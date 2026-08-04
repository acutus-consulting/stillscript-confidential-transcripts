import os
import sys
import json
import logging
import threading
import time
import webbrowser
from pathlib import Path
import customtkinter as ctk
from PIL import Image, ImageTk
import whisper
from tkinter import filedialog, messagebox
import anthropic

# Optional: use the OS keyring for secure API-key storage when available.
try:
    import keyring
    _HAS_KEYRING = True
except Exception:
    _HAS_KEYRING = False

# ─────────────────────────────────────────────
#  LEGACY USER-DATA MIGRATION (masterplan 4.1)
# ─────────────────────────────────────────────
# The product was renamed DanScribe -> StillScript in masterplan 4.1. Every
# on-disk location the old name owned is listed here, oldest name first, so
# an existing beta installation keeps its config, its API key fallback, its
# log, its transcripts, and — most importantly — its multi-GB downloaded
# models, instead of silently starting fresh.
#
# MOVE, not copy, deliberately:
#   * ~/.danscribe_models/ holds the Accurate model (5.75 GiB) and the
#     diarization model. Copying would need that much free space again, on
#     machines already measured at 81% full, and would leave two copies of a
#     6 GB tree that can silently diverge. os.replace()/rename on the same
#     filesystem is atomic and instant regardless of size.
#   * For the small files, a copy would leave the old and new config able to
#     drift apart, with no rule about which wins. One-way move keeps exactly
#     one source of truth.
# Nothing is ever deleted: if a new-name target already exists, the old path
# is left untouched rather than overwritten, so a partial or repeated
# migration can never destroy newer data.
#
# Runs at import, BEFORE logging is configured, because the log file itself
# is one of the things being moved — so it cannot use `logger`. Failures are
# swallowed per-item: a migration problem must never stop the app from
# starting, it just means that one item stays where it was.
_LEGACY_PATH_MIGRATIONS = [
    (Path.home() / ".danscribe_config.json", Path.home() / ".stillscript_config.json"),
    (Path.home() / ".danscribe.log", Path.home() / ".stillscript.log"),
    (Path.home() / ".danscribe_models", Path.home() / ".stillscript_models"),
    (Path.home() / "Documents" / "DanScribe_Transcriptions",
     Path.home() / "Documents" / "StillScript_Transcriptions"),
]


def migrate_legacy_user_data(migrations=None):
    """Move any surviving DanScribe-named user data to its StillScript name.

    Returns a list of (old, new) pairs actually migrated — empty on a fresh
    install, which is the normal case and explicitly not an error.

    Skips (leaving the old path alone) when the old path doesn't exist, or
    when the new path already exists. Never overwrites, never deletes.
    """
    migrated = []
    for old_path, new_path in (migrations if migrations is not None
                               else _LEGACY_PATH_MIGRATIONS):
        try:
            if not old_path.exists() or new_path.exists():
                continue
            new_path.parent.mkdir(parents=True, exist_ok=True)
            os.rename(old_path, new_path)
            migrated.append((old_path, new_path))
        except Exception:
            # Deliberately silent-but-safe: no logger yet (the log file is
            # itself mid-migration), and a failure here must not block
            # startup. The old data is still where it was.
            pass
    return migrated


_MIGRATED_LEGACY_PATHS = migrate_legacy_user_data()

# ─────────────────────────────────────────────
#  LOGGING
# ─────────────────────────────────────────────
# Log to a file in the user's home dir. A windowed .exe has no console, so
# print() output is invisible — logging is the only way to diagnose issues.
LOG_PATH = os.path.join(Path.home(), ".stillscript.log")
logging.basicConfig(
    filename=LOG_PATH,
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
)
logger = logging.getLogger("stillscript")

if _MIGRATED_LEGACY_PATHS:
    # Now that logging exists, record what the pre-logging migration did.
    for _old, _new in _MIGRATED_LEGACY_PATHS:
        logger.info("Migrated legacy user data: %s -> %s", _old, _new)

# ─────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────

def resource_path(relative_path):
    """Get path to files inside the EXE."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def _ensure_bundled_binaries_on_path():
    """Make ffmpeg/ffprobe bundled inside the frozen exe discoverable.

    The Windows build ships ffmpeg.exe and ffprobe.exe via PyInstaller
    --add-binary "...;.", which unpacks them into sys._MEIPASS (the onefile
    bundle root). Every audio call in this app shells out to a *bare*
    "ffmpeg"/"ffprobe" name — Whisper's own decoder (whisper/audio.py) as
    well as our diarization helpers (_probe_audio_duration_seconds,
    _diarize). Those resolve via PATH, so we prepend the bundle dir here.

    Prepend (not append) so the bundled copy wins over any stale system
    ffmpeg. No-op when running from source (sys._MEIPASS is absent), so the
    dev/Linux PATH is left exactly as-is.
    """
    base_path = getattr(sys, "_MEIPASS", None)
    if not base_path:
        return
    parts = os.environ.get("PATH", "").split(os.pathsep)
    if base_path not in parts:
        os.environ["PATH"] = base_path + os.pathsep + os.environ.get("PATH", "")


# Run at import so ffmpeg is resolvable before any transcription starts,
# regardless of entry point.
_ensure_bundled_binaries_on_path()

CONFIG_PATH = os.path.join(Path.home(), ".stillscript_config.json")
_KEYRING_SERVICE = "StillScript"
# Masterplan 4.1: the keyring entry is keyed by service name, so renaming the
# service would orphan an existing user's stored API key just as surely as
# moving a file would — it is user data, not a label. _get_api_key() below
# falls back to reading the old service name and re-homes the key under the
# new one on first use.
_LEGACY_KEYRING_SERVICE = "DanScribe"
_KEYRING_USER = "claude_api_key"

# Claude model used for AI summaries. Kept as a single constant so it can be
# updated in one place when Anthropic retires a model snapshot.
CLAUDE_MODEL = "claude-opus-4-8"


def _get_api_key(config):
    """Read the Claude API key from the OS keyring, falling back to the config file."""
    if _HAS_KEYRING:
        try:
            key = keyring.get_password(_KEYRING_SERVICE, _KEYRING_USER)
            if key:
                return key
        except Exception as e:
            logger.warning("Keyring read failed, falling back to config file: %s", e)
        # Masterplan 4.1 — nothing under the new service name; an existing
        # beta user's key is still filed under the old one. Re-home it so
        # this lookup only happens once, but return it either way: a failed
        # re-home must not cost the user their key.
        try:
            legacy_key = keyring.get_password(_LEGACY_KEYRING_SERVICE, _KEYRING_USER)
            if legacy_key:
                try:
                    keyring.set_password(_KEYRING_SERVICE, _KEYRING_USER, legacy_key)
                    keyring.delete_password(_LEGACY_KEYRING_SERVICE, _KEYRING_USER)
                    logger.info("Migrated the stored API key to the StillScript "
                                "keyring entry.")
                except Exception as e:
                    logger.warning("Could not re-home the legacy keyring entry "
                                   "(the key still works): %s", e)
                return legacy_key
        except Exception as e:
            logger.warning("Legacy keyring read failed: %s", e)
    return config.get("api_key", "")


def _set_api_key(config, api_key):
    """Store the API key securely. Prefer the OS keyring; never leave it in the JSON file."""
    config.pop("api_key", None)  # never persist the key in plaintext JSON
    if _HAS_KEYRING:
        try:
            if api_key:
                keyring.set_password(_KEYRING_SERVICE, _KEYRING_USER, api_key)
            else:
                try:
                    keyring.delete_password(_KEYRING_SERVICE, _KEYRING_USER)
                except Exception:
                    pass
            return
        except Exception as e:
            logger.warning("Keyring write failed, storing key in config file: %s", e)
    # Fallback only when no keyring backend is available.
    config["api_key"] = api_key


def load_config():
    # "Consistent" default (masterplan 1.4/2.9): fits the audit/provenance
    # promise Accurate mode is built around — a legal/clinical user should be
    # able to re-run a recording and get the same transcript back. Masterplan
    # 3.2 extended this same key to Fast mode too, so both modes share one
    # Settings choice and one persisted value rather than two parallel ones.
    defaults = {"api_key": "", "last_language": "Auto-Detect", "reproducibility": "Consistent"}
    if not os.path.exists(CONFIG_PATH):
        return defaults
    try:
        with open(CONFIG_PATH, "r") as f:
            config = json.load(f)
        if not isinstance(config, dict):
            raise ValueError("config is not a JSON object")
    except (OSError, ValueError, json.JSONDecodeError) as e:
        logger.error("Could not read config (%s); using defaults.", e)
        return defaults

    # Merge defaults so a partial/old config never has missing keys.
    merged = {**defaults, **config}

    # Validate persisted choices against known values. (A pre-Phase-2
    # "last_model" key may still be in the file; it is simply ignored now
    # that mode selection replaces the model dropdown.)
    if merged.get("last_language") not in LANG_CODES:
        merged["last_language"] = defaults["last_language"]
    if merged.get("reproducibility") not in ("Consistent", "Best effort"):
        merged["reproducibility"] = defaults["reproducibility"]
    return merged

def save_config(data):
    try:
        with open(CONFIG_PATH, "w") as f:
            json.dump(data, f)
        # Restrict to owner read/write so the config isn't world-readable.
        try:
            os.chmod(CONFIG_PATH, 0o600)
        except OSError:
            pass  # e.g. on filesystems that don't support chmod (some Windows setups)
    except OSError as e:
        logger.error("Could not save config: %s", e)

# ─────────────────────────────────────────────
#  CONFIGURATION DATA
# ─────────────────────────────────────────────

LANG_CODES = {
    "Auto-Detect": None, "Afrikaans": "af", "English": "en", "Sesotho": "st",
    "isiZulu": "zu", "isiXhosa": "xh", "Setswana": "tn", "Sepedi": "nso",
    "Siswati": "ss", "Tshivenda": "ve", "Xitsonga": "ts", "isiNdebele": "nr"
}

# ── Transcription modes ──────────────────────
# "Vinnig" (Fast) is the only mode with a working engine. It runs Whisper
# Medium through the transcribe_audio() seam — the path verified byte-for-byte
# in Phase 1. "Akkuraat" (Accurate) is presented in the UI but disabled until
# Phase 3 wires the large-v3 + Afrikaans-adapter engine behind it; nothing in
# this build must ever let Accurate silently fall back to Medium.
FAST_MODE_MODEL = "medium"

# Human-readable engine label recorded in each transcript's provenance footer.
# Phase 3 adds the accurate-mode engine (large-v3 + adapter revision SHA).
FAST_MODE_ENGINE_LABEL = "StillScript Fast — Whisper Medium"

# Masterplan 3.2 — Fast mode's reproducibility axis. openai-whisper's own
# transcribe() already runs a temperature-fallback ladder by default (retrying
# low-confidence segments at progressively higher, sampling temperatures) —
# verified (not assumed) to be the exact tuple below, openai-whisper 20250625's
# own default, identical in shape to accurate_engine.TEMPERATURE_BEST_EFFORT.
# transcribe_audio() never overrode it before this item, so Fast mode has
# silently run in "Best effort" mode on every transcription, with no way to
# pin it to a single, deterministic temperature=0.0 ("Consistent") — exactly
# the same axis Golf 2.8/2.9 investigated for Accurate mode, just unexposed
# here. Defined locally rather than imported from accurate_engine: Fast
# mode's runtime must never depend on that module (see
# transcribe_audio_accurate()'s docstring below) even though the values
# happen to be identical.
FAST_TEMPERATURE_CONSISTENT = 0.0
FAST_TEMPERATURE_BEST_EFFORT = (0.0, 0.2, 0.4, 0.6, 0.8, 1.0)


def _describe_fast_reproducibility(temperature):
    """Human label for the temperature value transcribe_audio() actually
    decoded with (masterplan 3.2) — mirrors
    accurate_engine._describe_reproducibility() exactly, including its
    "derive from the real value actually used, never re-derive from
    Settings" principle (masterplan 2.6/2.9): if a future caller ever passes
    some other temperature, provenance must say so honestly rather than
    mislabelling it as one of the two named Settings choices it isn't.
    """
    if temperature == FAST_TEMPERATURE_CONSISTENT:
        return "Consistent"
    if temperature == FAST_TEMPERATURE_BEST_EFFORT:
        return "Best effort"
    return f"Custom (temperature={temperature!r})"


# Masterplan 2.11 — mode-scope notices shown next to the Mode buttons (section
# 3 of the main window), swapped by _select_fast_mode()/_select_accurate_mode()
# so whichever mode is active always has its own scope reminder visible. Not a
# gate: both modes stay fully usable regardless of what's selected in the
# Audio Language menu above — this is disclosure, not enforcement, matching
# how AccurateConsentDialog already states limitations without blocking
# anything a user chooses to do.
#
# Evidence base (masterplan 2.8, all real, not assumed): a 3:20 noisy-bar
# clip (reclassified as an extreme case) and a 46-min mixed Afrikaans/English
# committee recording both real Accurate-mode runs. Fast mode has no engine
# of its own to re-test here — its Afrikaans/Dutch-drift issue was already
# known before 2.8 — but 2.8's finding that this product does not validate
# ANY mode against non-Afrikaans audio applies to Fast mode just as much.
FAST_MODE_SCOPE_NOTE = (
    "⚡ Fast mode is tuned and tested for Afrikaans. It's known to drift "
    "toward Dutch spelling occasionally, even on Afrikaans audio. For any "
    "other language, StillScript has not validated Fast mode's accuracy — "
    "treat those results as unverified."
)

# Short version for the main window, next to the Mode buttons — always
# visible once Accurate is selected, unlike AccurateConsentDialog's fuller
# explanation, which a user only sees once, on first activation.
ACCURATE_MODE_SCOPE_NOTE = (
    "🎯 Accurate mode always transcribes as Afrikaans, regardless of the "
    "Audio Language setting above. If your recording mixes in English or "
    "another language, expect it to run slower and carry a meaningfully "
    "higher chance of errors that need a full read-through, not a light one."
)

# Masterplan 1.4 / 2.9 — Settings copy for the reproducibility choice.
# Deliberately does NOT frame "Consistent" as simply the safer option: a real
# A/B test on one genuinely difficult recording (masterplan 2.8, ACCURATE
# MODE ONLY) found temperature=0 ("Consistent") produced a LARGER and more
# degenerate repetition-collapse than the fallback ladder ("Best effort") on
# that same clip, even though pre-collapse quality was slightly better and
# timing matched the ~3x baseline exactly. One data point, not a settled
# rule — but "Consistent sounds safer" is not what was actually measured, so
# the copy below says so plainly rather than defaulting to the
# reassuring-sounding claim. Written to stand alone for Wave 5.2 (User
# Manual), not just as an in-app aside.
#
# Masterplan 3.2 extended this same setting to Fast mode, which shares the
# identical underlying temperature-fallback mechanism (both modes ultimately
# call openai-whisper's transcribe() with the same two temperature values —
# verified, not assumed). The copy below is honest that the specific
# measurements quoted (the ~3x timings, the repetition-collapse severity
# finding) come from Golf 2.8's Accurate-mode-only test and have not been
# separately measured for Fast mode — do not silently extend those numbers to
# Fast mode without a real measurement, the same discipline this copy already
# applies to Accurate mode's own claims.
REPRODUCIBILITY_EXPLANATION = (
    "This setting applies to both Fast and Accurate mode. Either mode can "
    "decode audio two ways. This changes speed and exact repeatability — "
    "and, on genuinely difficult audio, it may also change how a rare "
    "failure shows up if the model struggles. Neither option is simply "
    "\"safer\"; read both before choosing.\n\n"
    "Consistent (default): the same recording always produces the exact "
    "same transcript, with no randomness. In Accurate mode this runs at the "
    "documented ~3x-real-time pace; on one real difficult Accurate-mode test "
    "recording, though, Consistent produced a MORE severe version of a rare "
    "failure (parts of the transcript degenerating into repeated words) "
    "than Best effort did on the same audio — so treat \"Consistent\" as "
    "meaning reproducible, not automatically safer on hard audio.\n\n"
    "Best effort: retries unclear parts of the audio before settling on an "
    "answer, which can occasionally give a slightly different transcript if "
    "you run the same recording twice, and (in Accurate mode) runs "
    "substantially slower (up to ~3x slower than Consistent was measured on "
    "one difficult clip). In that same test, its retries partly avoided the "
    "worst form of the failure above — but this is one data point, not a "
    "guarantee it will do so on other recordings, and Fast mode's own speed "
    "and failure behaviour under each option have not been separately "
    "measured.\n\n"
    "For most recordings — clear audio, one main speaker or language — "
    "either option works well in either mode. For long, difficult, or "
    "heavily overlapping/noisy recordings, consider trying both and "
    "comparing, especially if a transcript looks like it starts repeating "
    "itself. Always review the transcript against the audio for anything "
    "that matters."
)

# Masterplan 2.4 — real multipliers measured across golf 2.8/2.10's Accurate-
# mode test runs (transcription time only; diarization is separate, below).
# Kept as a RANGE per Reproducibility setting, not a single number: golf 2.8
# found a real ~3x spread depending on how difficult the audio actually is,
# and one "~3x" figure would misrepresent that spread as precision it doesn't
# have.
#
#   Consistent:  2.98x (clean, 2.10)     .. 3.08x (extreme noise, 2.8)
#   Best effort: 3.06x (clean, 2.8/2.10) .. 9.46x (extreme noise, 2.8)
#                                            (mixed-language, 2.8: 7.175x,
#                                             already inside this range)
#
# Consistent's range is narrower because every difficult-audio run measured
# under Consistent has landed close to 3x — Best effort's own fallback
# retries are what cause its much wider spread on hard audio. Consistent has
# NOT been tested yet on heavily multi-language audio specifically (golf
# 2.10's own remaining optional follow-up) — ACCURATE_TIME_ESTIMATE_EXPLANATION
# below says so explicitly rather than implying that narrow range is proven
# for that case too.
ACCURATE_TRANSCRIPTION_MULTIPLIER_RANGE = {
    "Consistent": (2.98, 3.08),
    "Best effort": (3.06, 9.46),
}

# Diarization (golf 2.12's pyannote backend) runs as a separate step AFTER
# transcription and adds roughly its own length again, regardless of the
# Reproducibility setting above — it never reads that setting at all (it has
# no temperature of its own to vary). Measured 0.95x-1.21x additional across
# golf 2.10's real run (1.13x) and golf 2.12's own spike benchmarks
# (0.948x-1.21x) on real audio.
ACCURATE_DIARIZATION_MULTIPLIER_RANGE = (0.95, 1.21)

# Shown in AccurateTimeEstimateDialog (masterplan 2.4), alongside the
# per-file numeric range _estimate_accurate_time_range() computes. A distinct
# step from AccurateConsentDialog (2.3, one-time model-download consent) and
# the persistent scope notices (2.9/2.11, mode/language scope) — this is
# about how long THIS specific file will take, shown on every Accurate-mode
# run, not just the first activation.
ACCURATE_TIME_ESTIMATE_EXPLANATION = (
    "Accurate mode is intentionally a \"start it and check back later\" "
    "step, not something you need to watch. Typically about 3x the "
    "recording's length with Consistent — real tests have stayed close to "
    "that even on very difficult audio. Best effort is usually similar, but "
    "can run significantly longer — up to roughly 9x or more — on very "
    "noisy, overlapping-speaker, or heavily multi-language audio, because it "
    "retries unclear passages before settling on an answer. Consistent "
    "hasn't been tested yet on heavily multi-language audio specifically, "
    "so treat that combination as a possible exception. Speaker "
    "identification, if enabled, adds roughly the recording's length again "
    "on top of this."
)


def _estimate_accurate_time_range(duration_seconds, reproducibility, diarize):
    """Real-multiplier time estimate for masterplan 2.4 — a range, not a
    single number (see ACCURATE_TRANSCRIPTION_MULTIPLIER_RANGE's own
    reasoning above for why false precision would misrepresent golf 2.8's
    own measured spread).

    Returns (low_seconds, high_seconds), or None if duration_seconds is None
    — _probe_audio_duration_seconds() already returns None for "couldn't
    determine the length" rather than raising, so this mirrors that and
    lets callers show a duration-less version of the estimate instead of
    crashing on unreadable audio metadata.

    An unrecognised `reproducibility` value (should never happen — load_config()
    already validates it against the same two choices SettingsWindow offers)
    falls back to "Consistent"'s range rather than raising, the same
    fail-safe convention load_config() itself uses for this exact key.
    """
    if duration_seconds is None:
        return None
    t_low, t_high = ACCURATE_TRANSCRIPTION_MULTIPLIER_RANGE.get(
        reproducibility, ACCURATE_TRANSCRIPTION_MULTIPLIER_RANGE["Consistent"]
    )
    low = duration_seconds * t_low
    high = duration_seconds * t_high
    if diarize:
        d_low, d_high = ACCURATE_DIARIZATION_MULTIPLIER_RANGE
        low += duration_seconds * d_low
        high += duration_seconds * d_high
    return (low, high)


# Third-party attribution shown in the "About" (Credits) dialog. Masterplan
# 2.7 appends the fine-tuned Afrikaans model + dataset entries (CC-BY-4.0)
# below — same list, same CreditsWindow, no UI change.
CREDITS = [
    {
        "name": "OpenAI Whisper",
        "detail": ('Radford et al., "Robust Speech Recognition via '
                   'Large-Scale Weak Supervision", arXiv:2212.04356'),
        "license": "Apache License 2.0",
        "url": "https://arxiv.org/abs/2212.04356",
    },
    # Masterplan 2.7. Attribution language and BibTeX taken verbatim from the
    # published HF model card (huggingface.co/DanieClar/stillscript-whisper-
    # large-v3-afrikaans, License/Attribution/Citation section) rather than
    # paraphrased, so this entry and the model card cannot drift apart. The
    # Afrikaans adapter is CC-BY-4.0; Accurate mode's merged model is only
    # the base model's weights (openai/whisper-large-v3, Apache-2.0, credited
    # above via the Whisper entry) plus this adapter — "nothing was retrained;
    # the only operation performed was merging the published adapter into the
    # published base weights", per the model card's own key statement. This
    # entry covers both the required creator/title/source/license attribution
    # and the "indicate if changes were made" requirement CC-BY-4.0 also asks
    # for, and applies regardless of which HF revision/layout (single-file vs
    # chunked) a given install downloaded, since both carry the same merge.
    {
        "name": "Whisper Large V3 Afrikaans (adapter)",
        "detail": (
            'André Oosthuizen, "Whisper Large V3 Afrikaans", 2025, HuggingFace. '
            "Merged into StillScript's Accurate-mode model without retraining — "
            "the only operation performed was merging this published adapter "
            "into the published openai/whisper-large-v3 base weights."
        ),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/andreoosthuizen/whisper-large-v3-afrikaans",
    },
    {
        "name": "afrikaans-30s (training dataset)",
        "detail": ("André Oosthuizen — training data used for the Whisper "
                   "Large V3 Afrikaans adapter above."),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/datasets/andreoosthuizen/afrikaans-30s",
    },
    # Masterplan 2.12. Same CC-BY-4.0 attribution shape as the two entries
    # above: title (name), author + modification notice (detail), source (url),
    # licence (license). Wording and the cited work are taken from the upstream
    # model card rather than paraphrased. The pipeline is redistributed
    # unmodified from StillScript's own mirror, because the upstream repository
    # is gated and a shipped desktop app cannot make each end user accept its
    # conditions; CC-BY-4.0 expressly permits that redistribution.
    {
        "name": "pyannote speaker-diarization-community-1",
        "detail": (
            "Hervé Bredin and the pyannote.audio team — the speaker-"
            "identification pipeline behind \"Identify different speakers\". "
            "Redistributed unmodified from StillScript's mirror; no "
            "retraining, fine-tuning or conversion was performed. Cites "
            "Plaquet & Bredin, \"Powerset multi-class cross entropy loss for "
            "neural speaker diarization\", Interspeech 2023, and Bredin, "
            "\"pyannote.audio 2.1 speaker diarization pipeline\", Interspeech "
            "2023."
        ),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/pyannote/speaker-diarization-community-1",
    },
]

# Model cache — prevents reloading every time
_model_cache = {}

def get_model(model_name):
    if model_name not in _model_cache:
        _model_cache[model_name] = whisper.load_model(model_name)
    return _model_cache[model_name]


def release_fast_mode_model():
    """Evict Fast mode's cached Whisper model(s) from memory (masterplan 2.5).

    Called before Accurate mode loads large-v3. Medium alone sits at roughly
    1.5 GB resident; large-v3 peaks at ~8.7 GB while generating. Having both
    resident at once is exactly the scenario that risks OOM on ordinary
    hardware, and it is avoidable — Fast and Accurate are never used in the
    same instant, only the same session.

    Clearing the dict (not just deleting individual entries) is what actually
    matters here: _model_cache is the ONLY thing holding a reference to the
    loaded `whisper.Whisper` object, so once every entry is gone CPython's
    refcounting frees the underlying C-allocated tensors immediately — no
    lingering reference elsewhere in this module keeps it alive (get_model()
    returns the object to its caller, but callers do not stash it; they call
    get_model() again next time). The explicit gc.collect() following it
    exists for the same reason PyTorch's own docs recommend it after freeing
    large tensors: reference cycles (e.g. autograd graph fragments, even
    though this app never calls backward()) can otherwise sit uncollected for
    a while under CPython's generational GC, and this is exactly the moment a
    plausibly multi-GB allocation is about to follow it.

    A no-op, not an error, if nothing is loaded yet — Accurate mode used first
    in a fresh session has nothing to release. Switching back to Fast mode
    afterward is unaffected by this: get_model() simply reloads from disk on
    its next call, the same cache-miss path it already always had.
    """
    if not _model_cache:
        return
    released = list(_model_cache.keys())
    _model_cache.clear()
    import gc
    gc.collect()
    logger.info("Released Fast-mode model(s) from memory before loading "
               "Accurate mode: %s", released)


def transcribe_audio(path, *, language, task, model_name, temperature=FAST_TEMPERATURE_BEST_EFFORT):
    """Run Whisper on `path` and return its result dict ({"text", "segments", ...}).

    This is the single seam through which all transcription flows. The three
    branches below are the exact language-dispatch logic (Afrikaans prompt,
    other forced language, auto-detect) that previously lived inline in run() —
    this call-site invariance holds regardless of which branch executes, or of
    where in the app this function is called from.

    `temperature` (masterplan 3.2) mirrors transcribe_audio_accurate()'s own
    parameter and the same division of responsibility: the caller
    (start_transcription()) reads the user's Settings choice and translates it
    to FAST_TEMPERATURE_CONSISTENT/FAST_TEMPERATURE_BEST_EFFORT before calling
    here — this function never reads Settings itself. Defaults to
    FAST_TEMPERATURE_BEST_EFFORT, openai-whisper's own default, so a caller
    that omits it (as every caller did before this item) sees unchanged
    behaviour. Only FAST_TEMPERATURE_CONSISTENT (plain 0.0, no sampling) is
    actually deterministic; the default fallback ladder can retry a
    low-confidence segment at a higher, sampling temperature, so — unlike the
    old docstring here claimed — output is not guaranteed byte-identical
    across runs unless "Consistent" is selected. The returned dict's
    "reproducibility" key reports the label for the value ACTUALLY used to
    decode (via _describe_fast_reproducibility()), not a re-derivation — the
    same source-of-truth principle accurate_engine.transcribe() uses.
    """
    model = get_model(model_name)

    # Whisper's initial_prompt biases the decoder toward the style/
    # spelling of the prompt text itself, not toward instructions it
    # "understands" — a short meta-sentence gives it little to latch
    # onto. Afrikaans and Dutch share a lot of vocabulary, so smaller
    # models especially can drift into Dutch orthography even with
    # language="af" forced. A longer, natural Afrikaans passage with
    # constructions Dutch doesn't have (like "nie ... nie" double
    # negation) gives the decoder a much stronger signal.
    af_prompt = (
        "Goeiemôre almal, baie dankie dat julle almal hier kon kom. "
        "Ons gaan vandag praat oor die begroting en die volgende "
        "stappe wat ons moet neem. Dit is nie maklik nie, maar ons "
        "sal dit saam reg kry. Het julle dalk enige vrae daaroor?"
    )
    if language == "af":
        # Tried condition_on_previous_text=False here to stop Dutch
        # drift from compounding across windows — real-world testing
        # showed it made things worse (hallucinated Unicode noise,
        # injected English words, repetition loops), since the
        # previous-window context also suppresses that kind of
        # hallucination. Reverted; back to Whisper's default.
        result = model.transcribe(path, task=task, language="af", initial_prompt=af_prompt,
                                   temperature=temperature)
    elif language:
        result = model.transcribe(path, task=task, language=language, temperature=temperature)
    else:
        result = model.transcribe(path, task=task, temperature=temperature)
    result["reproducibility"] = _describe_fast_reproducibility(temperature)
    return result


def transcribe_audio_accurate(path, *, language, task, model_dir=None, **engine_kwargs):
    """Accurate-mode counterpart to transcribe_audio() — the second entry point.

    Runs Whisper large-v3 + the Afrikaans adapter through a direct
    transformers `generate()` call (see accurate_engine.py; the CT2 /
    faster-whisper path for that model is known broken and must never be
    used). Returns the same {"text", "segments", ...} shape as the Fast seam,
    so diarization and the exporters are unaffected by which engine ran.

    Two things this deliberately does NOT do:
      * it is not wired into run() — activating the Accurate button is
        masterplan item 2.3;
      * it never falls back to Fast/Medium. If the engine is unavailable the
        AccurateEngineUnavailable error propagates, because silently
        substituting a weaker model is exactly the failure this product
        cannot ship.

    accurate_engine is imported lazily so that its heavy dependencies
    (torch/transformers, absent from the Fast-mode runtime) can never break
    startup or the Fast path.

    release_fast_mode_model() runs first (masterplan 2.5) — before large-v3's
    from_pretrained() call, not before this whole function, so a download that
    accurate_engine.load_engine() needs to wait on never costs Fast mode its
    warm cache for nothing. If this raises (e.g. loading large-v3 hits real
    memory pressure), Fast mode's cache is already empty either way; the next
    Fast-mode run just reloads Medium from disk, same as first launch.
    """
    release_fast_mode_model()
    import accurate_engine
    return accurate_engine.transcribe(
        path, language=language, task=task, model_dir=model_dir, **engine_kwargs
    )


def build_provenance(*, mode, language_label, task, diarized, num_speakers=None, engine=None,
                     model_id=None, model_revision=None, model_layout=None,
                     guard_verification=None, reproducibility=None):
    """Describe which engine produced a transcript, for the audit footer.

    Returns a plain dict so callers can extend it without any other caller
    changing — the same seam principle as transcribe_audio(). Callers pass
    display-level values (the language *label* the user picked, not the ISO
    code) so the footer reads the way the operator set it up.

    `engine` defaults to FAST_MODE_ENGINE_LABEL so the existing Fast-mode call
    site is unaffected. Accurate mode passes its own engine's label explicitly
    — accurate_engine.transcribe()'s result dict already carries it as
    result["engine"] (ACCURATE_ENGINE_LABEL), so the caller never needs to
    import accurate_engine just to get this string. Getting this wrong would
    mean an Accurate-mode transcript's audit footer claiming it was produced by
    Fast mode instead — a provenance product exists specifically to prevent
    that kind of quiet misattribution.

    model_id / model_revision / model_layout / guard_verification (masterplan
    2.6) are Accurate-mode-only and additive: they default to None and are
    only included in the returned dict when a caller actually passes them, so
    the Fast-mode call site (which never passes them) gets a provenance dict
    identical to before this wave — and format_provenance_lines() below only
    ever renders lines for the ones that are present.

    reproducibility (masterplan 2.9, extended to Fast mode by 3.2) follows the
    same additive convention as model_id/etc., but is NOT Accurate-mode-only:
    both call sites now pass it, sourced from their own engine's result dict
    (result.get("reproducibility")), never re-derived from Settings here.
    """
    from datetime import datetime
    provenance = {
        "engine": engine or FAST_MODE_ENGINE_LABEL,
        "mode": mode,
        "language": language_label,
        "task": task,
        "diarized": diarized,
        "num_speakers": num_speakers if diarized else None,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    if model_id is not None:
        provenance["model_id"] = model_id
    if model_revision is not None:
        provenance["model_revision"] = model_revision
    if model_layout is not None:
        provenance["model_layout"] = model_layout
    if guard_verification is not None:
        provenance["guard_verification"] = guard_verification
    if reproducibility is not None:
        provenance["reproducibility"] = reproducibility
    return provenance


def format_provenance_lines(provenance):
    """Render a provenance dict as an ordered list of 'Label: value' strings,
    shared by the .txt footer and the .docx footer so both stay in sync."""
    task_label = "Translate to English" if provenance.get("task") == "translate" else "Original language"
    if provenance.get("diarized"):
        spk = provenance.get("num_speakers")
        diar_label = f"Yes (up to {spk} speakers)" if spk else "Yes"
    else:
        diar_label = "No"
    lines = [
        f"Engine: {provenance.get('engine')}",
        f"Mode: {provenance.get('mode')}",
        f"Language setting: {provenance.get('language')}",
        f"Task: {task_label}",
        f"Speaker identification: {diar_label}",
    ]
    # Optional lines, only appended when present. model_id/model_revision/
    # model_layout/guard_verification (masterplan 2.6) remain Accurate-mode-
    # only — build_provenance()'s Fast-mode call site never passes them, so
    # they are simply absent here, unchanged since before this wave.
    # Reproducibility (masterplan 2.9, extended to Fast mode by 3.2 — which
    # decode strategy was actually used) is no longer Accurate-only, but is
    # still listed first among these, ahead of which model weights backed the
    # run, since it describes decode-time behaviour rather than which
    # model/weights were loaded.
    if provenance.get("reproducibility"):
        lines.append(f"Reproducibility: {provenance['reproducibility']}")
    if provenance.get("model_id"):
        lines.append(f"Model: {provenance['model_id']}")
    if provenance.get("model_revision"):
        lines.append(f"Model revision: {provenance['model_revision']}")
    if provenance.get("model_layout"):
        lines.append(f"Download layout: {provenance['model_layout']}")
    if provenance.get("guard_verification"):
        lines.append(f"Guard verification: {provenance['guard_verification']}")
    lines.append(f"Generated: {provenance.get('timestamp')}")
    return lines


# Below this duration, _diarize() uses the whole-file librosa.load path,
# which keeps speaker labels byte-identical to previous behaviour — the
# common case, where label accuracy matters most and memory was never a
# problem (measured delta ~1.1 GB at this length). At/above it, memory
# actually bites on long recordings (multi-GB peaks observed on hour-plus
# files), so _diarize() switches to the decode-once-to-temp-WAV path, which
# can shift a speaker label by a small margin — acceptable on long
# recordings where individual turn boundaries aren't scrutinized.
# ⚠ UNUSED as of masterplan 2.12 — kept, not deleted, deliberately.
#
# This gate had NO callers as of masterplan 2.12. It existed to bound memory
# while extracting one feature vector per Whisper segment: short files loaded
# whole via librosa.load, long ones decoded once to a temp WAV and seeked per
# segment. The pyannote backend does no per-segment feature extraction at
# all, so there is nothing left to gate — the branch was not "removed as a
# simplification", it lost its subject.
#
# Removing this constant, and the frozen-app PATH note at the top of this file
# that mentions it, is a tidy-up that was explicitly scoped OUT of 2.12
# (backend swap only). Do not assume the gate is live: it is not.
#
# _probe_audio_duration_seconds() below is a DIFFERENT story: it had no
# caller either, from 2.12 until masterplan 2.4 gave it its first real one
# (the pre-transcription time estimate) — same function, genuinely revived,
# not a coincidental name match.
DIARIZE_LONG_FILE_THRESHOLD_SEC = 20 * 60  # 20 minutes


def _probe_audio_duration_seconds(path):
    """Return audio duration in seconds via ffprobe container metadata only
    (no decode), so the check itself stays cheap even on very long files.
    Returns None if the duration can't be determined; callers should treat
    that as "unknown" and fall back to the safe default path.
    """
    import subprocess
    try:
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", path],
            capture_output=True, text=True, timeout=10,
        )
        return float(out.stdout.strip())
    except Exception:
        return None

# ─────────────────────────────────────────────
#  SETTINGS WINDOW
# ─────────────────────────────────────────────

class SettingsWindow(ctk.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("StillScript — Settings")
        # Grown from the original 520x320 (masterplan 2.9) to fit the
        # reproducibility explanation below without clipping it — this window
        # is not scrollable and not user-resizable, so the geometry must
        # genuinely be tall enough, not just plausible-looking. Measured (not
        # guessed) via winfo_reqheight() under Xvfb: real content needed
        # ~644px as of 2.9; masterplan 3.2 lengthened the explanation to cover
        # both modes, re-measured at ~670px. 700 still leaves real, checked
        # headroom (re-verify with winfo_reqheight() again if this copy grows
        # further) rather than an arbitrary round number.
        self.geometry("650x875")
        self.resizable(False, False)
        self.grab_set()

        config = load_config()

        ctk.CTkLabel(self, text="⚙️ Settings", font=("Arial", 25, "bold")).pack(pady=20)

        ctk.CTkLabel(self, text="Claude API Key:", font=("Arial", 17)).pack(pady=(10, 2))
        self.api_entry = ctk.CTkEntry(self, width=505, show="•", placeholder_text="sk-ant-...")
        self.api_entry.pack(pady=5)
        existing_key = _get_api_key(config)
        if existing_key:
            self.api_entry.insert(0, existing_key)

        # Clickable link
        link_label = ctk.CTkLabel(
            self,
            text="🔗 Get your API key at console.anthropic.com",
            font=("Arial", 14),
            text_color="#4da6ff",
            cursor="hand2"
        )
        link_label.pack(pady=4)
        link_label.bind("<Button-1>", lambda e: webbrowser.open("https://console.anthropic.com"))

        ctk.CTkLabel(
            self,
            text="Your key is stored locally on your device only.",
            font=("Arial", 13),
            text_color="gray"
        ).pack(pady=2)

        # Masterplan 1.4 / 2.9 — reproducibility choice. Originally
        # Accurate-mode-only; masterplan 3.2 wired Fast mode to this same key,
        # so one Settings choice now governs both modes.
        ctk.CTkLabel(
            self, text="🔁 Reproducibility (Fast and Accurate mode):", font=("Arial", 17, "bold")
        ).pack(pady=(20, 2))
        self.repro_var = ctk.StringVar(value=config.get("reproducibility", "Consistent"))
        ctk.CTkOptionMenu(
            self, values=["Consistent", "Best effort"],
            variable=self.repro_var, width=505,
        ).pack(pady=5)
        ctk.CTkLabel(
            self, text=REPRODUCIBILITY_EXPLANATION, font=("Arial", 13),
            text_color="gray70", justify="left", wraplength=550,
        ).pack(pady=(4, 10), padx=20)

        ctk.CTkButton(
            self, text="Save Settings", command=self.save,
            width=240, fg_color="#1f538d"
        ).pack(pady=20)

    def save(self):
        config = load_config()
        _set_api_key(config, self.api_entry.get().strip())
        config["reproducibility"] = self.repro_var.get()
        save_config(config)
        messagebox.showinfo("StillScript", "Settings saved!")
        self.destroy()

# ─────────────────────────────────────────────
#  CREDITS / ABOUT WINDOW
# ─────────────────────────────────────────────

class CreditsWindow(ctk.CTkToplevel):
    """Attribution surface for the models/datasets StillScript builds on.

    Driven entirely by the module-level CREDITS list — masterplan 2.7 added
    the fine-tuned Afrikaans model + dataset entries (CC-BY-4.0) by appending
    two dicts, with no change to this window.
    """
    def __init__(self, parent):
        super().__init__(parent)
        self.title("StillScript — About / Credits")
        # Grown from 560x420 to fit masterplan 4.2's logo lockup at the top
        # without clipping the Close button below — this window is not
        # scrollable/resizable, so the geometry must genuinely be tall
        # enough. Measured (not guessed) via winfo_reqheight() under Xvfb:
        # real content needs 566px; 600 leaves real, checked headroom.
        self.geometry("700x750")
        self.resizable(False, False)
        self.grab_set()

        # Masterplan 4.2 — the real mark, primary horizontal lockup, at the
        # brand guide's own "generous room" size (page 2). Same asset the
        # main window's banner uses, just a larger height, so the lockup's
        # own aspect ratio (never stretched) still holds.
        #
        # Masterplan 4.2 (2026-08-04 follow-up) — Danie reported this window rendering completely
        # empty (no logo, no text at all) on his real desktop, which never
        # reproduced under Xvfb. Investigated rather than guess-fixed: read
        # customtkinter 5.2.2's actual source (ctk_image.py, ctk_label.py).
        # CTkImage already caches its rendered PhotoImage objects in
        # self._scaled_light/dark_photo_images, and CTkLabel holds the
        # CTkImage as self._image — so the classic "PhotoImage GC'd because
        # nothing but a local var referenced it" bug does NOT appear to
        # apply to this code as written. `self._about_logo` below is kept
        # anyway, at zero cost, as defense-in-depth against that exact
        # failure mode (matches the same discipline _set_app_icon() already
        # uses for the raw ImageTk.PhotoImage case, where it truly matters).
        #
        # A more specific, code-confirmed candidate was found instead:
        # customtkinter's ScalingTracker.check_dpi_scaling() (scaling_tracker.py)
        # runs on a 100ms timer on Windows and, when it detects the process's
        # monitor DPI scaling has changed (e.g. the window moved to a
        # differently-scaled monitor, or Windows re-queries DPI after the
        # window settles), sets window.attributes("-alpha", 0.15) — making it
        # nearly invisible — while it rescales every widget, then restores
        # alpha to 1. If that rescale sequence is ever interrupted, the
        # window can be left stuck at alpha 0.15: visually indistinguishable
        # from "completely empty". This is a real Windows-HiDPI/multi-monitor
        # mechanism Xvfb (one virtual display, no real DPI variance) cannot
        # trigger or catch — consistent with Danie seeing it and Xvfb never
        # catching it. Not disabling customtkinter's own DPI system (too
        # blunt an intervention without being certain this is the cause);
        # instead a cheap, targeted safety net at the end of __init__ below
        # forces alpha back to 1 shortly after the window is built, in case
        # this exact sequence is what happened. This is the best evidence-
        # based fix possible without a real Windows machine to reproduce on
        # (see Golf 6.1) — NOT a confirmed root cause, and Danie's own
        # real-desktop confirmation is still needed; see the masterplan entry.
        try:
            img = Image.open(resource_path("stillscript_logo_horizontal_white.png"))
            logo_h = 110
            logo_w = round(logo_h * img.width / img.height)
            self._about_logo = ctk.CTkImage(light_image=img, dark_image=img, size=(logo_w, logo_h))
            ctk.CTkLabel(self, image=self._about_logo, text="").pack(pady=(20, 4))
        except Exception as e:
            logger.warning("About-window logo load failed (non-fatal): %s", e)

        ctk.CTkLabel(self, text="ℹ️ About StillScript", font=("Arial", 25, "bold")).pack(pady=(20, 4))
        ctk.CTkLabel(
            self,
            text="StillScript is built on the following open-source work:",
            font=("Arial", 16),
            text_color="gray",
        ).pack(pady=(0, 10))

        frame = ctk.CTkScrollableFrame(self, width=600, height=350)
        frame.pack(pady=5, padx=20, fill="both", expand=True)

        for entry in CREDITS:
            card = ctk.CTkFrame(frame)
            card.pack(fill="x", pady=8, padx=4)

            ctk.CTkLabel(
                card, text=entry["name"], font=("Arial", 18, "bold"), anchor="w"
            ).pack(fill="x", padx=12, pady=(10, 2))
            ctk.CTkLabel(
                card, text=entry["detail"], font=("Arial", 14), text_color="gray",
                anchor="w", justify="left", wraplength=550
            ).pack(fill="x", padx=12, pady=2)
            ctk.CTkLabel(
                card, text=f"License: {entry['license']}", font=("Arial", 14), anchor="w"
            ).pack(fill="x", padx=12, pady=2)

            url = entry.get("url")
            if url:
                link = ctk.CTkLabel(
                    card, text=f"🔗 {url}", font=("Arial", 14),
                    text_color="#4da6ff", cursor="hand2", anchor="w"
                )
                link.pack(fill="x", padx=12, pady=(2, 10))
                link.bind("<Button-1>", lambda e, u=url: webbrowser.open(u))

        ctk.CTkButton(
            self, text="Close", command=self.destroy,
            width=190, fg_color="#1f538d"
        ).pack(pady=(6, 16))

        # Masterplan 4.2 (2026-08-04 follow-up) safety net — see the long comment above the logo
        # block for why. Cheap and harmless if nothing was ever wrong.
        self.after(400, lambda: self.attributes("-alpha", 1.0))

# ─────────────────────────────────────────────
#  SPEAKER NAME ASSIGNMENT WINDOW
# ─────────────────────────────────────────────

class NameAssignWindow(ctk.CTkToplevel):
    def __init__(self, parent, speakers: list, callback):
        super().__init__(parent)
        self.title("Assign Speaker Names")
        self.geometry("565x525")
        self.resizable(False, False)
        self.grab_set()
        self.callback = callback
        self.entries = {}

        ctk.CTkLabel(self, text="🎤 Assign Names to Speakers", font=("Arial", 22, "bold")).pack(pady=20)
        ctk.CTkLabel(
            self,
            text="Leave blank to keep 'Speaker 1', 'Speaker 2', etc.",
            font=("Arial", 14),
            text_color="gray"
        ).pack(pady=(0, 15))

        frame = ctk.CTkScrollableFrame(self, width=455, height=265)
        frame.pack(pady=5, padx=20)

        for speaker in speakers:
            row = ctk.CTkFrame(frame, fg_color="transparent")
            row.pack(fill="x", pady=6)
            ctk.CTkLabel(row, text=f"{speaker}:", width=130, anchor="w", font=("Arial", 16, "bold")).pack(side="left", padx=5)
            entry = ctk.CTkEntry(row, width=275, placeholder_text="Enter name...")
            entry.pack(side="left")
            self.entries[speaker] = entry

        ctk.CTkButton(
            self, text="✅ Confirm & Continue",
            command=self.confirm, width=300, fg_color="#1f538d", height=55
        ).pack(pady=20)

    def confirm(self):
        name_map = {}
        for speaker, entry in self.entries.items():
            name = entry.get().strip()
            name_map[speaker] = name if name else speaker
        self.callback(name_map)
        self.destroy()


class AccurateConsentDialog(ctk.CTkToplevel):
    """First-activation consent screen for Accurate mode's model download.

    Shown once, before any network access, using real numbers from
    accurate_model_download.describe_download() — never a number remembered
    from an earlier investigation, which could be stale by the time a user
    actually clicks the button. `on_confirm`/`on_decline` follow the same
    callback-then-destroy convention as NameAssignWindow.confirm() above.
    """

    def __init__(self, parent, info, *, on_confirm, on_decline):
        super().__init__(parent)
        self.title("StillScript — Accurate Mode Setup")
        # 480x420 through masterplan 2.3; grown to fit the mixed-language
        # notice added in 2.11 without cramming the existing rows.
        self.geometry("600x775")
        self.resizable(False, False)
        self.grab_set()
        self._on_confirm = on_confirm
        self._on_decline = on_decline

        # Kept as attributes (not just packed and forgotten) so the real
        # numbers behind the dialog can be checked directly rather than by
        # parsing rendered label text.
        self.size_gb = info["total_bytes"] / (1024 ** 3)
        self.info_labels = []

        ctk.CTkLabel(
            self, text="🎯 Accurate Mode — one-time setup",
            font=("Arial", 21, "bold"),
        ).pack(pady=(20, 4))
        ctk.CTkLabel(
            self,
            text="Accurate mode uses a larger, more precise speech model that\n"
                 "is not included in the installer. It downloads once, the\n"
                 "first time you use Accurate mode.",
            font=("Arial", 16), justify="center",
        ).pack(pady=(0, 16))

        info_frame = ctk.CTkFrame(self, fg_color="gray20")
        info_frame.pack(fill="x", padx=24, pady=4)

        def _row(icon_text):
            label = ctk.CTkLabel(
                info_frame, text=icon_text, font=("Arial", 16),
                justify="left", anchor="w", wraplength=480,
            )
            label.pack(fill="x", padx=14, pady=8)
            self.info_labels.append(label)

        _row(f"📦  Size: about {self.size_gb:.1f} GB")
        _row("🌐  Requires an internet connection.")
        _row(
            "⏱  Time: this can take anywhere from tens of minutes to several\n"
            "hours, depending on your connection speed — on a slower line,\n"
            "expect it to take hours, not minutes. You can leave StillScript\n"
            "running in the background; if it's interrupted, it picks up\n"
            "where it left off rather than starting over."
        )
        # Masterplan 2.11, based on real Accurate-mode runs (2.8): a 3:20
        # noisy-bar clip and a real 46-min mixed Afrikaans/English committee
        # recording. Stated here, not just in the main window's smaller
        # reminder, because this is the one point where a first-time user is
        # deciding whether Accurate mode fits what they actually record —
        # someone whose work regularly involves code-switched audio should
        # know that going in, not discover it after a multi-hour transcription.
        _row(
            "🌍  Built for Afrikaans: this mode always transcribes as\n"
            "Afrikaans, no matter what Audio Language is selected above —\n"
            "its model is tuned specifically for Afrikaans. If your\n"
            "recording has real, significant English (or another language)\n"
            "mixed in, expect it to run slower, and expect a meaningfully\n"
            "higher chance of scattered errors that need a full, careful\n"
            "read-through rather than a light check. This doesn't mean the\n"
            "transcript is unusable — long stretches, including whole\n"
            "English passages, often come through correctly — just that\n"
            "mixed-language recordings need closer review than a\n"
            "purely-Afrikaans one would."
        )

        # Confidentiality point — stated plainly, in its own visually distinct
        # block, not folded into the paragraph above as fine print.
        ctk.CTkLabel(
            self,
            text="🔒 This download is one-way: it only fetches the language\n"
                 "model's weights onto your computer. It has nothing to do with\n"
                 "your recordings — your audio and transcripts are never\n"
                 "uploaded, now or ever.",
            font=("Arial", 14, "bold"), text_color="#8fd19e",
            justify="center", wraplength=505,
        ).pack(pady=(16, 4))

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=20)
        self.decline_btn = ctk.CTkButton(
            btn_frame, text="Not now", command=self._decline,
            width=180, height=50, fg_color="gray30",
        )
        self.decline_btn.grid(row=0, column=0, padx=8)
        self.confirm_btn = ctk.CTkButton(
            btn_frame, text="⬇ Download & Continue", command=self._confirm,
            width=265, height=50, fg_color="#1f538d", font=("Arial", 17, "bold"),
        )
        self.confirm_btn.grid(row=0, column=1, padx=8)

    def _confirm(self):
        self.destroy()
        self._on_confirm()

    def _decline(self):
        self.destroy()
        self._on_decline()


class AccurateDownloadProgressDialog(ctk.CTkToplevel):
    """Progress display for the Accurate-mode model download.

    Rendering is driven entirely by DownloadProgress objects handed to
    update_progress() — the same shape accurate_model_download already
    defines and emits, not a new one invented here. `on_cancel` is called once,
    from the Tk main thread (this dialog's Cancel button is a normal Tkinter
    command), and only sets a threading.Event; it does not stop anything by
    itself, cancellation happens in the download function on its own schedule.
    """

    def __init__(self, parent, *, on_cancel):
        super().__init__(parent)
        self.title("StillScript — Downloading Accurate-mode model")
        self.geometry("575x275")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", lambda: None)  # no window-close shortcut around Cancel
        self.grab_set()
        self._on_cancel = on_cancel

        self.phase_label = ctk.CTkLabel(
            self, text="Checking the download…", font=("Arial", 17, "bold"),
        )
        self.phase_label.pack(pady=(24, 8))

        self.bar = ctk.CTkProgressBar(self, width=455)
        self.bar.pack(pady=6)
        self.bar.set(0)

        self.detail_label = ctk.CTkLabel(
            self, text="", font=("Arial", 14), text_color="gray",
        )
        self.detail_label.pack(pady=(4, 10))

        self.cancel_btn = ctk.CTkButton(
            self, text="Cancel", command=self._cancel,
            width=170, height=45, fg_color="gray30",
        )
        self.cancel_btn.pack(pady=6)

        self._cancelling = False

    def update_progress(self, progress):
        """Render one DownloadProgress sample. Safe to call only on the main thread."""
        if self._cancelling:
            return
        self.phase_label.configure(text=progress.message)
        self.bar.set(max(0.0, min(1.0, progress.fraction)))

        if progress.total_bytes:
            done_mib = progress.downloaded_bytes / (1024 ** 2)
            total_mib = progress.total_bytes / (1024 ** 2)
            rate_mib = progress.bytes_per_second / (1024 ** 2)
            eta = "…" if progress.eta_seconds is None else _format_eta(progress.eta_seconds)
            self.detail_label.configure(
                text=f"{done_mib:,.0f} / {total_mib:,.0f} MiB   "
                     f"{rate_mib:.2f} MiB/s   ETA {eta}"
            )
        else:
            self.detail_label.configure(text="")

    def _cancel(self):
        # Chunk-level cancellation (see accurate_model_download's RESUME notes):
        # this takes effect after the chunk in flight finishes, not instantly.
        # Say so, and disable the button so a second click can't do anything odd.
        self._cancelling = True
        self.cancel_btn.configure(state="disabled", text="Cancelling…")
        self.phase_label.configure(
            text="Cancelling — finishing the current part first…")
        self._on_cancel()


def _format_eta(seconds):
    seconds = max(0, int(seconds))
    hours, rem = divmod(seconds, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours}h {minutes}m"
    if minutes:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


class AccurateTimeEstimateDialog(ctk.CTkToplevel):
    """Per-run time estimate, shown right before an Accurate-mode
    transcription actually starts (masterplan 2.4) — after any first-time
    model-download consent (masterplan 2.3) is already resolved, and distinct
    from it: that dialog is about whether to download a fixed-size model at
    all; this one is about how long THIS specific file will take, given its
    length and the currently-selected Reproducibility setting (masterplan
    2.9). Shown on every Accurate-mode run, not just the first activation,
    because the estimate genuinely differs per file — unlike the persistent
    mode/language scope notices (2.9/2.11), which say the same thing every
    time regardless of which file is selected.
    """

    def __init__(self, parent, *, duration_seconds, reproducibility, diarize,
                 on_start, on_cancel):
        super().__init__(parent)
        self.title("StillScript — Accurate Mode Time Estimate")
        # Measured via winfo_reqheight() under Xvfb, same discipline as
        # SettingsWindow's own geometry comment — not guessed. Real content
        # needs ~405px across every duration/diarize/reproducibility
        # combination tested; 460x430 leaves real, checked headroom.
        self.geometry("575x540")
        self.resizable(False, False)
        self.grab_set()
        self._on_start = on_start
        self._on_cancel = on_cancel

        # Kept as attributes, same reasoning as AccurateConsentDialog.size_gb:
        # so a test can check the real numbers behind the dialog directly,
        # rather than parsing rendered label text.
        self.duration_seconds = duration_seconds
        self.reproducibility = reproducibility
        self.estimate_range = _estimate_accurate_time_range(
            duration_seconds, reproducibility, diarize
        )

        ctk.CTkLabel(
            self, text="⏱  Accurate Mode — Time Estimate",
            font=("Arial", 21, "bold"),
        ).pack(pady=(20, 4))

        info_frame = ctk.CTkFrame(self, fg_color="gray20")
        info_frame.pack(fill="x", padx=24, pady=4)

        def _row(text):
            label = ctk.CTkLabel(
                info_frame, text=text, font=("Arial", 16),
                justify="left", anchor="w", wraplength=480,
            )
            label.pack(fill="x", padx=14, pady=8)
            return label

        if duration_seconds is not None:
            _row(f"🎵  Recording length: {_format_eta(duration_seconds)}")
        else:
            _row("🎵  Recording length: could not be determined.")

        _row(f"🔁  Reproducibility setting: {reproducibility}")

        if self.estimate_range is not None:
            low, high = self.estimate_range
            diarize_note = " (includes speaker identification)" if diarize else ""
            self.estimate_label = _row(
                f"⏳  Estimated time: {_format_eta(low)} – {_format_eta(high)}"
                f"{diarize_note}"
            )
        else:
            self.estimate_label = _row(
                "⏳  Estimated time: unknown (couldn't read this file's "
                "length) — expect the ranges described below once "
                "processing begins."
            )

        ctk.CTkLabel(
            self, text=ACCURATE_TIME_ESTIMATE_EXPLANATION, font=("Arial", 14),
            text_color="gray70", justify="left", wraplength=490,
        ).pack(pady=(12, 10), padx=20)

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=16)
        self.cancel_btn = ctk.CTkButton(
            btn_frame, text="Cancel", command=self._cancel,
            width=170, height=50, fg_color="gray30",
        )
        self.cancel_btn.grid(row=0, column=0, padx=8)
        self.start_btn = ctk.CTkButton(
            btn_frame, text="▶ Start Transcription", command=self._start,
            width=265, height=50, fg_color="#1f538d", font=("Arial", 17, "bold"),
        )
        self.start_btn.grid(row=0, column=1, padx=8)

    def _start(self):
        self.destroy()
        self._on_start()

    def _cancel(self):
        self.destroy()
        self._on_cancel()


# ─────────────────────────────────────────────
#  MAIN APPLICATION
# ─────────────────────────────────────────────

class StillScriptApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("StillScript Confidential Transcripts — v3.1.0")
        self.geometry("775x1150")
        self.resizable(False, False)

        self.config_data = load_config()
        self.current_transcript = None
        self.diarized_segments = None
        self.speaker_name_map = {}

        self._set_app_icon()
        self._build_ui()

    def _set_app_icon(self):
        """Window/taskbar icon (masterplan 4.2) — the real StillScript mark,
        symbol-only on its Ink Navy ground, exactly as the brand guide's own
        "App Icons" page specifies for 16-256px contexts (no wordmark at
        these sizes). `iconphoto(True, ...)` is used rather than the
        Windows-only `iconbitmap(".ico")` so this works identically on every
        platform this app runs on during development, not just the shipped
        Windows build; `True` also makes every Toplevel opened from this
        window (Settings, About, etc.) inherit the same icon automatically.
        Sizes 16/32/48/256 mirror the exact set the brand guide illustrates
        on its own "App Icons" page — Tk picks whichever fits each context.
        Best-effort: a missing/corrupt icon file must never stop the app
        from starting.
        """
        try:
            # Kept as an attribute, not a local: Tk does not itself retain a
            # reference to a PhotoImage passed to iconphoto(), so a
            # locals-only list would be garbage-collected and the icon would
            # silently vanish once __init__ returns.
            self._app_icon_images = [
                ImageTk.PhotoImage(Image.open(resource_path(f"stillscript_icon_{size}.png")))
                for size in (16, 32, 48, 256)
            ]
            self.iconphoto(True, *self._app_icon_images)
        except Exception as e:
            logger.warning("App icon load failed (non-fatal): %s", e)

    # ── THREAD-SAFE UI UPDATES ──────────────

    def _ui(self, func, *args, **kwargs):
        """Marshal a UI update onto the main thread.

        Tkinter is not thread-safe; worker threads must not touch widgets
        directly. `after(0, ...)` queues the call on the main event loop.
        """
        try:
            self.after(0, lambda: func(*args, **kwargs))
        except Exception as e:
            logger.error("UI update failed: %s", e)

    # ── BUILD UI ────────────────────────────

    def _build_ui(self):
        ctk.set_appearance_mode("dark")

        # ── Top banner: centered logo (masterplan 4.2, 2026-08-04 follow-up) ──
        banner = ctk.CTkFrame(self, fg_color="transparent")
        banner.pack(pady=(10, 0), fill="x", padx=20)

        # Logo — masterplan 4.2 (2026-08-04 follow-up): the real StillScript mark (brand
        # guide's primary horizontal lockup, white/reversed variant for this
        # dark-mode ground — see brand/StillScript_Brand_Identity.pdf page 2
        # "primary lockup ... any use with generous room" and page 5's
        # navy/black/cream/white-grounds-only rule; this dark CTk theme's
        # near-black ground qualifies). Height fixed, width computed from the
        # asset's own aspect ratio so it is never stretched off-ratio (a
        # brand-guide "Don't"). Masterplan 4.2 (2026-08-04 follow-up): the SA flag graphic that used
        # to share this banner is gone (Danie: it visually clashed with the
        # logo); the logo is centered in the freed space instead of pinned
        # left. No side= on pack() below — pack's default centers a widget
        # across the packing frame's cross axis, so this needs no manual
        # spacer, unlike the old left+spacer+flag layout.
        logo_loaded = False
        try:
            img = Image.open(resource_path("stillscript_logo_horizontal_white.png"))
            logo_h = 85
            logo_w = round(logo_h * img.width / img.height)
            logo_img = ctk.CTkImage(light_image=img, dark_image=img, size=(logo_w, logo_h))
            ctk.CTkLabel(banner, image=logo_img, text="").pack()
            logo_loaded = True
        except Exception:
            ctk.CTkLabel(banner, text="StillScript", font=("Arial", 30, "bold")).pack()

        # Made in SA text under banner — masterplan 4.2 (2026-08-04 follow-up): Danie asked to keep
        # this tagline even though the flag graphic above it is gone.
        ctk.CTkLabel(
            self,
            text="Made in South Africa  ·  for South Africans",
            font=("Arial", 13),
            text_color="#888888"
        ).pack(pady=(2, 6))

        # Settings button (top right) — masterplan 4.2 (2026-08-04 follow-up): placed with relx=1.0 +
        # anchor="ne" (relative to the window's right edge), not the old
        # hardcoded x=480. That value was tuned for the pre-follow-up 620px-wide
        # window with a small left-pinned logo; once the window widened and
        # the logo was centered and enlarged, a fixed x=480 landed on top of
        # the logo instead of the corner. Relative placement stays correct
        # regardless of window width, so this class of bug can't recur if
        # the window is ever resized again.
        settings_btn = ctk.CTkButton(
            self, text="⚙️ Settings", command=self.open_settings,
            width=145, height=35, fg_color="gray30"
        )
        settings_btn.place(relx=1.0, x=-15, y=15, anchor="ne")

        # Credits / About button (below Settings)
        credits_btn = ctk.CTkButton(
            self, text="ℹ️ About", command=self.open_credits,
            width=145, height=35, fg_color="gray30"
        )
        credits_btn.place(relx=1.0, x=-15, y=55, anchor="ne")

        # 1. Task
        ctk.CTkLabel(self, text="1. Select Task:", font=("Arial", 17, "bold")).pack(pady=(10, 2))
        self.task_var = ctk.StringVar(value="Translate to English")
        ctk.CTkOptionMenu(
            self, values=["Original Language", "Translate to English"],
            variable=self.task_var, width=455
        ).pack(pady=5)

        # 2. Language
        ctk.CTkLabel(self, text="2. Audio Language:", font=("Arial", 17, "bold")).pack(pady=(10, 2))
        self.lang_var = ctk.StringVar(value=self.config_data.get("last_language", "Auto-Detect"))
        ctk.CTkOptionMenu(
            self, values=list(LANG_CODES.keys()),
            variable=self.lang_var, width=455
        ).pack(pady=5)

        # 3. Mode — Fast vs Accurate. Two plain buttons rather than
        # CTkSegmentedButton because Phase 1/2 needed to render one segment
        # disabled while the other was live; Accurate is active as of Wave 2.3,
        # but the two-button layout is kept rather than reshuffled, since it
        # already matches the rest of the app's button styling.
        ctk.CTkLabel(self, text="3. Mode:", font=("Arial", 17, "bold")).pack(pady=(10, 2))
        self.mode_var = ctk.StringVar(value="fast")

        mode_frame = ctk.CTkFrame(self, fg_color="transparent")
        mode_frame.pack(pady=2)

        self.fast_btn = ctk.CTkButton(
            mode_frame, text="⚡ Fast (Vinnig)",
            command=self._select_fast_mode,
            width=220, height=50, fg_color="#1f538d", font=("Arial", 17, "bold")
        )
        self.fast_btn.grid(row=0, column=0, padx=(0, 8))

        self.accurate_btn = ctk.CTkButton(
            mode_frame, text="🎯 Accurate (Akkuraat)",
            command=self._select_accurate_mode,
            width=220, height=50, fg_color="gray30", font=("Arial", 17, "bold")
        )
        self.accurate_btn.grid(row=0, column=1)

        # Persistent, non-blocking scope notice (masterplan 2.11) — swapped by
        # _select_fast_mode()/_select_accurate_mode() below, same as the two
        # buttons' colours, so it always reflects whichever mode is actually
        # selected. This is disclosure, not a gate: neither mode is disabled
        # or blocked by anything stated here.
        self.mode_scope_label = ctk.CTkLabel(
            self, text=FAST_MODE_SCOPE_NOTE, font=("Arial", 13),
            text_color="gray70", justify="center", wraplength=550,
        )
        self.mode_scope_label.pack(pady=(4, 0))

        # 4. Speaker identification
        ctk.CTkLabel(self, text="4. Speaker Identification:", font=("Arial", 17, "bold")).pack(pady=(10, 2))
        self.diarize_var = ctk.BooleanVar(value=False)

        diarize_frame = ctk.CTkFrame(self, fg_color="transparent")
        diarize_frame.pack(pady=2)
        ctk.CTkSwitch(
            diarize_frame, text="Identify different speakers",
            variable=self.diarize_var
        ).pack(side="left", padx=10)

        # Number of speakers selector
        ctk.CTkLabel(diarize_frame, text="  Expected speakers:", font=("Arial", 14)).pack(side="left", padx=(20, 5))
        self.num_speakers_var = ctk.StringVar(value="2")
        ctk.CTkOptionMenu(
            diarize_frame,
            values=["2", "3", "4", "5", "6", "7", "8"],
            variable=self.num_speakers_var,
            width=85
        ).pack(side="left")

        # Status & progress
        self.status_label = ctk.CTkLabel(self, text="Ready", font=("Arial", 17), text_color="gray")
        self.status_label.pack(pady=(20, 5))
        self.progress_bar = ctk.CTkProgressBar(self, width=575)
        self.progress_bar.pack(pady=5)
        self.progress_bar.set(0)

        # Main button
        self.main_btn = ctk.CTkButton(
            self, text="▶  START PROCESSING",
            command=self.start_transcription,
            height=70, width=410,
            font=("Arial", 22, "bold"),
            fg_color="#1f538d"
        )
        self.main_btn.pack(pady=20)

        # Action buttons (after transcription)
        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=5)

        self.name_btn = ctk.CTkButton(
            btn_frame, text="🎤 Assign Speaker Names",
            command=self.open_name_assignment,
            width=240, height=50, fg_color="gray30",
            state="disabled"
        )
        self.name_btn.grid(row=0, column=0, padx=10)

        self.summarize_btn = ctk.CTkButton(
            btn_frame, text="🤖 Summarize with Claude",
            command=self.summarize_with_claude,
            width=250, height=50, fg_color="#2d6a4f",
            state="disabled"
        )
        self.summarize_btn.grid(row=0, column=1, padx=10)

        # Output area
        ctk.CTkLabel(self, text="Output:", font=("Arial", 16, "bold")).pack(pady=(15, 2))
        self.output_box = ctk.CTkTextbox(self, width=670, height=215, font=("Arial", 16))
        self.output_box.pack(pady=5)

    # ── MODE SELECTION ──────────────────────
    # Both buttons stay enabled and clickable; the colour swap is the only
    # feedback that a mode is "selected" (there is no separate indicator), so
    # both handlers must repaint both buttons — not just set mode_var — or a
    # click would change behaviour invisibly.

    def _select_fast_mode(self):
        self.mode_var.set("fast")
        self.fast_btn.configure(fg_color="#1f538d")
        self.accurate_btn.configure(fg_color="gray30")
        self.mode_scope_label.configure(text=FAST_MODE_SCOPE_NOTE)

    def _select_accurate_mode(self):
        self.mode_var.set("accurate")
        self.accurate_btn.configure(fg_color="#1f538d")
        self.fast_btn.configure(fg_color="gray30")
        self.mode_scope_label.configure(text=ACCURATE_MODE_SCOPE_NOTE)

    # ── SETTINGS / CREDITS ──────────────────

    def open_settings(self):
        SettingsWindow(self)

    def open_credits(self):
        CreditsWindow(self)

    # ── TRANSCRIPTION ───────────────────────

    def start_transcription(self):
        file_path = filedialog.askopenfilename(
            title="Select Audio File",
            filetypes=[("Audio Files", "*.mp3 *.wav *.m4a *.flac *.mp4 *.ogg *.webm")]
        )
        if not file_path:
            return

        # Validate the selection before doing any heavy work.
        if not os.path.isfile(file_path):
            messagebox.showerror("Error", "The selected file could not be found.")
            return
        if os.path.getsize(file_path) == 0:
            messagebox.showerror("Error", "The selected file is empty.")
            return

        mode = self.mode_var.get()

        # Save last used settings. Masterplan 2.9 found and fixed a real bug
        # here: this used to write self.config_data directly — the snapshot
        # load_config() took once at app startup — which meant clicking Start
        # Processing silently clobbered any change made via the Settings
        # window (e.g. the reproducibility choice, or the API key's config-
        # file fallback) back to whatever was on disk when the app launched,
        # unless the user restarted the app first. Re-reading fresh here
        # means an in-session Settings change survives the very next run, not
        # just the next app launch.
        self.config_data = load_config()
        self.config_data["last_language"] = self.lang_var.get()
        save_config(self.config_data)

        task_choice = self.task_var.get()
        language_label = self.lang_var.get()
        lang_code = LANG_CODES[language_label]
        whisper_task = "translate" if task_choice == "Translate to English" else "transcribe"
        do_diarize = self.diarize_var.get()
        num_speakers = int(self.num_speakers_var.get())

        self.output_box.delete("1.0", "end")

        # Accurate mode has its own flow — first a possible download (with
        # consent), then transcription. It is never allowed to fall back to
        # the Fast/Medium engine, so this branches out completely rather than
        # sharing run() below.
        if mode == "accurate":
            self._start_accurate_transcription(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )
            return

        model_name = FAST_MODE_MODEL
        self.status_label.configure(text=f"Status: Loading model ({model_name})...")
        self.progress_bar.set(0.1)
        self.main_btn.configure(state="disabled")
        self.name_btn.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")

        def run():
            try:
                self._ui(self.progress_bar.set, 0.3)
                self._ui(self.status_label.configure, text="Status: Processing audio...")

                # Read fresh from disk, not self.config_data (loaded once at
                # startup) — same reasoning as _transcribe_with_accurate_model()'s
                # read (masterplan 2.9, extended to Fast mode by 3.2): a
                # Settings change must reach the very next transcription, not
                # require an app restart.
                repro_choice = load_config().get("reproducibility", "Consistent")
                temperature = (
                    FAST_TEMPERATURE_CONSISTENT if repro_choice == "Consistent"
                    else FAST_TEMPERATURE_BEST_EFFORT
                )

                result = transcribe_audio(
                    file_path,
                    language=lang_code,
                    task=whisper_task,
                    model_name=model_name,
                    temperature=temperature,
                )

                self._ui(self.progress_bar.set, 0.7)

                if do_diarize:
                    self._ui(self.status_label.configure, text="Status: Identifying speakers...")
                    segments = self._diarize(result, max_speakers=num_speakers, audio_path=file_path)
                    self.diarized_segments = segments
                    transcript_text = self._segments_to_text(segments, self.speaker_name_map)
                else:
                    self.diarized_segments = None
                    transcript_text = result["text"]

                # current_transcript feeds the on-screen preview AND the Claude
                # summary prompt, so it must stay pure transcript. Provenance is
                # attached only at the export layer (below), never here.
                self.current_transcript = transcript_text

                provenance = build_provenance(
                    mode="Fast (Vinnig)",
                    language_label=self.lang_var.get(),
                    task=whisper_task,
                    diarized=do_diarize,
                    num_speakers=num_speakers,
                    # Masterplan 3.2 — sourced from the result dict (the
                    # ACTUAL temperature transcribe_audio() decoded with), not
                    # re-read from Settings here. Same principle Accurate mode
                    # already uses (2.6/2.9): if the two ever disagreed,
                    # provenance must report reality, not the request.
                    reproducibility=result.get("reproducibility"),
                )
                output_path = self._save_transcript(
                    transcript_text, file_path, provenance=provenance
                )

                preview = transcript_text[:2000] + ("..." if len(transcript_text) > 2000 else "")
                self._ui(self.progress_bar.set, 1.0)
                self._ui(self.status_label.configure, text="✅ Done!")
                self._ui(self.output_box.insert, "1.0", preview)

                if do_diarize and self.diarized_segments:
                    self._ui(self.name_btn.configure, state="normal")
                self._ui(self.summarize_btn.configure, state="normal")

                self._ui(messagebox.showinfo, "StillScript", f"Transcription complete!\nSaved to folder:\n{output_path}\n\nBoth .txt and .docx files created.")

            except Exception as e:
                logger.error("Transcription failed: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"An error occurred:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")

            self._ui(self.main_btn.configure, state="normal")
            self._ui(self.progress_bar.set, 0)

        threading.Thread(target=run, daemon=True).start()

    # ── ACCURATE MODE (Wave 2.3) ─────────────
    # First activation: consent -> cancellable/resumable download -> full
    # verification -> transcription. Subsequent activations: straight to
    # transcription. Both routes converge on _transcribe_with_accurate_model()
    # for the actual engine call, so Wave 2.4's transcription-progress /
    # time-warning UI only has one call site to extend, not two.
    #
    # accurate_model_download is imported lazily inside these methods, the
    # same discipline accurate_engine.py and transcribe_audio_accurate() already
    # follow — Fast mode's runtime must never acquire this import just because
    # the module exists on disk.

    def _start_accurate_transcription(self, file_path, language_label, lang_code,
                                       whisper_task, do_diarize, num_speakers):
        import accurate_model_download as amd

        self.main_btn.configure(state="disabled")
        self.name_btn.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")

        if amd.is_model_ready():
            # Subsequent activation — skip consent AND the download UI
            # entirely, per masterplan 2.3. ensure_accurate_model() still runs
            # (a local stat + JSON read; no network) so the directory handed to
            # the engine is always the one it actually verified, never assumed.
            self.status_label.configure(text="Status: Loading Accurate-mode model...")
            self.progress_bar.set(0.1)

            def already_ready_worker():
                try:
                    model_dir = amd.ensure_accurate_model()
                except Exception as e:
                    logger.error("Accurate model check failed unexpectedly: %s", e, exc_info=True)
                    self._ui(self._on_accurate_setup_failed, e)
                    return
                # Masterplan 2.4 — cheap (ffprobe metadata, no decode), safe to
                # call on this background thread before handing off to the main
                # thread for the dialog itself (Tk widgets are main-thread only).
                duration_seconds = _probe_audio_duration_seconds(file_path)
                self._ui(
                    self._show_accurate_time_estimate,
                    model_dir, file_path, duration_seconds, language_label,
                    lang_code, whisper_task, do_diarize, num_speakers,
                )

            threading.Thread(target=already_ready_worker, daemon=True).start()
            return

        # First activation. Fetch the REAL size/file-count before asking for
        # consent — never a number remembered from an earlier report, which
        # could be stale — off the main thread like every other network call.
        self.status_label.configure(text="Status: Checking Accurate-mode download details…")

        def fetch_info():
            try:
                info = amd.describe_download()
            except Exception as e:
                self._ui(self._accurate_consent_check_failed, e)
                return
            self._ui(
                self._show_accurate_consent, info,
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )

        threading.Thread(target=fetch_info, daemon=True).start()

    def _on_accurate_setup_failed(self, exc):
        self.status_label.configure(text="Status: Ready")
        self.progress_bar.set(0)
        self.main_btn.configure(state="normal")
        messagebox.showerror("StillScript", f"Accurate mode could not start:\n{exc}")

    def _accurate_consent_check_failed(self, exc):
        self.status_label.configure(text="Status: Ready")
        self.main_btn.configure(state="normal")
        messagebox.showerror(
            "StillScript",
            "Could not check the Accurate-mode download details. This usually "
            f"means there is no internet connection right now.\n\n{exc}\n\n"
            "Fast mode works offline."
        )

    def _show_accurate_consent(self, info, file_path, language_label, lang_code,
                                whisper_task, do_diarize, num_speakers):
        self.status_label.configure(text="Status: Ready")

        def on_confirm():
            self._download_then_transcribe_accurate(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )

        def on_decline():
            # The user said no. Accurate mode simply isn't used this time —
            # reset and don't ask again until they click Accurate again.
            self.main_btn.configure(state="normal")
            self.progress_bar.set(0)

        AccurateConsentDialog(self, info, on_confirm=on_confirm, on_decline=on_decline)

    def _download_then_transcribe_accurate(self, file_path, language_label, lang_code,
                                            whisper_task, do_diarize, num_speakers):
        import accurate_model_download as amd

        cancel_event = threading.Event()
        progress_dialog = AccurateDownloadProgressDialog(self, on_cancel=cancel_event.set)

        def worker():
            try:
                model_dir = amd.ensure_accurate_model(
                    progress_callback=lambda p: self._ui(progress_dialog.update_progress, p),
                    cancel_event=cancel_event,
                )
            except amd.AccurateModelDownloadCancelled:
                # Catch this BEFORE AccurateModelDownloadError (its own base
                # class) — a deliberate cancel is not a failure and must never
                # show an error dialog. Whatever chunks landed stay on disk.
                self._ui(self._on_accurate_download_cancelled, progress_dialog)
                return
            except amd.AccurateModelDownloadError as e:
                self._ui(
                    self._on_accurate_download_failed, progress_dialog, e,
                    file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
                )
                return
            except Exception as e:
                logger.error("Accurate model download failed unexpectedly: %s", e, exc_info=True)
                self._ui(
                    self._on_accurate_download_failed, progress_dialog, e,
                    file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
                )
                return

            self._ui(progress_dialog.destroy)
            # Masterplan 2.4 — same reasoning as the already-ready path above:
            # cheap, background-thread-safe duration probe, then hand off to
            # the main thread for the dialog.
            duration_seconds = _probe_audio_duration_seconds(file_path)
            self._ui(
                self._show_accurate_time_estimate,
                model_dir, file_path, duration_seconds, language_label,
                lang_code, whisper_task, do_diarize, num_speakers,
            )

        threading.Thread(target=worker, daemon=True).start()

    def _on_accurate_download_cancelled(self, dialog):
        dialog.destroy()
        self.status_label.configure(text="Status: Ready (download cancelled)")
        self.progress_bar.set(0)
        self.main_btn.configure(state="normal")
        # Deliberately no error dialog — the user asked for this, and nothing
        # already downloaded was lost (see accurate_model_download's
        # cancellation notes: chunks are only ever removed after the guard
        # passes, never on cancel).

    def _on_accurate_download_failed(self, dialog, exc, file_path, language_label, lang_code,
                                      whisper_task, do_diarize, num_speakers):
        dialog.destroy()
        self.status_label.configure(text="Status: Ready")
        self.progress_bar.set(0)

        # AccurateModelDownloadError's message is already written for a
        # non-technical user (masterplan 2.1a.2) — shown as-is, no wrapping.
        can_retry = getattr(exc, "can_retry", False)
        if can_retry and messagebox.askretrycancel("StillScript", str(exc)):
            self._download_then_transcribe_accurate(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )
            return

        self.main_btn.configure(state="normal")

    def _show_accurate_time_estimate(self, model_dir, file_path, duration_seconds,
                                      language_label, lang_code, whisper_task,
                                      do_diarize, num_speakers):
        """Masterplan 2.4 — the seam Wave 2.3 deliberately left: both Accurate-
        mode routes above now converge here instead of calling
        _transcribe_with_accurate_model() directly, so the time estimate is
        shown exactly once per run regardless of which route got here. Runs
        on the main thread (Tk widgets require it) — the actual heavy work is
        deferred to a fresh background thread only if the user clicks Start,
        never blocking the UI either way.

        Reads Settings fresh (masterplan 2.9's own discipline: never
        self.config_data, which is a startup snapshot) so the dialog shows
        whatever Reproducibility choice is actually in effect right now.
        """
        repro_choice = load_config().get("reproducibility", "Consistent")

        def on_start():
            threading.Thread(
                target=self._transcribe_with_accurate_model,
                args=(model_dir, file_path, language_label, lang_code,
                      whisper_task, do_diarize, num_speakers),
                daemon=True,
            ).start()

        def on_cancel():
            # Same reset shape as declining the golf 2.3 consent dialog —
            # the user said no, that is not a failure.
            self.status_label.configure(text="Status: Ready")
            self.progress_bar.set(0)
            self.main_btn.configure(state="normal")

        AccurateTimeEstimateDialog(
            self, duration_seconds=duration_seconds, reproducibility=repro_choice,
            diarize=do_diarize, on_start=on_start, on_cancel=on_cancel,
        )

    def _transcribe_with_accurate_model(self, model_dir, file_path, language_label, lang_code,
                                         whisper_task, do_diarize, num_speakers):
        """Runs on a background thread — the shared tail for both Accurate-mode
        routes above, reached only after masterplan 2.4's time-estimate dialog
        (see _show_accurate_time_estimate()) has already been shown and
        confirmed. Wires masterplan 2.4's real transcription-progress UI
        through the progress_callback= seam Wave 2.3 left here; see
        accurate_engine.transcribe()'s docstring for exactly what that
        callback reports and how often."""
        self._ui(self.status_label.configure, text="Status: Loading Accurate-mode model...")
        self._ui(self.progress_bar.set, 0.3)
        try:
            self._ui(self.status_label.configure,
                     text="Status: Processing audio (Accurate mode)...")

            # Read fresh from disk, not self.config_data (loaded once at
            # startup) — same reasoning as summarize_with_claude()'s API-key
            # read: a change made in Settings mid-session must take effect on
            # the very next transcription, not require an app restart.
            # Masterplan 2.9. Translated to the engine's own constants here
            # (not a bare number) so accurate_engine.py never needs to know
            # about Settings' label strings.
            import accurate_engine
            repro_choice = load_config().get("reproducibility", "Consistent")
            temperature = (
                accurate_engine.TEMPERATURE_CONSISTENT if repro_choice == "Consistent"
                else accurate_engine.TEMPERATURE_BEST_EFFORT
            )

            # Masterplan 2.4. This callback runs INSIDE transformers'
            # generate() loop, on THIS background thread — not marshaled
            # through self._ui() itself (only the widget updates it triggers
            # are). An exception escaping it would propagate up through
            # generate() and abort a run that can take over an hour, so
            # every line that can fail is inside its own try/except that
            # only logs, never raises — a UI bug must never be able to do
            # that to a long transcription. Silently doing nothing on a
            # malformed/unexpected `p` is deliberate: the alternative
            # (crashing the transcription over a progress-display glitch)
            # is strictly worse for this product's actual purpose.
            run_start = time.monotonic()

            def on_progress(p):
                try:
                    seek_frames = float(p[0][0])
                    total_frames = float(p[0][1])
                    if total_frames <= 0:
                        return
                    fraction = max(0.0, min(1.0, seek_frames / total_frames))
                    elapsed = time.monotonic() - run_start
                    processed_s = seek_frames / accurate_engine.PROGRESS_FRAMES_PER_SECOND
                    total_s = total_frames / accurate_engine.PROGRESS_FRAMES_PER_SECOND
                    # Scaled into the existing 0.3-0.7 "processing audio" band
                    # (0.1 = model loading, 0.7 = transcription done, 1.0 =
                    # fully done) rather than inventing a new range — the
                    # milestones before/after this call are unchanged.
                    self._ui(self.progress_bar.set, 0.3 + fraction * 0.4)
                    self._ui(
                        self.status_label.configure,
                        text=(
                            "Status: Transcribing (Accurate mode)... "
                            f"{_format_eta(processed_s)} / {_format_eta(total_s)} "
                            f"of audio processed (elapsed {_format_eta(elapsed)})"
                        ),
                    )
                except Exception as e:
                    logger.warning(
                        "Accurate-mode progress update skipped (non-fatal, "
                        "transcription continues): %s", e
                    )

            result = transcribe_audio_accurate(
                file_path, language=lang_code, task=whisper_task, model_dir=model_dir,
                temperature=temperature, progress_callback=on_progress,
            )

            self._ui(self.progress_bar.set, 0.7)

            if do_diarize:
                self._ui(self.status_label.configure, text="Status: Identifying speakers...")
                segments = self._diarize(result, max_speakers=num_speakers, audio_path=file_path)
                self.diarized_segments = segments
                transcript_text = self._segments_to_text(segments, self.speaker_name_map)
            else:
                self.diarized_segments = None
                transcript_text = result["text"]

            self.current_transcript = transcript_text

            provenance = build_provenance(
                mode="Accurate (Akkuraat)",
                language_label=language_label,
                task=whisper_task,
                diarized=do_diarize,
                num_speakers=num_speakers,
                # accurate_engine.transcribe()'s result already carries its own
                # label (ACCURATE_ENGINE_LABEL) — reuse it rather than import
                # accurate_engine here just for a string.
                engine=result.get("engine"),
                # Masterplan 2.6 — likewise sourced from the result dict, not
                # re-derived here, so accurate_engine stays the single source
                # of truth for what backs a given transcription.
                model_id=result.get("model_id"),
                model_revision=result.get("model_revision"),
                model_layout=result.get("model_layout"),
                guard_verification=result.get("guard_verification"),
                # Masterplan 2.9 — likewise sourced from the result dict (the
                # ACTUAL temperature the engine decoded with), not re-read
                # from Settings here. If the two ever disagreed, provenance
                # must report reality, not the request.
                reproducibility=result.get("reproducibility"),
            )
            output_path = self._save_transcript(
                transcript_text, file_path, provenance=provenance
            )

            preview = transcript_text[:2000] + ("..." if len(transcript_text) > 2000 else "")
            self._ui(self.progress_bar.set, 1.0)
            self._ui(self.status_label.configure, text="✅ Done!")
            self._ui(self.output_box.insert, "1.0", preview)

            if do_diarize and self.diarized_segments:
                self._ui(self.name_btn.configure, state="normal")
            self._ui(self.summarize_btn.configure, state="normal")

            self._ui(
                messagebox.showinfo, "StillScript",
                f"Transcription complete!\nSaved to folder:\n{output_path}\n\n"
                "Both .txt and .docx files created.",
            )

        except Exception as e:
            logger.error("Accurate transcription failed: %s", e, exc_info=True)
            self._ui(messagebox.showerror, "Error", f"An error occurred:\n{e}")
            self._ui(self.status_label.configure, text="Status: Error")

        self._ui(self.main_btn.configure, state="normal")
        self._ui(self.progress_bar.set, 0)

    # ── DIARIZATION ─────────────────────────

    def _diarize(self, whisper_result, max_speakers=4, audio_path=None):
        """Attach speaker labels to a transcript.

        Returns exactly the shape it always has — an ordered list of
        {"speaker": "Speaker N", "text": ...} with consecutive same-speaker
        turns merged — so _segments_to_text(), the name-assignment dialog and
        both exporters are unaffected by what happens inside.

        The speaker timeline comes from pyannote (diarization_engine), computed
        over the raw audio and NOT from Whisper's segment boundaries. That
        distinction is the entire point of masterplan 2.12: the old backend
        averaged one feature vector per ~30s Whisper segment, which made a
        speaker change inside a segment unrepresentable and collapsed a
        confirmed four-voice stretch into a single speaker.

        Falls back to the pause-based heuristic below if diarization is
        unavailable — logged at ERROR, not warning, because silently shipping a
        near-one-speaker transcript is precisely the failure this product
        cannot afford.
        """
        segments = whisper_result.get("segments", [])
        if not segments:
            return [{"speaker": "Speaker 1", "text": whisper_result["text"]}]

        try:
            if audio_path is None:
                raise ValueError("no audio path provided")

            import diarization_engine

            turns = diarization_engine.diarize(audio_path, max_speakers=max_speakers)
            if not turns:
                raise ValueError("diarization returned no speaker turns")

            # Referenced via the class, not self. _diarize() has always been
            # callable with an unbound/None self — the existing test harnesses
            # do exactly that — and _attach_speakers is a @staticmethod, so
            # going through self would silently reintroduce an instance
            # requirement and turn every such call into a fallback.
            labelled = StillScriptApp._attach_speakers(segments, turns)
            if not labelled:
                raise ValueError("no text could be attached to the speaker timeline")
            return labelled

        except ImportError as e:
            logger.error("Speaker-diarization libraries unavailable (%s); falling back to "
                         "pause-based detection, which is much less accurate.", e)
        except Exception as e:
            logger.error("Speaker diarization failed (%s); falling back to pause-based "
                         "detection, which is much less accurate.", e, exc_info=True)
        # ── Fallback: pause-based detection ─────────────
        pauses = []
        for i in range(1, len(segments)):
            gap = segments[i].get("start", 0) - segments[i - 1].get("end", 0)
            pauses.append((i, gap))

        sorted_pauses = sorted(pauses, key=lambda x: x[1], reverse=True)
        num_changes   = max_speakers - 1
        change_points = set(idx for idx, gap in sorted_pauses[:num_changes] if gap >= 0.8)

        result = []
        current_speaker = 1
        for i, seg in enumerate(segments):
            text = seg.get("text", "").strip()
            if not text:
                continue
            if i in change_points and current_speaker < max_speakers:
                current_speaker += 1
            speaker_key = f"Speaker {current_speaker}"
            if result and result[-1]["speaker"] == speaker_key:
                result[-1]["text"] += " " + text
            else:
                result.append({"speaker": speaker_key, "text": text})

        return result if result else [{"speaker": "Speaker 1", "text": whisper_result["text"]}]

    @staticmethod
    def _attach_speakers(segments, turns):
        """Map a speaker timeline onto Whisper's segments.

        `turns` is [(start, end, raw_speaker)], non-overlapping and time-sorted.

        A Whisper long-form segment can be ~30s and span several speaker turns,
        so a segment is split where the timeline says the speaker changed, and
        its words are apportioned across the parts in proportion to how long
        each speaker held the floor. Word-level timings would make this exact;
        segment-level timings are all Whisper gives us here, so this is an
        approximation — but a far better one than handing a whole 30s block to
        whichever speaker happened to dominate it, which is what the old
        backend effectively did.

        Raw pyannote labels (SPEAKER_00, ...) are renumbered to "Speaker N" in
        order of first appearance, matching the previous backend's convention
        so the name-assignment dialog keeps working unchanged.
        """
        raw_to_label = {}
        out = []

        for seg in segments:
            text = (seg.get("text") or "").strip()
            if not text:
                continue

            s0 = float(seg.get("start", 0.0) or 0.0)
            s1 = float(seg.get("end", s0) or s0)

            overlaps = []
            if s1 > s0:
                for t0, t1, spk in turns:
                    o = min(s1, t1) - max(s0, t0)
                    if o > 0:
                        overlaps.append((max(s0, t0), o, spk))
                overlaps.sort(key=lambda x: x[0])

            if not overlaps:
                # No speech detected here by the diarizer (music, noise, a gap).
                # Keep the text with whoever was last speaking rather than
                # inventing a speaker or dropping the words.
                carry = out[-1]["_raw"] if out else (turns[0][2] if turns else "SPEAKER_00")
                parts = [(carry, text)]
            elif len(overlaps) == 1:
                parts = [(overlaps[0][2], text)]
            else:
                words = text.split()
                total = sum(o for _, o, _ in overlaps)
                parts, idx = [], 0
                for n, (_, o, spk) in enumerate(overlaps):
                    if n == len(overlaps) - 1:
                        take = len(words) - idx
                    else:
                        take = int(round(len(words) * (o / total))) if total > 0 else 0
                        take = max(0, min(take, len(words) - idx))
                    chunk = words[idx:idx + take]
                    idx += take
                    if chunk:
                        parts.append((spk, " ".join(chunk)))
                if not parts:
                    parts = [(overlaps[0][2], text)]

            for spk, chunk in parts:
                if spk not in raw_to_label:
                    raw_to_label[spk] = f"Speaker {len(raw_to_label) + 1}"
                label = raw_to_label[spk]
                if out and out[-1]["speaker"] == label:
                    out[-1]["text"] += " " + chunk
                else:
                    out.append({"speaker": label, "text": chunk, "_raw": spk})

        for entry in out:
            entry.pop("_raw", None)
        return out

    def _segments_to_text(self, segments, name_map=None):
        lines = []
        for seg in segments:
            speaker = seg["speaker"]
            name = name_map.get(speaker, speaker) if name_map else speaker
            lines.append(f"{name}: {seg['text']}")
        return "\n\n".join(lines)

    # ── NAME ASSIGNMENT ──────────────────────

    def open_name_assignment(self):
        if not self.diarized_segments:
            return
        # Preserve insertion order (no sorting that breaks numbering)
        seen = []
        for seg in self.diarized_segments:
            if seg["speaker"] not in seen:
                seen.append(seg["speaker"])

        def on_names_confirmed(name_map):
            self.speaker_name_map = name_map
            updated_text = self._segments_to_text(self.diarized_segments, name_map)
            self.current_transcript = updated_text
            self.output_box.delete("1.0", "end")
            self.output_box.insert("1.0", updated_text[:2000] + ("..." if len(updated_text) > 2000 else ""))
            self.status_label.configure(text="✅ Names assigned!")

        NameAssignWindow(self, seen, on_names_confirmed)

    # ── AI SUMMARY ──────────────────────────

    def summarize_with_claude(self):
        config = load_config()
        api_key = _get_api_key(config).strip()

        if not api_key:
            messagebox.showwarning(
                "API Key Missing",
                "Please enter your Claude API key via ⚙️ Settings."
            )
            return

        if not self.current_transcript:
            messagebox.showwarning("No Transcript", "Please complete a transcription first.")
            return

        self.summarize_btn.configure(state="disabled")
        self.status_label.configure(text="Status: Claude is generating summary...")
        self.progress_bar.set(0.3)

        def run():
            try:
                client = anthropic.Anthropic(api_key=api_key)

                prompt = f"""You are a professional minutes writer. The following is a transcription of a meeting or conversation.

Please provide a concise summary that includes:
1. Main points discussed
2. Decisions made (if any)
3. Action items (if any)
4. Participants (if names are available)

Transcription:
{self.current_transcript}

Write the summary in the same language as the transcription."""

                message = client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=2000,
                    # Disable extended thinking: a summary doesn't need it, and it
                    # keeps latency and token cost down (Sonnet 5 thinks by default).
                    thinking={"type": "disabled"},
                    messages=[{"role": "user", "content": prompt}]
                )

                summary = message.content[0].text
                output_path = self._save_summary(summary)

                self._ui(self.progress_bar.set, 1.0)
                self._ui(self.status_label.configure, text="✅ Summary complete!")
                self._ui(self.output_box.delete, "1.0", "end")
                self._ui(self.output_box.insert, "1.0", summary)

                self._ui(messagebox.showinfo, "StillScript", f"Summary complete!\nSaved to folder:\n{output_path}\n\nBoth .txt and .docx files created.")

            except anthropic.AuthenticationError:
                logger.warning("Claude authentication failed")
                self._ui(messagebox.showerror, "Invalid API Key", "Please check your Claude API key in Settings.")
                self._ui(self.status_label.configure, text="Status: API error")
            except anthropic.RateLimitError:
                logger.warning("Claude rate limit hit")
                self._ui(messagebox.showerror, "Rate Limited", "Too many requests. Please wait a moment and try again.")
                self._ui(self.status_label.configure, text="Status: Rate limited")
            except anthropic.APIError as e:
                logger.error("Claude API error: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"Claude API error:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")
            except Exception as e:
                logger.critical("Unexpected error during summary: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"An unexpected error occurred:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")

            self._ui(self.summarize_btn.configure, state="normal")
            self._ui(self.progress_bar.set, 0)

        threading.Thread(target=run, daemon=True).start()

    # ── FILE SAVING ──────────────────────────

    def _get_output_dir(self):
        import platform
        if platform.system() == "Windows":
            base = Path.home() / "Downloads"
        else:
            # Linux/Mac: prefer ~/Documents, fall back to ~/Downloads.
            docs = Path.home() / "Documents"
            base = docs if docs.exists() else Path.home() / "Downloads"
        path = str(base / "StillScript_Transcriptions")
        os.makedirs(path, exist_ok=True)
        return path

    def _make_docx(self, text, docx_path, doc_type="transcript", provenance=None):
        """Convert text to a formatted .docx file using python-docx.

        `provenance`, when given, is rendered as a small gray audit footer at
        the very bottom (below the transcript body) — additive metadata that
        records which engine produced the document.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title_text = ("StillScript Confidential Transcripts — Meeting Summary"
                      if doc_type == "summary"
                      else "StillScript Confidential Transcripts — Transcript")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
        run.font.name = "Calibri"

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%d %B %Y"))
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_run.font.name = "Calibri"

        doc.add_paragraph()  # spacer

        # Parse and add content lines
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Speaker line: "Name: text"
            if doc_type == "transcript" and ":" in stripped:
                parts = stripped.split(":", 1)
                if len(parts[0]) <= 40:
                    p = doc.add_paragraph()
                    speaker_run = p.add_run(parts[0] + ": ")
                    speaker_run.bold = True
                    speaker_run.font.name = "Calibri"
                    speaker_run.font.size = Pt(11)
                    text_run = p.add_run(parts[1].strip())
                    text_run.font.name = "Calibri"
                    text_run.font.size = Pt(11)
                    continue

            # Bold markdown headings **text**
            if stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                r = p.add_run(stripped.strip("*"))
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
                r.font.name = "Calibri"
                continue

            # Numbered headings like "1. **Main Points**"
            if stripped[:3].rstrip(". ").isdigit() and "**" in stripped:
                clean = stripped.replace("**", "").lstrip("0123456789. ")
                p = doc.add_paragraph()
                r = p.add_run(clean)
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
                r.font.name = "Calibri"
                continue

            # Bullet points
            if stripped.startswith("- ") or stripped.startswith("• "):
                bullet_text = stripped.lstrip("-• ").strip()
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(bullet_text)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.name = "Calibri"
            r.font.size = Pt(11)

        # Provenance footer — small gray audit block at the very bottom.
        if provenance:
            doc.add_paragraph()  # spacer
            sep = doc.add_paragraph()
            sep_run = sep.add_run("─" * 40)
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            sep_run.font.name = "Calibri"

            heading = doc.add_paragraph()
            h_run = heading.add_run("Transcription provenance")
            h_run.bold = True
            h_run.font.size = Pt(9)
            h_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h_run.font.name = "Calibri"

            for line in format_provenance_lines(provenance):
                pp = doc.add_paragraph()
                pr = pp.add_run(line)
                pr.font.size = Pt(9)
                pr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                pr.font.name = "Calibri"

        doc.save(docx_path)

    def _save_transcript(self, text, source_path, provenance=None):
        from datetime import datetime
        base = os.path.splitext(os.path.basename(source_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = self._get_output_dir()

        # The transcript body is written verbatim; the provenance footer, if
        # present, is appended AFTER it, clearly separated — additive metadata,
        # not part of the transcript content itself.
        txt_body = text
        if provenance:
            footer = "\n".join(format_provenance_lines(provenance))
            txt_body = (
                text
                + "\n\n" + ("─" * 40) + "\n"
                + "Transcription provenance\n"
                + footer + "\n"
            )

        txt_path = os.path.join(out_dir, f"{base}_transcript_{timestamp}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(txt_body)
        docx_path = os.path.join(out_dir, f"{base}_transcript_{timestamp}.docx")
        try:
            self._make_docx(text, docx_path, "transcript", provenance=provenance)
        except Exception as e:
            logger.warning("Could not create .docx: %s", e)
        return out_dir

    def _save_summary(self, text):
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = self._get_output_dir()
        txt_path = os.path.join(out_dir, f"summary_{timestamp}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        docx_path = os.path.join(out_dir, f"summary_{timestamp}.docx")
        try:
            self._make_docx(text, docx_path, "summary")
        except Exception as e:
            logger.warning("Could not create .docx: %s", e)
        return out_dir


# ─────────────────────────────────────────────
#  START APPLICATION
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import multiprocessing
    multiprocessing.freeze_support()  # Required for the Windows .exe build
    ctk.set_appearance_mode("dark")
    app = StillScriptApp()
    app.mainloop()
