import os
import sys
import io
import json
import base64
import logging
import threading
import webbrowser
from pathlib import Path
import customtkinter as ctk
from PIL import Image
import whisper
from tkinter import filedialog, messagebox
import anthropic

# Optional: use the OS keyring for secure API-key storage when available.
try:
    import keyring
    _HAS_KEYRING = True
except Exception:
    _HAS_KEYRING = False

# ─────────────────────────────────────────────
#  LOGGING
# ─────────────────────────────────────────────
# Log to a file in the user's home dir. A windowed .exe has no console, so
# print() output is invisible — logging is the only way to diagnose issues.
LOG_PATH = os.path.join(Path.home(), ".danscribe.log")
logging.basicConfig(
    filename=LOG_PATH,
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
)
logger = logging.getLogger("danscribe")

# ─────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────

def resource_path(relative_path):
    """Get path to files inside the EXE."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def _ensure_bundled_binaries_on_path():
    """Make ffmpeg/ffprobe bundled inside the frozen exe discoverable.

    The Windows build ships ffmpeg.exe and ffprobe.exe via PyInstaller
    --add-binary "...;.", which unpacks them into sys._MEIPASS (the onefile
    bundle root). Every audio call in this app shells out to a *bare*
    "ffmpeg"/"ffprobe" name — Whisper's own decoder (whisper/audio.py) as
    well as our diarization helpers (_probe_audio_duration_seconds,
    _diarize). Those resolve via PATH, so we prepend the bundle dir here.

    Prepend (not append) so the bundled copy wins over any stale system
    ffmpeg. No-op when running from source (sys._MEIPASS is absent), so the
    dev/Linux PATH is left exactly as-is.
    """
    base_path = getattr(sys, "_MEIPASS", None)
    if not base_path:
        return
    parts = os.environ.get("PATH", "").split(os.pathsep)
    if base_path not in parts:
        os.environ["PATH"] = base_path + os.pathsep + os.environ.get("PATH", "")


# Run at import so ffmpeg is resolvable before any transcription starts,
# regardless of entry point.
_ensure_bundled_binaries_on_path()

CONFIG_PATH = os.path.join(Path.home(), ".danscribe_config.json")
_KEYRING_SERVICE = "DanScribe"
_KEYRING_USER = "claude_api_key"

# Claude model used for AI summaries. Kept as a single constant so it can be
# updated in one place when Anthropic retires a model snapshot.
CLAUDE_MODEL = "claude-opus-4-8"


def _get_api_key(config):
    """Read the Claude API key from the OS keyring, falling back to the config file."""
    if _HAS_KEYRING:
        try:
            key = keyring.get_password(_KEYRING_SERVICE, _KEYRING_USER)
            if key:
                return key
        except Exception as e:
            logger.warning("Keyring read failed, falling back to config file: %s", e)
    return config.get("api_key", "")


def _set_api_key(config, api_key):
    """Store the API key securely. Prefer the OS keyring; never leave it in the JSON file."""
    config.pop("api_key", None)  # never persist the key in plaintext JSON
    if _HAS_KEYRING:
        try:
            if api_key:
                keyring.set_password(_KEYRING_SERVICE, _KEYRING_USER, api_key)
            else:
                try:
                    keyring.delete_password(_KEYRING_SERVICE, _KEYRING_USER)
                except Exception:
                    pass
            return
        except Exception as e:
            logger.warning("Keyring write failed, storing key in config file: %s", e)
    # Fallback only when no keyring backend is available.
    config["api_key"] = api_key


def load_config():
    defaults = {"api_key": "", "last_language": "Auto-Detect"}
    if not os.path.exists(CONFIG_PATH):
        return defaults
    try:
        with open(CONFIG_PATH, "r") as f:
            config = json.load(f)
        if not isinstance(config, dict):
            raise ValueError("config is not a JSON object")
    except (OSError, ValueError, json.JSONDecodeError) as e:
        logger.error("Could not read config (%s); using defaults.", e)
        return defaults

    # Merge defaults so a partial/old config never has missing keys.
    merged = {**defaults, **config}

    # Validate persisted choices against known values. (A pre-Phase-2
    # "last_model" key may still be in the file; it is simply ignored now
    # that mode selection replaces the model dropdown.)
    if merged.get("last_language") not in LANG_CODES:
        merged["last_language"] = defaults["last_language"]
    return merged

def save_config(data):
    try:
        with open(CONFIG_PATH, "w") as f:
            json.dump(data, f)
        # Restrict to owner read/write so the config isn't world-readable.
        try:
            os.chmod(CONFIG_PATH, 0o600)
        except OSError:
            pass  # e.g. on filesystems that don't support chmod (some Windows setups)
    except OSError as e:
        logger.error("Could not save config: %s", e)

# ─────────────────────────────────────────────
#  CONFIGURATION DATA
# ─────────────────────────────────────────────

LANG_CODES = {
    "Auto-Detect": None, "Afrikaans": "af", "English": "en", "Sesotho": "st",
    "isiZulu": "zu", "isiXhosa": "xh", "Setswana": "tn", "Sepedi": "nso",
    "Siswati": "ss", "Tshivenda": "ve", "Xitsonga": "ts", "isiNdebele": "nr"
}

# ── Transcription modes ──────────────────────
# "Vinnig" (Fast) is the only mode with a working engine. It runs Whisper
# Medium through the transcribe_audio() seam — the path verified byte-for-byte
# in Phase 1. "Akkuraat" (Accurate) is presented in the UI but disabled until
# Phase 3 wires the large-v3 + Afrikaans-adapter engine behind it; nothing in
# this build must ever let Accurate silently fall back to Medium.
FAST_MODE_MODEL = "medium"

# Human-readable engine label recorded in each transcript's provenance footer.
# Phase 3 adds the accurate-mode engine (large-v3 + adapter revision SHA).
FAST_MODE_ENGINE_LABEL = "DanScribe Fast — Whisper Medium"

# Third-party attribution shown in the "About" (Credits) dialog. Masterplan
# 2.7 appends the fine-tuned Afrikaans model + dataset entries (CC-BY-4.0)
# below — same list, same CreditsWindow, no UI change.
CREDITS = [
    {
        "name": "OpenAI Whisper",
        "detail": ('Radford et al., "Robust Speech Recognition via '
                   'Large-Scale Weak Supervision", arXiv:2212.04356'),
        "license": "Apache License 2.0",
        "url": "https://arxiv.org/abs/2212.04356",
    },
    # Masterplan 2.7. Attribution language and BibTeX taken verbatim from the
    # published HF model card (huggingface.co/DanieClar/stillscript-whisper-
    # large-v3-afrikaans, License/Attribution/Citation section) rather than
    # paraphrased, so this entry and the model card cannot drift apart. The
    # Afrikaans adapter is CC-BY-4.0; Accurate mode's merged model is only
    # the base model's weights (openai/whisper-large-v3, Apache-2.0, credited
    # above via the Whisper entry) plus this adapter — "nothing was retrained;
    # the only operation performed was merging the published adapter into the
    # published base weights", per the model card's own key statement. This
    # entry covers both the required creator/title/source/license attribution
    # and the "indicate if changes were made" requirement CC-BY-4.0 also asks
    # for, and applies regardless of which HF revision/layout (single-file vs
    # chunked) a given install downloaded, since both carry the same merge.
    {
        "name": "Whisper Large V3 Afrikaans (adapter)",
        "detail": (
            'André Oosthuizen, "Whisper Large V3 Afrikaans", 2025, HuggingFace. '
            "Merged into StillScript's Accurate-mode model without retraining — "
            "the only operation performed was merging this published adapter "
            "into the published openai/whisper-large-v3 base weights."
        ),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/andreoosthuizen/whisper-large-v3-afrikaans",
    },
    {
        "name": "afrikaans-30s (training dataset)",
        "detail": ("André Oosthuizen — training data used for the Whisper "
                   "Large V3 Afrikaans adapter above."),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/datasets/andreoosthuizen/afrikaans-30s",
    },
    # Masterplan 2.12. Same CC-BY-4.0 attribution shape as the two entries
    # above: title (name), author + modification notice (detail), source (url),
    # licence (license). Wording and the cited work are taken from the upstream
    # model card rather than paraphrased. The pipeline is redistributed
    # unmodified from StillScript's own mirror, because the upstream repository
    # is gated and a shipped desktop app cannot make each end user accept its
    # conditions; CC-BY-4.0 expressly permits that redistribution.
    {
        "name": "pyannote speaker-diarization-community-1",
        "detail": (
            "Hervé Bredin and the pyannote.audio team — the speaker-"
            "identification pipeline behind \"Identify different speakers\". "
            "Redistributed unmodified from StillScript's mirror; no "
            "retraining, fine-tuning or conversion was performed. Cites "
            "Plaquet & Bredin, \"Powerset multi-class cross entropy loss for "
            "neural speaker diarization\", Interspeech 2023, and Bredin, "
            "\"pyannote.audio 2.1 speaker diarization pipeline\", Interspeech "
            "2023."
        ),
        "license": "CC-BY-4.0",
        "url": "https://huggingface.co/pyannote/speaker-diarization-community-1",
    },
]

# Model cache — prevents reloading every time
_model_cache = {}

def get_model(model_name):
    if model_name not in _model_cache:
        _model_cache[model_name] = whisper.load_model(model_name)
    return _model_cache[model_name]


def release_fast_mode_model():
    """Evict Fast mode's cached Whisper model(s) from memory (masterplan 2.5).

    Called before Accurate mode loads large-v3. Medium alone sits at roughly
    1.5 GB resident; large-v3 peaks at ~8.7 GB while generating. Having both
    resident at once is exactly the scenario that risks OOM on ordinary
    hardware, and it is avoidable — Fast and Accurate are never used in the
    same instant, only the same session.

    Clearing the dict (not just deleting individual entries) is what actually
    matters here: _model_cache is the ONLY thing holding a reference to the
    loaded `whisper.Whisper` object, so once every entry is gone CPython's
    refcounting frees the underlying C-allocated tensors immediately — no
    lingering reference elsewhere in this module keeps it alive (get_model()
    returns the object to its caller, but callers do not stash it; they call
    get_model() again next time). The explicit gc.collect() following it
    exists for the same reason PyTorch's own docs recommend it after freeing
    large tensors: reference cycles (e.g. autograd graph fragments, even
    though this app never calls backward()) can otherwise sit uncollected for
    a while under CPython's generational GC, and this is exactly the moment a
    plausibly multi-GB allocation is about to follow it.

    A no-op, not an error, if nothing is loaded yet — Accurate mode used first
    in a fresh session has nothing to release. Switching back to Fast mode
    afterward is unaffected by this: get_model() simply reloads from disk on
    its next call, the same cache-miss path it already always had.
    """
    if not _model_cache:
        return
    released = list(_model_cache.keys())
    _model_cache.clear()
    import gc
    gc.collect()
    logger.info("Released Fast-mode model(s) from memory before loading "
               "Accurate mode: %s", released)


def transcribe_audio(path, *, language, task, model_name):
    """Run Whisper on `path` and return its result dict ({"text", "segments", ...}).

    This is the single seam through which all transcription flows. The three
    branches below are the exact language-dispatch logic (Afrikaans prompt,
    other forced language, auto-detect) that previously lived inline in run().
    Whisper is deterministic at temperature=0, so for fixed inputs this returns
    byte-identical output regardless of where it is called from.
    """
    model = get_model(model_name)

    # Whisper's initial_prompt biases the decoder toward the style/
    # spelling of the prompt text itself, not toward instructions it
    # "understands" — a short meta-sentence gives it little to latch
    # onto. Afrikaans and Dutch share a lot of vocabulary, so smaller
    # models especially can drift into Dutch orthography even with
    # language="af" forced. A longer, natural Afrikaans passage with
    # constructions Dutch doesn't have (like "nie ... nie" double
    # negation) gives the decoder a much stronger signal.
    af_prompt = (
        "Goeiemôre almal, baie dankie dat julle almal hier kon kom. "
        "Ons gaan vandag praat oor die begroting en die volgende "
        "stappe wat ons moet neem. Dit is nie maklik nie, maar ons "
        "sal dit saam reg kry. Het julle dalk enige vrae daaroor?"
    )
    if language == "af":
        # Tried condition_on_previous_text=False here to stop Dutch
        # drift from compounding across windows — real-world testing
        # showed it made things worse (hallucinated Unicode noise,
        # injected English words, repetition loops), since the
        # previous-window context also suppresses that kind of
        # hallucination. Reverted; back to Whisper's default.
        result = model.transcribe(path, task=task, language="af", initial_prompt=af_prompt)
    elif language:
        result = model.transcribe(path, task=task, language=language)
    else:
        result = model.transcribe(path, task=task)
    return result


def transcribe_audio_accurate(path, *, language, task, model_dir=None, **engine_kwargs):
    """Accurate-mode counterpart to transcribe_audio() — the second entry point.

    Runs Whisper large-v3 + the Afrikaans adapter through a direct
    transformers `generate()` call (see accurate_engine.py; the CT2 /
    faster-whisper path for that model is known broken and must never be
    used). Returns the same {"text", "segments", ...} shape as the Fast seam,
    so diarization and the exporters are unaffected by which engine ran.

    Two things this deliberately does NOT do:
      * it is not wired into run() — activating the Accurate button is
        masterplan item 2.3;
      * it never falls back to Fast/Medium. If the engine is unavailable the
        AccurateEngineUnavailable error propagates, because silently
        substituting a weaker model is exactly the failure this product
        cannot ship.

    accurate_engine is imported lazily so that its heavy dependencies
    (torch/transformers, absent from the Fast-mode runtime) can never break
    startup or the Fast path.

    release_fast_mode_model() runs first (masterplan 2.5) — before large-v3's
    from_pretrained() call, not before this whole function, so a download that
    accurate_engine.load_engine() needs to wait on never costs Fast mode its
    warm cache for nothing. If this raises (e.g. loading large-v3 hits real
    memory pressure), Fast mode's cache is already empty either way; the next
    Fast-mode run just reloads Medium from disk, same as first launch.
    """
    release_fast_mode_model()
    import accurate_engine
    return accurate_engine.transcribe(
        path, language=language, task=task, model_dir=model_dir, **engine_kwargs
    )


def build_provenance(*, mode, language_label, task, diarized, num_speakers=None, engine=None,
                     model_id=None, model_revision=None, model_layout=None,
                     guard_verification=None):
    """Describe which engine produced a transcript, for the audit footer.

    Returns a plain dict so callers can extend it without any other caller
    changing — the same seam principle as transcribe_audio(). Callers pass
    display-level values (the language *label* the user picked, not the ISO
    code) so the footer reads the way the operator set it up.

    `engine` defaults to FAST_MODE_ENGINE_LABEL so the existing Fast-mode call
    site is unaffected. Accurate mode passes its own engine's label explicitly
    — accurate_engine.transcribe()'s result dict already carries it as
    result["engine"] (ACCURATE_ENGINE_LABEL), so the caller never needs to
    import accurate_engine just to get this string. Getting this wrong would
    mean an Accurate-mode transcript's audit footer claiming it was produced by
    Fast mode instead — a provenance product exists specifically to prevent
    that kind of quiet misattribution.

    model_id / model_revision / model_layout / guard_verification (masterplan
    2.6) are Accurate-mode-only and additive: they default to None and are
    only included in the returned dict when a caller actually passes them, so
    the Fast-mode call site (which never passes them) gets a provenance dict
    identical to before this wave — and format_provenance_lines() below only
    ever renders lines for the ones that are present.
    """
    from datetime import datetime
    provenance = {
        "engine": engine or FAST_MODE_ENGINE_LABEL,
        "mode": mode,
        "language": language_label,
        "task": task,
        "diarized": diarized,
        "num_speakers": num_speakers if diarized else None,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    if model_id is not None:
        provenance["model_id"] = model_id
    if model_revision is not None:
        provenance["model_revision"] = model_revision
    if model_layout is not None:
        provenance["model_layout"] = model_layout
    if guard_verification is not None:
        provenance["guard_verification"] = guard_verification
    return provenance


def format_provenance_lines(provenance):
    """Render a provenance dict as an ordered list of 'Label: value' strings,
    shared by the .txt footer and the .docx footer so both stay in sync."""
    task_label = "Translate to English" if provenance.get("task") == "translate" else "Original language"
    if provenance.get("diarized"):
        spk = provenance.get("num_speakers")
        diar_label = f"Yes (up to {spk} speakers)" if spk else "Yes"
    else:
        diar_label = "No"
    lines = [
        f"Engine: {provenance.get('engine')}",
        f"Mode: {provenance.get('mode')}",
        f"Language setting: {provenance.get('language')}",
        f"Task: {task_label}",
        f"Speaker identification: {diar_label}",
    ]
    # Accurate-mode-only lines (masterplan 2.6). Only appended when present,
    # so Fast-mode's footer is unchanged — build_provenance()'s Fast-mode
    # call site never passes these keys, so they are simply absent here.
    if provenance.get("model_id"):
        lines.append(f"Model: {provenance['model_id']}")
    if provenance.get("model_revision"):
        lines.append(f"Model revision: {provenance['model_revision']}")
    if provenance.get("model_layout"):
        lines.append(f"Download layout: {provenance['model_layout']}")
    if provenance.get("guard_verification"):
        lines.append(f"Guard verification: {provenance['guard_verification']}")
    lines.append(f"Generated: {provenance.get('timestamp')}")
    return lines


# Below this duration, _diarize() uses the whole-file librosa.load path,
# which keeps speaker labels byte-identical to previous behaviour — the
# common case, where label accuracy matters most and memory was never a
# problem (measured delta ~1.1 GB at this length). At/above it, memory
# actually bites on long recordings (multi-GB peaks observed on hour-plus
# files), so _diarize() switches to the decode-once-to-temp-WAV path, which
# can shift a speaker label by a small margin — acceptable on long
# recordings where individual turn boundaries aren't scrutinized.
# ⚠ UNUSED as of masterplan 2.12 — kept, not deleted, deliberately.
#
# This gate and _probe_audio_duration_seconds() below now have NO callers. They
# existed to bound memory while extracting one feature vector per Whisper
# segment: short files loaded whole via librosa.load, long ones decoded once to
# a temp WAV and seeked per segment. The pyannote backend does no per-segment
# feature extraction at all, so there is nothing left to gate — the branch was
# not "removed as a simplification", it lost its subject.
#
# Removing these two, and the frozen-app PATH note at the top of this file that
# mentions them, is a tidy-up that was explicitly scoped OUT of 2.12 (backend
# swap only). Do not assume the gate is live: it is not.
DIARIZE_LONG_FILE_THRESHOLD_SEC = 20 * 60  # 20 minutes


def _probe_audio_duration_seconds(path):
    """Return audio duration in seconds via ffprobe container metadata only
    (no decode), so the check itself stays cheap even on very long files.
    Returns None if the duration can't be determined; callers should treat
    that as "unknown" and fall back to the safe default path.
    """
    import subprocess
    try:
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", path],
            capture_output=True, text=True, timeout=10,
        )
        return float(out.stdout.strip())
    except Exception:
        return None

# ─────────────────────────────────────────────
#  SETTINGS WINDOW
# ─────────────────────────────────────────────

class SettingsWindow(ctk.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("DanScribe AI — Settings")
        self.geometry("520x320")
        self.resizable(False, False)
        self.grab_set()

        config = load_config()

        ctk.CTkLabel(self, text="⚙️ Settings", font=("Arial", 20, "bold")).pack(pady=20)

        ctk.CTkLabel(self, text="Claude API Key:", font=("Arial", 13)).pack(pady=(10, 2))
        self.api_entry = ctk.CTkEntry(self, width=420, show="•", placeholder_text="sk-ant-...")
        self.api_entry.pack(pady=5)
        existing_key = _get_api_key(config)
        if existing_key:
            self.api_entry.insert(0, existing_key)

        # Clickable link
        link_label = ctk.CTkLabel(
            self,
            text="🔗 Get your API key at console.anthropic.com",
            font=("Arial", 11),
            text_color="#4da6ff",
            cursor="hand2"
        )
        link_label.pack(pady=4)
        link_label.bind("<Button-1>", lambda e: webbrowser.open("https://console.anthropic.com"))

        ctk.CTkLabel(
            self,
            text="Your key is stored locally on your device only.",
            font=("Arial", 10),
            text_color="gray"
        ).pack(pady=2)

        ctk.CTkButton(
            self, text="Save Settings", command=self.save,
            width=200, fg_color="#1f538d"
        ).pack(pady=20)

    def save(self):
        config = load_config()
        _set_api_key(config, self.api_entry.get().strip())
        save_config(config)
        messagebox.showinfo("DanScribe AI", "Settings saved!")
        self.destroy()

# ─────────────────────────────────────────────
#  CREDITS / ABOUT WINDOW
# ─────────────────────────────────────────────

class CreditsWindow(ctk.CTkToplevel):
    """Attribution surface for the models/datasets DanScribe builds on.

    Driven entirely by the module-level CREDITS list — masterplan 2.7 added
    the fine-tuned Afrikaans model + dataset entries (CC-BY-4.0) by appending
    two dicts, with no change to this window.
    """
    def __init__(self, parent):
        super().__init__(parent)
        self.title("DanScribe AI — About / Credits")
        self.geometry("560x420")
        self.resizable(False, False)
        self.grab_set()

        ctk.CTkLabel(self, text="ℹ️ About DanScribe AI", font=("Arial", 20, "bold")).pack(pady=(20, 4))
        ctk.CTkLabel(
            self,
            text="DanScribe is built on the following open-source work:",
            font=("Arial", 12),
            text_color="gray",
        ).pack(pady=(0, 10))

        frame = ctk.CTkScrollableFrame(self, width=500, height=290)
        frame.pack(pady=5, padx=20, fill="both", expand=True)

        for entry in CREDITS:
            card = ctk.CTkFrame(frame)
            card.pack(fill="x", pady=8, padx=4)

            ctk.CTkLabel(
                card, text=entry["name"], font=("Arial", 14, "bold"), anchor="w"
            ).pack(fill="x", padx=12, pady=(10, 2))
            ctk.CTkLabel(
                card, text=entry["detail"], font=("Arial", 11), text_color="gray",
                anchor="w", justify="left", wraplength=460
            ).pack(fill="x", padx=12, pady=2)
            ctk.CTkLabel(
                card, text=f"License: {entry['license']}", font=("Arial", 11), anchor="w"
            ).pack(fill="x", padx=12, pady=2)

            url = entry.get("url")
            if url:
                link = ctk.CTkLabel(
                    card, text=f"🔗 {url}", font=("Arial", 11),
                    text_color="#4da6ff", cursor="hand2", anchor="w"
                )
                link.pack(fill="x", padx=12, pady=(2, 10))
                link.bind("<Button-1>", lambda e, u=url: webbrowser.open(u))

        ctk.CTkButton(
            self, text="Close", command=self.destroy,
            width=160, fg_color="#1f538d"
        ).pack(pady=(6, 16))

# ─────────────────────────────────────────────
#  SPEAKER NAME ASSIGNMENT WINDOW
# ─────────────────────────────────────────────

class NameAssignWindow(ctk.CTkToplevel):
    def __init__(self, parent, speakers: list, callback):
        super().__init__(parent)
        self.title("Assign Speaker Names")
        self.geometry("450x420")
        self.resizable(False, False)
        self.grab_set()
        self.callback = callback
        self.entries = {}

        ctk.CTkLabel(self, text="🎤 Assign Names to Speakers", font=("Arial", 18, "bold")).pack(pady=20)
        ctk.CTkLabel(
            self,
            text="Leave blank to keep 'Speaker 1', 'Speaker 2', etc.",
            font=("Arial", 11),
            text_color="gray"
        ).pack(pady=(0, 15))

        frame = ctk.CTkScrollableFrame(self, width=380, height=220)
        frame.pack(pady=5, padx=20)

        for speaker in speakers:
            row = ctk.CTkFrame(frame, fg_color="transparent")
            row.pack(fill="x", pady=6)
            ctk.CTkLabel(row, text=f"{speaker}:", width=110, anchor="w", font=("Arial", 12, "bold")).pack(side="left", padx=5)
            entry = ctk.CTkEntry(row, width=230, placeholder_text="Enter name...")
            entry.pack(side="left")
            self.entries[speaker] = entry

        ctk.CTkButton(
            self, text="✅ Confirm & Continue",
            command=self.confirm, width=250, fg_color="#1f538d", height=45
        ).pack(pady=20)

    def confirm(self):
        name_map = {}
        for speaker, entry in self.entries.items():
            name = entry.get().strip()
            name_map[speaker] = name if name else speaker
        self.callback(name_map)
        self.destroy()


class AccurateConsentDialog(ctk.CTkToplevel):
    """First-activation consent screen for Accurate mode's model download.

    Shown once, before any network access, using real numbers from
    accurate_model_download.describe_download() — never a number remembered
    from an earlier investigation, which could be stale by the time a user
    actually clicks the button. `on_confirm`/`on_decline` follow the same
    callback-then-destroy convention as NameAssignWindow.confirm() above.
    """

    def __init__(self, parent, info, *, on_confirm, on_decline):
        super().__init__(parent)
        self.title("StillScript — Accurate Mode Setup")
        self.geometry("480x420")
        self.resizable(False, False)
        self.grab_set()
        self._on_confirm = on_confirm
        self._on_decline = on_decline

        # Kept as attributes (not just packed and forgotten) so the real
        # numbers behind the dialog can be checked directly rather than by
        # parsing rendered label text.
        self.size_gb = info["total_bytes"] / (1024 ** 3)
        self.info_labels = []

        ctk.CTkLabel(
            self, text="🎯 Accurate Mode — one-time setup",
            font=("Arial", 17, "bold"),
        ).pack(pady=(20, 4))
        ctk.CTkLabel(
            self,
            text="Accurate mode uses a larger, more precise speech model that\n"
                 "is not included in the installer. It downloads once, the\n"
                 "first time you use Accurate mode.",
            font=("Arial", 12), justify="center",
        ).pack(pady=(0, 16))

        info_frame = ctk.CTkFrame(self, fg_color="gray20")
        info_frame.pack(fill="x", padx=24, pady=4)

        def _row(icon_text):
            label = ctk.CTkLabel(
                info_frame, text=icon_text, font=("Arial", 12),
                justify="left", anchor="w", wraplength=400,
            )
            label.pack(fill="x", padx=14, pady=8)
            self.info_labels.append(label)

        _row(f"📦  Size: about {self.size_gb:.1f} GB")
        _row("🌐  Requires an internet connection.")
        _row(
            "⏱  Time: this can take anywhere from tens of minutes to several\n"
            "hours, depending on your connection speed — on a slower line,\n"
            "expect it to take hours, not minutes. You can leave StillScript\n"
            "running in the background; if it's interrupted, it picks up\n"
            "where it left off rather than starting over."
        )

        # Confidentiality point — stated plainly, in its own visually distinct
        # block, not folded into the paragraph above as fine print.
        ctk.CTkLabel(
            self,
            text="🔒 This download is one-way: it only fetches the language\n"
                 "model's weights onto your computer. It has nothing to do with\n"
                 "your recordings — your audio and transcripts are never\n"
                 "uploaded, now or ever.",
            font=("Arial", 11, "bold"), text_color="#8fd19e",
            justify="center", wraplength=420,
        ).pack(pady=(16, 4))

        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=20)
        self.decline_btn = ctk.CTkButton(
            btn_frame, text="Not now", command=self._decline,
            width=150, height=42, fg_color="gray30",
        )
        self.decline_btn.grid(row=0, column=0, padx=8)
        self.confirm_btn = ctk.CTkButton(
            btn_frame, text="⬇ Download & Continue", command=self._confirm,
            width=220, height=42, fg_color="#1f538d", font=("Arial", 13, "bold"),
        )
        self.confirm_btn.grid(row=0, column=1, padx=8)

    def _confirm(self):
        self.destroy()
        self._on_confirm()

    def _decline(self):
        self.destroy()
        self._on_decline()


class AccurateDownloadProgressDialog(ctk.CTkToplevel):
    """Progress display for the Accurate-mode model download.

    Rendering is driven entirely by DownloadProgress objects handed to
    update_progress() — the same shape accurate_model_download already
    defines and emits, not a new one invented here. `on_cancel` is called once,
    from the Tk main thread (this dialog's Cancel button is a normal Tkinter
    command), and only sets a threading.Event; it does not stop anything by
    itself, cancellation happens in the download function on its own schedule.
    """

    def __init__(self, parent, *, on_cancel):
        super().__init__(parent)
        self.title("StillScript — Downloading Accurate-mode model")
        self.geometry("460x220")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", lambda: None)  # no window-close shortcut around Cancel
        self.grab_set()
        self._on_cancel = on_cancel

        self.phase_label = ctk.CTkLabel(
            self, text="Checking the download…", font=("Arial", 13, "bold"),
        )
        self.phase_label.pack(pady=(24, 8))

        self.bar = ctk.CTkProgressBar(self, width=380)
        self.bar.pack(pady=6)
        self.bar.set(0)

        self.detail_label = ctk.CTkLabel(
            self, text="", font=("Arial", 11), text_color="gray",
        )
        self.detail_label.pack(pady=(4, 10))

        self.cancel_btn = ctk.CTkButton(
            self, text="Cancel", command=self._cancel,
            width=140, height=36, fg_color="gray30",
        )
        self.cancel_btn.pack(pady=6)

        self._cancelling = False

    def update_progress(self, progress):
        """Render one DownloadProgress sample. Safe to call only on the main thread."""
        if self._cancelling:
            return
        self.phase_label.configure(text=progress.message)
        self.bar.set(max(0.0, min(1.0, progress.fraction)))

        if progress.total_bytes:
            done_mib = progress.downloaded_bytes / (1024 ** 2)
            total_mib = progress.total_bytes / (1024 ** 2)
            rate_mib = progress.bytes_per_second / (1024 ** 2)
            eta = "…" if progress.eta_seconds is None else _format_eta(progress.eta_seconds)
            self.detail_label.configure(
                text=f"{done_mib:,.0f} / {total_mib:,.0f} MiB   "
                     f"{rate_mib:.2f} MiB/s   ETA {eta}"
            )
        else:
            self.detail_label.configure(text="")

    def _cancel(self):
        # Chunk-level cancellation (see accurate_model_download's RESUME notes):
        # this takes effect after the chunk in flight finishes, not instantly.
        # Say so, and disable the button so a second click can't do anything odd.
        self._cancelling = True
        self.cancel_btn.configure(state="disabled", text="Cancelling…")
        self.phase_label.configure(
            text="Cancelling — finishing the current part first…")
        self._on_cancel()


def _format_eta(seconds):
    seconds = max(0, int(seconds))
    hours, rem = divmod(seconds, 3600)
    minutes, secs = divmod(rem, 60)
    if hours:
        return f"{hours}h {minutes}m"
    if minutes:
        return f"{minutes}m {secs}s"
    return f"{secs}s"


# ─────────────────────────────────────────────
#  MAIN APPLICATION
# ─────────────────────────────────────────────

class DanScribeApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("DanScribe AI — Professional v3.1.0")
        self.geometry("620x920")
        self.resizable(False, False)

        self.config_data = load_config()
        self.current_transcript = None
        self.diarized_segments = None
        self.speaker_name_map = {}

        self._build_ui()

    # ── THREAD-SAFE UI UPDATES ──────────────

    def _ui(self, func, *args, **kwargs):
        """Marshal a UI update onto the main thread.

        Tkinter is not thread-safe; worker threads must not touch widgets
        directly. `after(0, ...)` queues the call on the main event loop.
        """
        try:
            self.after(0, lambda: func(*args, **kwargs))
        except Exception as e:
            logger.error("UI update failed: %s", e)

    # ── BUILD UI ────────────────────────────

    def _build_ui(self):
        ctk.set_appearance_mode("dark")

        # ── Top banner: Logo + SA Flag side by side ──
        banner = ctk.CTkFrame(self, fg_color="transparent")
        banner.pack(pady=(10, 0), fill="x", padx=20)

        # Logo (left)
        logo_loaded = False
        try:
            img = Image.open(resource_path("logo.jpg"))
            logo_img = ctk.CTkImage(light_image=img, dark_image=img, size=(130, 130))
            ctk.CTkLabel(banner, image=logo_img, text="").pack(side="left", padx=(10, 0))
            logo_loaded = True
        except Exception:
            ctk.CTkLabel(banner, text="DanScribe AI", font=("Arial", 24, "bold")).pack(side="left", padx=10)

        # Spacer
        ctk.CTkLabel(banner, text="", fg_color="transparent").pack(side="left", expand=True)

        # SA Flag (right) — embedded as base64
        try:
            FLAG_B64 = "/9j/4AAQSkZJRgABAQEAlgCWAAD/4QBsRXhpZgAASUkqAAgAAAADADEBAgAHAAAAMgAAABICAwACAAAAAQABAGmHBAABAAAAOgAAAAAAAABHb29nbGUAAAMAAJAHAAQAAAAwMjIwAqAEAAEAAAD6BQAAA6AEAAEAAAB5AwAAAAAAAP/bAEMAAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAf/bAEMBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAf/AABEIA3kF+gMBEQACEQEDEQH/xAAfAAABBQEBAQEBAQAAAAAAAAAAAQIDBAUGBwgJCgv/xAC1EAACAQMDAgQDBQUEBAAAAX0BAgMABBEFEiExQQYTUWEHInEUMoGRoQgjQrHBFVLR8CQzYnKCCQoWFxgZGiUmJygpKjQ1Njc4OTpDREVGR0hJSlNUVVZXWFlaY2RlZmdoaWpzdHV2d3h5eoOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4eLj5OXm5+jp6vHy8/T19vf4+fr/xAAfAQADAQEBAQEBAQEBAAAAAAAAAQIDBAUGBwgJCgv/xAC1EQACAQIEBAMEBwUEBAABAncAAQIDEQQFITEGEkFRB2FxEyIygQgUQpGhscEJIzNS8BVictEKFiQ04SXxFxgZGiYnKCkqNTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqCg4SFhoeIiYqSk5SVlpeYmZqio6Slpqeoqaqys7S1tre4ubrCw8TFxsfIycrS09TV1tfY2dri4+Tl5ufo6ery8/T19vf4+fr/2gAMAwEAAhEDEQA/AOM/az+I/wDwUQ/Zy8SG+sv20/2v9d+GGu3ki+GfEx/aK+MTy2Ez75l8OeImi8XrHb6xbxK5tbkLHa63axPeWixTxahYWH8V/RV+lfkP0huHFgcbWo5J4mZHhKcuJOHFXlClj6UOSlLiHh5VKjqYjKcRVlFYrDOVTE5NiqscLi5VaFbAY7H/AMyfTH8AvG76MfFLzDAcfeIPEHhNxDjaseFeKpcS57Otl1afPWjwxxPKljI0sNnWGoxm8Ji1GlhM+wlGeNwcaOJo5ll2XfH3/DfX7df/AEep+1p/4kd8Yv8A5sq/rj21X/n7U/8AA5f5n8T/APESvEX/AKL7jX/xKs9/+bw/4b6/br/6PU/a0/8AEjvjF/8ANlR7ar/z9qf+By/zD/iJXiL/ANF9xr/4lWe//N4f8N9ft1/9Hqftaf8AiR3xi/8Amyo9tV/5+1P/AAOX+Yf8RK8Rf+i+41/8SrPf/m8P+G+v26/+j1P2tP8AxI74xf8AzZUe2q/8/an/AIHL/MP+IleIv/Rfca/+JVnv/wA3h/w31+3X/wBHqftaf+JHfGL/AObKj21X/n7U/wDA5f5h/wARK8Rf+i+41/8AEqz3/wCbw/4b6/br/wCj1P2tP/EjvjF/82VHtqv/AD9qf+By/wAw/wCIleIv/Rfca/8AiVZ7/wDN4f8ADfX7df8A0ep+1p/4kd8Yv/myo9tV/wCftT/wOX+Yf8RK8Rf+i+41/wDEqz3/AObw/wCG+v26/wDo9T9rT/xI74xf/NlR7ar/AM/an/gcv8w/4iV4i/8ARfca/wDiVZ7/APN4f8N9ft1/9Hqftaf+JHfGL/5sqPbVf+ftT/wOX+Yf8RK8Rf8AovuNf/Eqz3/5vD/hvr9uv/o9T9rT/wASO+MX/wA2VHtqv/P2p/4HL/MP+IleIv8A0X3Gv/iVZ7/83h/w31+3X/0ep+1p/wCJHfGL/wCbKj21X/n7U/8AA5f5h/xErxF/6L7jX/xKs9/+bw/4b6/br/6PU/a0/wDEjvjF/wDNlR7ar/z9qf8Agcv8w/4iV4i/9F9xr/4lWe//ADeH/DfX7df/AEep+1p/4kd8Yv8A5sqPbVf+ftT/AMDl/mH/ABErxF/6L7jX/wASrPf/AJvD/hvr9uv/AKPU/a0/8SO+MX/zZUe2q/8AP2p/4HL/ADD/AIiV4i/9F9xr/wCJVnv/AM3h/wAN9ft1/wDR6n7Wn/iR3xi/+bKj21X/AJ+1P/A5f5h/xErxF/6L7jX/AMSrPf8A5vD/AIb6/br/AOj1P2tP/EjvjF/82VHtqv8Az9qf+By/zD/iJXiL/wBF9xr/AOJVnv8A83h/w31+3X/0ep+1p/4kd8Yv/myo9tV/5+1P/A5f5h/xErxF/wCi+41/8SrPf/m8P+G+v26/+j1P2tP/ABI74xf/ADZUe2q/8/an/gcv8w/4iV4i/wDRfca/+JVnv/zeH/DfX7df/R6n7Wn/AIkd8Yv/AJsqPbVf+ftT/wADl/mH/ESvEX/ovuNf/Eqz3/5vD/hvr9uv/o9T9rT/AMSO+MX/AM2VHtqv/P2p/wCBy/zD/iJXiL/0X3Gv/iVZ7/8AN4f8N9ft1/8AR6n7Wn/iR3xi/wDmyo9tV/5+1P8AwOX+Yf8AESvEX/ovuNf/ABKs9/8Am8P+G+v26/8Ao9T9rT/xI74xf/NlR7ar/wA/an/gcv8AMP8AiJXiL/0X3Gv/AIlWe/8AzeH/AA31+3X/ANHqftaf+JHfGL/5sqPbVf8An7U/8Dl/mH/ESvEX/ovuNf8AxKs9/wDm8P8Ahvr9uv8A6PU/a0/8SO+MX/zZUe2q/wDP2p/4HL/MP+IleIv/AEX3Gv8A4lWe/wDzeH/DfX7df/R6n7Wn/iR3xi/+bKj21X/n7U/8Dl/mH/ESvEX/AKL7jX/xKs9/+bw/4b6/br/6PU/a0/8AEjvjF/8ANlR7ar/z9qf+By/zD/iJXiL/ANF9xr/4lWe//N4f8N9ft1/9Hqftaf8AiR3xi/8Amyo9tV/5+1P/AAOX+Yf8RK8Rf+i+41/8SrPf/m8P+G+v26/+j1P2tP8AxI74xf8AzZUe2q/8/an/AIHL/MP+IleIv/Rfca/+JVnv/wA3h/w31+3X/wBHqftaf+JHfGL/AObKj21X/n7U/wDA5f5h/wARK8Rf+i+41/8AEqz3/wCbw/4b6/br/wCj1P2tP/EjvjF/82VHtqv/AD9qf+By/wAw/wCIleIv/Rfca/8AiVZ7/wDN4f8ADfX7df8A0ep+1p/4kd8Yv/myo9tV/wCftT/wOX+Yf8RK8Rf+i+41/wDEqz3/AObw/wCG+v26/wDo9T9rT/xI74xf/NlR7ar/AM/an/gcv8w/4iV4i/8ARfca/wDiVZ7/APN4f8N9ft1/9Hqftaf+JHfGL/5sqPbVf+ftT/wOX+Yf8RK8Rf8AovuNf/Eqz3/5vD/hvr9uv/o9T9rT/wASO+MX/wA2VHtqv/P2p/4HL/MP+IleIv8A0X3Gv/iVZ7/83h/w31+3X/0ep+1p/wCJHfGL/wCbKj21X/n7U/8AA5f5h/xErxF/6L7jX/xKs9/+bw/4b6/br/6PU/a0/wDEjvjF/wDNlR7ar/z9qf8Agcv8w/4iV4i/9F9xr/4lWe//ADeH/DfX7df/AEep+1p/4kd8Yv8A5sqPbVf+ftT/AMDl/mH/ABErxF/6L7jX/wASrPf/AJvO8+GP/BTn9vf4W+PvC/j+x/av+PXjGfwxqkWonwr8TPi18QviF4E8Q24DRXekeJvCPibxLf6Tqum39rJNbyFoYr+yaRNQ0i+03VrWy1C2ca9aMlL2k3bpKUmn6pvU9LKPFzxJyfMsHmdPjTiTHywlZVfqeb53mmZ5dio6qdDF4LF4upRrUqkHKL0jVptqrQqUq0KdWH92X/BPj/goX8IP+CgXwmj8Y+CpovDXxI8NwWVr8VPhPfX0dxrvgvWbhCq3lnIVgk1zwdq80U8nh3xLBbxR3UaS2GowadrdlqGm23q0a0a0brSS+KPVP9U+j/W5/ot4ZeJ2R+JmSrH5fKOFzXCxpwznJalRSxOX15L44O0XiMDXkpPC4uMVGaTp1Y0sRTq0off1bH6WFABQAUAFAFPUbCHVNPvtMuZL2K31GzubG4l07UdQ0jUI4buF4JZLHVtJubLVNMvESRmttQ028tL+zmCXNncwXEccq5VqUcRRrUKjqxhWpVKU5Ua1bDVlGpFwk6WIw9SliKFVKTdOtQq061KdqlKpCcYyXRhMTUwWKw2Moxw862ExFHE0oYvCYXH4WdShUjVhHE4HHUcTgsbh5SilWwmMw9fC4mm5UcRRq0pzhL+TD/go58Nf28/2MPGL+KfC37Vv7VHir9nzxXqckfhLxXL8c/ilcX3hS/uPMnj8FeNJYvEypFqMMayf2LrJSK08R2ULSxiDU7fULC2/zl8cMh8X/C7M3mOXeIniHmPBeY13HLsxlxbxBOtltafNOOVZrKOPSjXjFP6rimo0sfSi5RUMRCvRp/7q/RF4x+jD9Ifh+OSZ54HeCOSeK2R4OM89yOHhpwTSwue4Wly058R8OwqZM5TwlSbg8xy5SqYjKMTUUJOrg6uFxVb8wP8AhtD9sT/o7D9pb/w+3xR/+amvwH/iKXib/wBHF47/APEu4g/+eB/aP/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkD/htD9sT/o7D9pb/wAPt8Uf/mpo/wCIpeJv/RxeO/8AxLuIP/ngH/Eu/wBH/wD6MZ4O/wDisuCv/nIH/DaH7Yn/AEdh+0t/4fb4o/8AzU0f8RS8Tf8Ao4vHf/iXcQf/ADwD/iXf6P8A/wBGM8Hf/FZcFf8AzkD/AIbQ/bE/6Ow/aW/8Pt8Uf/mpo/4il4m/9HF47/8AEu4g/wDngH/Eu/0f/wDoxng7/wCKy4K/+cgf8Noftif9HYftLf8Ah9vij/8ANTR/xFLxN/6OLx3/AOJdxB/88A/4l3+j/wD9GM8Hf/FZcFf/ADkD/htD9sT/AKOw/aW/8Pt8Uf8A5qaP+IpeJv8A0cXjv/xLuIP/AJ4B/wAS7/R//wCjGeDv/isuCv8A5yB/w2h+2J/0dh+0t/4fb4o//NTR/wARS8Tf+ji8d/8AiXcQf/PAP+Jd/o//APRjPB3/AMVlwV/85A/4bQ/bE/6Ow/aW/wDD7fFH/wCamj/iKXib/wBHF47/APEu4g/+eAf8S7/R/wD+jGeDv/isuCv/AJyB/wANoftif9HYftLf+H2+KP8A81NH/EUvE3/o4vHf/iXcQf8AzwD/AIl3+j//ANGM8Hf/ABWXBX/zkP7l/wBknWtZ8Sfsqfsy+IvEWranr/iDX/2fPgxrWu67rV/darrOtazqvw48N32p6tq2p30s97qOp6jezz3l/f3k811eXU0txcSyTSO5/wBafDfFYrHeHnAWNxuJr4zG4zgvhfFYvF4qtUxGKxWKxGR4GtiMTicRWlOrXr16s51a1arOVSrUlKc5SlJt/wDNP475dl+T+OHjLlOU4HB5XlWV+K3iHl2WZZl2FoYLL8uy/BcXZxhsHgcDg8NClhsJg8JhqVPD4XC4enToUKFOFKlCFOEYr+dfxd4R8N+PPDer+EPF+kWmu+HNdtHstU0y+QvDcQuQysrKVlt7m3lWO4s7y3kiu7K6ihurWaG4hjkX/lE4U4s4i4G4iynizhPNsXkfEOSYuGNy3M8FPkrUK0LxlGUZKVKvh69OU8Pi8JiIVcLjMLVrYXFUa2HrVKcv9S+M+DOF/EPhfOuDONMlwXEHDPEGCqYDNspx9NzoYmhNqUZRlFxrYfFYerGnicFjcNUpYvA4ulRxeEr0cTRpVY/zm/tZ/sl+JP2cvEhv7D7Xrvwv128kXw14laPfPp8z75V8O+IzEixQatBErm1uwkVrrdrE93apDcRX9hYf9Cv0VfpV8O/SG4d+oY/6pkfibkmEhLiPhyNRwo5hRhyUpcQ8PKrOVWvlVerKKxWFc6uKybE1Y4XFTrUKuBx+O/5kPpkfQ34o+jHxQ8yy767xD4TcQY2pDhbimdPnr5ZXnz1o8M8TyowjSw2cYelGbweMUKWEz7CUp4vBwo4mjmOXZf8AHlf10fxKFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAe3/s7ftFfFv9lf4s+GfjR8FfFNz4V8beGZzskXdPpOu6TO8Z1Pwz4n0syRwa34b1mKJIdS0y4IyUgvbOaz1Ozsb61uE5U5KUXZr8fJ90+p9DwvxTnfB2dYTP8gxk8HmGElo9ZUMTRk17bCYyjdRxGExCSjVpT7RqU5U61OnUh/oPf8E6/wDgo38Iv+CgvwtXX/DElv4U+LfhWysY/ir8Jry8STVPDeoTIkTa3oTuyza74G1W7Eg0fXI41lgcjTNZgsdUTyZfXo1o1o3Wkl8Uez7run0f3n+m3hd4p5J4m5P9Zwjjgs7wdOms5yWpUTrYSrJJfWMO3aWJy6tO/sMQleL/AHNeNOsuV/onWx+ohQAUAFABQByXjzwH4P8Aif4P8Q+APH/h7TfFfg7xXps+ka/oGrwefY6jY3AG5HAKywzwyLHc2d7bSQ3thew297ZXFvd28M8fm5vlGWZ/lmNybOcFQzHLMxoTw2MweJhz0q9Ge6eqlGcZJTpVYSjVo1Ywq0pwqwjNe7wxxPxBwZxBlPFPC2bYzI+IcjxlLH5VmuAq+yxOExVJu0otqUKlOpCU6OIw9aFTD4rD1KuGxNKrh6tWnP8Ai1/4KNf8E5/GH7F3jJvEnhpdS8Vfs/8AizUpI/B/i+WMz3vhi+n8ydPBPjWSGNYodWgiSQ6Pq+yGz8S2MD3EC2+o2+p6bZf5a+N/gjmfhbmjx+AVfMeDMxryjlmZSXPVy+tPmmsqzWUUoxxEIqX1XE2jSx9GDnBQrwr0KX/RN9Ef6W/D/wBIjh9ZPnEsHkfipkWDhPP8hhP2WGzrDUuSnPiPh2FSbnVwNWpKP9oYDmqYjJsTVjSqyq4Stg8ZifzGr8DP7NCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Qa/Yv/wCTO/2T/wDs2n4E/wDqrvC1f7O+Fv8AybLw6/7IThH/ANZ/Lz/lR+kR/wApAeOf/Z4vE3/1tc7PwNr/AJKj/WY5zxb4S8N+O/Dmr+EfF2kWeu+HddtJLLVNLvoy8FxC+GVlZSssFzBKqXFpeW8kV1Z3UUN1azQ3EMcq/Q8KcV8RcD8RZTxZwpm2LyPiHI8XDG5ZmeCnyVsPWheMlJSUqdfD16cp0MXhcRCrhsZhqtbDYqlVoValOXzPGXBvDHiDwxnXBnGeS4LiHhniDBVMBm2U5hTdTD4nD1LSjKMouNXD4nD1Y08Tg8bhqlHGYHF0qOLwlejiaNKrD+c/9rT9krxJ+zn4jOo6d9s134Xa7eOnhvxI8e+fTZ5N8q+HPEbRIsUGqwRq5tLsLFa63axPdWqQ3EV/YWP/AEJ/RT+lZw99IXh7+z8weEyPxNyTCQnxFw7GfJQzKhBwpS4h4eVWcqlfLK1SUFi8I5VMTk2Kqxw2JlVw9bA47G/8yP0yfoa8UfRj4neaZYsbxD4ScQY2cOGOKZ01UxGV4ip7StHhjieVGEKWHzehShN4LGqFLCZ7hKU8XhIUcVRzDLsB8dV/Xh/EYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB7D8Bfj38Vf2Z/in4X+MnwZ8V33g/wAd+E7wXFjqFqfMtL+zkKrqGha7pzn7LrPh/WLcG01bSL5JLW8t25VJkhmjqE5QkpRdmv6s+6fU93hviTOeEs5wefZDjamBzHBT5qdWGsKtOWlXD4mk/cxGGrx9ytQqJwnF9JKMl/oKf8E2/wDgpV8Kv+Cg/wAMvt2mmx8G/G/whYWv/C0/hPLeeZcadKxjt/8AhLPCT3D/AGnWvAuq3bBba7/e3ug3k0eia7i4fTr/AFf16FeNaPaa+KP6run+Gz8/9M/CnxYybxNyj2lH2eA4hwNOH9s5K53nSk7R+u4JyfPiMurTdoVNamGqSWHxPvOlVr/pXW5+sBQAUAFABQBx/j/wB4N+KXg3xF8PviD4d03xX4N8V6bPpOv6Bq0Pn2V/ZT4ODgpLb3NvKsd1Y31rLBfaffQW99Y3FveW8M8fm5xk+WcQZZjcmzrBUMxyvMaE8NjcHiY89KtSnZ2eqlCcJKNSlWpyhWo1oQrUZwqwhNe/wtxTxDwVxDlPFfCubYzI+IcjxlPH5XmmBqezxOFxFO6urqUKtGrCU6OJw1eFTDYvDVKuGxNKrh6tSnL+LL/gov8A8E6vGX7FnjQ6/wCHxqXiv4A+LNSlj8GeMpY/OvPDt5N5lwngrxrJBGkNvrdtCkh0rVNkFl4nsIJLu0S3v7bVdL07/LLxu8Es08LM0eNwSr5jwbmVeUcrzOUeergasuaayrNZRSjDFQipfVsRaFLH0YSqU1CtDEUKP/RR9En6WvD/ANInh1ZVmrweR+KeRYOE+IeH4T9nh83w9PlpS4j4dhVnKpVy6tUlD69gearicmxVWNCvKrha2CxmK/MyvwY/skKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/0Gv2L/8Akzv9k/8A7Np+BP8A6q7wtX+zvhb/AMmy8Ov+yE4R/wDWfy8/5UfpEf8AKQHjn/2eLxN/9bXOz8Da/wCSo/1mCgDnfFnhPw5468Oav4S8W6RZ674d12zksdU0u+j8yC5gchgQQVkguIJVS4tLu3eK6s7qKG6tZobiGOVfoOFeKuIeCeIcp4r4UzbF5HxBkmLhjcszPBVOSvh68LppqSlTrUK1OU6GKwuIhVw2Lw1WrhsVSq4erUpy+a4x4O4Z8QOGc54N4yybBcQcM8QYKpgM3ynMKftMPisPUakmnFxq0MRQqxp4nB4zD1KWLwWLpUcXhK1HE0aVWH86P7W37JPiP9nPxEdT037Zrvwt128dPDniN033Glzyb5V8OeI2iRY4dUhjVzZ3gWK11q2ia5tliuYr6xs/+hL6KX0reHvpCcPrLcyeEyTxPyTCQlxDw9Gfs6GaUIclOXEPDsas5VK2W1qko/XMG51cTk2JqRw+IlWw1XBY7Gf8yf0yvoacT/Rk4mebZUsbxD4R8QY2pDhniedP2mJyjE1OerHhjiiVKEaVDNaNKM3gceoUsJnuFpTxOFjRxdHMMvwPxvX9fH8QBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHqnwU+NnxP8A2d/ib4V+L/we8Waj4L8f+Dr9b/R9a05wQysDHeaZqdnKHtNW0TVbVpbHWNG1CG40/U7Caa0u4JYpCKqMpQkpRdmuv+fdPr3PZ4f4gzfhbN8HnmRY2rgMywNVVKFek7pp6VKNanK8K+Hrwbp16FWMqVanKUJxaZ/oE/8ABM3/AIKc/DH/AIKC/Dnyj/Z3gn9oHwdptu/xN+Fv2ptroGitW8beBjdSvd6v4J1K7kjSRHe41PwpqFzFouuyTJPout6/69CvGtHtNfFH9V3T/DZ9G/8ATDwl8XMo8Tcqt+6y/ibA0YvN8n53qvdg8wy7nk518vq1Gk03Otgqs44fEuXNQxGK/UOtz9fCgAoAKACgDjfiD8PvBfxW8F+I/h58Q/Dum+LPBnizTZtJ1/QdVh86zvrOYqwOVZJra7tp0iu7C/tJYL7Tr+C2v7C5t7y3gnj8zOcmyviHK8dkudYKhmOV5jQnhsZg8THmpVqUrPo1KFSElGpRrU5QrUK0IVqM4VYQmvoOFeKuIuCOIsp4s4TzfGZHxFkWMp47K80wNT2eIw2Ip3T0alTrUK1OU6GKwteFXDYvDVa2FxVGth61SnP+K7/gop/wTt8afsV+NjreiDUvFfwD8WalLH4J8ayRedd6FdyiS4XwV41eCNIbXXrWBJDpmpBILHxRYW8l9ZJb3ttq2laX/lj42+CeaeFma/W8J7fMeDsyryjlWayjzVMHVlzTWVZrKEVCnjKcFJ4evaFLMKMJVaShVhiMPQ/6K/ol/S04d+kVw6stzN4PI/FHIsJCfEfDsJ8lDNKEOSlLiPh2NWcqlfLK1SUFjMHz1cTkuKqxw+JlVw1bA47GfmlX4Qf2KFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/8md/sn/8AZtPwJ/8AVXeFq/2d8Lf+TZeHX/ZCcI/+s/l5/wAqP0iP+UgPHP8A7PF4m/8Ara52fgbX/JUf6zBQAUAc94r8KeHfHPh3V/Cfi3SLPXfDuu2cljqml30fmQXMEmCDkFZIZ4ZFSe1uoHiubS5iiuraaK4ijkX3+F+KeIOCuIMq4q4VzbF5JxBkmLp43LMzwNT2eIw1eF091KnWo1qcp0MVhq8KmGxeGq1cNiaVWhVqU5fN8X8IcM8fcM5zwdxjk2C4g4Z4gwVTL83yjMKftMNi8NUtJapxq0MRQqxp4jCYvD1KWLwWLpUcXhK1HE0aVWH86X7W/wCyR4j/AGdPER1XSxea78LNcvHTw94idPMuNJuJN8q+HPEjRIscOpRRq5sr4JFa61bRNcW6w3UV7Y2n/Qj9FH6V3D/0hOH1leaPCZJ4oZJhIT4g4fhP2eHzbD0+WnLiHh2NWcqlbLq1SUfruCc6uJybE1I0K8q2Fq4LG4v/AJlPpl/Qz4m+jLxK84yhY3iDwi4gxs4cNcTTh7XE5NiavPVhwxxROlCNKjmlGnGf1DHqFLCZ7haU8Rh4UcXRx+AwfxrX9gH8PBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHpHwi+L3xI+A/xG8K/Fn4S+LNU8E+P/AAZqUeqaB4g0mUJPbzKGjuLW6glWS01LStRtZJrDV9I1GC50zVtNubnT9RtbmzuJoXqMpQkpRdmndP8Arc9XJM8zXhzNcHnWSY2tl+Z4Cqq2GxVF2lGVmpQnFpwq0asHKlXoVYzpV6U50qsJwlKL/v8AP+CX/wDwVH+HH/BQLwANH1X+y/A/7R/g3S4ZfiJ8NluGS11i1iMVtJ48+H/2uWS61Hwne3Mka6hpzy3Oq+D9RuYtL1eW7s7nRNe1316FeNZdprePfbVdbXfy/F/6W+EPjBlXiZlnsK3scv4qwFGMs0ylSahXgrQeZZZzyc6uCqTaVWk5TrYGrNUa7nCeHxOJ/Vytz9mCgAoAKACgDi/iJ8O/BXxZ8E+I/h18RfDmneLPBfizTZtK17QdViMtre2spV1ZXRkntLy0nSK807UbOWC/02/gtr+wube8t4J4/LzvJcq4jyrHZJneCoZjleZUJ4bGYPER5qdWnKzTTTU6dWnNRq0a9OUK1CtCFajOFWEJr6LhPiziPgXiPKOLeEs3xeR8RZFjKeOyvNMFU5K+HrwumpRkpUq+Hr05ToYvCYiFXC4zC1a2FxVGth61SnL+Kv8A4KI/8E8PG37FXjg6vpA1LxX8BfFuozJ4G8cSRCW50e6kElwvgrxo9vGkFp4is4EkbT9QEcFh4p0+CTUNPjt7y21jSNJ/yw8bPBTNfCvNvrOG9vmPB+ZV5LKc2lHmqYWpLmmsqzRwioU8bTgpOjW5YUcwowlWoqFWGJw2H/6Lfom/Sy4c+kXw2sBmDwmR+KGRYSnLiXhuNRwo5hRhyUpcR8OxqzlVr5TiKsorF4Vzq4rJcVVhhcVKrh62Ax+O/Nevwo/sEKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP8AQa/Yv/5M7/ZP/wCzafgT/wCqu8LV/s74W/8AJsvDr/shOEf/AFn8vP8AlR+kR/ykB45/9ni8Tf8A1tc7PwNr/kqP9ZgoAKACgDn/ABV4V8PeN/D2r+E/FmkWeu+Hdds5LHVdLv4/Mt7q3kwecFZIZopFSe1uoHiubS5ihurWaG4hjlX3uGOKOIOC+IMq4p4WzbGZJxBkmLp47LM0wNT2eIw2Ip3V1dShVo1acp0MThq8KmGxeGqVcNiaVWhVqU5fOcXcI8Nce8NZzwfxhk2C4g4a4gwVXL83yjMaXtcNi8NVs7OzjUo16NSMMRhMXh6lLFYPF0qOLwlajiaNKrD+dP8Aa4/ZH8Rfs6+ITq2ki8134V65eOnh/wAQunmXOj3Em+VfDniRokWOLUIkV/sN+EittZto2mhWG7ivLK1/6D/oofSwyD6QeQLKc2eEyTxQyTBwnn2Qwn7PD5vh6fLTnxDw7GpOVSrgKs5R+vYFzqYnJ8TUjRrSrYWrg8bif+Zb6Zv0MuJfozcSSznJljuIPCHP8bOHDnEk4e1xOSYqrz1YcMcTzpQjTo5lSpxn/Z+Y8lLC55haUq9CNHGUcdgcJ8Z1/YZ/DYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB3vwv8Aih8QPgv4/wDC3xS+FvirVvBXj7wXqsOs+G/Euiz+RfafewhkdWV1kt7yxvLeSax1TTL6G503VtNubvTdStLqwuri3lcZOMlKLs1qmelk+cZnkGZ4POMnxlbL8ywFaNfC4uhLlqUqkbppp3jUp1IuVOtRqRnSrUpzpVYTpzlF/wB9v/BLP/gqp8P/ANvvwQnhPxS2k+B/2mvB+kxzeN/Accxg07xdYWwjgm8e/DxLqWSe70OeVkOtaE0tzqfhG9nS2u5LzS7jS9Z1H16GIVZWdlUW67+cfLut1+L/ANKPB3xkyzxKy9YLGOjl/FuBoqWYZany0sbThaMsyytTk5Tw85NPEYdylWwVSXJN1KMqOIq/rtXQftwUAFABQAUAcT8R/hx4I+LngjxJ8OPiP4c07xZ4L8WadLpeu6FqkRktru2kKukkciMlxZ31ncJFe6bqVnLBf6ZqFvbX9hc295bwzJ5WeZJlXEmU47I88wNDMcrzKhLD4zCYiPNCpTlZqSaanTq0pqNWhXpShWoVoQrUakKsITX0fCPF3EnAfEmUcXcI5vi8i4iyLGQxuWZngp8lahWheMozjJSpYjDYilKph8ZhMRCrhcZhatbC4qjWw9apTl/FP/wUO/4J5+N/2KPHP9pab/aPiv4E+LdRnTwJ46khElxpdy4luR4L8ZtbxpBZ+JbK3SRrK9WO3sPFGn28up6bHb3VtrGkaP8A5X+NfgrmvhXm3t8P7fMeEMyryWUZvKPNPDzfNP8AsrNHBKFLHUoJulVUYUcwowlXoKFSGKw2G/6L/on/AEsOG/pGcNfU8b9UyPxOyLCU5cT8MxqONLG0ouFJ8RcOqtOVXEZPiKsoLE4Zzq4rJcVVhg8ZOrRrZfj8w/N2vww/r4KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBBr9i//kzv9k//ALNp+BP/AKq7wtX+zvhb/wAmy8Ov+yE4R/8AWfy8/wCVH6RH/KQHjn/2eLxN/wDW1zs/A2v+So/1mCgAoAKACgDA8U+FvD3jbw9q3hTxXpNnrvh7XLOSx1TS7+Pzba6t5MHBwVeKaKRUntrmF47m0uY4rm2liuIo5F93hnifP+DM/wAq4p4XzXGZJxBkmLp47K80wNT2eJwuIp3V02pQq0qsJTo4nDVoVMPisPUq4bE0qtCrUpz+d4t4S4b474bznhDi/JsDxBw1xBgquX5vlGY0va4XGYWrZ2kk41KValUjCvhcVQnSxWDxVKji8JWo4mjSqw/nU/a5/ZG8Q/s6+IG1jRxea78K9cvHTQPEDp5tzo1zLvkTw54jeNFSK/jRX+wX+2O21m3jaWJYruK7s7f/AKDvon/SxyD6QeQrKM3lg8k8UclwkZ57kMJ+zw+c4ely058Q8PRqTc6mCqTlH6/gOaricnxFSNOrKthKuExeI/5l/pnfQx4k+jNxJLO8kWO4g8IM/wAdOHDvEc4e1xWRYqrz1YcM8UTpQjTpZhSgp/2dmXJSwueYalKtSjRxtLG4HDfGNf2KfwyFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAdn8O/iJ44+Evjfwx8Sfht4o1fwZ468G6tb634Z8T6FdNaanpOpWxOyaGQBo5YZo3ktb2yuY57HUbGe5sL+2ubK5uIJGm4tSTaad01vc78rzTMclzDCZtlOMr4DMcBXjiMJi8PNwrUasNpReqlGSbhUpzUqdWnKdKrCdOcov8Avd/4JUf8FYfA37eXg+2+H/j2XSfBX7UfhTSBN4o8JxstlpHxD0+yQJc+N/h9HNM7yRbAtx4k8Mb5L7w5cSPNB9s0Nor5PWoYhVVZ6VEtV/N5r9V0/E/0k8GvGjLvEfAwyzMnRy/jDBUObGYJNQoZnSpq08wyxSk242tLFYS7q4WTco+0w/LVX7IV0n7qFABQAUAFAHDfEr4a+BvjB4G8SfDb4k+HNP8AFngvxZp0uma5oepxl4LmByrxTQyoyXFlqFlcJFe6ZqdlNBqGmahb21/YXNvd28MyeTnuRZTxLlOOyPPMDQzHKsxoSw+LwmIjeFSDs4yjJNTpVqU1GrQr0pQrUK0IVqM4VIRkvpeD+MOJeAeJco4v4QzfF5FxFkWLhjctzPBzUatGrG8ZwqQkpUsThcRSlUw+MweJp1cLjMLVrYXFUatCrUpy/ik/4KF/8E9vHP7E/jv7dY/2j4s+Bfi3UJ18A+PXhDz2E7CS5Hgzxk1vGlvY+KLG3SRrW7WO3sPFGnwSappcdvcW+saPo3+VvjT4L5t4VZv7aj7fMeEcyrzWT5w43nRm+aosrzRwioUswpQUnTqJQo5hRhLEYeMJwxOGwv8A0YfRQ+ldw19I3hn6rivqmR+JmRYSnLijhiNRxpYunFwoviLh5VZyq4nJcTWlFV6Ep1cVkuKqwwWNnVpVcBmGY/nHX4ef12FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/wDJnf7J/wD2bT8Cf/VXeFq/2d8Lf+TZeHX/AGQnCP8A6z+Xn/Kj9Ij/AJSA8c/+zxeJv/ra52fgd8pWN0khnimhhube5tp4bq0u7W5iS4tbyzu7aSW2vLK7t5YrmzvLWWa2u7aWK4t5ZYZUkb/k+zzJM34azjMuH8/y3GZRneT43EZdmmWY+jPD4zA43C1JUq+HxFGolKFSnOLT3UlaUXKMlJ/6tZfmGCzXA4XMstxVDHYDHUKeKwmLw1SNWhiKFaKnTq0qkW1KMotPundNJpoSvLOwKACgAoAKAMHxR4X8P+NPD+reFfFWk2eueHtcs5bDVdKv4/Ntru2lwSrDIeOWNwk1vcwvHc2tzHFc20sVxFHKvucNcS59wdn2VcUcL5rjMkz/ACTGU8dleaYGp7LE4TE0r2lFtShUp1ISnRxGHrQqYfFYepVw2JpVaFWpTl89xZwnw3x1w3nPCPF+T4LiDhviDBVcuzfKMxpe2wuMwtazcZK8Z0qtKcYV8NiaE6WJwmJpUcVha1HE0aVWH8637XX7IviD9nbxA2taKt5rnwq1y8ZNB1518250S5lLyJ4d8RvGoSO9jQN/Z+obY7fWLeMyRiK8iu7SD/oN+id9LLIfpBZDHJs5lg8k8UslwcZ55kcJeyw2d4aly058Q8PRqTlOpg6k5ReYZfzVMRlGIqKE3VwdXC4uv/zM/TP+hfxJ9GfiOWe5EsdxB4QZ/jpQ4f4hnD2uLyDF1uerDhniedKEYU8dThGf9m5ny0sNneGpyqU40cdRxmCw/wAXV/Y5/CwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB1Xgfxx4w+Gni/w54/8AeJNX8IeNPCOrWmu+GvE2g3kthq2jatYyCW2vLO6hYOjqwKyI26GeF5Le4jlglkjdpuLUk2mndNb3O3Lsxx2U47C5nlmLr4HMMDXhiMJi8NUlTr0K1N3hUhOOqaejTvGUW4yUoyaf8AeV/wSd/4K2+EP25vDFp8K/ipc6R4O/an8MaUX1PR4/K0/RfizpenQZuvGPge3Zljh1aGFGuvFfg+EmXTf32s6Kk2gfaodC9XD4hVVyy0qL/ybzXn3XzWl7f6O+C/jXgfETBwyfOJUMDxjg6N61BctLD51RpR9/H5fFu0a0YpzxuBjd0fer0FLDOccN+1ldR+/BQAUAFABQBwfxN+GXgX4x+BPEnw1+JXhvT/ABZ4K8WafJput6JqUZaGeFiskNxbzRslzY6jY3CRXumanZTW+oaZqEFvfWNxBdQRSp5GfZDlHE+UY7Is9wNHMcqzGjKhi8JXTcZxbUozhJNTpV6NRRq0K9KUK1CtCFalOFSEZL6fg3jLibw/4myfjHg/N8VkfEeRYuGMy3MsJJKpSqRTjUpVac1KjicJiaUqmGxmDxMKuFxmFq1sNiaVWhVqQl/FD/wUI/4J8+O/2JvHn2m2OoeK/gf4t1Cdfh/4/eANLaysJLkeDfGLW8aW1j4rsLZJHguEjt7DxPp9vJq2lRQSwavpGi/5WeNHgxm/hVnHtKftsx4TzKtNZNnDheVKT5p/2XmbglClmFGCk4TShRx9GDxGHjCcMThsL/0Z/RS+lZwz9I3hj2Nf6pkfiVkWFpvirhaNVqFeCcKL4hyBVpyrYnJMVWlFVaTnVxWTYqrDA46dWnVwGPzD86a/ET+tgoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/0Gv2L/wDkzv8AZP8A+zafgT/6q7wtX+zvhb/ybLw6/wCyE4R/9Z/Lz/lR+kR/ykB45/8AZ4vE3/1tc7P8zz/glJ/wVgk+CreH/wBmX9p3X7i5+CU80el/DT4mak895ffBa6upsQ6Dr0oE13f/AAluriUnaizXngCeV9Q0uOfQHv8ASov5R+mH9DvJ/H3KK3F3CNLB5N4uZRg+XCYyXJhsFxjg8NTfssjz2rZRhjYQiqWTZ1Uu8K+TAY+cstlSq5d9N4HeOWO8NsbDJM6nXx/BWOr81eguariMir1pe/mOXQ1lKhKTc8fgI6VvexOGSxftIYv+tBWjkjhmhmt7m3ube3vLS7tLiC8sr2yvII7qyvrG9tZJrW+sL61mhu7G+tJprS8tJobq2mlgljkb/ntzzI844ZzjMuH+IMtxmT53k+Mr5fmmV5hQqYbG4HG4abp18PiKFRKcKkJp7q0lacXKMlJ/6aZfmGBzbA4TM8sxdDHZfjqFPE4PGYapGrQxFCrFSp1aVSLalGSfqndSSkmha8o7AoAKACgAoAwvE/hjQPGegat4W8U6TZ654f1yzlsNV0q/i822u7aXGVYZDxyI4Wa3uIXjuLa4jiubaWKeKORfb4b4kz7g/Psr4n4YzXGZJn+S4unjsrzTAVXSxWExNK9pxesJ05xc6VehWjUw+Jw9Srh8TSq0KtSnL5/ivhXhzjjhzOOEeLsnwOf8N5/gauXZvlGY0lWwmNwlZLmhNXU6dSnNQrYbE0Z08ThMTTpYrC1qOIo0qsP51/2u/wBkPxB+zvr7a5oi3mu/CnXLxk0PXXUy3Wg3UxeRPDviN41CpdooYadqW2O31eBCyiK9iurWL/oL+ib9LPIfpA5FHJM7lg8k8U8lwcZ51kkJeywue4akownxDw9CpJynhZycXmOXc1TEZTXqKMpVcHVw2Kq/8zX00PoXcR/Ro4iln+QRx2f+D+f46UMgz+cXWxfD2LrOdSnwzxPOnBQp4yEVNZZmjjTw2dYem5RVHHUsXhKXxZX9lH8KBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHReEvFvijwF4n0Dxr4K1/VvCvi7wtqtlrnhzxHoV9cabrGi6xp06XNjqOnX9q8c9tdW06LJHJG4ORg5UsC02mmm007prdM6sDjsZluMw2YZfia+Dx2DrU8ThcXhqkqVfD16UlOnVpVINShOEkmmn66Nn92v/AAST/wCCu/hj9tfQNP8Agx8Z73SfCn7U/h3SzmNRb6Zonxo0vTbcvdeJ/Cdsvl29l4rtbaJ7vxZ4NtlVEjSfxF4aibQxqemeF/Vw+IVVcstKi/8AJvNefdfNaXt/ox4KeN2D8QMNTyHPqlHBcY4Wjqly0sPn1GlG88Zgo6RhjYRTnjcBBJJKWKwkXh/bUsH+5NdR/QwUAFABQAUAcD8UPhf4D+M/gPxJ8M/iZ4b0/wAV+CvFmnyadrWjaihaOWNiJILq1njKXNhqen3KRX2l6pZSwX+mX8FvfWVxBcwRyr4+f5BlHFGUY7Ic+wNHMcqzGjKhisLXTcZRdnGpTnFqpRr0ZqNXD4ilKFahWhCrSnCpCMl9RwXxpxP4ecT5Pxlwdm+KyPiPIsXHF5dmOEklOE0nGpRrU5KVLFYPFUZVMNjcFiYVcLjMLVq4bE0qtGrOEv4nv+Cgv/BPzx5+xN4+EsJ1DxX8EvFl/cD4e/EFoAXhciS4/wCEP8YG3jS2sPFun2yO8UqpBY+JrCCTV9Jjhkh1fSNF/wAq/GfwZzjwqzjmj7bMeFMxrT/sbOXC8oP3p/2ZmfIlCjmNGCbjNKFHH0YPE4aMJRxOGwv/AEa/RU+lVwx9I7hfkqfVcj8SMiwtJ8V8KxqvlqRThS/1gyBVpyrYrIsXWlFThKVXFZNiqkcBj51I1MBj8x/O6vxM/rMKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/0Gv2L/wDkzv8AZP8A+zafgT/6q7wtX+zvhb/ybLw6/wCyE4R/9Z/Lz/lR+kR/ykB45/8AZ4vE3/1tc7P8T2vuz8dP30/4JS/8FXpfgZJoP7NP7TGu3N58C7m4TT/h18Rr83F9qHwSvLycsmk6syCa81H4S3t1M0lzaRJPfeBbqaXWNDhn0qTVdGuf4a+mF9DzJvH/ACerxXwpTweTeLeT4PlwOPnyYfBcXYPDQfssiz6qkowxMIr2WT5zUvLBy5cHjJTy6UJ4H+hvA7xxx/hrjoZNnM6+P4Kx1e+Iw65qtfJK9WXv5jlsLtulKT58dgI6V1zYiglilKOJ/rajkhnht7m2uLa8tLy2tb6xvrG6t76w1CwvreO7sNR0+/tJZ7PUNO1CzmgvLC/s557O+s54bu1nmt5o5G/56s+yHOeF85zPh3iLLMZk2eZNjK+X5pleYUJ4fGYHGYebhWoV6M0pRlGSumrwnBxqU5ShKMn/AKa5dmOAzfAYTNMsxdDH5fj6FPE4PGYapGrQxFCrHmhUpzi2mmnqvijJOMkpJpOryTtCgAoAKACgDD8TeGdA8ZaBq3hfxTpNnrnh/XLOWw1XSr+ITWt5azD5kdchkkRgssE8TR3FtcRxXFvLFPFHIvt8OcR57whnuV8TcM5rjMlz/JMZSx+V5pgKrpYrCYqk3yzhLWM4Ti5Uq9CrGpQxNCpVw+Ip1aFWpTl4HFPC3DvG/DuccJ8W5Pgc/wCG8/wNbLs3yjMaKrYTG4SsvehUjdThUhJRrYfEUZ08ThcRTpYnDVqOIpUqsP52P2vP2Qtf/Z415td0JbzXPhRrd4yaLrbqZrrQLqYs6eHfETxqFS5RQRpupFY4NWgQkCK9juLZP+gn6Jn0tMi+kBkcciz2WDyXxTyXBxnnOTQkqOFz/DUVGE+IOHoTk5Tw8pOLzLLVKpXyqvNXdXBVMPiZ/wDM59NH6FvEX0aeIZ8RcPRx3EHg9n2OlDIs+qRdbGcN4yu51KfDPE86cFGGJhFSWV5q408NnNCm2lRx9LE4WHxVX9mH8IhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAG14b8SeIPB3iDRfFfhPW9V8N+J/DmqWWt6B4g0O+udM1jRdY024ju9P1PTNRs5IbqyvrK6ijntrm3ljlhlRXRgwBpptNNNpp3TW6Z0YXF4rA4nD43BYithMXha1PEYbE4epOjXoV6UlOlWo1YOM6dSnNKUZxalFpNM/uf/AOCRX/BYHw/+2Lo+l/Af496lpXhn9qLQ9O8vTdRYW2l6J8cNM062Lz6zoUCeVaaf48tLWF7rxP4TtUjt76GO48SeF4F0xdW0bwx6mHxKq+7Oyn9yl5rz7peq0vb/AER8EvHHC8dUKPDnElWjhOMMPStSqvko4fiGlSjeWIw0VywpZjCEXPGYKCUakVLFYOPsVXoYP9566z+jwoAKACgAoA8++Kfws8A/GrwD4k+GPxN8N2HivwV4ssH0/WNH1BCVdSRJb3lncRlLnTtU0+5SK+0vVLKWC/02/ggvLOeG4hjkXxuIeH8n4qyfHZDn2Bo5jlWY0XRxWFrJ2avzQq05pqpRxFGoo1cPiKUoVqFaEKtKcZxUl9VwTxtxR4dcUZPxnwbm+KyPiPI8VHF5fmGFkuaMrOFWhXpTUqOLwWLoyqYbG4LEwq4bGYWrVw+IpVKVSUX/ABOf8FAv+Cf/AI9/Yl+IHH2/xV8FfFl/cD4d/ENoAWBxJcf8Ij4uNvGltp/i7TrZHZWVILHxJYQSaxpEcRh1bStF/wAqvGbwazjwqzi69tmPCuY1p/2LnThqn70/7NzLkioUcyowTaaUKOOowlicMouOIw+F/wCjf6K30qOF/pHcLWl9VyPxGyLC0v8AWzhRVXaSvCl/b+QqrOVbFZDi60oqScquJyjFVY4DHzqKpgcdmP551+Kn9XhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6DX7F/8AyZ3+yf8A9m0/An/1V3hav9nfC3/k2Xh1/wBkJwj/AOs/l5/yo/SI/wCUgPHP/s8Xib/62udn+J7X3Z+OhQB+9X/BKf8A4KuT/AObQv2bf2lNbu774CXd0LHwD4/vftOoal8Dr++uGcWN8I1nvdT+E1/eTPNqWlwR3F94Lu55/EHhyCe0k1rQdW/h76YH0Psl+kDk1Tijheng8l8WsmwfJl2ZSUcPg+K8Jh4N0sh4gqJJKqkvZ5RnE+apgJuOGxTqZfJfVf6D8D/HDH+GmPjlGbyr4/gvHV+bFYVXq18mr1ZWnmWWxb1g2+bG4GNo4lJ1aXLik/bf1xQzW91b215Z3Vpf2N9aWuoafqGn3dtqGnalp1/bxXmn6lpuo2Us9lqOm6jZzwXun6jZTz2V/ZzwXdpPNbzRyt/zz5/kGdcLZ1mnDnEWWYzJs9yXG18vzXK8fRlQxmCxmGm4VqFalLVSjJXjKLlCpBxqU5zpzjOX+nGW5lgM4wGEzTK8XQx+XY+hTxWDxmGqKpQxFCrHmhUpzW6admnaUZJxmoyTSkryDtCgAoAKACgDE8SeG9B8YaDqvhjxRpVlrmga3Zy2Gq6VqEQmtby1mGGSRThldWCywzxMk9tOkdxbyxTxRyL7XDvEWecJZ5lfEvDWaYzJc+yXGUsfleaYCq6OKweKou8alOWsZRlFyp1qNSM6GIoTqUMRTq0atSnLweKOF+HuNeHs44T4syfA5/w5n+BrZdm+UZlRVfB47B1179OpB2lGcZKNWhXpSp4jDYinSxOGq0sRSp1I/wA7X7Xv7IOvfs8a6/iDQFvNc+E+t3jJo2supmu/Dt3MWePw74idFCrOqhhpmplUg1WFCCIr6Oe3H/QT9Ev6WuR/SAyOOQ59LB5L4qZLg1POMnhJUcLxBhaKjCpxBw9Ccm5UZNqWZZYpTr5XWndOrgqlCu/+Z76af0LOIfo1cQz4j4chjs/8Hc+x0oZHnk4uvjOGcZXlKdPhriapCKjGvFKUcqzZxhh84oQs1SzCliMOviev7NP4OCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKANTRNb1nw1rOleIvDuraloWv6FqNlrGia3o97c6bq2katptzHeafqemajZyQ3djf2N3DFc2l3bSxXFvPHHNFIkiKwabTum0+63+82w+Ir4SvRxWFrVcNicNVp18PiKFSdKtQrUpqdKtSqwcZ06lOcVOE4SUoySkmmrn9w/8AwSD/AOCxWjftY6bo37PH7Rmr6boH7TGlWK2vhrxLP9m0zRvjlp9hAS1xaRoIbPTfiRa20TT634et0itdfijm13w3BGi6no2jeph8SqloTf7zo+k/+D32v0P9CvA/x0w/GlHD8L8U16WG4to0+TC4uXJSocRU6cdZ00uWFLNYQTliMLFKGJSliMLFL2tCh/QLXWf0wFABQAUAFAHnnxX+FPgD43eAPEvww+J/hyx8VeC/Fdg9hq+k3ynkbhJbX1jcxlbnTtV065SK+0vVLKWG+06+ggu7SaOaJWHi8RcO5NxXk2OyDP8AA0cxyrMaLo4nDVk9deaFWlUi1UoYihUUauHxFKUK1CtCFWnOM4pn1nA/HHFPhxxTk/GfBmb4nI+IsjxUcVgMfhpK6dnCthsTRmpUcXgcXRlPDY3BYmFTDYvDVatCvTnTnKL/AImP2/8A9gLx/wDsS/EHy3N94q+DPiu+uf8AhXXxENuB5gAe4PhPxZ9nRbbTvF+m2ysxCrDZeIbGF9Y0iOIR6ppejf5U+Mvg3nPhTnNv32Y8LZjWn/Ymdciu95/2dmXJFQo5lQgm9FCljaUXicMo2xGHwv8A0cfRY+lNwt9I/hXmX1bJPEPI8NS/1t4T9q/d1jS/t3I/azlWxeQ4yrKK1lUxOVYmpHL8fKfNgsbmH581+Ln9WBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/8md/sn/8AZtPwJ/8AVXeFq/2d8Lf+TZeHX/ZCcI/+s/l5/wAqP0iP+UgPHP8A7PF4m/8Ara52f4ntfdn46FABQB+8H/BKn/gq1c/s+XGifs4/tIa1eaj+z9fXf2PwP45uxdajqnwL1HULlpGgnSJLi/1P4T6hezyXOs6JbR3F/wCELy4uPE3ha3nWXXvD3iL+IvpffRAyP6QeS1OJeG4YPJPFnJsG4ZXms1Ghg+J8JQi5U8g4hqRjvvDKs2kpVcuqSVGs6mAnKFH+gfBHxvzHwzx8cqzWVfMOC8fX5sZg03Ur5RWqtKeZZZFv/t7GYJNQxUU6lPlxKUp/11wT213bWl9ZXdnqGn6hZ2epabqWm3lrqOmappmo2sV9puqaXqVjLcWOp6XqdjcW99pupWNxcWOoWNxBeWdxPbTxSv8A88vEPD2ecJ55mnDXEuV4zJc+yXGVsvzXKswoyoYvBYuhLlqUqtOW99J06kHKlWpShWozqUqkJy/04yzNMvzrL8Hm2VYyhj8ux9CGJweMw01Uo16NRXjOEl84zjJKdOalTqRjOMoqWvGO8KACgAoAKAMXxH4c0LxfoWq+GfE+lWet6BrdnLYarpWoQie0vbSYYeORDghlIWSGaNkmt5kjngkjmjjkX2eHuIs84TzzK+JeGs0xmS59kuMpY/K80wFV0cVg8VRd4VKc1dSjJOVOrSqKdGvRnUoV6dSjUqQl4XE/DHD/ABpw/m/CnFeUYHPuHc+wNbLs3yjMqKr4PHYOurTpVYOzjKMlGrRrU5Qr4evCniMPVpV6VOpH+dv9r79kHXf2eddfxD4eS91z4T63eMuj6w6tPdeG7qZiyeHvEUqKFWUZK6VqjBIdVhQq3l30U8Nf9BH0Svpa5J4/5HHIOIJ4PJfFTJcJGWbZRCUaGF4iwtKKjPP+H6c5uUqbdpZplsXOrldaalF1MFVo1V/zP/TU+hXxB9GviCfEvDUMdn/g7n2NlDJc7nGWIxnC+Mryc6fDXE1WEFGNVK8cozaShRzehBwkqeYUq9F/Etf2efwYFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAXtM1TUtE1LT9Z0bUb7SNY0m+tNT0rVdMu7iw1LTNSsJ47qx1DT761kiurK+s7qKK5tLu2ljuLeeOOaGRJEVgeZpRrVcPVpYihVqUK9CpCtRrUZyp1aNWnJTp1aVSDU6dSnNKcJwkpRklKLTSZ/bf/wAEff8Agsjpn7TVl4f/AGav2mtbstH/AGibC0j07wV46vnt7HSvjfbWsZWK0ucLDa6b8TYbdB9psECWvi/y5L/SUh1RrjSj6mGxPtLQm/f6P+f/AO2/M/0D8DvHajxdTw3CfFuIp0OKacFSwGY1HGnR4hjBaQl8MKWbqK9+krQx1nUoKNZyon9Eddh/UIUAFABQAUAec/Fn4TfD/wCOPw+8SfC/4oeHLLxV4L8V2LWOq6VeqQVIIktdQ0+6jK3Om6vptysd7peqWUkN7p97DDc20ySxg14nEfDmTcWZNjsgz/A0sxyvMaTpYnD1U+/NTrUaitUoYmhUUauHxFKUatGrGNSnKMopn1vAvHXFXhtxVlHGnBeb4nJOIsjxKxOBx2GkmmmnGthcVRlejjMDjKMp4fG4LEQqYfF4epUo1oShNo/iW/b7/YF+IP7EnxD+zz/bvFPwc8VXtyfhz8Rvs4C3CAPcHwr4p+zottpvjHTLZWaSNVhs9fsoW1nR0RF1LTdH/wAqPGPwcznwpzrll7XMeGMxqz/sTO+TSa1n/Z+Y8iVOhmdCCbaSjSxlKLxOGSSr0MN/0d/Rb+lJwr9I/hT2tP6tkniDkeHori7hL2zbpSbjSWd5J7WTrYzIMZWaUZt1MRleIqRy/HylKWExmP8Az+r8ZP6oCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i//AJM7/ZP/AOzafgT/AOqu8LV/s74W/wDJsvDr/shOEf8A1n8vP+VH6RH/ACkB45/9ni8Tf/W1zs/xPa+7Px0KACgAoA/db/glX/wVZu/2dLnRv2dv2i9ZvtT/AGedQvDbeDfGdwt1qerfAnUtRunlkxFClxf6p8KNRvriW78QeHbSK4v/AAre3F14s8I2s803iLw34s/ib6Xn0Qci+kJkc+IeHoYPJPFjJcG4ZRnE0qOE4kwtCMpU+H+Ipwi3Km9Y5XmjjOvldWfJP2uBnVor9+8EfG7MfDLMI5Zmcq+YcGY+upY3BRbqV8qrVGlLM8sjJ2Ulo8Zg0408XCPNHkxEYTf9eltc2l9aWeoafe2Op6bqVlZappeqaXe2up6Vqulanaw3+matpOqWE1xYappWqWFxb3+manYXFxY6jY3Fve2dxPbTxSv/AM8fEfDme8IZ7mvDPE+VYzJM/wAkxlXAZrlWYUnRxeDxdF2nTqQd1KMk41KNanKdHEUZ08RQqVaNSnUl/p3lea5dneXYPNspxlDMMtzChDE4PGYaaqUa9GorxlGW6ad4zhJRqUqkZU6kY1ISipq8Q7woAKACgAoAxfEfh3QvF2har4Z8TaVZ63oGt2c2n6rpWoQie0vbScYeKVDyCDiSKVGSaCZI54JI5o0kX2OH+Ic74UzvLOJOG80xmS57k2MpY/K80wFaVDF4PFUXeFWlUW6avCrTmpUq9KdSjWhUpVJwl4fE3DPD/GfD+b8K8VZRgc+4dz7A1suzfKMyoxxGDx2DxEeWpSrU5app2qUqtOUK2HrQp16FSnWp06kf53v2v/2P9d/Z61x/EfhxLzW/hNrV4U0nVnDT3fhm7nZmj8P+IZFXhuq6VqrBIdTiXy5PLv0kik/6Bvok/S3yTx+yWHDvEU8Hk3irk2DU80yqDjQwnEmEoRUamf5BTlK7W0s0yuLnVy6rL2tP2mBnCpT/AOaH6av0KeIPo2Z/U4o4Yhjs+8HM9xzhk+cTUsRjeFcZiJSnS4b4lqRjZS3hlGcSUaOaUoeyq+yzGFSlU+Iq/tI/gkKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAtWV7eabeWmo6dd3VhqFhdQXthf2VxLa3lleWsqT2t3aXUDxz211bTxpNBcQyJLDKiSRurqCAunUqUqkKtKc6VWlONSnVpylCpTqQkpQnCcWpQnCSUoyi1KMkmmmrn9qX/BHf8A4LL2f7QMHhz9l79qnxDbaf8AHW3ht9I+G/xM1SaK1sfjJFCixWugeILhzHBZ/FBI1VLW4Ypb+O8Yj8rxUfI1/wBPD4nntCo/f+zL+byf978/Xf8AvvwM8eIcTxwvCHGWJhS4jjGNDKs2rSjCnn0Yq0MNiZO0YZwkrQlpHMulsb7uJ/pErtP6qCgAoAKACgDzb4u/CP4e/HX4e+JPhb8UfDll4p8F+KrJrPU9MvFKyRuD5lpqWm3abbnTNY0y5WO90vVLOSK8sLyKK4glV158PiThvJeLslx3D/EGBpZhleY0nSxFCrdNPenXoVFapQxNColVw+IpSjVo1YxnCSaPr+A+POK/DPivKONeCs3xGScRZJiViMFjcO04zi0418Ji6Er0cZgMZRc8PjcFiIzw+Kw9SdKrCUZH8Sn7fH7BHxC/Yk+In2O7+2+KfhB4pvblvhx8RhbBY72JQ858M+JhAgttM8Y6XbAmeACK01u0ibWNIVYhf2Glf5T+MXg9nXhTnfs6ntcw4azCrUeR53yaVYq8/qGP5EoUMzw8NZxSjSxdOLxOGSiq1HD/APR59Fz6UXCn0j+E/b0Pq+S8f5Jh6K4v4S9s3KhNuNJZ1k3tZOtjMgxtZr2dRupXy2vNZfj5Sm8Lisd8B1+Nn9ShQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6DX7F//Jnf7J//AGbT8Cf/AFV3hav9nfC3/k2Xh1/2QnCP/rP5ef8AKj9Ij/lIDxz/AOzxeJv/AK2udn+J7X3Z+OhQAUAFABQB+5v/AASt/wCCq97+zbd6P+z1+0Nq1/q37OmpXzQeEvFsyXeqax8CNU1O6eeaaCCBLjUNV+FWp39xLeeJvDFnFc6h4avri78YeDLWe9m8SeGPGn8V/S6+iHkP0hsiln2QxweR+K+SYOUMlzucVRwvEGFoqU4cPcRzpxcp4eUnJZbmbjUxGU15uyq4KpiMNP8AfPBPxtzLwxzFZdmLr5hwbmFdSx+XxfPWy2tUajLNMrjJpKqkk8XhLxpY2nHeGIjSqx/r7s7yy1GysdT0y+sNV0vVbCx1XSdW0q+tNU0nV9J1S0hv9L1fSNUsJriw1TSdUsLi3v8ATNTsLm4sdRsbi3vbO4ntp4pX/wCd/iXhrP8Ag7Ps14X4oyrGZHxBkeMq4DNcqzCk6OKweKpNc0Jxu4zhOLjWoV6UqmHxWHqUsThqtWhVp1J/6eZTm2W57luCzfKMZQzDLMwoQxODxmGnz0a9Ge0ovSUZRacKlOajVpVYzpVYQqQnFWK8M9AKACgAoAKAMbxD4e0PxZoeqeGvEul2etaDrdnNp+q6VqEKz2l7aTrtkimjb8HjkQrLDKqTQvHKiOvsZBn+dcK51lnEfDmZ4zJs8ybGUcfleaYCtKhi8Hi6EuanVpVI/ONSE1KnWpynSrQnSnOEvE4l4ayDjHIM24W4pynA57w9nuBr5bm+UZlQjiMFj8FiI8tWjWpy+U6dSDjVo1YwrUalOtThOP8AO/8Atgfsfa5+z3rcniXw0l5rfwl1q8KaXqrhri88L3c7Fo/D/iGRV6HlNJ1ZgsWoxqIZjHqCMk3/AECfRI+lxkvj7k0OHOI54PJvFXJ8Gp5nlkHGhhOJsJQio1M+yCnKW60nmuVRcquX1JOtR9pgJxnR/wCaP6a30KM/+jbn1Tirhanjs+8G89xzhlObTU8RjeEsbiZylS4c4kqxjte8MnzmajSzOnFUK7p5lCUK/wAQV/ah/AwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQBNb3FxaXEF1azzW11bTR3Ftc28rw3FvcQuJIZ4Jo2WSKaKRVkjljZXR1DqwYA0FRnKEozhKUJwkpwnFuMoyi7xlGSacZRaTTTTT1Tuf2a/8Ec/+CzsHxkj8Mfsq/taeJYrX4vRpa6H8K/i5rVykNt8VFRVgsPCPjK+mZY4PiSFCW+j61O6xePjssrx08aGGTxd6eGxPPanUfv/AGZP7Xk/73/pXrv/AHl4F+PMc+WE4N41xcYZ4lDD5NneImowzlK0aeCx9SVlHNto0MRJpZlpCo1j7PG/0y12n9aBQAUAFABQB5n8YPg/8PPjx8O/Enws+Kfhyz8UeDPFNmbXUdOugVmgmU77PVNLvExcaZrOl3IS80vVLN47uyu4o5onGGDeFxLw1kvF+S47h/iDA0swyvMKTp16NRWlCS1pYjD1V79DFYedquHxFNxqUqkVKL3v9jwDx9xZ4Y8WZRxtwVm+IyXiHJMQq+ExdB3hVhL3cRgsbh5XpY3L8bRcsPjcFiIzoYmhOdOpF3TX8SP7ev7BfxE/Yk+Iv9nah9s8UfCXxRd3T/Df4kLbbIdRgTdO3hvxIIFFvpnjHSrfm7tR5drrFrGdY0hRbm8stM/ym8YfB/OvCnO/Y1va5hw3mFWo8jzxQtGtBXm8DjuRclDM8PD+JD3YYmnF4nDLk9rSof8AR99F/wCk/wAJ/SP4S+t4X6vkvHeS0KMeL+EXW5qmEqStTWb5Q6kva43IMbV/g1/frYCvNYDHv2v1fEYz4Jr8dP6hCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Qa/Yv8A+TO/2T/+zafgT/6q7wtX+zvhb/ybLw6/7IThH/1n8vP+VH6RH/KQHjn/ANni8Tf/AFtc7P8AE9r7s/HQoAKACgAoAKAP3D/4JX/8FVL/APZmvdJ/Z/8A2gtV1DV/2cdUv3i8MeJ5Eu9V1j4Earql289xeWVvAlxqGq/C7U7+4mvfFXhOyhub/wAP31xd+M/BVpPqc3ibwv46/i/6XH0ReH/pEZDLO8lWDyLxWyTByhkefTj7LC55hqXPUhw7xJKnCU6mDnOUv7OzHlqYnJ8RUdSmq2Dq4vB1/wB58FPGzM/DDMll+YOvmPBuYV1LMcujLnrZdVnaMs0ypTkoxrxik8VhbxpY6nFRk6deNGvT/sDsb6w1Sw0/VtJ1DT9X0jV9PsdX0fWNIvrTVdI1jSNUtIb/AEvV9I1XT5rnT9V0nVLC4t7/AEzU7C5ubHUbG4gvLO4ntpopX/53eJ+GOIODOIM24W4qynGZHxDkeMqYDNcqx9L2WJwmJpWbjJXlCpSqQlCthsTRnUw2Lw1SlicNVrYetTqz/wBPspzfLc+y3BZxk+NoZjlmYUIYnB4zDT56VelPqnpKM4yUoVaVSMKtGrGdKtCFWE4K1XgnohQAUAFABQBj+IPD+ieK9E1Tw34k0uz1rQtas5tP1XStQhW4s72zuFKywzRN1B4ZHUrJFIqSxOkqI6+vkOfZ1wvnWWcRcO5njMmzzJsZRx+WZpgK0qGLwWLoS5qdajVi7pp3jOEuanVpynSqxnTnOMvE4k4byHjDIc24X4oynA57w/nuBr5dm+UZlQhicFj8FiIuNWhXpTTTT0lCcXGrRqxhWozhVhCcf54f2wf2Pdb/AGfNak8T+GI7zWvhLrN4V03U3D3F54VvLhyYtA8QSAZKEkx6Rq77Y9QQC3uCmoIRcf8AQF9Ef6XOTePmTU+GuJamDybxWyfBqeY5bFxoYPijCUIpVM+yGnKWk0rTzbKYOVTAzk8Rh1PASvh/+aX6a/0J8++jdntXizhOnjs98Gs8xzhleazUsTjuEMbiZt0uHOI6sY3dNtunk2dTUaeY01HDYlwzKFsV8O1/a5/AgUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQA+OSSKRJYneKWJ1kjljZkkjkRgyOjqQyOjAMrKQysAQQRmgabi1KLakmmmm0007pprVNPVPe5/Yr/wAEcf8AgtDH8Rh4X/ZP/a98UpF8QwLTQPhD8aNfu1ji8fY2W2m+B/iBqVw4WPx1/q7Tw74ou5AnjU+VpesSjxiba88XelhsTzWp1H720ZPr5Sffz69Xff8AunwK8e1mv1PgvjjGKOae5hskz/Ezssz2jSy/M6s3ZZjtDC4ybtmGlGu1juSpjv6iK7j+vgoAKACgAoA8w+Mfwc+HXx8+HXiT4V/FTw5aeJ/Bvii0Ntf2NyNlxa3CZey1bSb1B9o0vWtLuNl3pmp2jpc2dzGroxUuj+BxPwxknGOSY7h7iHA08wyvMKfJWo1NJ05rWlicNVXv4fFYedqmHxFNqpSqJSTtdP7Pw+8QeLfC7i3KON+Cc3r5LxDkuIVbC4qi+alWpS93EYHHYeX7rG5djaXNQxuCrxlRxFGcoySfLKP8SH7eX7B3xF/Yl+Ix0vU/tfif4UeJ7u6k+G/xIS12W+q2ybpm8P8AiAQr9n0vxhpUGPttl8lvqlun9r6SDatcW1h/lL4v+EGd+FOefV6/tcfw7j6lSWR54qdoYiCvJ4LG8vuYfM8PD+LS0hiIL6zhl7Nzp0f+j/6MP0nuEvpH8I/XcH7DJuOcloUY8X8ISrc1XBVpWprNcqdR+1xuQY6rf6viPerYKrL6hj2qyo1sV8G1+QH9PBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+g1+xf/yZ3+yf/wBm0/An/wBVd4Wr/Z3wt/5Nl4df9kJwj/6z+Xn/ACo/SI/5SA8c/wDs8Xib/wCtrnZ/ie192fjoUAFABQAUAFABQB+3X/BLH/gqlqP7MF/pfwD+PuqajrP7NurahInh7xA8d3qus/AjV9Uu3uLnUtMtoEuNQ1X4Zapf3E194u8H2MNxfaLfXF3418E2k2sTeJvDHj7+M/pb/RH4e+kTw+83yhYPIvFTI8HOGQcQTh7PDZzhqfPUhw5xJKlCVSrl9WpKX1DMFCrislxNSVajCthK2NwOL/dvBXxrzPwwzNYLHOvmPB2YV1LMssjLnq4GrO0ZZplSnJRhiYRS+s4bmhSx9KKhUcK8MPiKP9hen6hp2sadp2s6PqWnazo2s6dYaxo2s6Pf2mq6PrOj6raQ6hpWsaPqthNcWGqaTqlhcW9/pup2FxcWWoWVxBd2k81vNHI3/O1xTwtxDwTxDm3CnFeUYzIuIsjxlTA5rlWPp+zxOFxNO0rOzlTrUa1OUK+FxVCpVwuMwtWjisLWrYetSqz/ANQMnzjLOIMswWc5NjaGY5ZmNCOIweMw0+elWpSuutpQqQmpU61GpGFahWhUo1oQqwnBW68A9IKACgAoAKAMjX9A0XxVouqeHPEemWetaFrVnPp+q6XqEK3FnfWdwpSWGeJwQQQcqwxJHIFljdJEVx62RZ9nPDGc5bxDw9meMyfO8nxlHH5ZmmArTw+MwWMw81OlXo1YNNSTVpRd4VIOVOpGdOcovxeI+Hci4uyLNuGOJ8qwOe8P57ga+W5vlGZUIYnBY/BYmDhWoYijNNSjJO8ZLlqU6ihVpThVhCcf54v2wv2PNa/Z91qXxT4XjvNa+Ems3m3TtRffcXnhS8uHJi0DX5QCWiYny9I1iTCXyBbW6ZNRUfav+gH6I30ucm8e8np8McT1MHk/ivk+D58wy+PLh8HxVhMPBKrnuQ020o1or95m+UQbngpuWKwqnl8msL/zT/TY+hNnv0b89q8XcJU8dnvg1nmO5MtzOanicdwdjsTNulw7xHVSblRlJ+zyXO6ijTzCCjhMXKGZwTxnw3X9sn8AhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFACglSGUlWUghgSCCDkEEcgg8gjkHmgLtO6dmtU+t+5/Xp/wAEb/8AgtGPEn/CK/sl/theKwviT/Q/D3wd+N/iK9wviP7lrpfgL4k6rdP8viL/AFVn4Y8ZXsoXxD+50jxDMPEH2XU/EHo4bE3tTqPXaM318pPv2fXrrv8A3B4E+Paxf1PgrjnGWxf7vC5FxDiqmmL2hRy3Na03pivhp4THVJf7V7tDEy+s8lbE/wBVdd5/ZAUAFABQAUAeXfGb4M/Dn4//AA58R/Cr4qeHbXxN4O8T2hgvLOceXdWV0mWsdY0e+UGfS9b0q4K3WmanalZ7WdAfnieWKT5/ijhfJOMskx3D3EOBp4/LMfT5KtKelSlUV3SxWGqr38Pi8PO1ShiKbU6c11i5Rl9r4e+IfF3hZxdlHG/BGb18m4gyauquHxFJ81HEUZWWJy/MMM37LG5bjqV6GMwddSpV6Un8M4wnH+I39u/9hL4jfsS/EhtH1cXXiX4W+Jrq6l+G3xIjtdlrrNnGTK2h66IgYNL8X6VAyjUdPJSG/hUarpW+ylkjtP8AKXxe8Ic78Kc8+rYn2mP4fx9SpLI88jTtTxNNXk8JjOX3MPmeHhb21G6jWh/tGHvSlKNL/pA+jH9JzhH6R/CKx+AdDJ+Nsmo0IcYcITr89fL687QWZ5Y6jVXG5BjqibwuKtKphajeBx3LiIQnX+Eq/IT+mwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i//AJM7/ZP/AOzafgT/AOqu8LV/s74W/wDJsvDr/shOEf8A1n8vP+VH6RH/ACkB45/9ni8Tf/W1zs/xPa+7Px0KACgAoAKACgAoAKAP2x/4Jaf8FT9U/ZZ1HTPgT8d9S1LXP2atZ1GRdF1ox3eq618CtX1W7e4utX0a1gS4v9V+G+p39xNfeMfBljFPe6Xe3F5418EWkuuzeJfDPxA/jf6Wv0SeHfpFcPPNMrWDyLxTyPBzhw9xHODp4fNsPT56seHOJJUoSqV8srVJTeBxyhVxeSYqrLEYaNfC1sfl+O/dPBbxqzTwvzT6njHXzHg/Ma6lmmVxlz1cFVlywea5UpyjGni4RUfrGHcoUcwowVKq4VoYbE4f+xLTNT0zW9M0zXND1PTdb0PW9NsNZ0TW9Gv7TVdG1rRtVtIr/StY0fVbCa4sNU0rVLC4t77TtSsbiezvrOeG6tZpYJY5G/52OK+FOIuB+Is34T4tyjGZFxFkWMqYHNcrx1PkxGGxELSWsXKlXoV6UqeIwmLw9SrhcbhatHF4StWw1alVn/qDk2c5XxDleCzrJcbQzHK8xoRxODxmHlzUq1KV09GlOnUpzUqVejVjCtQrQqUa9OnWpzhG7Xzx6YUAFABQAUAZGvaDovinRdT8O+ItMs9Z0PWbOfT9V0vUIEuLO+s7hCk0E8TghlYHIYYdHCyRssiqw9XIs9zjhnOMt4g4fzLGZPneT4yjj8szTAV54fGYLGYeanSr0K0GpRlFqzTvGcHKFSMoSlF+NxFw9kfFuR5rwzxNlWBzzIM8wNfLc3yjMqEMTgcfgcTB062HxFGonGUZRd1JWnTmo1KcoVIRkv55P2w/2O9a/Z+1mXxV4VjvNa+Ems3myw1B99ze+Ery5cmLQdflALNA7Hy9H1iTCXqhbS7ZNRVWvf8AoA+iN9LvJ/HrKKfC3FFXB5P4r5Rg+bHYCPJh8HxZg8PBe1z3IqbajHERivaZvk9O88HJyxeEU8vlOOD/AOaj6bP0Jc9+jhnlbjDhCljs88Gc8x3Jl2Yz58VjuDMdipt0eHeIqqTlPDTlL2WSZ5UtDHxUcFjZQzOMJY/4Zr+2z/P8KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP60/8Agjf/AMFpQw8K/slftieKzvzaeHvg78c/Ed9ww/d2uk+AfiZqt2+Q4Oyz8M+Nr6Y78waR4kmDC31eT0MNitqdV/4Zv8pN/g/vP7V8CvHy/wBS4K46xuvuYXIuIsVU32hRy3N603vtTwmYVJa+7Qxcr8teX9Ztegf2mFABQAUAFAHlfxq+C3w4/aD+G/iP4U/FXw7beJfB/iW1MN1bS4jvNPvIwzWOt6JfBWm0rXNKnIudN1K2IlglBRxLbyzwS/PcU8LZHxnkeO4e4iwVPH5Zj6fLUpy92rRqxu6OKwtZLnw+Lw837ShXhaUJaPmhKcZfb+HXiLxf4VcX5RxxwPm1bJ+IMnre0o1oXnh8Vh52WKy7McM2qeOy3HU06OMwla8KsHzRcKsKdWH8Rf7df7CvxH/Yl+JJ0PWxc+JPhj4luLuf4bfEiG1aOy1yxiYyNo2srGDBpfi/SYWQapphfyrmPZqmmNLYz4g/yk8XfCPPPCnPfqmK58fkGPnUnkWeRpuNLF0ou7wuKUbww+ZYaLX1ihflqRtiKDlSnaH/AEg/Rl+kzwj9I7hBZllzo5Pxnk9KhT4w4QqV1PEZbiZrlWY5c5tVcbkGOqKTwWM5eejPmwWMUMTTvV+F6/Iz+mAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/0Gv2L/8Akzv9k/8A7Np+BP8A6q7wtX+zvhb/AMmy8Ov+yE4R/wDWfy8/5UfpEf8AKQHjn/2eLxN/9bXOz/E9r7s/HQoAKACgAoAKACgAoAKAP2n/AOCW3/BUzVv2U9T074HfHLUNT179mjW9Sk/svUxHdarrXwM1nVbt57vXdBtIEnv9U+Hup308t/4z8FWMc93Y3k93408FWkniCXxH4c8f/wAd/Sz+iXw39Ivh15llyweReKWR4OcOG+JZwcKGZUIOdWPDnEjownVxGU16kpvB4xQrYvI8VVni8JCthq2YZdmH7l4L+NOa+F2afVcV7fMeEMxrxlmuVRlzVMJUlywea5UpyUKeNpwUfb0HKFHMaMFRrSp1YYbFYb+xrSdW0nX9J0rX9A1XTdd0HXtMsNa0LXdFv7XVdG1vRtVtYr7S9X0jVLGWex1LS9Ssp4Lywv7Oea1vLWaK4t5ZIpFY/wDOrxbwlxHwLxHm/CPF2UYzIuI8ixk8DmmV46ChXw1eCUoyUoudLEYfEUp08Tg8ZhqlbCY3CVqOLwlevhq9KrP/AFCyXOsq4iyvBZ3kmOoZjleY0I4jB4zDycqdWnJtNNSUZ06tOalSr0KsYVqFaFSjWp06sJwjfr509QKACgAoAKAMnXtC0bxRo2p+HvEWmWes6HrNnPp+qaXqECXNlfWdyhSaC4hkBV0dTweGRgroyuqsPVyPPM44azjLeIMgzLGZPnWUYyjj8szPAV54fGYLGYeaqUa9CtTalCcZLXVxnFyhNShKUX4/EPD+R8WZHmvDXEuVYHPMgzzA4jLc3ynMsPDFYHMMDioOnXw+JoVE4zhOL0ekoSUakJRnGMl/PL+2J+x1rP7P+sy+LPCcV7rPwj1m822N8++5vfCF7cufK0HX5QCz2zu3l6NrMmFvF22V6y6kqPf/AO/30Rfpd5P49ZRS4V4qq4PJ/FfKMHzY3BR5MPg+LcHhofvc8yKldRhioRXtc4yeneWEbljcFGWXOpDAf81X02voSZ59HHO63GPB1LHZ54M53juTAZhPnxWP4Kx2KqP2PD3EVVJynhKk5eyyPPKlo41cmAx84ZoqVTMfhev7dP8APwKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP6rv+CN//BaM+Hv+EV/ZL/bC8VlvD3+h+H/g78b/ABFekt4f+5a6X4C+JOq3T5bQP9VZ+GPGV7Lu0H9zpHiGY6EbXU9C9DDYranUflGb/BSf5S+/uf2X4E+Pbw31PgrjnG3w3uYbIuIcVU1w20KOW5tWm9cNtDB4+o74b3aGKl9X5K2H/rxBDAMpDBgCCDkEHkEEcEEcg969A/t3fUWgAoAKACgDyj42/BL4bftDfDbxF8Kfit4dt/EfhDxJb+XPBJtiv9Mv4g5sNd0K/wBjzaVrulTP9o07Ubf54n3wzJPaT3NtN87xXwpkfGuRY7h3iLBQx2W46HLODtGth60U/Y4vCVrOWHxmHk+ehXhrF3jJTpzqQl9z4ceI/F/hRxhlPHHA+bVsoz/J63PSqxvPDYzDTcfrWWZnheaNPHZZjqa9li8JV92pG04Sp16dGtT/AIif26P2GPiT+xN8Sm8P6+tx4j+G3iO4u7j4bfEiC1aKw8Q6fE299K1VY98OleLdJieNNX0lpNsilNS05p9OuYpF/wAo/Fvwkzzwqz14PGKeOyLHTqTyLPI03GjjaMXd4fEJXjh8yw0WlicO3aStXoOdCpFr/pD+jR9JfhD6R3B6zXK3SyjjDKKVClxhwhUrqeKynF1Fyxx2Bc+Wpjshx9SM5YDHKN4PmweMVLGUpwl8OV+Sn9KBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+g1+xf/wAmd/sn/wDZtPwJ/wDVXeFq/wBnfC3/AJNl4df9kJwj/wCs/l5/yo/SI/5SA8c/+zxeJv8A62udn+J7X3Z+OhQAUAFABQAUAFABQAUAFAH7Pf8ABLn/AIKk6x+ydq1h8FPjZfap4g/Zm13U5DZXix3Oq638ENZ1W6aa88ReHLSJZr3UvAeo3s8t9408E2SS3MF1Nd+MvBtq/iOTxB4f8efx/wDSx+ibw19I3hz69gfqeQ+KGRYScOGeJ503Ghj6MHOsuHOJHRhOticnxFWc5YXFRhWxmRYurPGYKFbD1sxy3Mv3DwY8aM18Lc1+r4j2+Y8I5jXjLNsojLmqYapJKDzXKlUlGFLHUoKKrUXKFDMaMFQryp1YYXF4X+yHRtZ0fxHo+k+IvDur6Z4g8PeINMsNb0HX9Ev7XVdF1zRdUto73TNX0jVLKWaz1HTNRs5obuyvrSaW3ubeWOaKR0YE/wDOnxfwhxLwFxJm/CHF+T4zIeI8ixc8FmmV46CjWoVopSjOE4ynRxOFxFKVPE4PG4apWwmOwlaji8JXrYatSqz/ANRMkzvKuI8qwOd5JjqGY5XmNCOIweMw8nKnVptuMk1JRnSrUpxlSr0KsYV8PXhUoV6dOrTnCOjXzZ6gUAFABQAUAZWuaHo/ibR9T8P+IdNs9Z0TWbOfT9U0vUIEubK+srlDHPb3EEgKvG6k9eVOHUq6qw9TJM7zfhvN8uz/ACHMcZlGdZRjKGPyzM8BXnhsZgsZh5qpRxGHrU2pQnCST3tJXjJSjKSfkZ/kGScVZJmnDfEmV4HO8hzvBYjLc2ynMsPTxWBzDA4qDp18NiaFVShUp1ISe65oytODjOMZL+eb9sX9jnWfgBrE3i7wlFeaz8I9Yvdtnetvub3wde3Uh8rQtel+Z3tJHbytG1qX5bsbLG/ddSEUmo/7+fRE+l5lHjxlNLhPiyrg8o8V8pwfNi8JHkw2D4vweGh+9zvI6V1GGMpwXtc4yenrhnz47Axll7q08v8A+az6bf0I87+jnnVbjPgyjjs78Gc7x3LgsbL2mKx3BGOxVR+x4f4hrWlOpgqk5Klkee1dMYuTL8wnHNFRq5n8K1/cB/nyFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9Rf8AwRx/4LQyfD0+Fv2Tv2vfFLS/D8m00D4Q/GnX7svL4DyUttM8D/EHUrly0ngf/V2nh3xTdyF/BmYtL1qVvCAtr3wn34bE2tTqPTaMn08pPt2fTrpt/YHgV49vLPqfBfHGMby33MNkef4mbcsu2hSy7M603d5ftDC4ybbwOlHES+o8lTBf2JRyJKiSxOkkciLJHJGwdJEcBkdHUlXR1IZWUkMCCCQa9E/udNSSaaaaTTTumnqmn1T3T6j6BhQAUAFAHkvxw+B/w1/aK+GviL4UfFfw9B4i8JeIrfbLG22LUdJ1GJX/ALP1/QNQ2PLpWvaVM5n0/UIQShMlvcR3NjcXVrP85xZwnkXG2RY3h3iLBQxuW46FpRdo18NXin7HGYOtZyw+Mw0nz0a0b2fNCcalKdSnP7vw28SeMPCXjDKeOeBs2q5Tn2UVeaE1ephMdhKjj9byvNMLzRhjssx1OPssXhajSkuSrSnRxNGhXpfxD/tyfsOfEn9if4mP4b8RpP4h+HfiGe7ufht8SILRotO8S6ZCwdtO1FULxaV4s0mKSJNZ0d5D96PUNPe50y6t5z/lF4teEue+FWfPA41TxuSY2dSeRZ7Cm40MdQi7uhXSco4fMsNFxWKwrk9416LqYepCb/6Q/o0/SV4P+kdwdHN8plSyni7KadCjxhwhVrqpi8nxlROMcZhJS5Z47IsfOE5ZdmEYLaeExcaONoVaR8Q1+Tn9JBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6DX7F//Jnf7J//AGbT8Cf/AFV3hav9nfC3/k2Xh1/2QnCP/rP5ef8AKj9Ij/lIDxz/AOzxeJv/AK2udn+J7X3Z+OhQAUAFABQAUAFABQAUAFABQB+yv/BLz/gqNrf7JGsWXwa+Mt5qniP9mbxBqcjwyolzqmt/BTWdUuWlvPE3he0jE13qPgvULyaS98aeCLNJJvtEt14u8I258SPrmi+OP5D+lf8ARP4a+kbw19bwn1TIfE7IsJUjwxxTOm40cZRi51lw5xI6MJ1sTkmJrTnLDYmMK2MyLF1p47AwrUa2ZZbmf7d4M+M+beFua+wr+3zHhHMa8ZZvlCnedCcuWDzXKlUkoUswpQjFVaTlChmNGEcPiJU6lPC4vCf2T6HrmieJ9E0fxN4Z1jS/EXhvxFpdjrnh/wAQaJfW2qaLrmi6pbx3mm6tpOpWcktpf6df2ssdxaXdtLJDPDIroxBr/nQ4w4P4m4A4lzjg/jDJ8XkPEeRYueCzPLMbBRq0asUpwqU5wlOjicJiaM6eJwWNw1SthMdhK1HF4SvWw1alVn/qJkeeZTxLlOBzzI8dQzHKsxoqvhMXh5NwqQbcZRlGSjUpVqVSMqWIw9aMK+HrwqUa9OnVhOC1K+aPWCgAoAKACgDK1zQ9H8S6PqXh/wAQabZ6xomsWc+n6ppeoQJc2V9ZXKGOe3uYJAySRyIxBBGQcMpDAEenkudZtw5m2XZ9kOY4zKM5yjGUMflmZ4CvPDYzBYzDTVShiMPXptThUhNJpp2avGScW0/Iz7Icl4pyXNOHOI8rwOd5DneBxGW5tlOZYenisDmGBxVOVLEYbFYeqpQqUqkJNNNXTtKLUkpL+en9sb9jjWPgDq83i/whFe6z8I9YvNtpduZLq+8G3tzJ+60PXZcF5LKR28rRdal4ufksNQcakIZtS/37+iH9L3KfHfKqXCPF1XB5R4r5ThObE4WPJhsFxhg8NT/e53ktLSFPHU4RdXOMmpf7v7+PwEXl7rUcu/5rvpufQizr6Oec1+NeCqOOzvwYzrHcuExkvaYrH8D47F1LUcg4gre9Orl9WpJUciz2s7Yr3MuzKcc0VCvmvwlX9xH+e4UAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf00f8Ec/+Czs/wAHpPDH7Kn7WviaW5+Ekj2uh/Cn4u61cvNc/C1nK2+n+EPGl9OzST/Dckpb6Nrk7tL4CzHZXzv4LEMvhHuw2J5bU6j93aMn9nyfl59D+tvAvx6lkbwfBvGuLlPJG4YbJs7rycp5O3aNLA5hUk255Ve0cPiJXllulOo3gOV4H+zOC4guoIbq1miuba5ijuLe4gkSaC4gmQSRTQzRs0csUsbLJHIjMjowZWIINekf3hGUZxjOEozhOKlGcWpRlGSvGUZJtSjJNNNNpp3TJaCgoAKACgDyH46fAv4aftG/DTxD8J/iv4fh8QeE/EMGGGVh1PRtThV/7P8AEPh/UDHJJpWvaTLIZrC+iVh80trdw3Wn3V5Z3HzfFvCWRcb5FjeHeI8FDG5bjYarSNfC14p+xxuCr2lLDYzDyk5Ua0U95U6kalGpVpT+98NPEvjHwj4xynjrgbNamVZ7lNW8X71TB5hg6ko/W8qzXC80IY7K8dCKp4rC1Gm7Qr0KlDFUcPiKX8Q37cP7D/xL/Yo+JsnhjxMk/iD4f+IJru6+G/xIt7R4dM8UaXC4ZrK+VTJFpXirSY5Io9c0R5WaNnjv7GS60u7tLuT/ACh8WfCfPfCrPngMep43JsbKpUyPPIU3GhmGHi7ulWSco4fMcPGUVi8I5NpuNajKph6tOpL/AKRfo2fST4O+kbwbHOcmlTyrirKqdCjxhwhWrxqY3JcbUi1HE4aT5Z47I8fOFSeW5lGCU1GeFxUKGOoV6EPiavyk/o4KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i//AJM7/ZP/AOzafgT/AOqu8LV/s74W/wDJsvDr/shOEf8A1n8vP+VH6RH/ACkB45/9ni8Tf/W1zs/xPa+7Px0KACgAoAKACgAoAKACgAoAKACgD9i/+CX/APwVD179kPW7P4Q/F+71XxL+zL4h1R5NsaXGp638GtZ1O5Ml54q8JWqeZc33hS+upXvPGngq1Vnnle48UeFoV8SHV9L8Y/yN9K36KPDH0juGvrOHeEyHxMyLCVI8LcVSptUsTSi6lZcO8R+xhOtisixNac5Ua8YVsZkWLrTx+XwrUq2ZZZmn7Z4NeMubeFubeyq+3zLhLMa0XnGTqac6U2oweaZV7SSp0cxpQUVUpuUKGY0IRw2JlCdPCYvB/wBl+ga/oPizQtF8U+Fta0vxJ4Z8SaXZa54e8Q6HfW+p6NrmjalAl1p+qaXqNq8lte2N5byJNBcQyMjq3Zgyj/nP4y4N4n8PuJs44O4xyfF5DxHkWLng8zyzGxSqUqiSnTq0qkJTo4rB4qjOnisDjsLUrYPHYStRxeEr1sPWp1Jf6jZFnuU8TZTgc8yPHUcxyrMaKr4TF0JNwnBtxlCcZKNSjXo1IypYjD1oQr4evCpRr04VYTgtavmD1goAKACgAoAy9b0XSPEmkaloOv6bZ6xousWc+n6ppeoQR3VlfWV1G0U9tcwShkkikRiCCOOGBDAEenk2c5tw9m2XZ7kWY4zKc5ynGUMflmZ4CvUw2MwOMw1RVaGJw9em4zp1ac4qSknrqmnFtPyc9yLJuJ8mzTh3iLK8FnWRZ1gsRlubZTmWHp4vA5hgcXTlSxGFxWHqxlTq0qtOTjKMlppJNSSa/nr/AGx/2N9X+AerTeMfB0N5rHwj1e8221y3mXN74LvbqTEWia5L80kmnyyMItF1qU4uMpp+oONSEE2p/wC/P0QvpfZT47ZVR4Q4vrYTKfFjKcJzYjDx5MNguMcHhqd6udZLS92FPH04RdXOMmpa0LTzDL4vLnXo5b/zYfTe+hDnP0dc5r8b8EUMbnXgxnWO5cNipe0xeP4Fx2LqWo5Dn9Z81SrltWpJUciz6s39YvDLMzms0+r4jNvg+v7kP89AoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/pH/4I7/8ABZi8+AM/hz9l39qrxDc6h8DLiW30f4a/E7VZpbq++DkszrDaeHvEdzIXnu/hezssdndMZLjwFkIvmeEsReHe3DYnktTqP3dlJ/Z8n/d/L0P6s8DPHifDUsJwfxlipVOHpSjQynN60pTqZE5PlhhcXJ3lUyhtpU6jvLLdvewXu4X+1CyvbPUrO01HT7u2v9Pv7aC9sb6yniurO9s7qJZ7a7tLmB5Ibm2uIZEmgnhd4ponWSN2RgT6Z/fNOpCrCFWlOFSnUhGpTqU5KcKkJpShOE4txnCcWpRlFtSTTTaZZoLCgAoAKAPHvjv8CPhn+0j8M/EPwn+LHh+HX/Cuvw5BGyHVdD1WFJBp3iLw7qJjlk0rXtKkkaWyvY1dWV57K9gvNNvL2yufmeL+Ech45yHG8OcR4OOMy7Gx8o4jCYiKl7DG4Ku4ylh8ZhpScqNaKas50qsKtCpVpT+/8MfE7jLwh4yyrjrgbNamV55ldSzT5qmBzLA1JQeLynNsIpwhjsrx0IKGJw05RkpRp4jD1KGMoYbE0f4hP23f2I/iZ+xT8TpPCniqObXvAmvS3d38N/iPa2jw6V4r0mFwWtbpQZY9K8U6VHLDHruhSTO9u8kV7ZS3mlXllez/AOUPix4UZ94V59LL8wUsZk+MlUqZHnlOm44fMcPFpunUSco4fMMMpRjjMJKTcG41qUqmGq0qs/8ApG+jd9JDg76RnBsM8ySVPK+J8rhQocX8I1q8amOyPH1ItRr0JNQnjskx04VJ5ZmcacY1Yxnh8RChjsPicPT+K6/Kz+iwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/AEGv2L/+TO/2T/8As2n4E/8AqrvC1f7O+Fv/ACbLw6/7IThH/wBZ/Lz/AJUfpEf8pAeOf/Z4vE3/ANbXOz/E9r7s/HQoAKACgAoAKACgAoAKACgAoAKACgD9g/8AgmF/wVA1/wDY/wBdtvhP8WbnVvE/7M/iTU2kltohNqOs/B/WdSuQ934u8I2uJJ73w7eTSSXPjDwXbFftsjy+I/DyJ4iXULDxR/JP0rPopcL/AEj+Gfb0XhMh8S8iwlWPCvFcqUlTr0051lw7xH7GE62KyHFV5zlSrRhWxmR4utUzDL4VqdbMstzT9q8G/GXN/CzNvZz9tmXCmY1oPOcmU05U5NRg80yv2klTo5lRpqKnByhQzGjCOGxUoShhMXg/7NPDniPw94x8P6J4t8Ja5pXibwt4m0uz1zw74i0O9h1HRtb0fUYVuLHUtNvrdnhubS5hYOkiNlTujkVJUkRf+c3jTgvijw84ozjg3jLJ8XkXEeRYuWEzLLcZFKdOaSnSrUasHOji8Hi6E6eKwOOwtSthMdhK1HF4WtWw9anUl/qPkOfZRxPlGBz3IsdRzHKsxoqvhcXQk3GcbuM4TjJKpRr0akZ0cRh60YV8PXhUo1qcKsJRWzXy564UAFABQAUAZmtaLpPiPSdR0HXtOs9X0XV7O40/VNM1CCO6sr+yuo2iuLa5t5Q0csUsbFWVgeuRggEelk+cZrw/muX55keYYvKc4ynF0MfluZYCvUw2MwOMw1RVaGJw1elKNSnVp1IqUZRfk7ptPys8yPJ+JsnzPh7iHLMFnWR51gsRlubZTmWHp4vAZhgcXTlSxGFxWHrRlTq0atOTjKMk97q0kmv57P2yP2N9X+AmrT+M/BkN5rHwi1e8xBcMZLq+8FXt1J+60TW5Tukl02WRhFoutyk+edmnak41L7Pcar/vv9EH6X+VeOuV0eDuMa2EynxYyrCc1aivZ4bBcZ4PDU71c5yakuWnSzGlTi6uc5NSS9iufMcug8v+sUMs/wCbL6b30IM5+jvnGI454GoY7OvBnOcao0MRJ1MXj+BMdi6lqORZ7WfNUrZXVqyVHI89rN+2bp5ZmlRZmsNic2+Da/uc/wA8QoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/oo/4I+f8FktR/Zou/D/7NH7Tmt3urfs83t1Hp3gbx7evcX+q/BO5upVSGwvWzNc6h8MJJnJltI0kvPB7SPeaYk2jibTYOzDYnktCb9x7P+X/AO1/I/qPwN8d6vCc8LwlxdiKlfhipNUsuzKo5VK3D8pu0aVT4p1cocnrBXqYG7nRUqF6Mf7a9M1PTda07T9Y0fULLVtI1aytdS0vVdMu4L/TtT06+gjurLUNPvrWSW2vLK8tpYri1ureWSC4gkjmhkeN1Y+pvqf6A0qtLEUqdehVp1qFanCrRrUpxqUqtKpFTp1adSDlCpTqQkpwnGTjKLUotp3LtBoFABQAUAeN/Hv4CfDL9pT4ZeIPhN8WNAi13wvr0W+ORNkOr6Bq8KSLp3iTw5qLRyvpWvaVJK8lneRo8ckb3FjfQXmmXt9ZXPzHGHB+Q8dZDjeHOI8HHGZfjI3TVoYnB4mKl7DHYGu4yeHxmHlJypVUpJpzpVoVaFWrSqfoPhf4ocZeD/GeVcdcDZpPLM7yypaUZc1TAZpgKkoPF5Rm+EU4Rx2WY6MIxxGHnKMozjSxOGq4fGYfDYmj/EF+21+xN8Tf2KvifL4R8Wxy674I12S7vPhx8RrS0kh0jxdo8Mi7oZ13SppfifSklgi1/QZJ5JbSWSK7tJb3SL3T9Quv8oPFbwpz7wrz+WW5lGWLynGSq1cjzynTccNmWGhJXhNXksPmGHUoRxmDlJypylGrSlVw1WjWqf8ASP8ARx+kdwb9IzgyGfZFOGWcSZZChQ4v4Rr141MfkOYVIvlq03aEsbk2OlCrPK80hTjCvCFTD14YfH4bF4Wj8YV+Wn9EBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6DX7F//Jnf7J//AGbT8Cf/AFV3hav9nfC3/k2Xh1/2QnCP/rP5ef8AKj9Ij/lIDxz/AOzxeJv/AK2udn+J7X3Z+OhQAUAFABQAUAFABQAUAFABQAUAFABQB+vP/BMb/gp54j/Y68Q23wu+KNxqvin9mjxLqjS3+nxCXUNZ+FOsajMPtXjHwZbljJcaTcSt9o8XeD4SseqqJNZ0ZYfEKTx61/Jv0qvoqcL/AEj+GFUg8LkPiTkWFqx4U4slSfJUgnOt/q9xF7GE6+LyDFV5znTqRhVxeSYutUzDLoVYVsyy7Nf2jwc8Y838LM35Ze2zHhXMa0HnWSqa5oyajT/tPLPaSVOhmVGCipRbhRzChCOFxUoShhcXg/7O/DPibw5418OaF4x8Ha7pXijwn4o0q01zw54k0O8i1DR9b0e/jEtpqOnXkJKTQTISCPllhlWW3uI4bmGaJP8AnL424J4o8OuKM44M4zyfFZFxHkWKlhMxy7FxXNCVlOlXoVYOdDF4LF0Z08VgcdhalXCY3CVaOKwtarQqwnL/AFIyDP8AKOKMowOfZDjqOY5VmNFV8LiqLdpK7jOnUhJKpQxFGopUcRh60YV8PXhUo1oQqQlFblfKnsBQAUAFABQBm6zo2k+ItK1HQtd06z1fRtXs59P1PTNQt47qyv7K6jaK4trq3lVo5YZY2ZXVgQc56gGvRyjOM04fzTL88yTMMXlWcZVi6GPy3MsBXqYbG4HG4apGrQxOGxFKUalKrSqRUozjJO61um0/LzvJMn4lyfM+H+IMswWc5JnOCxGXZrlWZYeli8BmGBxdOVHE4XF4atGdKtRrU5SjOE4tNO+9mfz3ftk/sbat8BdVn8aeCoLzV/hHq15iKZjJdX3gm9upMRaNrUp3STaZNI4i0XW5STKSmm6nINRFtc6t/vt9ED6YGV+OmWUeDeM6+EyrxYyrCN1aSVPC4LjTB4anernGT0ly06WZ0qcXWznJqSSppTzHLYPL/rOHyv8A5tPpv/Qfzj6PGb4jjvgTD43OfBjOMbajWbq4vH8B47F1bUcjz2s+arWymtVmqORZ9WbdVunlea1Fmf1XFZx8FV/dJ/ncFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9BH/BIL/gsVrH7KGpaL+zv+0bq+o6/wDs06rfLaeGfE85udS1n4HX99PkzW0aia81L4b3NzK02s6BbrLdeH5JJtb8OwyE6jo2rdmHxLptQm7wez6wf6p9V03XVP8AprwP8dK/BlXD8L8VV6uJ4TrVOTCYyXPVr8PVastZRS5qlXKpzblXw0VKeFbliMLGV6tCt/cLoutaP4k0fSvEPh7VdO13QNd06y1fRdb0e9ttS0nV9K1G3ju9P1PTNRs5JrS+sL61miubS7tppYLiCWOaKR43Vj6l76p3T1T3vfqf6E0MRQxVCjisLWpYjDYilTr4fEUKkatGvRqxVSlWpVYOUKlOpCUZwnCTjOLUotppmnQahQAUAFAHi/x++APwx/aY+GOv/Cb4saDHrfhnXI/MgnjMcOs+HtYhSRdO8SeG9RaKV9L13S3ld7W6VJIZopLjT9Qt73S72+sbn5fjHg7IOO8gxnDnEeDji8Bi480ZK0MVgsTFSVDHYGu4yeHxmHcm6dRKUZRlOjWhVw9WrSqfofhb4pcZ+DnGeV8dcC5pPLc5y2fLUpy5qmX5tl9SUHi8ozjCKcI43LMbGEY16EpRqU6kaWKwtXD43D4bE0f4gv21/wBir4nfsV/FCbwb4wik1vwZrcl3e/Dn4i2lpJDo3jHRYZF3JIu6VNL8S6UssEPiHw/LPJNYzyRXVrLe6RfabqV7/k/4q+FefeFmfyyzM4yxeV4uVWrked06bjhszwsZK8ZK8lh8fh1KEcbg5SlKlOUalOVXDVaFer/0kfRz+kZwb9IvgunxBkE4ZbxFlsaGH4u4Sr141MwyDMakZWnF2hLG5PjpU6tTKs0hThTxNOFSjXhh8fhsZhMP8aV+Xn9ChQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/8AJnf7J/8A2bT8Cf8A1V3hav8AZ3wt/wCTZeHX/ZCcI/8ArP5ef8qP0iP+UgPHP/s8Xib/AOtrnZ/ie192fjoUAFABQAUAFABQAUAFABQAUAFABQAUAFAH64f8Eyv+CnHib9jbxHB8NfiTNqviv9mnxRqpm1bR4d99rHww1e/kVbnxr4Igd90lpK5E3ivwlGyW+uQo2o6eLfxBEst9/J/0qPorcLfSP4XTTwuReI+RYWquE+LXSfK43nW/1f4g9jCVfGcP4utKcoyjGri8lxdaeY5dCrGpmOXZp+zeDvjFnHhZm7T9tmPC2Y1oPOskU1dO0af9p5Z7SSp0Mzo04pNOUKOYUYRwuKlBwwuKwf8AaF4V8VeGfHXhnQPGngvX9K8VeEfFWlWuueG/Euh3aX2ka3pF8m+2v7C6j4kifDRyxyLHcWtzHPZ3kNveW88Ef/OVxxwPxT4ccU5xwXxpk+KyLiPIsVLC5hl+KirxdlOjiMPWg50MZgcZQlTxWBx2FqVcJjcLVpYnDVqtGrCb/wBSeH+IMn4pyfA59kOOo5jlWY0VWw2Kot2au41KVWnJKpQxFCopUcRh60YVsPWhOlWhCpCUVvV8oeyFABQAUAFAGbrGj6V4g0rUdD1zTrPV9H1azuNP1PTNQt47qyv7K6jaK4tbq3mVo5oZo2ZXR1IIPrzXo5Tm+aZDmmAzvJcwxeVZvlWLoY/LcywFephcbgcbhqkauHxOGxFKUalKtSqRjOE4STTXqeZnWS5RxJlGZZBn+W4LOckznBYnLs1yrMsPSxeAzDAYulKjicJi8NWjOlWoVqU5QqQnFpp99T+fD9sr9jXVfgNqs/jbwTBeav8ACPVrwCOQmS6vvBF7dSYi0bWZTulm0qaRxFoutzEmQlNM1OT+0Ra3Wr/76fRA+mDlnjlllDgvjSvhMq8V8rwl5wXs8NguNcHhad6ub5RSXLTo5rSpxdbOcmpJKCVTMsth/Z/1nC5V/wA2304PoPZv9HnNsRx7wFh8bnPgznGNShUftcXj+Acdi6tqOSZ3WfPVr5PWqzVDIs9rNyqSdPKs1qf2l9UxecfBFf3Wf51hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+9P/BIr/gsFr/7HesaV8Bvj5qWqeJf2Xtb1HytL1Jhc6prfwP1PUbgvNrGhwIJrzUfAV3cyvdeJvCdqktxYSyXHiPwvA2pNq2jeJuvD4l03yTu4N6Pdwb6/wCHq181rdP+kPBHxyxPA1ejw3xLVrYvhDEVbUaz5q2I4erVZ3lXw8VzVKuWznJzxeCheVKTlisHH2rr0MX/AHPeHPEegeL9A0XxV4V1rS/EnhnxHpllregeINEvrbU9H1rSNSt47vT9T0zUbOSa1vbG9tZY7i2ureWSGaKRXR2Vga9RNNXTTT1TTun53P8ARHC4rDY7DUMbgsRRxeExVGniMNicPUhWoYihVip0q1GrTcoVKdSElKE4ycZRaabubNM3CgAoAKAPFP2gv2fvhj+038MNe+E3xY0JNZ8N61H5ttdReXDrXhvWoI5V07xL4a1F4pm0zXdLeV2trlUkguIJLnTtRtr3Sr6+sbn5XjPgzIOPcgxnDnEeDWKwOKXNCpG0cVgcVFSVDHYGu4yeHxeHcm6dRKUZxlOjXhVw9WtSqfo3hX4qcZ+DXGmV8dcC5nLLs4y6XJWo1Oepl2b5dVnCWLyfOMJGcFjcsxsYRVai5Qq0qkKOLwlbDY7DYbE0f4gP20/2Lfif+xZ8UJvBXjSF9Z8I6y93e/Dr4iWdpJBovjPRIZVB+UtMumeI9MWaCHxD4eluJZ9OuJYbi3mvtIvtM1O+/wAn/FPwsz/ws4glleaReKy3FOrVyTO6VOUMLmmEhJXuryWHx2HUoRxuClOU6M5RnCVbDVqGIq/9JH0dfpFcGfSK4Lp8RcPVI5dn+XxoYfi3hLEV4VMx4ezKpBta2pyxuUY1wq1MpzanShSxdKFSjWp4bH4bG4LDfHNfmJ/QQUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/APJnf7J//ZtPwJ/9Vd4Wr/Z3wt/5Nl4df9kJwj/6z+Xn/Kj9Ij/lIDxz/wCzxeJv/ra52f4ntfdn46FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfrP/wAEzv8Agpp4p/Yy8TQ/Dz4hy6r4s/Zq8U6r52u6DCWvdX+G+rXzolx438CwSyKCGwsnijwqskVp4jtozc2xtdegtryT+U/pTfRY4V+kfwslL6tkfiJkeFqrhLi72T91XnW/sHP/AGMJV8Zw9jK8pS92NXFZPiqs8xy6FT2mYYDM/wBk8HvGHOfCzOG17bMeF8xrQedZJzrV2VP+0st9o1ChmdCmkndwo4+jCOFxUo8mFxOE/tI8I+LvC3j/AMLeH/HHgfxDpXizwd4s0q21zw14l0O6W90nWtJvFJgvbK4UK2NyyQXFvMkN5Y3kNxYX9va31rc20X/ORx1wLxV4bcVZxwXxrk+KyLiPI8S8Nj8BiknulOjisLXg5UMZgMZRlDE4HHYapVwuMwtWliMPVqUqkZP/AFK4e4iyfivJ8Dn2Q46jmOV5hSVXD4mi33tUo1qcrVKGJoVFKliMPWjCtQrRnTqwjOLR0NfJHtBQAUAFABQBnavpGl6/peoaJren2eraPq1ncafqemahbx3Vlf2V1G0Nxa3VvMrxzQzRuySI6kMCa9DKc2zPIczwGdZLj8XlWb5Vi6GPy3MsBXqYXG4HG4WpGrh8VhsRSlGpRrUakYzhOEk00eZnOTZTxFlOZZDn2W4LOMlzjBYnLs1yrMcPSxeAzDAYulKjicJi8NWjOlXoV6U5QqU6kXGUWz+fL9sr9jTVPgRqlx448DwXmrfCPVbwBWJkur7wPe3UmItI1iY75Z9ImlcQ6LrUxLMxj0vVJDqP2W71b/fL6H30wcs8cctocE8bV8JlfivleEbcf3eFwXG2DwtO9XNsppLlp0c2o0oOtnOTUkkoqpmeWQ+ofWsNlX/Nx9OH6DubfR7zXE8f8A4fG5x4M5vjUoybq4vH8AY7GVbUcmzqtLnq18lr1pqhkWe1nKUpOnlOb1f7SeDxmc/Atf3cf50hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+5n/BJL/grx4n/AGKdf0/4L/Gm91bxX+yz4i1Q4UC41PW/gvqupXBe68TeFLcebc3vhO6uZXu/Fng62DOZHuPEfhqEa42qaX4o6sPiHTfLJ3pv1bj5ry7r5+v9EeCnjfjOAMTSyDP6lfG8HYqtt71bEZBWqzvPF4KOs6mCnOTnjcBG7u5YrCRWI9tRxn92fhPxZ4Y8eeGdB8aeC9f0nxT4S8U6VZa54d8R6FfW+p6PrWj6lAlzY6jp1/avJb3VrdQSJJFLE7Ag4OGBA9VNNJp3T1TXU/0WwWNwmZYTDY/AYmjjMFjKNPEYXFYapGtQxFCrFTp1aVSDcZwnFpppv7zoaZ1BQAUAFAHiP7Qv7PXww/ae+F+u/Cf4saEmseHdYTzrO8h8uDXPDOtwxypp3iXw1qLxTNpmt6Y0rmCcRy291by3Om6lbX2lX19Y3PynGnBeQcfZBjOHOI8IsVgcUuanVjyxxeAxcYyVDH4Cu4ydDF0HKThO0oVISqUK8KuHq1qVT9I8KPFfjTwY40yzjrgXM5Zfm+Xy9niMPU56mW5zltScJYvJs5wkZ01jMtxqpx9rScoVaNaFHGYOthsdhsNiaP8AD9+2f+xj8T/2LvijceB/G8D6v4V1d7u++HfxDs7WSHRPGuhQyqDIgLTLpviDTFmgg8ReHpp5bnTLmWKaGW+0i+0rVL//ACe8UvC3P/C3iCeVZtB4nLsS6lXJc7pU5Rwma4SMlqruSoY2gpQjjcFKcp0JyjOMquGq4fEVv+kn6O/0iOC/pE8FU+JOHKkcBnmAjQw3FnCeIrwq5lw7mVSEmoyaVN4zKsY6dWplObU6UKWMowqU6tPDY/DY3BYX4+r8zP38KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Qa/Yv/wCTO/2T/wDs2n4E/wDqrvC1f7O+Fv8AybLw6/7IThH/ANZ/Lz/lR+kR/wApAeOf/Z4vE3/1tc7P8T2vuz8dCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD9Xv+CaP/BTDxZ+xf4pj8B+PH1Xxd+zZ4r1UT+I/DcDG81f4f6peNHFP468BxTSoguAqo/iXwx5sFj4os4QyvZ65bWGpQ/yt9KT6LXCn0j+FfZ1Pq2R+IWSYar/AKo8X+xbdN3nW/sPPfZRlWxvD2MrSlKUVGpicpxNSeY5dGbnjsFmP7F4QeMGc+Fmcc0fa5jwzmFWDzvJOde/oof2hlzqSVOhmdCmkk24UsbSjHDYppRw+Iwv9p3g7xl4T+IfhPw7478CeItK8W+DPFulW2t+GvE2h3Iu9K1nSrvd5V3azbUkUrIkttd2lzFBfadfQXWnaja2moWl1aw/84/HnAfFfhnxXm/BXGuT4nI+IskxLw+NwWJSalFrnoYvCV4OVHG4DGUXDEYLHYadTDYvD1IVqNSUJJn+pfDvEeTcWZNgc/yDHUswyvMKSq4fEUm7p3tUo1qcrVKGJoVFKliMPVjGrRqxlCpFSR0lfIHthQAUAFABQBn6tpOma9pmoaLren2eq6Rqtncafqem6hbx3Vlf2V1G0Nza3VtMrxTQTROySRupVlJBFd+VZrmWR5lgM5ybH4vK82yvF0Mfl2ZYCvUwuNwONwtSNbD4rC4ijKFWjXo1YRnTqQkpRkk0zzc5ybKeIcpzLIs9y7BZxkucYLE5dmuV5jh6WLwGYYDF0pUcVhMXhq8Z0q+Hr0pyp1KdSMoyjJpo/nz/AGy/2M9U+BOp3HjnwNBeat8I9VvAASZLq+8DXt1JiLSdXlO+WfR55XEOi61MSxZo9L1WQ6h9ku9X/wB8foffTCy3xvy3D8Ecb4jCZX4r5ZhH/wA+8LguN8Hhad6ua5VSXJSo5xRpQlWzjJqSUeVVMzyyn9QWLwuVf83X04voOZt9H3NcT4geH+GxuceDOb41dauMx/h/jsXVtRyfOasuetXyOvWmqGR57WlKXM6eU5vV/tF4PGZz8B1/eB/nMFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH7X/8EnP+Ct3i/wDYa8T2fwp+Ktzq/jD9ljxPqpk1LSY/O1HW/hLquoz7rrxj4ItyzST6PPM7XXizwdBiPUSZtb0RIvEH2y31/qw+IdJqMtabevVxv1Xz3Wvda7/v/gv42Y7w8xdPJs5nXx/B2LrXq0VzVcRklarK88dl8dXKhKTc8bgY6VfexGHUcT7SOJ/vI8EeN/CHxK8I+HfH3gHxHpHi7wX4u0mz13w14l0G9i1DSNZ0m/iE1re2V3CzJJG6na6nbLDKskE6RzxyRr6qakk0009U1qmf6O5fmGBzbA4XMstxVDHYDG0YYnCYvDVI1aFejUXNCpTnFtNNbreMk4ySkml1NM7AoAKACgDw39on9nf4X/tQ/C7XfhP8WNDXVtA1dPPsNQt/Kh13wtrsMUqad4m8M6i8UzabremtM5il2S2t5bS3Wmapa32lX19Y3PyXG3BWQcf8P4vhziPCLE4PErno1ocscXl+LhGSoY/AV3GToYqg5S5ZWlCpTlUoV4VcPWq0p/pXhN4s8aeC3GuWcdcDZk8DmuAl7LE4arz1MtzrLKk4SxeTZzhIzprGZdjFTj7SHNCtQrQo4zB1sNjsNhsTS/h+/bM/Y0+KH7F/xRuPAnjqBtV8Naq13ffD34hWVpLDoXjfQYZVXzodzzLp2vacJreHxF4emuJbrSbqWKWOW90m+0rVdQ/ye8UPC/iDwt4gnlGbweJwGIdSrkudUqcoYTNcHGS96N3JUMZQUoQxuClOVTD1JRlGVbDVcPiK3/SV9Hr6QvBf0iOCqXE3DVVYHOcCqGG4s4UxNeFTM+HM0qQk/Z1LKDxeV4x06tTKc2p0oUcdQhOE4YbH4bHYLC/IVfmp++BQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/oNfsX/wDJnf7J/wD2bT8Cf/VXeFq/2d8Lf+TZeHX/AGQnCP8A6z+Xn/Kj9Ij/AJSA8c/+zxeJv/ra52f4ntfdn46FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+q3/BNX/gpb4u/Yr8WL4J8atqvi/9m/xbqqz+KfCsD/atV8DapdGKCbx54BiuJY4o9QSNIv8AhIvDhmttO8W2NukcsllrNppWsWP8tfSi+i5wn9I/hT2Nf6vknH+SYet/qjxeqN50JNyq/wBi50qUXWxvD+MrSbqU0p4jLMRUlmGXpzli8Jj/ANg8IfF/OvCzOfaU/a5hw1j6sP7byRzsqi0h9fwDm1DD5nQgvdk3Gli6cVhsU1FUa+G/tT8FeNfCPxI8I+HPH3gHxHpXi/wV4v0q31vwz4m0S4+1aXrGl3O9UuLeQrHLFJFNHNaX1jdw2+o6XqNtd6XqlpZ6lZ3dpD/zi8fcA8WeGPFmb8E8bZPiMk4iyXEOhi8JXV4VIS97D43BYiN6WNy/G0nHEYLG4eU6GJoTjUpzabS/1M4b4kybi3JsFn+QY6lmGWY+n7ShXpu0oyTtVoV6b9+hiaE06WIw9VRqUqkZRnFNa9PXxx7gUAFABQAUAZ+q6Vpmu6Zf6NrVhaarpOq2lxYalpt/bx3VlfWV1G0Nza3VtMrxTwTxO0ckcisrKxBFd+V5pmWSZlgc4yfHYvLM1yzF0Mfl2Y4GvUwuNwWNwtSNbD4rC4ijKNWjXo1YxqU6kJKUZJNM87OMoyriDKsxyPPMuwWb5Nm+CxOXZpleY4eli8BmGAxdKVDFYTF4WvGdGvh69Gc6dWlUjKM4yaaP59v2zP2MtT+Bep3HjvwJb3mq/CTVbsZGZbu+8C3t1JiLStWlO+WfRZ5XEOi61MWYM0el6rJ9v+x3mr/73fQ9+mHlvjdl2H4H45xOEyvxXyzCvlf7vC4LjjB4Wm5Vc0yukuSlRzmjShKtnGTUUouKqZnldP6isXhMp/5vfpx/QbzX6P2aYrxC8PcNjc38Gs2xi543q4zH+H2PxlXlo5Rm9aXPWxGQ4itONDI89rylNTlSyjN6rzF4LG518AV/eZ/nEFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH7Kf8Epv+Csfjf9g/xdb/AA8+IMmr+Nf2XPFer+d4j8LRM15rHw51O/kVbrxv4AjmkUAk4n8SeFRJFZeIIke6tTaa4q3Vz04fEOk7Su4N6rrF91+q6+p+7+DXjTmPhzjYZXmbr5hwfja/NisGm6lfKqtVpTzDLVJ9/exeCuqeJSc4cmJ9+f8Ae18PfiH4I+LHgnwz8SPhv4n0jxn4F8Y6Tba34Z8TaFdLeaZq2m3QJjmglGHjlikWS3vLO4SG9sL2G4sb63t7y3ngj9ZSUkpRaaeqa6n+kWWZpl+dZfhM1yrGUMfl2PoxxGExeGmqlGvSntKMt0004VISUalKpGdOpGNSEorsqZ3hQAUAFAHhP7Rn7Ofwu/ak+F2ufCf4saIuqaFqi/adN1K28qHXvCmvQxSpp3ibwzqMkUx07WdOM0gRyktre2kt1peqW19pV9e2Vx8jxvwRw/4gcP4vhziPCLEYPELnoV4cscZl2MjGSoY/AV3GToYqg5OztKnVpyqYfEU6uHq1aU/03wk8W+NfBTjXLeOeBsyeCzPBS9ljMHWc6mWZ5llScJYvJs5wsZwWLy/FqEeaPNCvh68KONwVbD43DYfEUv4ff2yP2OPij+xj8UrrwF48tm1Pw9qbXV98PviBZWssOg+ONAilVRc225pRp+t6eJYIPEXh6aeW70i7lidJb3Sb3StV1H/J7xP8MOIPC7iCpk+cQeIwWIdStk2c0qco4PNsHGSXPC7kqOLoc0IY3BSnKphqkotSq4eth8RX/wCkr6Pn0g+CvpD8FUeKOGKywWbYNUcNxVwria8KmacN5pUg5exrWjTeLy3FuFWrlObU6UKGPoQnGUMPjsNjsDhPkevzY/eQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Qa/Yv/AOTO/wBk/wD7Np+BP/qrvC1f7O+Fv/JsvDr/ALIThH/1n8vP+VH6RH/KQHjn/wBni8Tf/W1zs/xPa+7Px0KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/U3/AIJs/wDBSnxj+xR4uXwh4vOq+Mf2cfFuqpceLvB0EgudU8G6nc+Vby+Pvh/HczRQQ6xFDHCNf8PvPa6Z4x062jtbuWx1ez0XXdJ/l36UH0XuEvpH8KfVsV7DJOPMlw9Z8I8Xxo81TC1HzVXlGcKnH22O4fxlZt1qK5q+X15vH4Be1eJw+N/XvCLxdzrwszr21H2mYcOY+rTWd5I6lo1oq0Pr2Bc3yYfM6ENIVNKeKpxWGxL5PZVcP/az4G8c+Dvib4P8N/EH4feJNK8YeCfF+lQa14Z8TaJObjTNX024Los0LOkU8E8E8U9lqOn3sFrqek6na3mlatZ2Wp2V3aQf84niD4fcW+F3Fub8E8b5RiMl4hyWu6OKwtZc1KtSl72Hx2BxEb0sbl2NpctfBY2hKdHEUZKUZX5ox/1M4a4lyXi/JcFxBw/jqePyzH0/aUa0HacJrSrh8RSfv4fFYed6eIoVEqlKommrWb6qvjD3goAKACgAoAoarpWm65pt/o2s2Fpqmk6paXFhqWm39vFd2V9ZXUbQ3Nrd20yvFPBPE7RyxSKyurEEHNd2WZnmOS5jgc3yjHYvLM1yzF0Mdl2Y4GvUwuNwONwtSNbD4rC4mjKFWjXo1YRqU6tOUZwnFNO55+b5Tlef5XmOSZ3l+DzbJ82weJy7NMrzHDUsZgMwwGMpSoYrB4zC14zo4jD4ijOdKtSqwlCcJOMk0z+fj9s39jHU/gZqVz498BW93qvwk1S7y6Zlur7wJe3UuItL1SVi8s+hzyusOja1MzOrtHpWqyfbvsd5q/8Avb9Dz6YmXeNuX4fgXjrE4TLPFbLMK+SdqeFwXHODwtPmq5nllJctKhndClCVbOMnoqMJwVTNMrprBLF4PKf+b76cn0Gs08AMzxXiJ4d4bGZv4NZtjF7Sneri8f4e47GVeWllObVZOpWxGQYitONHI88rSlOE5U8ozeo8weCx2dfn7X96n+cAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfrz/wSy/4KreP/ANgXxvH4R8Vtq3jj9mTxhq0c3jXwJHN9o1Hwff3TRwz+Pfh6lzKkFrrMESo2t6CZbfTPFlnAsFzJZ6rBpur2PRQxEqTs7uD3Xbzj59119dV+3+DvjLmXhtmCwONdfMOEcdXUswy5S5quBqztGWZZWpyUYYiKSeIwzlCljacVGbp1o0a9P++v4YfE/wAAfGfwD4W+KPwu8VaT418BeNNKh1nw34l0Wfz7HUbKYsjAh1juLS9tLiOay1PTL6G21LStRt7rTdStbW+tbi3j9eMlJKUXdPVNdf6/Pc/0nyjN8tz7LcHnGT42hmGW4+jGvhMXh5c1OrTldPe0oVISUqdajUjGrRqxnSrQhUhOK7ymekFABQAUAeDftIfs4fC79qf4W638KPivoo1LRdSButK1W2EUWv8AhLX4YpY9P8TeGdQkjlNhq+nmaQAlJbS/tJbrS9Utr3S728s5/kOOOB+H/ELh/F8O8RYVV8LXXtMPiKfLHGZdjIxkqOPwFeUZOjiaPM9WpU61OVTD4inVw9WrSn+n+EPi9xr4JcbZdxzwNmLweY4N+xx2BrOdTK89yupOEsXk2c4WM4LFYDFKEW0pQr4XEQo43BVsPjcNh8RT/h8/bE/Y7+KX7GfxTuvh/wCP7U6joWom6v8AwB4+sbWWLQPHPh+KVUF5Zl2lFjrNgJYIPEXh6aeW70a8ljIlvdLvNK1XUf8AJ3xO8MuIPC7iGpk2cQdfB13UrZNnNKnKODzbBxkl7SndyVHFUeaEcbgpTlUwtSUXzVcPVw+Ir/8ASX9H/wCkDwV9IbgmjxVwtWWEzPCKjhuKeFsTXhUzThrNZwcnh8Ryqm8Vl+KcKtXKs2p0oUMww8Jpww+Nw+OwOE+S6/OD92CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/0Gv2L/APkzv9k//s2n4E/+qu8LV/s74W/8my8Ov+yE4R/9Z/Lz/lR+kR/ykB45/wDZ4vE3/wBbXOz/ABPa+7Px0KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP1F/wCCbn/BSbxp+xL4wHhbxR/avjH9nPxfqsdx408FQSLPqXhTUbgQ20vj/wCHyXU0Vvba/BbxQLrmhvPa6X40021hsNQmsdVstB8QaH/MP0nfow8I/SP4TeDxqoZLxzk1CtLhHjCFHmrYKrK9R5Vmypr2uO4fxtX/AHjD3lWwVWTx2A5a3tqWK/XPCTxczvwszr2+H9pj+HsfUgs7yOVS1PEQVo/XMG5Pkw+ZUIfwqtlCvBfV8Tem4To/2veA/Hvgz4o+DPDXxE+HfiXSvGPgfxhpcOteGvE2iztPp2radOzx+ZH5iRXNtc21zFcWGp6Zf29rqmj6pa3ukavZWOqWN5Zwf84XiH4ecXeFnF2b8EccZRXybiDJq/s8Rh6vvUcTQld4bMMvxKXssdluNp2rYPGUHKlWpveM4zhH/U3hjifJOMMlwXEHD+Np4/LMdT56dWGlSlUWlXDYmk3z4fFYed6dehUSnTmusXGUutr4o98KACgAoAKAKOqaXput6bf6PrFjaappWqWlxYajp1/BFdWV9ZXUTQ3NrdW0yvFPBPE7RyxSKyOjEMCDXdlmZ5jk2Y4HN8ox2LyzNMsxdDHZfmOBr1cLjcFjcLVjWw+KwuJoyhVoV6FWEalKrTlGcJxUotNHn5tlWWZ7lmYZLnWX4PNsozbB4nLszyzMcPSxmAzDAYylOhisHjMLXjOjiMNiKM50q1GrCUKkJSjJNNn8/P7Z37GGpfA7Ubnx94BtrvVPhJqd3mWIGW7vvAd7dS4i03U5W3zT6FPM4h0bWZmZ43aPStWkN8bK91f/AHq+h39MXLvGvL8NwJx5icLlnitluFtSqtUsLguOcHhabdTMctprlpUM8oUoOtm+UUlGFSEamaZXT+pLGYPKf+cH6cv0Gcz8AszxfiN4c4XGZt4N5rjL1qKdXGY/w8x2Lq8tLK81qy562I4exFacaGSZ5XlOpTnKllGcVXj3gcdnX591/fB/m6FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH6v/APBL7/gqR8R/+Cf3j4aLq/8Aanjj9m/xlqsM3xC+G6XAku9FuphFbS+PPh8LuaO10/xXZ20cQ1HTXlttK8YafbRaXqstpe2+i67onRQxEqLs9YN+8u395efl169Gv2jwh8Yc18M8y+r1/bZjwrj60ZZnlSleeHnLlg8yyznkoUsZCCXtaTlCjjqUFRrOFSNDE4f+/r4R/F34cfHf4deFfiz8JfFml+NvAHjTTI9V8P8AiHSJWeC5gZmjntrmCVY7vTtU066jmsNX0jUYLbU9I1K2utO1G1tr22mhT1oyjNKUXdPZ/wDD6/fqf6V5JneVcR5Xg86yXG0cwyzH0VWwuKotuM4ttSjOMkp0q1KalTr0KsYVqFWE6VWEKkJRXo9UeqFABQAUAeBftKfs2fC39qr4Wa18KPitowv9I1AG70fWLQRReIPCHiCKKWPT/E3hq/kilNjqtiZXUhkls9Qs5brS9UtrzTby6tZvjuOuBuH/ABD4fxXDvEWF9tha/wC8w2Jp8scZluNjGUaOPwNaUZeyxFHma1UqdalKph8RCrQq1Kcv1Hwf8YONfBDjbLuOeB8weFzDCP2GPwFdzqZXn2VVKkJ4vJs4wsZw+s4HEqEWmpQxGFxEKONwVbD4zD0K9P8Ah7/bB/Y/+KX7GvxTu/h58QrU3+j35ur/AMBePLG2li8P+OvD0Uyot/YM7S/YtWsvNgg8Q+H555bzRb2SMGS80280vVdR/wAnvE3wz4g8L+IamS5zTdbC1vaVsnzilTlHB5tgoySVWk25eyxNLmhHG4OU5VcLVlH3qtCrh8RW/wCkzwB8fuCvpCcE0OK+Fa6wuY4VUcNxRwvia0J5rw1ms4OTw2JUVB4nA4lwqVcqzWnThh8xw8Je5h8Zh8bgsJ8oV+cn7mFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+g1+xf/yZ3+yf/wBm0/An/wBVd4Wr/Z3wt/5Nl4df9kJwj/6z+Xn/ACo/SI/5SA8c/wDs8Xib/wCtrnZ/ie192fjoUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH6ff8ABOD/AIKReNv2I/Gf/COeIhqvjL9nbxfqsc/jjwLBKs2o+GtQnWG1k8f/AA+S7mhtrTxLa20UCazosk9ppPjfS7SDS9VnsNTsvDviPw9/Mv0m/oycIfSP4ReAzBUcm41yejXnwhxhToc+Iy6vO83luZKFqmOyDG1UvreEcnUw1R/XcC4YiM41/wBa8JfFvO/CzO/rOGdTHZBjqlNZ3kcqlqWKpq0freEcm44fMqEL+xr25asV9XxHNScXT/ti+H/xA8E/FbwV4a+I/wAOPE2leMfA3jHS4dZ8NeJdGmabT9UsJnkiZlEqQ3VpeWl1Dcafqulahb2mraLq1pfaPrFjY6rY3lnB/wA4PiL4dcX+FPF+bcD8cZRWyfP8nrclajUvPD4vDzbeGzLLcUkqeOy3HU17XCYyi3CpBuMlCtCrSh/qdwvxRknGWSYLiDh7G08dluOhzQnHSrRqq3tsLiqTbnh8Xh5Pkr0J+9GVmnKEoTl2FfDn0AUAFABQAUAUdT0zTta06+0jV7G01PStTtLiw1HTr+3iu7K+srqJobm0u7aZXint54neOWKRGSRGZWBBNduW5lmGT5hgs2ynG4rLc0y3FUMdl+YYGvVwuMwWMwtWNbDYrC4mjKFWhiKFWEalKrTnGcJxUotNHBmuVZZnmW5hk2c5fg81yjNcHiMvzPLMxw1LGYDMMDi6U6GKweMwteFSjiMNiKM50q1GrCVOpCUoyi02j+fz9s/9jDUfgfqN18QPh/bXep/CXU7sGeEGW6vfAd7dS7Y9O1KRt80+g3EziHR9YmZnido9K1WQ3psr3V/96fodfTFy/wAasBhuAuPcThct8VcuwrVCs/Z4XB8dYPC03KpmGX01y0qGe0KUJVs3yikowqwjUzTK6awixmDyr/nD+nN9BjM/ATMsX4j+G+ExmbeDeaYy+Iw6dXGY/wAO8di6vLSyzM6snOvieHcRWnGjkud15TqUakqWUZxVeOlgcdnP58V/fR/m0FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH6jf8Ey/+CnXxO/4J9fEbyX/tLxt+z74y1O3f4m/C77UN0UjCK1bxv4FN1KlppPjbTrWONZY3eDTPFmn20Wi65JBJBout6BvQryoy6uD+KP6rz/P8T9g8JPF3N/DLNeV+1zDhnH1ovN8n59Yt2g8wy7nkoUMwpQSUk3GljaUI4fEOLjh8Rhv9Af4K/Gv4Y/tD/DPwr8X/AIPeLdN8aeAPGOnrf6NrWmuchgTHeabqdnKEu9J1vSrpZbDWNG1GG31HS7+Ce0vLeKaNlr14yjOKlF3T6/5+fc/0wyDP8o4oynB55keNpY/LcfSVWhiKTfe06VWErVKOIozUqdehVjGrRqxlCpFSTR6nVHsBQAUAFAHz9+0v+zR8Lf2rfhZrPwp+Kuji90u+DXeia3aLDH4h8HeIYopI7DxL4av5I5DZ6lZmV0kjdZLPUrKS50vU7e7067ubaT43jvgTh/xE4exXDvEWG9th6373C4qnyxxuWY2MZRo4/A1pRl7LEUuZppqVKvSlUoYiFShUqQl+qeDvjFxt4HcbZfxxwPmH1bHYZqhmOXV3UnlWf5VOcJYrJ84wsJw+sYPEckZRlGUMRg8RCjjcFWoYuhRrQ/h7/a9/ZC+KX7G/xUvPh18RLQ3ul3hub/wL46sbaaLw/wCOvDscwRNS05naX7JqVp5kMGv6BPPJe6JfSIjSXen3WmanqH+TviX4a8QeGHENXJM6purhqvtK2UZvShKOCzfBRkkq9Fty9liKXNCGMwc5yq4WrJJupQqYfEV/+kzwD8fOCfpB8E4fi3hPEfV8dh/Y4XibhnE1qc814ZzadNylhMWoqH1jB4jkqVMrzSnThh8xw0JSUaGKo4zB4X5Ur87P3AKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBBr9i//kzv9k//ALNp+BP/AKq7wtX+zvhb/wAmy8Ov+yE4R/8AWfy8/wCVH6RH/KQHjn/2eLxN/wDW1zs/xPa+7Px0KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/TX/gnJ/wAFHvHH7EPjVtC1tdU8Zfs9eMNUin8eeAYZlkvtBv5UhtH8f/D9buaG1svFVpawwR6tpUs9ppPjfSrSDSNYnsdQsfDfiXw1/NH0mfozcH/SO4QlluZxo5RxnlFGvU4Q4wp0FPE5ZiZpzeX5goWqY7IcdVUVjcE5OdGTWMwTp4qnep+seE3iznnhbnaxWEc8dkWOqU455kkqnLSxdKL5frOGcrxw+ZYeLboYhK1RfuMRzUZe7/bR8O/iJ4H+LXgjwz8Sfht4m0vxj4G8Y6XFrHhvxJo8ry2OpWUjyQuCk0cN3ZX1ldw3OnatpOo29pq2i6taX2j6xY2OqWN5Zwf84HiR4b8YeE/GGbcD8c5TWyjPsoq8tSnK88LjcLNy+q5nlmK5VTx2WY6EXVwmLpe7Nc1OpGliKVajT/1O4V4qyPjTI8FxDw9jYY7LcbC8ZL3a2HrRt7bCYujdyw+Lw8nyVqM9U7Ti50p06k+zr4U+iCgAoAKACgClqem6frOnX2kavY2mp6XqdpcWGo6df28V1ZX1ldRNDc2l3bTq8M9vPC7xzRSoySIzKwIJrsy7McflGPwWa5VjcVluZ5biqGOy/MMDXq4XGYLGYWrGthsVhcTRlCtQxFCtCFWlWpzjOnOKlGSaTOHNMry3O8tx+T5xgMHmuU5rg8Rl+ZZbmGHpYvA4/A4ylOhisHjMLXhUo4jDYijUnSrUasJ06lOcoTi02j+f79s/9i/UPgjf3XxC+H1td6n8JtSu83FuPNur3wFe3Uu2PT9QlYvNceH7iZ1h0jWJmaSCRo9K1WQ3hsr3Vv8Aef6HX0xsB40YDC8A8f4rC5d4q5dhWsPiGqeFwfHWDwtJyqY/AQjyUaGf0KUJVs2ymkowrwjUzTK6awixmDyr/nF+nP8AQYzLwGzHGeJPhthMZmvg5mmM5sVhV7bGY/w7x2Lq8tPLsyqydSvieG8RWmqOS51XlKph6k6WUZxVljXgcfnP561/fp/muFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH6Xf8E2f+ClnxV/4J8fEz7ZpxvvGfwN8X6han4p/CeS82QX8YEdt/wlvhB7h/s2i+O9LtEVYLo+VY+ILKCPRNezCmmajo29GvKjLq4P4o/quzX47M/WfCnxZznwyzbnpe0x/D2Oqw/tjJXUtGotIfXcC5vkw+Y0YfDP3aeKhFYfEvlVKrQ/0EvgP8efhX+0v8LfC3xk+DPiux8YeA/FtmLmw1G0Pl3Vjdx4XUND1zT5D9q0bxBo9wWs9W0i+SO7srlCrq0bxSyevCcakVKLun/Vn59z/TPhziTJuLMnwefZDjaeOy3G0+elVhpOnNaVcPiKT9+hiaE7wr0KiU6c1Zpppv2CqPcCgAoAKAPnv9pv8AZl+Fv7WHwr1j4VfFTSPtenXga80HX7NYY/EPg3xFHDJHYeJfDd9LHJ9l1C08xkmhdZLLVLGS50zU7e5sLqeF/i+POA+H/EXh7FcPcQ4b2tCrerg8ZTUVjcsxqjKNHHYGtJP2danzNSi70sRSlOhXhUo1Jwf6t4NeMvG3gZxtgOOOCMf7DF4e2HzTK8RKpPKeIcpnUjPE5RnGGhOHt8LX5FKnUi44jB4mFLGYOrRxVGnUj/D1+11+yN8U/wBjn4qXvw4+I1n9q0+6+03/AIH8cWNvNH4e8d+HY5hHHqmlvIZPst/a+ZDBr2hTTSXuiX0ixSPc2Nzpupah/k74leG3EHhjxDVyPO6XtKFTnrZTm1GElgs3wSlZYig3f2danzRhjMJKUquFrNRk6lKdCvW/6TfAbx54J+kFwRhuLuEcR7DF0fZYXiXhrE1qc824ZzaVNyngsbGPL7bC1uWpVyzM6dOOHzLDRc4Ro4mjjMHhflivz0/bQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i//AJM7/ZP/AOzafgT/AOqu8LV/s74W/wDJsvDr/shOEf8A1n8vP+VH6RH/ACkB45/9ni8Tf/W1zs/xPa+7Px0KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP0u/4J0f8ABRrx1+w/43bSNWXU/Gf7PvjDU4Z/H/w+hnR73R72RIbRvH3w/F5PDaWHi6ytIYItS02ae00jxxpVpBout3FjfWXhrxN4X/mz6S30aeDvpHcIPKs2VPKOLsppV6vCHGFGgqmLyjF1FzSweNjFxnj8ix1SMFj8vlNNPlxeEnRxlKnVP1bwo8WM98Lc8WMwTnjckxs6cM8yOdRxoY2jF2Veg3eOGzHDxcnhsUou+tGuqlCcoH9tnw3+JHgX4v8AgXwz8S/hp4n0zxj4F8Y6ZHq/h3xHpErvZ39o7yQSo8U8cN5YahYXkNzp2r6RqVtaatomrWl9pGr2Vjqlld2kP/N94l+GvGPhJxjm3A3HOU1cpz7KanvRfNUweYYOpKX1TNcqxfLGGPyvHQg6mFxVNK9qlCvChi6GIw9L/U/hTivI+NcjwfEPD2NhjcuxsdHpGvhq8UvbYPG0bylh8Zh5SUa1Gbe8alOVSjUpVZ9tXwZ9GFABQAUAFAFLUtO0/WNPvdJ1aytdS0zUrW4sdQ0++t4rqyvrK6iaG5tLu2nV4bi3uIXeKaGVGjkjZldSCRXZl+YY/Kcfg80yvG4rLsyy7FUMdgMfgq9XC4zBYzC1Y1sNisLiaMoVqGIoVoQq0q1KcalOpGMoyUkmcOZ5Zl2dZdj8ozfA4TNMqzTCYjL8yy3MMPSxmBx+BxdKdDFYPGYWvGpRxOGxNGpOlXoVoTp1ac5QnFxbR+AH7aH7F2ofBO/u/iH8PLW61P4Tajdbrq1Xzbq98A3l1KFjsb+Ri81x4euJpFh0nV5i0lvI0elarIbtrK91X/eX6HP0x8B4zYHC+H/iBisLl3ipl+F5cLipezwuD47wmGpuVTG4GEeSjQ4goUoSq5rlVJRhiIRqZnldNYaONweV/wDOR9Of6C2ZeA+YYzxL8NMJi808HczxfNjMHF1cXj/DrG4uqo08vzCrJ1K+J4axFeoqOTZ1WlOphqk6WUZxVeLeBx+cfnnX9/n+agUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfox/wTn/AOCjvxa/4J9fFIa74ce68W/CHxXeWcfxV+E11etDpviKyiIhGvaBJL5kOh+ONJti39l6zHH5V7Eo0nWY7rTnQW21GtKjK61i/ij3812fn95+peFvipnfhlnH1nCueNyPG1ILOclnU5aWKpx91YnDSleOHzChFv2NdK1RL2OIU6TXJ/oNfs8ftEfCX9qb4T+GfjP8FfFVr4q8E+J4CY5U2waromqwJGdT8NeJtLLvPoniTRpZVh1LTLkll3Q3drLd6dd2V7c+xCcakVKLun/Vn59z/TbhfijJeMclwmf5BjIY3L8ZHSStGth60UvbYTF0buWHxdByUa1GeqvGcHOlOnUn7dVH0AUAFABQB87/ALT37MPws/az+Fer/Cv4p6T9psrnde+HvEVkkKeIvBfiKOGSOy8R+HL2WN/s97b72jubaQPZarYyXGm6lBcWdxJGfiuPuAeH/Ebh7E8PcQ4b2lGpergsbSUVjcrxqi40sdgask+SrC7jODvSxFGU6FeE6VSUT9Y8GPGfjbwL43wHG/BOO9jiaNsPmuVYiVSWU8Q5TOpGeJyjN8PCUfbYatyqdKtFxxOCxMaWMwlWliKUJr+Hn9rb9kr4p/sd/FS/+G3xIsTcWc/2i/8ABPjWxt5k8O+O/DiTCOLV9JlkL+ReQb4oNc0SaV77RL5xBOZ7Sew1C+/yc8SPDfiDwy4hrZFnlL2lKfPWyrNaMJLBZvgVK0cTh5Pm5KsLxhi8JOTq4Ws+WTnTnRrVv+k3wH8d+CvpAcEYXjDhDE+yxNL2WF4j4cxNWnLNuGc3lT554HHQio+1w9XlnVy3MqdOOGzLDJ1KapV6eKwuG+Xa/Pz9rCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/AEGv2L/+TO/2T/8As2n4E/8AqrvC1f7O+Fv/ACbLw6/7IThH/wBZ/Lz/AJUfpEf8pAeOf/Z4vE3/ANbXOz/E9r7s/HQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP0n/AOCdf/BRbx5+w946bTtQXU/GXwB8YanBN8Q/h3FOjXem3bpDaN488BfbJobTT/GVhaQwRX9jLPaaT420m0t9D12eyu7Lw14l8Lfzh9JP6NfBv0jeDpZPnMaeU8V5VTr1uEOMKNCNTG5Ljakbyw2KinCePyLHTjCOZZZOpFTUYYrCzw+OoUMRD9T8KfFbPfC3PFjsA5Y3JsZOnDPMjqVHHD4+hFtKtRbUlhsxw8ZSeFxcYtpt0a0auHqVKcv7b/hn8TPAfxk8B+Gfib8MvE+m+MfAvjDTY9V8P+IdKkdra8tmeSGaGaGdIbzT9S0+8huNO1jR9St7TVtF1a0vdJ1ays9Ss7q1i/5vPE3wy4y8IeMs14F46ymplWe5VUT+1UwWZYGpKawmb5Ri3CEMflWOjCU8NioRi1KNXDYmnh8bh8ThqP8AqjwlxbkXG+RYPiLh3GxxmXYyPlHEYXERUXXwWNoc0pYfGYdyUa1GTas4VaU6tCrSq1O6r4E+kCgAoAKACgCnqOn2Gr2F7pWq2VrqWmala3FjqGn30EV1ZX1ldRNBc2l3bTq8Nxb3ELvFNDKjxyRuyOpUkV15fmGOyrHYPNMsxmKy/MsvxVDG4DH4KvVwuMwWMwtWNbDYrC4mjKFahiKFaEKtGtSnGpTqRjOMlJJnFmWW5fnOX47KM3wOEzPKszwmIwGZZdj8PSxeBx+BxdKdDFYPGYWvCpRxOGxNGpOlXoVoTp1ac5QnGUZNH4Bfto/sW3/wUv7v4ifDu1utS+E+o3W68s1826vfAN5dSgR2V7Ixea48OzzOItJ1aZmktpGj0vVZGuTZXuqf7x/Q5+mRgfGXBYTw+8QsXhcv8VMvw3Lg8ZL2eFwfHeEw1JynjMHCPJRw/ENClCVXNcqpKNPEwjUzPK6aw6xmCyz/AJyvp0fQWzHwJzDG+Jnhng8Xmng9meL58dgoOri8d4dY3F1VGngcfUk518TwziK9SNHJ84rSnUwlSVPKc3qvFSwGOzb88K/0CP8ANEKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP0A/4J7/APBQz4v/APBP34sJ4v8ABc03iX4beJLiytfip8J76+kg0PxnpEDlVvrJys0eieMtIhkmk8PeJIbeSS3d5LDUYNQ0W8vtPuNqNaVGV1qn8UejX6NdH+lz9M8MfE/PPDPOljsBKWLynFypwznJalRxw+PoRdvaU3aSw+PoRcnhcXGLcW3Sqxq4epUpS/0JP2aP2mPhB+1t8I/Dnxp+CfiaHxH4R1+Mw3NvJ5dvr3hfXYIoX1Twp4s0pZppNG8R6Q88a3llJJJDPBLa6npl1qGj6hp2o3frwnGpFSi7p/en1T7Nf8HZn+nHCfFuR8bZJhc/4fxccVgcSuWcXaOJweJioutgsbR5pOhi6DklUptyjKMoVqM6tCrSqz98qz6QKACgAoA+c/2ov2XfhZ+1t8K9V+FnxS0vzrWfffeGvEtkkK+IvBXiNIXjsvEXh68lRvJuoN5ivLOTdY6vYPPp2owzWs7KPiOP+AOH/Efh7E8P8QYfnpzvVwOOpKKxuV41RcaWNwVWSfLUhflq05XpYii50a0ZU5tH634LeNPG3gTxvgeNuCsd7OvS5cNnGUYiVR5TxHlEqkZ4jKc2oQkvaUavKp0MRC2JwOJjTxeEqU61NN/w8/tZfsm/FP8AY9+KmofDT4l2HnW8vn3/AIM8Z2MEy+HfHXhxZvLh1nRppN/lXEW6OHWdGmle+0S/b7NcGa3lsr29/wAnfEfw54h8M+Ia2RZ7R5qcuetleaUoSWCzfAqfLHFYaUr8tSN4xxWFlJ1sJWfJPmhKlVq/9J3gV46cE+P/AARheMeDsVyVoeywvEPD2Kq03m3DObunz1MvzCnG3tKVS06mXZjTgsNmWGXtaXJWhicNh/mGvgD9nCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i//AJM7/ZP/AOzafgT/AOqu8LV/s74W/wDJsvDr/shOEf8A1n8vP+VH6RH/ACkB45/9ni8Tf/W1zs/xPa+7Px0KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/R/8A4J3/APBRHx9+w748a0ul1Lxj8BvGGpQS/EX4cx3CNcWdwyQ2h8d+BBdzQ2eneNNPs4YYrq1lmtNK8aaVaQaFr09ncWfhzxJ4X/nP6SP0beDPpG8HSyTPIwyrijK6detwhxhQw8auPyLHVIpyoV4qVOWPyPHThThmmVTqwjWjGGJw1TDZhh8Li6P6j4V+Kue+FuerMMvcsZlGMlTp55kdSq4YfMcPFu1Sm7SWGzDDqU5YPGRhJwblSqxq4arWo1P7c/hf8UfAPxp8A+Gfif8AC/xPpvjDwN4w05NT0HXtLdzBcwl3huLe5t50ivNN1TTryK407WNH1K3tNV0bVLW80vVLO01C0uLeP/m88UPC/jPwf4zzXgXjvKp5XneWTUoyi5VcBmmAqSmsHnGT4xwpxx+VY+NOcsNiYxjOM4VsLi6WGx2GxWFof6pcI8XZFxxkWD4i4dxkcZl+Mi007QxODxMFF18DjqHNKWGxmHckqtKTlGUZQr0Z1sNWo1qnfV+fH0oUAFABQAUAU9Q0+w1awvdL1SztdR03UbWex1DT76CK6s72zuomhubW7tp1eG4t7iF3imhlRo5I2ZHUqSK68Bj8dleOweZ5ZjMTl+Y5fiaGNwGPwVerhsZg8Zhqsa2GxWFxNGUK1DEUK0IVaNalONSnUjGcJKSTOLMsuy/OMvx2U5tgcJmeV5nhMRgMxy7H4eli8Dj8Di6U6GKweMwteFSjicNiaNSdKvQrQnTq05yhOMoyaf4CftpfsW33wWvrz4jfDm0utR+FGo3W+9sl826vfAN5dS4SzvJGMk1x4cnmcRaVqsrNJaSPHpeqSNcGzvdT/wB4focfTIwPjHgsJ4e+IeLw2X+KeAw3JgsbP2eGwfHeEw1NueLwkFyUqHEVClB1c0yukowxdONTM8spqhHG4PLf+c36dP0Fcw8C8fjfE7wxwWLzPwezLF8+PwEHWxeO8Osbi6vLDB42pJ1K+J4XxFepGjlGcVpTq4OpOnlOb1ZYh4HH5r+d1f6CH+ZwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfdf7A37fvxj/YD+LkHj74e3L674J12Sysvih8LNRvZoPDnj7QbeVyAxVJ10fxTpSTXE3hnxTb2013pVzLNbXMGpaHf6vo+o7Ua0qMrrVP4o91+j7P80fo3ht4l574a53HMsrm8Tl+JdOnnGTVako4XMsNFvraXsMZRUpSwmMhCU6MnKE41cPUr0Kv+hT+yz+1R8HP2xPhBoHxo+CniNNa8Oauv2XVdKuvJt/EvgzxHDDFJqfhLxdpUc07aVr2mGaMyR+bNZ39nLaavo95qOjahYahc+vCpGpHmi7r8U+z8z/Tng7jHIuOsjw2f8P4pYjCV/crUZ8scXgMVGMXWwWOoqUnRxNHmV1eUKtOUK9CpVoVadWf0ZVn1IUAFABQB83/tS/stfCz9rn4V6p8LvijpnmQy+ZfeGPE9lHCPEXgnxGsLx2niDw/dyqfLmj3eVf2MpNjq9g02n38UkMoKfD+IPh9w94k8PYjh/iDD80Zc1XAZhSjH69lWOUXGljcHUkvdnG/LWoybpYmi5Ua0ZQlp+veCfjZxt4D8b4LjbgrG8lWHLhs6ybEzqPKeI8olUjPEZVmtCElz0529phsTC2JwGKjTxWFnCpD3v4ef2r/2Ufin+yB8VNR+GXxM0/fG/nX/AIP8YWMMw8O+OfDgmMdvreizybtkiZSHV9Jmka+0W/LWt0Hja1urr/JzxF8OuIfDTiGvkOfUeaL5q2WZnRhNYLN8DzWhisLOV7SV1HE4aUnVwta9OpzRdOpU/wCk3wN8cuCfH3gjCcZcG4vlnH2eG4gyDE1KbzbhrN3T56uXZjThbmhK0qmAx9OKw2Y4ZKvRcZxr0KHzLXwR+yhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6DX7F/8AyZ3+yf8A9m0/An/1V3hav9nfC3/k2Xh1/wBkJwj/AOs/l5/yo/SI/wCUgPHP/s8Xib/62udn+J7X3Z+OhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+i/8AwT0/4KG/ED9hzx80Ui6j4w+Bfi/UbeT4j/DZLhPNilKQ2h8b+BzdzRWmmeNtOs4oYpopZrXS/GGmWsGg6/NbS2vh7xD4Y/nf6R30cODPpGcGzyHPoQyziXLIYivwhxhh8PGrmGQZhViuanUjzU5Y/JMdKnShm2U1KsKeJhCniMPVwuY4bB43D/p/hZ4p594XZ6sxy2UsXlWLlTp53kdWq4YXMsNBu04u01hsfh1OcsHjYwlKlKUqdWFbC1a9Cr/bv8Kvir8Pvjd8P/DPxR+FvifTvGHgbxfp66jomuaa7+XKm5obmzvLaZIrzTNX0y7jn0/WNH1GC21PSNTtrrTtRtba8t5oU/5uvFHwt408HOM804E47yqeWZ3lklOnODnVy/NsuqzqRwec5NjJQpxx+VY9U5vD4iMIVKdSnXweMo4XMMLjMJQ/1S4P4wyHjrIcHxFw7jFi8Bi04yjK0MVgsVBRdfAY+gpSeHxmHcoqpTcpRnCVOvQqVsNWoV6noVfnh9OFABQAUAFAFS/sLHVbG80zU7O11DTtQtZ7K/sL2CK6s72zuomhubW6tp1eG4t7iF3imhlR45Y3ZHVlYg9WBx2NyvG4TMstxeJwGY4DE0Mbgcdg69XDYzB4zDVY1sPisLiaMoVqGIoVoQq0a1KcalOpGM4SUkmceYZfgM3wGNyrNcFhMyyzMsJiMBmOX4/D0sXgsdgsXSnQxWExmFrxqUMThsTQqTo16FaE6VWlOUJxlGTT/Ab9tL9iy++DF7efEf4cWd1qPwp1C5339gnm3V54BvLqXCWt07F5rjw3PM4i0zVJWeSzkaPTNUkaY2d7qX+7/wBDf6ZOC8YcHhPDvxExmGwHilgMPyYDHz9nhsHx3hMNTcp4nDQXJRw/EdClB1Myyyko08ZTjUzLLKapLGYLLv8AnP8Ap1fQUx/gbjsd4oeF+CxeZeD+Y4r2mZZdD2uLx3hzjcXV5YYTFzk6lfE8LYivUjRyrNqsp1cDUnTyrNqsq8sDjs0/Ouv9Bz/MsKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAPtf8AYW/bs+Mv7BXxftfiX8ML46n4d1VrPT/iV8NNTu5ovC/xF8NwTM5sb9UWb+zdd04TXM/hnxRbW8t/oN7NMDFf6Pf6zo2qa0qsqUuaOt910a/z7Pp8z9A8OvEbPvDfPIZtlFT22FrOFLNsprTksHmmFjJv2dVK/ssTS5pywmMjGVTDVJSuqtCpXoVv9Cz9kj9rj4NftpfB7RPjL8F9eGoaTe7bHxH4dvmgh8U+BPE8cEUuoeFPFumRTTGw1Wz81ZIZo3m07V7CS21fR7u+0y8trqT2KdSNWKlF6PdPdPqn/Wu6uf6dcFcbZDx9kWHz7IcT7WhUtTxWFqcscZl2LUU6uDxtGMpezrU73jJOVKvTca1CdSlOM39O1Z9cFABQAUAfNf7VP7K/ws/a7+Fep/C/4n6blW82+8K+KrGKH/hIvBHiMQtHa69oVzIvDLkRajp0rfYdYsDJY3qFWjlh+F8QvD3h7xJ4exGQZ/QunzVcuzGlGP17KcdyuNPGYOpLZr4a9CT9liqPNRqppqUf2DwR8buNvAXjfB8acF4y0lyYbPMjxM6n9k8SZQ6inXyvNKMHrF61MJi4L6zl+KUMThpKSnCp/Dz+1X+yp8U/2QvipqXwx+J2nZ/1t94S8W2MM3/CO+OPDnnGO213Q7mQf7sOqaZK5vtGv99neKf3E9x/k74ieHnEPhrxDXyDPqF/irZbmVKM/qWbYHmahi8JOXXaOIw8n7XC1r0qifuTn/0neB/jhwT4+cEYPjPgzF6+5hs+yLE1Kf8Aa3DWb+zU62WZlSg9etTBY2nH6tmGG5cRh5X9rSpfNNfBn7EFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+g1+xf/wAmd/sn/wDZtPwJ/wDVXeFq/wBnfC3/AJNl4df9kJwj/wCs/l5/yo/SI/5SA8c/+zxeJv8A62udn+J7X3Z+OhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfoh/wT4/4KEfEP9hz4gMVW/wDF/wAEfF2oW7/En4aC5UF32RWv/CZ+DDdSJaaV430y0jjjO+S303xXp1tDoWvyQm30PW/Dv89fSM+jlwX9IzgyfD/EMI5bxFlscRiOEOMMPh4Vcy4dzGrGPNGUXKnLH5LjpU6VPN8oqVqdLF04U69CrhMywuAx+F/TfC3xSz7wuz6OZ5ZJ4vK8XKnTzzI6tWUMLmmFhJ2adprDY/Dqc5YHHQhKdCcp06kK+FrYnDVv7evhN8Wvh58c/h74Z+Kfws8T2Hi7wP4tsFv9H1mwZhnDGK70/ULSUJd6XrOl3SS2GsaPqENvqOl6hBPZXtvFPE6D/m58VPCvjTwa40zTgTjvK5ZbnOXSVSjWpudXLc4y2rOpHB51kuNlTpLH5VjlTqOhXUKdWlVp18FjaGEzHCYzB4f/AFT4P4xyHjvIcHxFw7jFi8Bi1yzhLlhisDi4Ri6+AzCgpzeHxuGc4qpTcpQnCVPEYepXwtahXq+jV+dH1AUAFABQAUAVL+wstUsrzTdTs7bUNO1C2nsr+xvYIrqzvbO6iaG5tbq2mV4bi3uIXeKaGVHjkjZkdSpIPVgcdjcsxuEzHLsXicBmGAxNDGYHHYOvVw2LweLw1SNbD4rC4ijKFahiKFaEKtGtSnGpTqRjOElJJnJmGX4HNsDjcrzTBYXMctzHC4jA5hl+Ow9LF4LHYLF0p0MVhMZha8KlDE4bE0Kk6NehWhOlVpTlCpGUZNP8Cf20/wBiu++DV7efEj4b2d1qPwqv7nzNR09PNurzwDd3MuEt7lyXmuPDU8rrFpupyl5LKRk03U5DK1ne6h/u79Df6ZWC8YMHg/DrxGxmGwHijgcP7PLswqezw2E48wmGptzxGHiuSjh+JKFKEqmZZbTUaeOpxnmWWwVNYzB5f/zo/Tr+gpj/AAPx2O8UfC7BYrMfCDMcV7TM8sp+1xeO8Osbi6toYbEzk6lfE8K4ivONLK81qynVy+rOnlWa1JVZYHH5l+dFf6En+ZAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfYf7E37bnxn/YU+MNh8VfhNqf2iwuza6f8AED4fanczr4U+I/hiKdpJNG1yCLf9mvrbzZ5/D3iK2hbUvD1/K89v9osbnU9M1HWlVlSlzR/7eT2kvPz7Pp83f7rw/wDEHP8Aw6zynnOS1uenPkpZnllaUvqWa4SMuZ0MRFX5akLylhcVBOrhaknKPNTnWo1f9C/9jj9sn4MftvfB3Svi/wDB3WfNhfydP8YeDtRkgTxX8PfFHkLLd+G/E9hFI/lTJlpdN1ODfpmuaf5eo6ZPNC7rF69OrGrHmi/VdU+z/wA+p/p3wLx3kPiDkVHPMixHNF8tPHYGq4rG5ZjOXmnhMZTi3yyWsqVWN6WJpWq0ZSi3b6urQ+zCgAoAKAPmf9q39lP4WftffCvUvhj8TdOww86+8I+LrGGE+IvA/iPyTHba5odxKBlT8sOq6XK4sdasN9ndqGFvcW/wniJ4ecPeJfD1fIM+oa+9Wy3MqUY/Xspx3K1DF4ScvlHEYeT9liqN6VRX5Jw/Y/A/xw428A+N8Hxnwbi7r93hs+yHE1Kn9k8S5R7RTrZbmVKD0fxVMFjYReJy/FcuIoNr2tKr/Dz+1R+yx8U/2Rfipqnwv+J+m7ZE82+8K+KrKKY+HfG/hwzNHaa/oN1IvzRvgRajp8rfbtHvxLY30ayIjy/5OeIXh7xD4bcQ4jIM/oWa5quXZjSjL6lm2B5nGnjMHUlun8NehJ+2wtbmo1UmlKX/AEm+CPjdwT49cEYPjTgvGXjLkw2d5JiZ0/7W4bzdU1OvleaUYPSUbueExcF9WzDDOGJw0nGUoU/m2vhT9gCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Br9i/wD5M7/ZP/7Np+BP/qrvC1f7O+Fv/JsvDr/shOEf/Wfy8/5UfpEf8pAeOf8A2eLxN/8AW1zs/wAT2vuz8dCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP0K/wCCfn/BQT4ifsO/EMzQC+8W/Bbxbf2x+Jfw0+0qBcKFjtv+Eu8IG5kS10rxvpdoipHI7wWHiSxgi0PXXSOLSdV0P+ffpFfR04L+kXwXU4d4jpxy7P8ALo4jE8IcX4bDwq5nw5mVWEVL3XKm8fk2OdKjTzjJ6tanRx1GnSrUauDzPCZfmOD/AEzwv8Uc/wDC7Po5plcnistxTp0s7ySrVlDCZrhYSdtbTWHx+HU5ywOPhCVTDznOnOFfCV8Vha/9vvwh+L3w7+O/w78NfFX4VeJrHxZ4I8WWQvNK1WyYq8ciny7zTNTs5Nt1pWt6VciSy1fSL6OG+069ilt7iJWUE/8ANx4reFPGvgxxrmnAnHeVyy7OMvaq0K9Jzq5ZnWWVZ1I4PO8kxsqdJY7KscqU/Y1uSnWoVqeIwOOw+EzLB4zB4f8A1T4N4yyDjzIcJxFw7jFisDik4VKc+WGLwGLhGLr5fmFBSm8PjMO5x9pT5pQqU50sTh6lfC16Fer6TX5wfUhQAUAFABQBVvrGy1Oyu9O1G0tr/T7+2ns76xvYIrqzvLO6iaG5tbq2mV4bi3uIXeKaGVHjljdkdWViD04LG4zLcZhcxy/F4nA4/A4mjjMFjcHXqYbF4PF4apGth8VhcRRlCtQxFCtCFWjWpTjUp1IxnCSkkzkx+AwOa4HGZZmeDwuY5bmOFxGBzDL8dh6WLwWOwWLpToYrCYvC14VKGJw2JoVJ0a9CtCdKrSnKnUjKMmn+Bf7an7FV78HLy9+JXw1srm/+Fd/c+ZqWmR+bdXngG7uZcLBO7F5rjwzPM4j07UpWeWwkePTdTkZzZ3t//u39Db6ZWD8XsJg/DjxHxmHwPihgcP7PLcyqezw2E48wuGp3lWoxXJSw/EtClCVTMMupqNPH041Mxy2CgsZg8D/zqfTr+gnjvBHG47xT8LMDisx8IswxTq5rlVP2uLxvhzjMVVtGhXlLnr4nhTEVqipZZmdWU6uXVJ08rzWpKbwWOzH85a/0MP8AMUKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAPrH9jb9sz4z/ALD/AMYtL+Lvwe1jY/7jT/GfgzUpZ28KfEPwuJxLdeHPEtjE670OXl0rVYNup6FqBW/06ZH86KfSnVlSlzRfqukl2f8An0PtOBOPM+8Pc9o53kVez92nj8BVlN4LNMJzXnhcXTi1dauVGtH97hqtqlKSfMpf6F/7FH7bPwY/bq+D2n/Fb4Sap5N5bfZtP8feANTuIG8V/DrxPJAZZdF123iK+fZ3HlzT6B4gt4103xBYRtcWphu7fUdP0/2KVWNWPNF+qe8X2f8An1+8/wBPOAPEDIfEXIqWdZJW5Zx5KWZZbVnF43KsY43lh8TBfFCVpSw2JivZYqmnOHLONWlT+wK0PuAoAKACgD5k/aw/ZP8AhZ+2B8K9R+GfxM0/y5k86/8AB3jGxghbxH4G8RmHy4Na0WeXb5kUmEh1jSJpFsdbsAbW58uZLS7tPgvEXw64e8TOHq+Q59RtJc1bLMzowi8dlGO5bQxWFnK14vSOJw0pKli6N6dS0lTqU/2XwM8c+NvALjfCcZcHYrnhL2eG4g4fxNWosp4myhVOerl2Y04X5ZxvKpgMfTg8Tl2Kar0eaEq9Cv8Axq/GD/gm7+2P8JviJ4k8CQ/Af4o/Eey0W8ZNM8b/AA1+H/i3xh4Q8S6ZL+9sNU03VdF0m9hhe4t2RrvSrx4tT0u58yzvYEkjDP8A5g8TeBnidw5neOyiPCHEGeUsLVaw+bZFk2Y5lluPw8veo4ihiMLh6sYOcGnUw9WUa+Hqc1KrBSjd/wDQjwD9L/6PvHXCeUcT1PE7gvhHE5jh1LG8N8YcU5HkOf5PjYPkxWCxmCzDHYepUjSrKSoY7DxqYPG0eTEYarKE7R8z/wCGJf2yf+jUf2jf/DLfEb/5na8H/iFPif8A9G743/8AEWzv/wCYj7H/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Q/4Yl/bJ/6NR/aN/8ADLfEb/5naP8AiFPif/0bvjf/AMRbO/8A5iD/AImQ+j3/ANHx8JP/ABYnCf8A89Rr/sT/ALZCKXb9lH9o8gYzs+CXxJkbkgcJH4bZzyecKcDJPAJpPwp8T0rvw744+XCueN/csC2yo/SO+j5J2Xjl4RJv+bxG4Rgu+sp5vGK+b19SD/hi/wDbE/6NP/aW/wDDE/FH/wCZap/4hZ4nf9G646/8RLP/AP5gNP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+g/4Yv/AGxP+jT/ANpb/wAMT8Uf/mWo/wCIWeJ3/RuuOv8AxEs//wDmAP8AiYn6P/8A0fPwe/8AFmcF/wDz6D/hi/8AbE/6NP8A2lv/AAxPxR/+Zaj/AIhZ4nf9G646/wDESz//AOYA/wCJifo//wDR8/B7/wAWZwX/APPoP+GL/wBsT/o0/wDaW/8ADE/FH/5lqP8AiFnid/0brjr/AMRLP/8A5gD/AImJ+j//ANHz8Hv/ABZnBf8A8+j+5f8AZJ0XWfDf7Kn7Mvh3xFpOp6D4g0H9nz4MaLruha1YXWl6zous6X8OPDdjqek6tpl9FBe6dqenXsE9nf2F5BDdWd1DLb3EUc0boP8AWnw4wuKwPh5wHgsdhq+DxuD4M4XwuLwmKo1MPisLisPkeBo4jDYmhVjCrRr0KsJ0q1GrGNSnUjKE4qSaP+ajx3zHL848cPGXNspx2DzTKs08VvEPMcszPLsVRxuX5jl+O4uzjE4LHYHGYadXD4vB4vDVaeIw2Kw9SpRxFGpCrSnOE4yf+JVX2Z+UhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+gX7AP7f8A8R/2HfiIbyx+2eLPg74rvbUfEv4ZvdBIdQjQJbr4o8LtO32fSPGulWwC212fLtNbtIl0bWybcWN7pX4B9Ij6O/BX0i+C6nDfEtNZfnmXrEYnhHi/DYeFXNOGszrQipSjFypPHZRjnSo085yarWp0Mwo0qVWlVweZ4PLsywX6V4YeJ+f+F+fxzXKpvFZfiXTpZ3klWrKGEzbCQk2k2lNYfHYdTnPA4+FOdTDVJThOFfCV8VhcR/cB8HPjH8OPj78OfDXxW+FHiaz8V+CfFVp9p07UrUlJ7a4jwl9pGr2Ln7TpOu6TcFrTVtJvUju7K5Qq6tG0Usn/ADb+LHhPxr4Lca5nwJx3ljy/N8varYbE0XOrled5XVnUhg87yTGzp0ljsqxypVPZVfZ0q+Hr08RgMww+DzLB4zB4f/VTgzjPIOPcgwnEfDuLWJwWJThVpT5YYzL8ZCMZV8vzCgpT+r4zDuceeHNOnVpzpYnDVa+Er0K9X06vzY+qCgAoAKACgCte2VnqVnd6dqNpbX9hf209nfWV5BHc2l5aXMbQ3NrdW8yvDPb3ELvFNDKjxyxuyOrKxB6cHjMXl2LwuYYDFYjA47A4iji8FjcJWqYbF4TF4apGth8ThsRRlCtQxFCtCFWjWpTjUp1IxnCUZJM5cdgcFmeCxmW5lhMLmGXZhhcRgsfgMbQpYrB43B4qlOhisJi8NXjUo4jDYmjUnRr0K0J0qtKcqdSMoyaf4GftqfsVXnwdvL34l/DSyub/AOFl9cmXVNLj8y5u/AN3cy4WKZjvmn8MTzOI9P1CQvJp8jJp2oyMzWl5e/7s/Q1+mXhPFzCYPw38SMbh8D4n4LDqnlmaVHTw2E48wuHp3lVpRXJSocTUKUHUx+ApqNPMKcZ5hl0IpYvB4P8A51vp2fQSxvgnjcf4qeFeBxWYeEeYYp1c3yikquKxvhzjMVVtGlWk3Ur4nhPEVpqnl2ZVXOrllSdPLM0qSlLBY7H/AJx1/ocf5hBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9Qfsjftd/Gb9iv4w6L8Y/gxrxsNTtNlj4l8NX7TzeFvHnhh545b/wALeLNNiliF7pt35Ykt7iN4tR0i/S31XSLuz1G1huF0p1JUpKUXr1XRrqn/AF5rU+v4I43z7gDPaGe5DifZ1YWp4vCVXKWDzLCOSlUweNoxlH2lKdrwmnGrQqKNahUp1YRkv9Cz9hj9ur4M/t6fB+0+Jnwuvhp2v6YLTT/iR8NdTu4JfFPw68Szws50/UkjWI6homoGG4n8M+J7a3j0/XrKGbEdjq1jrGj6X69KrGrHmi9ftRe6f+XZ9fW5/p34d+IuQ+JGRwzfJ6nssTS5KWa5TWnF4zK8XKN3SqpW9rh6rUpYTGQiqWJpxlpTr069Cj9qVqffBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf4U9ABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAffX7A/7fXxJ/Yd+I39paX9r8VfCTxTeWqfEv4ZS3ZjtdWt49sK+IvDrzFrfSPGmlW+RYagFFvqdun9kawJbNoJrH8D+kN9Hngn6RXBVThniemsBnWAWIxPCXF2Fw8Kua8M5pVhFSqU1KdJ43Kcc6VGlnWS1a1KhmNClSqU6uDzPB5bmWB/SfDHxP4g8L8/jm2UT+s4DEunSzvJK1SUMHm2EhJtRk0p/V8bh+epPAY+EJ1MLUnOM4V8JXxeExP8AcJ8GPjR8Nv2gvhv4b+LHwn8S2vinwV4otfPsb6D91d2V3GFF/omt2DMbjSNf0idja6rpV0BPazgMDLby29xN/wA23i14S8beCnGuZcCcd5Y8DmuBftsJi6LnVyrPMrq1KkMHneSYydOmsblmNVKfs6jhSxGGr08RgMww+EzLCYzB0P8AVTgvjTIOPsgwnEXDuLWJwWJXJWo1OWGMy/GQjGVfL8woRnN4fGUOePNHmnTq05U8ThqtfCV6Fer6lX5ofVhQAUAFABQBWvbKz1GzutP1C1tr6wvrea0vbK8gjubS8tLmNobi2ureZXhnt54neKaGVHjljZkdWViD0YTGYvL8XhsfgMViMFjsFiKOLweMwlaphsVhMVh6ka2HxOGxFGUK1DEUKsIVaNalONSnUjGcJRkkzlx2BwWZ4LF5dmWEw2YZfmGGr4LH4HG0KWKweNweKpToYnCYvDV4zo4jDYijOdKvQrQnSq0pyhUjKMmn+B/7av7FN58ILy++Jnwysbm++Ft9cGbVtJi8y5u/AN1cSYEchO+a48LzyuEsL+QvJpsjJp+ou2bS8vP91/oafTLwni1hcF4beJWNw+C8TsFh1SyrNavs8PhePMLh6bbnBLkpUOJqFKDqY7AwUKeZU4zzDL4JrF4TCf8AOx9Oz6CON8FsZmHir4U4DE5h4SY7EutnOT0fa4rG+HWLxVWyhUk3Ur4nhKvWmqeX5jVc6uV1J08tzOpLmweNxn5wV/oif5gBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9H/srftV/GT9jn4waB8Z/gp4ifRvEOlMLTV9Iu/OuPDXjTw5NNFLqXhLxfpUc0C6roWpCGMsglhvdOvIrTV9HvNP1iwsb+2unUlTlzRdn17Ndn/W+p9Vwbxnn3AmeYbPsgxToYqi+SvQnzSwmYYWUoyq4LHUVKPtsPV5VdXjUpVFCvQqUq9OnVj/oU/sE/t9fBz9vv4RQfED4dXSaJ4z0OOysvih8LdRvYZ/EngDX7iJyqSFUgbV/DGqvBczeGPFNtbQ2mr20M0FxBp2t2GsaNpvsUq0a0bx3XxR6r/gefU/068N/EnIvErJI5nlc1h8fh1Tp5vk9WpGWLyzEyT0lpH2+ErOMpYTGQgoV4KUZRpYilXoUfuitT9ECgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/CnoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA+8/2C/wBvX4l/sPfEgaxopuvFHwr8T3VpF8TPhlPeGKx12zi/dLrmhvLvh0fxnpMDOdL1VU8q8iB0rVkuNOmxB+DfSE+j3wT9IngqrwvxTSWBzfArEYrhPi3C0IVc14YzWrCMXWoqU6X13K8a6VGlnOS1a1PD5lh6dKcKuDzLCZbmeA/R/DLxN4g8MM/hm+UTeIwOIdOlnWS1qsoYPN8HCTfJNpT+r4yhz1J4DHwpzq4SrOcZQr4Svi8Jif7hvgl8bfhp+0R8NfDnxZ+EviS28T+DfE1v5ltdRYivtNv4gn9oaDr2nlmn0jxBpEzi31PTLn95C5SaF7iyuLW7uP8Am18XPCLjbwS42zLgXjvLXgs0wT9vg8ZQdSrlWfZVVqVIYPO8kxk6dL65luM9nNQm4UsThcRTxGAzDD4TMcJi8JR/1U4K42yDj/IMLxFw7i/rGDxH7uvQqcsMZl2MhGMq+X5hQjKfsMXQ54uUeadKtSnSxOGq18LXo16nq9fmR9aFABQAUAFAFa8s7TUbS6sL+1t76xvrea0vbO7hjubW7tbmNobi2ubeZXint54neKaGVGjljZkdWViD0YTF4rAYrDY7A4nEYLG4PEUcXg8Zha1TD4rC4rD1I1sPicNiKUoVaFehVhCrRrUpxqU6kYzhJSSZzY3BYPMsHi8uzHCYbH4DH4avg8dgcbQpYrB4zB4qlKhicLisNXjOjiMNiKM50q9CtCdKrSnKnUjKMmn+CP7av7FF38Ibq++J3wxsri++F17cGbWNIiElzd+Abq4kwFYkvNceF55XCWV85aTTJGSw1B2U2l3df7qfQ0+mbhfFjDYHw08S8bh8F4m4OgqOUZvVdPD4XjzDYem22klClh+J6NKDnjMFBRp5nTjPHYCEZrFYTDf87f07voI4zwYxmYeK/hRl+Jx/hLjsS6+d5JRVXFYzw6xeJq7Sb562I4Rr1pqngcfUlOrlNScMuzKcoSweNxX5vV/oof5eBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB7/wDszftN/GD9kf4u+HfjT8E/E03h7xZoMnk3drL5txoHivQZ5YZNU8J+LdJWaGPWfDmrrBGt3aPJFcW1xFa6ppd3p+s6fp2o2lwqSpyUouz/AAa7PyPpuEuLs84JzzC5/wAP4t4XG4Z8s4SvPDY3DSlF1sFjaKlFV8LXUUpwbjOElCtRnSxFKlVh/oSf8E+f+ChXwf8A+CgPwmj8ZeCZovDfxG8OQWVr8VPhRf30VxrvgrWZ0ZVu7R9sMmt+DtYlink8O+JYLeOK6jSSx1CDTtbstR0y19ejWjWjdaSXxR6p/qn0f63P9OPDLxOyPxMyVY/L5RwuaYWNOGc5LUqKWJy+vJO04O0XiMDXlGTwuLjFKok6dWNLEU6tGH37Wx+lBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+FPQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfdf7CH7eHxM/Yf+JQ17w+1x4m+GPiW5tIPiZ8Mri8aHTvEVhCfLTV9IkkEkWj+MdJheRtH1lIisq79M1SO60u5lhH4T9IL6PvBP0h+Ca3C3FVH6nmmD9viuFOK8LQp1c24YzWrCMXiMPzSp/XMtxnsqNLOcmq1qeHzPD06UlVwmY4TLcywP6L4aeJnEHhjn8M4yep7fB1/Z0c5yatUlHBZvg4Sb9nVsp+wxdHmnPA4+EJVcJVnNONbC1sXhcT/cV8Dfjl8M/wBo34Z+Hfi18JPEdv4l8HeJIC0My7YdS0jUoVjOo+HvEWneZJNo/iHSJZFh1LTZ2YpuhurWW7067sr25/5tfF/wg438EONsx4G46y14PMsJ+/wOOoe0q5Tn+VVKk4YTOskxk6dP63l+L9nNJuFPE4TE06+Ax+HwuYYXFYWj/qrwRxvw/wCIGQYXiLh3Fe3wtf8Ad4jD1OWGNy3GxjGVfAZhQUpexxVHmi3aU6ValKnicNVrYatRrVPW6/Lz64KACgAoAKAK93aWuoWtzYX9tb3tje281peWd3DHc2t3a3MbQ3Ftc28yvFPbzxO8U0MqPHLG7I6srEHfC4rFYHFYbG4LE18HjcHXo4rCYvC1qmHxWFxWHqRrUMThsRSlCrQr0KsIVaNalONSnUjGcJRlFM5sZg8JmOExWX5hhcNjsBjsNXweNwWMoU8ThMZhMTSlRxOFxWGrRnRxGHxFGc6VejVhOnVpzlCpGUZNP8E/21v2KLr4R3V98T/hhY3F78MLy4abWdGhElzdeAbm4k9SXmn8LTyuEtLxy8mlSMljfu0bWt1P/ul9DP6ZuG8VsNgfDPxMx1DB+JeEoKjk2c1nTw+G47w1CD0fwUqHFFGlDnxeEgo081hGeNwMI1I4nC0f+d76d30D8X4NYrMPFnwmy/E4/wAJ8biJV89yKiqmJxnh1i8TUSvdudbE8I4itNQweNm51cnqShgMwqSpSwmMr/m3X+i5/lyFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHuH7Ov7Rfxb/ZW+LPhn40fBXxTc+FvGvhmc7XG+fSNe0md4zqfhjxRpXmRwa34b1mKJYdR024ZcskF7ZTWeqWdhf2twnKnJSi7Nfc11T7p/8Hc+h4W4pzvg3OsJn+QYyeDzDCS31lQxNCTTrYTGUbqOIwldJRq0pdVGpTlTrU6dSH9Sui/8HPPw6/sfSv8AhIv2UvGo1/8As6y/tv8AsX4jaE2j/wBrfZ4/7ROlG+0Bb3+zjd+abIXY+0rbmMTlpQzHuWOjbWm79bNWuf2Hh/pdZV7Cj9a4MzD617Kn9Y+r5phnQ9vyr2vsfaYdVPZOfN7Pn99RspNu7NP/AIiefhD/ANGp/Ej/AMOJ4Y/+UVP69H+SX3o2/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlAf8RPPwh/6NT+JH/hxPDH/wAoqPr0f5Jfeg/4m6yL/ojc2/8ADng//lAf8RPPwh/6NT+JH/hxPDH/AMoqPr0f5Jfeg/4m6yL/AKI3Nv8Aw54P/wCUB/xE8/CH/o1P4kf+HE8Mf/KKj69H+SX3oP8AibrIv+iNzb/w54P/AOUB/wARPPwh/wCjU/iR/wCHE8Mf/KKj69H+SX3oP+Jusi/6I3Nv/Dng/wD5QH/ETz8If+jU/iR/4cTwx/8AKKj69H+SX3oP+Jusi/6I3Nv/AA54P/5QH/ETz8If+jU/iR/4cTwx/wDKKj69H+SX3oP+Jusi/wCiNzb/AMOeD/8AlAf8RPPwh/6NT+JH/hxPDH/yio+vR/kl96D/AIm6yL/ojc2/8OeD/wDlAf8AETz8If8Ao1P4kf8AhxPDH/yio+vR/kl96D/ibrIv+iNzb/w54P8A+UB/xE8/CH/o1P4kf+HE8Mf/ACio+vR/kl96D/ibrIv+iNzb/wAOeD/+UB/xE8/CH/o1P4kf+HE8Mf8Ayio+vR/kl96D/ibrIv8Aojc2/wDDng//AJQH/ETz8If+jU/iR/4cTwx/8oqPr0f5Jfeg/wCJusi/6I3Nv/Dng/8A5QH/ABE8/CH/AKNT+JH/AIcTwx/8oqPr0f5Jfeg/4m6yL/ojc2/8OeD/APlB/Rv8D/ifZ/G74K/CD4z6fpVzoNh8Xfhd8P8A4n2Oh3tzFeXmjWfj7wnpPiu20q7u4I4obq506HVks57mGKOKeWF5Y40Rgo7Iy54xltzRUvvV/wBT+puHs3p8QZBkefUqM8NSzvJ8szenh6klOpQp5lgqONhRnOKUZzpRrqEpRSUpRbSSZ/h8VR7AUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH3J+wr+3X8Tv2IPiYviTw08/iP4ceI57O3+Jnwyurx4dL8UaZCxRNR06RhLHo/i/SIpJn0PXY4XKlpNO1KK90i7u7OX8M8f/ADgj6Q3BNbhTiyh9UzHCe3xXC3FWFo06mbcL5tUhGP1rC80qf1rL8V7OlSzjJ6tWnhszw1On+8wuPw2XZjgf0Pw18SuIPDHP6ec5NU9tha3s6WcZPWqSjgs3wUZN+yrWUvY4qjzzngcdCE6uEqzl7tbDVsVhcT/cZ8Cfjv8Mf2kfhl4d+Lfwk8RweI/CPiKEgN8kOq6JqsCRnUvDfiTTRJLLo/iLSJJUi1DT5ncbXgvrKe90u9sL+6/wCbTxh8HuN/A7jbMeBuOcu+q4/C3xGXZjh/aVMp4gympUnDCZ1kuLnCn9awGK5JRalGnicHiadfAY+hhsdhsRh6X+qvA/HPD/iDkGG4h4exXtsNW/d4rDVeWONy3GxjGVbAY+jGU/Y4mjzRekpUq9KVPE4apWw9alVn7BX5afYBQAUAFABQBXu7S1v7W5sb62gvLK8gmtby0uoY7i1urW4jaK4trmCZXingnid45oZUaOSNmR1ZWIO+FxWJwWJw+NwWIr4TGYSvSxWExeFrVMPicLicPUjVoYjD16UoVaNejVhGrSrU5xqU6kYzhJSSZz4zCYTMMJisBj8Lh8dgcdh6+ExuCxlGnicJi8JiacqOJwuKw9aM6OIw+IoznSr0asJ06tOcoVIyjJp/gt+2v+xPdfCa5v8A4o/C6wuL34Y3k7T63okAkuLrwFc3En3l+/LP4VmlcJa3blpNJkZbK9cwtbXUv+6H0M/pm4bxVw2B8MvE3HUMJ4lYSgqGS51WdPD4bjrD0Kfwy+GlQ4oo0oueJwsVGnm0IyxmCjGssThaf/PB9O/6B+L8HMVmPi14S5fiMd4UY3ESxGf5BQVTE4vw7xWIqfFH462I4Qr1pqGExc3OrktScMDj5yoPC4ur+bFf6Mn+WwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av8AyYp+xX/2aX+zl/6p7wbXuUv4VL/r3D/0lH+ufht/ybrgH/siuFf/AFRYA/xW60PtQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD7e/Ya/bn+KH7EPxNXxT4UeXxF8PvEM1na/Ev4ZXd5JBpHi3SoHZUvbOQrNHo/i3SI5Z5NA8QxQSPbvJNYX8N/o19qGn3P4d4++AXBH0hOCa/CfFtD6tj8L7bFcL8UYWjTqZvwxm1SnGP1vCOcofWcDifZ0qebZRVqww2Z4enBOeHxuGwGPwX6F4b+JPEHhln9POslq+1w9b2dHN8orVJRwWb4KMnJ0aySl7LEUuac8FjYQlWwlWUrKrh6uJw1f+474CfHv4X/tLfDDw98W/hH4ih8Q+E/EERRgwSDWNA1iBIm1Lwz4m0xZZpNH8RaQ80cd9YSySRvHLbajp9zf6Rf6dqN3/wA2vjJ4N8b+BvG2YcD8cZe8NjcPfEZZmeHVSeU8Q5TOpOGFznJsVOEPrGDxHJKM4SjDE4PEwrYLHUcPjMPWow/1V4F464f8Qsgw3EPD2K9rh6v7vF4SryxxuWY2MYyrYHH0Yyl7KvS5k1JOVKvSlDEYepVoVKdSXslflR9iFABQAUAFAEF1a219bXFle28F5Z3cE1rd2l1FHcW11bXEbRT29xBKrxTwTxO8c0UqtHJGzI6srEHfDYnE4LE4fGYPEV8Ji8JXpYnC4rDValDE4bE0KkatDEYevSlGrRr0asY1KVWnKNSnUjGcJKSTOfF4TC4/C4nA47DYfG4LG4ethMZg8XRp4jC4vC4mnKjiMNicPWjOlXw9elOdKtRqwnTq05yhOMoyaf4Mftr/ALE1z8KLnUPil8LbCe8+Gd3O1xruhQLJcXXgK4uJOZIx88s/hSaV9tvcsWk0d2W0vGa2Nvct/ub9DL6Z2H8UsPgPDHxPx9HCeJOFoqhkeeV3ToYbjrD0IaU6j92lR4ppU4uVfDxUaebwjLF4SKxKxGHX/PH9O/6B2K8H8VmPi34R5diMb4VYyvLE8Q8PYeNTE4rw7xOIqa1aa9+tiOEK9WfLh8TNzqZJUlHB42csLLDYp/mrX+jh/lmFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/Jin7Ff/Zpf7OX/AKp7wbXuUv4VL/r3D/0lH+ufht/ybrgH/siuFf8A1RYA/wAVutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAPtf9h/8Abi+KP7EXxPj8XeEJJNf8Ca/LZ2nxL+Gl7eSQaN4x0eCRtlxA+ydNH8V6Qk1xL4d8Rw28s1lLLNZXsGoaJf6ppd7+I+PXgLwR9IPgnEcI8X4f2GMw/tsVwzxNhaVOeb8MZtOmorGYKUnD2+Er8lOnmuVVakcNmeGhGM5UMXQwWNwn6B4ceI/EHhnn9LO8kq+0oVPZ0s2ymtUlHBZvgoybdCulzezr0+acsHjIwlVwlWTklUo1MRh6/wDcj+z/APtAfC79pv4X+H/i58IvEMeveFddQxTRSiO31rw7rVvHC+p+GPFGmLLM+keItIeeJb2yeWaGaGa11PTLrUNG1DTdSvP+bXxm8GeN/AvjbH8EccYD2GKo82IyrNcOqk8o4iymVScMNm+T4mcIe2wtblcatKahisFiY1cHjaNHE0alNf6q8Ccd8P8AiJkGG4g4fxPtKNS1PGYOq4xxuV41RUquBx1KMpezrU73hOLlRxFJwr4epUo1IzftFfk59mFABQAUAFAEFzbW17bXFneW8F3aXcEttdWtzFHPbXNtPG0U9vcQSq8U0E0TtHLFIrRyRsyOrKxB2w2JxGDxFDF4SvWwuLwtalicLisNVnQxGGxFCcatGvQrUpRqUa1GpGNSlVpyjOnOMZwkpJMwxWFwuOwuJwWNw1DGYPGUK2FxeExVGniMLisLiKcqWIw2JoVozpV6FelOdKtRqwlTq05ShOMoyaf4N/tsfsTXPwquNQ+Kfwr0+e7+Gl3O1xr+g26yXFz4DuJ3y00Q+eWbwpNI22Gdi0mjSMtrdu1q1vcD/cz6GX0zsP4o4fAeGHihj6OF8R8LRjh8iz7ESp0MPxzh6MNKNZ+7To8U0qcearRiowzmEZYnDRji1XoS/wCeX6eH0DsV4Q4nMvF3wiy6vjPCzGV5YniLh3DRqYjFeHmJxE7yr0UuerX4PrVZctGvJyqZHUnDCYuUsG8PiV+aVf6PH+WAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+nZ+wL/AMmKfsV/9ml/s5f+qe8G17lL+FS/69w/9JR/rn4bf8m64B/7IrhX/wBUWAP8VutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD7Q/Yj/bd+KX7EnxQi8Z+C5X13wVrslnZ/Er4aX15Jb6H410W3kfY6uEnXSPFGkrPcTeG/EsFvNcabcSz2t3BqOh6hq+j6j+KePHgPwR9IHgnE8IcYYb2WJo+1xPDfEmFpU5Zvwxm0oKMcdgZzcfa4ety06eZ5ZVqRw2ZYaKhUdLE0sJi8L994deI3EHhpn9LO8jrc9KfJSzXKq05LBZvglLmlh8QlfkqwvKWExcIutharcoqdKdehW/uT/Z7/aF+Fv7T/wALtA+Lnwi8QJrnhjW0MF1bTCO313wzrlvHC+p+FvFWlpNO2keIdJaeIXVo0s1vc201pqulXeo6JqOm6nef82vjT4Lcb+BPG2O4J42wPsq9LmxGUZvho1JZRxHlMqkoYfNspxE4x9rQqW5K9CfLicDiVUwmLp069OUX/qrwFx7w/wCIvD+H4gyDEc9OdqWOwVVxWNyvGqKlVwWNpRb5KkL81OpG9LEUnCvQnOnNM9tr8kPtQoAKACgAoAhuba3vLee0u4Ibq0uoZba6tbmJJ7e5t50aKaCeGVWjmhmjZo5YpFZJEZkdSpIO2HxGIwmIoYvCV62FxWFrUsRhsTh6s6OIw+IozjUo16FanKNSlWpVIxqUqtOUZ05xjOMlJJmGKw2GxuGxGDxuHoYvB4uhVw2LwuKpU8RhsVhsRTlSr4fEUKsZ0q1CtSnOnVpVIyp1KcpQnGUZNP8AB39tj9iW4+FlxqPxU+FWnzXXw2upmufEHh+3V57nwJcTvlri3UbpJvCksjYjlJaTRHZbe5ZrIwzx/wC5P0Mvpn4fxPoZf4X+KOYUcL4jYajHD5Bn+IlCjh+OaFGFo4fESfLTo8U0qcb1Ka5aedQjLEYdRxqrUKn/AD0fTx+gbifCTE5n4veEGW18Z4XYutPFcScN4WNTEYnw9xFed54nDRXPVr8H1qkrU6rcqmQzksNipSwDoYil+Z9f6Qn+VoUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/ACYp+xX/ANml/s5f+qe8G17lL+FS/wCvcP8A0lH+ufht/wAm64B/7IrhX/1RYA/xW60PtQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD7K/Yo/bY+Kf7E3xSh8b+B5n1rwhrT2dl8SfhtfXksGg+ONCt5XKqzBJ10nxNpKz3M/hnxPBbTXWk3U09vcwaloWo61oup/i/jr4FcEeP8AwTiuD+McLyVqftcVw7xFhqdN5vwzm0qfJDH5fUnbno1OWFPMcuqTWFzLDRVKtyVqeGxOH+88O/ETiDw14gpZ5kVbmhLkpZpllaclgs2wSlzSw2KjG/LON5SwuKinWwtV88OaEqtKr/cp+zt+0V8LP2pPhboPxc+EWvprXhvWFNve2VwIrfXvC2vQRQyan4W8VaYk07aVr+lGeL7RbmWa1u7Wa01bSLzUtE1HTdSu/wDm18a/BPjfwH42xvBfGuC5KkOfE5NnWGjUeUcR5S6koUM1yqvOK56c7KGJw07YrAYnnw2KhCpFOX+qvAHH/D/iNw/Qz/IMRzRdqWPwFVxWNyvGqKlUweMpxb5ZK/NSqxvSxNJxrUZSi3b3KvyE+3CgAoAKACgCG4t7e7t57S7ghurW6hlt7m2uIkmt7i3mRo5oJ4ZFaOaGaNmjlikVkkRmV1Kkg7YfEV8LXo4rC162GxOGrU8Rh8Th6k6NfD16M1Uo16NanKNSlWpVIxqU6kJRnCcVKMlJJmOJw2HxuHxGDxmHo4vCYujVw2KwuJpQr4fE4evCVKvh8RQqxnSrUa1KcqdWlUjKFSEpQnFxbT/CD9tn9iW4+F0+o/Fb4U6dNdfDe5le68ReHbZZJ7jwJPM+XurZBukm8JyyN8rndJobsIJydPMU0H+4/wBDH6Z9DxMoZf4XeKWYUsN4iYelHDcP8Q4mUKNDjihRhaGFxU2406XFVOEbyj7sM7jF1qCWPVWjX/56vp4/QNxHhNiMz8X/AAfyyvi/C/FVp4vibhnCwnXxPh9iK9S88XhILmq1uDqtSfuzfNU4fnJUMRKWWujXw/5mV/pGf5VBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+nZ+wL/yYp+xX/2aX+zl/wCqe8G17lL+FS/69w/9JR/rn4bf8m64B/7IrhX/ANUWAP8AFbrQ+1CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD7E/Yt/bU+Kv7E/xTh8d+A521jwtrDWdj8R/hxf3ksHh/wAd6DbyuywzMqTjSvEelie5n8M+J7e2nu9Gu57iGaDUtD1HW9E1X8a8cvA3gjx94JxXBvGeEtOPtMVkGf4WnT/tfhrNnT5KeY5bVn8UJWjTx+BqS+q5jhk6NdKcaFeh934e+IfEHhtn9HPcir3i+WlmWW1pS+pZtglLmnhcVBbSV5Sw+JivbYWq/aU21KpTqf3Lfs4/tHfCr9qj4WaH8XPhFry6v4e1YG21HTroRW/iHwn4ggihk1Lwr4r0yOac6Xr2mGeIzQiaezvrSaz1fR73UtE1HTtSu/8Am18b/BDjfwF42xnBnGmDs/fxOR55hoVP7I4kyn2jhRzPK601qnpDGYOo/rWX4nmw+JimoTqf6q+H/iBw/wCI/D9DPshr3+GlmGX1ZR+u5VjeXmqYTF04vfeVCvFeyxVK1Wk2uaMfdq/Hj7gKACgAoAKAIri3gu4J7W6hiuba5ikt7m3uI0mguIJkaOaGeGQNHLFLGzJJHIrI6MysCpIOtCvXwtejicNWq4fE4erTr4fEUKk6NehXpTVSlWo1abjUpVaVSMZ06kJRnCcVKLUkmY4jD4fGYevhMXQo4rC4qjVw+Jw2IpQr4fEYevCVOtQr0asZU61GtTlKnVpVIyhUhKUZxcW0/wAIf22v2JZ/hhPqXxX+FGnS3Pw4uJXuvEnhu1V5rjwLPM+XvLRBukl8JyyN1+Z9BdhFMTppjltf9xfoYfTQoeJVHLvCzxTzGlh/EOhShheHeIsTOFGhxvRpRtDB4ub5adLiqnCP92Gexi6lJLMlUpYr/nt+nl9AzEeFFfNPGHweyytivDHE1p4vijhfCQnXxHh9XrTcqmNwUFzVK3B1WpLf3qnD05KlWcsrdKthPzKr/SY/ynCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBOz9gX/kxT9iv/ALNL/Zy/9U94Nr3KX8Kl/wBe4f8ApKP9c/Db/k3XAP8A2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfX37GP7Z/xV/Yp+Klv4/8Ah/cnVfDmqm0sPiL8OtQu5oPD3j3w/BK7i2uiiTjTNf0wT3Nx4Y8T29tNeaHeTzxyQajomo63omrfj3jf4IcEePfBOM4M40wd/jxOR57hoU/7X4azb2bhSzPLK01vtDGYOo/q2YYbmw+Ii/3dSl9z4feIPEHhxxBQz7Ia/wDLSzHLq0pfUs1wXMpVMJi4J+sqFeK9thqtqtJ/HGf9y/7Nn7Sfwq/as+FeifFv4R64NU0LUx9k1XSrvyoPEXhDxDBDDLqXhXxVpsc0503W9NM8TOqyz2WoWc1pq+j3mo6NqGn6hdf82njj4G8b+AXG2L4N4zwl4y9picgz/DU6n9kcS5SqjhSzLLas1pKN408dgakvrWXYluhXTjKjWrf6q+HviFw/4kcP0c9yKvZrlpZlltaUfruVY3l5p4XFwi9U9ZYfERXssVStUpu6qQp+91+Nn3QUAFABQAUARTwQXUE1tcwxXFtcRSQXFvPGk0E8EyGOWGaKQNHLFLGzJJG6sjoxVgQSK1oV62GrUcThq1XD4jD1adehiKFSdKtQrUpqpSrUatNxqU6tOpGM6dSEozhNKUWmkzHEYehi6FfC4qhRxOFxNGrh8ThsRThWoYihWhKnWoV6NRSp1aNWnKVOrTqRlCpCUoyTi2n+En7bX7Ek/wAMptS+LHwn06W5+HVxK934l8NWqPNP4FmlctJe2UY3SS+EpHbkfM+gMwSQnSyktn/uH9DD6aFHxIo5d4V+KmY0sP4g0KcMLw5xJipxpUeN6VKPLTwWNqPlhS4qpwjZNtQz6Kc6ds0U6eM/58Pp5fQLxHhVXzTxi8HcrrYnwzxFWpjOKuFcHCdbEeH9etNyq5hgKceapW4Nq1JXkvenw5OXs6rllDp1cD+Y9f6UH+UgUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av8AyYp+xX/2aX+zl/6p7wbXuUv4VL/r3D/0lH+ufht/ybrgH/siuFf/AFRYA/xW60PtQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD64/Y2/bK+K37FnxUtviF8PLo6loWpG0sPiF8PNQu5oPDvj7w9BK7/Yr7Yk/9na3p3n3Nx4Z8TW9tNfaFezTBodQ0bUNb0TVvyHxs8EuCPHngnG8F8a4LnhLnxOS51hoU1m/DebKm4UM0yqvOL5akb8mKws74bH4Zzw2KhOEk4/b8Acf8QeHPEFDPsgxHLJctLH4Cq5PBZrguZSqYPGU09Yv4qNaNq2Gq2rUZKSaf9zH7NH7THwp/aw+FWi/Fr4Sa2NR0bUP9D1nRrzyYPEfg3xFDDFLqPhbxVpsc0x0/WNP86NwUkmsdSsZrTV9HvNQ0i/sb64/5tPHTwK438AONsVwdxlheelU9rieHuIcLTqLKOJspjU5IZhl9Sd+SrT5oU8xy+pN4rLcTL2VbnpVMNicR/qr4eeInD/iVkFHPMircs48lLM8srTi8blONceaWGxUVbmhK0pYXFRXscVSXPDlnGrSpfQFfjB94FABQAUAFAEU8ENzDNbXMMVxb3EUkFxBPGksM8MqGOWGaKQMksUqMySRurI6MVYEEitaNathq1LEYerVoYihVhWoV6NSdKtRrUpqdKrSqwcZ06tOcYzhUhJThNKUWmkzKvQoYqhWw2Ko0sThsTSqUMRh69OFahXoVoSp1aNalUUqdWlVpylCpTnGUJwlKMk02j8Jv22/2JJvhpNqXxZ+E2my3Hw7nlku/E/hi0R5ZvA00rlpb+wjXdJJ4SkdiXUbn8Psdrk6SUew/3A+hh9NGj4jUcu8K/FXMqeH8QKFOGE4a4lxc40qPG1KlFRpYDH1JOMKfFcIK0Ju0M/iuaNs25oY//ny+nn9Auv4W1818Y/BvK6uJ8NMRVqY3ivhTBwnWxHAFarNyq5jl1KKlUrcG1JybqQXNU4blLlnfJnGpl35i1/pWf5QhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+nZ+wL/AMmKfsV/9ml/s5f+qe8G17lL+FS/69w/9JR/rn4bf8m64B/7IrhX/wBUWAP8VutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAPrL9jv9sX4r/sX/FW1+Ivw4vPt+j6h9lsPH/w/wBQupovDfj/AMOwzPIdP1ERrN9g1ew864n8N+Jbe3lv9AvppiIr7Sb/AFrRtW/JPGnwV4I8d+CcbwTxvgfa0KnNicnzjDRpxzfhzNlTlDD5tlGJnGXsq9Pm5MRQmpYXH4aVTCYylVoVJRPteAuPuIPDriDD5/w/iOSpG1LHYGq5PBZpgnJSqYLG0otc9OVuanVi1Ww9VRrUJxqRTP7mf2Y/2nfhT+1p8KtH+LPwm1n7dpV7iy13Qr0ww+JPBfiOKGKXUPC/ijTopZvsOq2XmxyRyJJNYapYy2uraTd3ul3tpdzf82njv4Ecb/R+42xPB/GGG9rh63tcVw5xJhaVSOUcT5TGooRx2AnNy9liKPNCnmeWVZyxWW4mSp1HVw9bCYvFf6q+HXiLw/4l8P0s8yOryVYclLNcqrTi8blGNlFylh8TFW56U7SnhMXGKpYuknKKhVhXo0foWvxU++CgAoAKACgCOaGG4hlt7iKOeCeOSGeCZFlhmhlUpLFLG4ZJI5EZkkRwVdSVYEEitKNarh6tLEUKtShXoVIVqNajOVOrRq05KdOrSqQanTqU5pThOElKMkpRaaTMq9CjiaNbDYmjSxGHxFKpQxFCvThVo16NWDp1aNalUUoVaVWEpQqU5xlCcJOMk02j8KP22/2I5vhtNqfxa+EumyXHw8nkkvPFHhezjaWbwNLKxaXUdOiXc8vhKR2LSxqGfw6xw2dHKvp3+330L/po0vESllvhT4rZlTocfUadPCcM8TYucadHjWlTio0svzCrJqFPiqEFalVdoZ/FaWzZOOYf8+v09PoF1vDCtmvjJ4NZVVxHhviKtTG8W8JYKnOrW4BrVZOdbM8spRUp1eDqk5OVaklKfDUpWd8lcZ5Z+YVf6XH+TwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+nZ+wL/AMmKfsV/9ml/s5f+qe8G17lL+FS/69w/9JR/rn4bf8m64B/7IrhX/wBUWAP8VutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD6t/Y//AGwfiv8AsY/FWz+JHw2vftmm3v2Ww8eeA9QuZovDfj/w5DM0jaXqqRrL9j1Oz824n8O+IreCW/0C+mleNLzTb3V9J1T8n8Z/Bjgjx14Jx3BHHGA9vha3NiMqzXDKnDN+Hc2jTnDDZxk+KnCfsMVR53GrSmp4XHYaVXB42jXwtapTf2fAfHnEHh3xBh+IOH8T7OtC1LGYOq5SwWaYJyUquBx1GMo+0o1LJwmnGth6qhiMPUp1qcZr+5n9l79qH4Uftb/CrSPit8J9Y+16fdbbLxD4evmhi8SeCvEccMct94a8TWEUsv2TULXzFkguI3lsNVsZLfVNKurvT7qC4f8A5tPHrwF43+j7xtiOEeL8P7fCYj22K4Z4mwtGpDKOKMphUUVjcFKbn7DF4fnpU81yqrUnissxNSEZyr4TEYHHYz/Vbw48R+H/ABM4fpZ3klX2denyUc2ymtOMsblGNlFt0MQly+0o1OWc8HjIRjSxdKLlFU61PEYeh9FV+JH34UAFABQAUARzQxXEUsFxFHPBPG8M0MyLLFNFKpSSKWNwySRyIzI6OCrqSrAgkVpSq1aFWnXoValGvRqQq0a1KcqdWlVpyU6dWnUg1OFSnNKcJxkpRklKLTVzOtRo4mjVw+IpU6+Hr06lGvQrQjVo1qNWLhVpVaU1KFSnUhKUKkJxlGcZOMk02j8K/wBtv9iKX4cy6n8W/hHpkk/w/mkkvPFPhWzRpZvBEsrF5dS0yJQzyeEndi00Khn8Ose+jENpn+3f0L/ppUvEKnlvhT4r5lTo8eUqdPB8McUYupGnS40p04qNLLcyqy5YU+KYRSjRrSajxAlZ2zdNZl/z7/T0+gVW8Mq2a+MvgzlVXEeHFepUx3F3CGBpzq1uAqtWTnWzXKqMFKdXg2pOTlXoRUp8MybeuRtSyn8v6/0xP8mwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBOz9gX/kxT9iv/ALNL/Zy/9U94Nr3KX8Kl/wBe4f8ApKP9c/Db/k3XAP8A2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH1R+yJ+158V/2NfirZfEr4aX/ANosrn7NYeOPA+oXEyeG/HvhyOYySaTq8UYk+zX1t5k0+g69BE9/oV9I80K3Flc6npuo/lXjJ4N8EeOfBOP4H45y/wCs4LEXxGWZnh1Tp5vw9m0Kc4YXOclxc4VPq2Nw/PKM4TjUwuNw062Bx9DE4LEV6FT7LgXjriDw84gw3EPD2J9liKX7rF4SrzSwWZ4KUoyrYHH0Yyj7WhV5VKMlKNahVjTxGHqUq9KnUj/cz+y1+1N8KP2u/hVpXxU+FOr/AGi0m2WXiXw1fPDH4k8E+I1hSW88O+JLGKSTyLuHf5lpeRNJYavYtDqOm3FxazK4/wCbTx88A+N/o98bV+EuLqH1nA4n2+K4X4owtGpTyjijKadSMfreDc5VPq2Ow3tKVLN8oq1amKyvE1Kac8TgcTl+YY7/AFV8N/Enh/xNyCnnWSVPZYil7Ojm+UVqkZY7J8bKLfsa6Sj7XD1uWc8DjoQjRxlKMmo0sRSxWGw/0fX4efoIUAFABQAUARyxRTxSQTxxzQzRvFNDKiyRSxSKUkjkjcMkkciMVdGBVlJVgQTWlKrVoVadajUqUa1GpCrSq0pyp1aVWnJTp1KdSDU4VITSlCcWpRklJNNXM61GliKVWhiKVOvQr050a1GtCNWlWpVYuFSlVpzUoVKdSEpQnCacZxbjJNNo/C79tz9iKX4eS6p8XPhHpkk/gGaSS98V+FLONpJfBMkjF5tU0uFdzyeE3cl7iBQW8OklgDopzpf+3P0LvppUuP6eW+FHixmdOjx1ShTwfC3FOMqRp0uM6dOPLRyzM602oU+KYRSjh8RNqPEKSi7Zzpmn/P19PT6BNbw1q5t4zeDGVVK/h3WqVMdxhwfgac6tbgSrUk5183yihBSnU4OqTk54nDRUp8MyblFPInfKPy8r/TM/yYCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Oz9gX/AJMU/Yr/AOzS/wBnL/1T3g2vcpfwqX/XuH/pKP8AXPw2/wCTdcA/9kVwr/6osAf4rdaH2oUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfUX7JP7W/xX/Y5+Kth8TfhjqPmQTeRYeM/BeoTzDw3488OLMZJdG1qCMny7iHfLNoutQIb/Q75zcWxkt5r6yvfyzxi8HeCPHHgnMOBuOsu+tYDFf7Rl2Y4f2dPN+H82p06kMJnWSYydOp9Vx+F9pOLUoVcLjMNUr4DMMPisBisThqv2HA3HPEHh7xBhuIeHcV7HE0v3eKwtXmngszwUpRlWwGYUIyh7bDVeVNNSjWoVY08ThqtHE0qVWH9zP7Kn7Vnwo/a++FWmfFL4WarvjfyrLxV4VvpYV8SeB/EYhWS60DxBaRu2yVMmXT9QiDWGs2Jjv7CV4ndY/8Am18f/ADjf6PPG1bhPi2h9by7F+3xXCvFWEoVKeUcU5TSqRj9awvPKp9UzHCe1o0s5yarWqYrK8TUp/vMXl+Ly3Msw/1V8NfErh/xO4fhnWS1PY4qj7OjnGT1qkZY3J8bOLfsa1lH22Frcs54DHwhGljKUZ+7RxVHF4TDfStfhh+hhQAUAFABQAyWKOeOSGaNJoZkeKWKVFkjljkUpJHJG4KujqSrowKspIYEE1dOrUo1KdajUnSrUpxq0qtOcoVKdSElOFSnOLUoThJKUZxalGSTTTVzOrSpYilVoV6dOtQrU50q1GrCNSlVpVIuFSnVpzUoVKdSEnGcJpxlFuMk02j8L/23P2IZfh9Jqnxd+EOlvN4DleS98WeE7KNpJfBUkjF59V0qFQXk8JuxL3NsgL+HGJdAdEJ/sn/bf6F300qfHtPLPCfxZzOFLjilCnguFeKsbUjCnxjTglCjlWa1ptRhxTGKUMNiZtR4hSUJtZ1b+1f+f76ev0CavhxUzbxn8F8pqVvD2rUq4/jHg3A0pVKvA1SpJzr5xk1CCc6nB85tzxWFgpS4Zk3UgnkLf9j/AJdV/psf5KBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/Jin7Ff/Zpf7OX/AKp7wbXuUv4VL/r3D/0lH+ufht/ybrgH/siuFf8A1RYA/wAVutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD6c/ZO/ax+K37HnxV074n/C/Usq3k2Pi/whfzTf8ACN+OvDgmElxoeu20TcMuXm0rVYV+36LfEXdm5Vrm3ufy/wAX/CDgjxv4JzHgXjvLvrmXYz/aMBj8P7Olm/D+bU6dSGDzvI8ZOnV+qZjhPaTinKFXDYvDVcRl+YYfF5di8XhK31/BHHHEHh9xBheIuHcV7DFUf3eJw1TmngszwU5RlXy/MKEZR9vha3LF6SjVoVo0sVhatHFUaNan/cz+yf8AtY/Cj9sL4Vad8T/hdqfI8mx8XeEb+WH/AISTwN4jMIluND122jP+/LpmqQr9g1mx23lk/E8MH/Nr9IH6P3G/0eONq3CvFdH65lmM9viuFOLMLQqUsp4pymlUjF4nD806v1PM8J7SjSznJataricrxNWlJVcZluLyzM8w/wBVfDTxL4f8T8ghnOTVPYYuh7OjnOTVqkZ43KMbOLfsqtlH2+Er8lSeAx8IQpYylCacKGKoYvCYb6br8JP0QKACgAoAKAGSxxzRyQzRpLFKjxyxSqskcscilXjkRwVdHUlXVgVZSQQQTV06lSjUhWpVJ0qtKcalKrTlKFSnUhJShUhOLUoThJKUZRalGSTTTVyKtKnXp1KNanCtRrQnSq0qsI1KdWnUi4VKdSnNOM4Ti3GcJJxlFtSTTZ+Gf7bv7EMngCTVPi98IdLeXwLK8t94t8I2UbSSeDJHYvPq+kQICz+FHYl7q0QFvDjEyRg6GSNH/wBtPoXfTTp8dwyzwm8WszhS43pwp4LhTizG1Iwp8YQglChlObV5tRhxRGKUMLiptR4hSVOo1ndnm3+AH09foEVfDqpm/jR4LZTOt4f1Z1cw4y4MwFKVSrwPUqSdTEZzkuHgpSqcITk3UxmDgnLhmTdSknkDayX8tq/05P8AJEKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Oz9gX/kxT9iv/s0v9nL/wBU94Nr3KX8Kl/17h/6Sj/XPw2/5N1wD/2RXCv/AKosAf4rdaH2oUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9Lfsp/tWfFb9kD4q6Z8UfhbqmyRfKsfFfhW+kmbw5448OGZZbrQdftY2G5Gw0unajEBf6PfbL2xkVxJHN+Y+LvhFwT428E5lwJx3lv13K8bavgsbQdOlm2Q5tShUhg87yPGTp1fqeZ4P2tRRm4VcPisNVxGX5hh8ZluMxeEr/XcE8b8QeH/EGF4i4dxfsMXQ/d4jD1OaeCzLBTlGVfL8xoRnD2+Er8kW0pQq0asKWKwtWhi6FCvT/uZ/ZK/a1+FP7Yvwq0/wCJvwx1EJPGILHxl4Nv54T4k8CeI2h8yfRdat4yC8Mm2SbR9YhQWGt2K/abVklju7S0/wCbb6Qf0e+Nvo78bVeFuKaTx2U4518VwnxbhaFSllPFGVUqkYutQUpVfqWa4L2tGlnWS1a1XE5ZiKtKcauMyzGZZmeYf6qeGfibw/4n5BDOMnn9XxuH9nRzrJa1WM8blGMnFtU6jSh9YwdfkqTwGYQpwpYulCacKGLoYzCYb6hr8GP0YKACgAoAKAGSRxzRyRSoksUqNHLFIqvHJG6lXSRGBV0dSVZWBVlJBBBNXTqVKVSFWlOdOrTnGpTqU5ShUp1ISUoThOLUozjJKUZRalGSTTuiKlOnWp1KNanCrSqwnTq0qkY1KdSnOLjOnUhJOM4Ti3GcZJxlFtNNNn4a/tu/sQyeA5NU+L/wg0t5fA8ry33i/wAIWMbPJ4NkcmSfWNHgQFn8KuxL3lmgLeHGJliB0Isujf7Y/Qt+mpDjmGWeEvi3mcKfGlOFLA8J8W42qoQ4vhBKFDKM4rzajHiiMVGGDxk2lxCkqVV/23yyzf8AwD+nr9Aip4e1M38afBXKZ1eAqs62YcacF5fRlUqcETm3UxGd5Jh6acp8Iyk5VMbgqacuGW5VqKeQcyyX8tK/08P8jgoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Ts/YF/5MU/Yr/7NL/Zy/8AVPeDa9yl/Cpf9e4f+ko/1z8Nv+TdcA/9kVwr/wCqLAH+K3Wh9qFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH0j+yz+1P8AFb9kT4q6V8U/hXqxhuYvLsvE3hm9kmbw5438OtMkt34e8Q2cbr51vLt82yvY9t9pF8Ir/T5op4/n/NPFvwk4J8a+Ccy4E47yxY/Kscvb4TF0XClmuRZrSp1IYPPMjxs6dV4LNMF7WoqdRwq4fE4eriMvzDD4zLcZjMHiPrOCuNc/4Az/AAnEXDuL+r4zD/u69CpzTweY4Ocoyr5fmFCMofWMJiOSLlHmhVpVYUsThqtDF0KFen/eb+zZ8c9I/aT+CPw++Nmh+HPEnhPTvHmiJqkegeKtPuLDU7CZJpbS5WKSaGGLVtJluIJZ9E1+yX7BrukyWmp2wjW5MEX/ADKeOPhNj/BHxO4l8N8wzvKeIqmRV6LoZtlFelOGKwONowxeBlj8FCtXrZNmywtWn/aOT4upOrg8Rd0K+NwFXB5hi/8AWjw940w/iBwllXFWGy/G5XHMac1UwWOpzjKjiKFSVHELDYiVOnTx+CdaE/quOoxjCvTt7Snh8TCvhqHuVfkp9oFABQAUAFADJI0lR4pUSSORGjkjkUOkiOCro6MCro6kqysCGBIIINXCpOlOFWnOdOpTnGpTqQk4ThODUozhOLUozjJKUZRaaaTTuRUp061OdKrCFWlVhKnUp1IqdOpTmnGcJwknGcJxbjKMk1JNpppn4U/t3fsa6f8AC8Xvxl+GiWdj4Ev9Qgj8S+FGnhtv+EY1XVLlYbe68PRyun2jQ9Qu5VjOkQB7nRriQGzik0ZjHo/+3v0HvpgY/wAS3g/B/wASJ4zHccYLAV6nDfFKo1sS+Jcsy3DyrYjC8QVKUJ/V86wGFpyqRzavyYfOMPTaxlWnnEVUzf8A59v2hH0Hst8J1j/HHwrhgsv8PswzLD0+KuD5V6GFXCmbZtiY0MPi+GqVadP6zkGY4yrGlLJcN7TFZHiaqeBo1Mjk6WSfmBX+mB/k2FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/p2fsC/wDJin7Ff/Zpf7OX/qnvBte5S/hUv+vcP/SUf65+G3/JuuAf+yK4V/8AVFgD/FbrQ+1CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD97/8AglV/wSoufjfcaD+0b+0doU9p8GbaaHU/h/8AD/U4ZILn4r3ELiS21zXLeQLLD8OIpFD21u4WTxs65G3wuGfXv84/ppfTVwnhBhMd4Z+GWNw2P8U8bh/ZZrm1P2WKwfAGFxNNNVKsZKdHE8UV6M1Uy/LqsZ0stpzp5lmdOSeDwOO/qXwF8A6/HFbD8WcW4erh+D8PV58Hgp89GvxLWpS1jCScalLKKc4uOJxUHGeKkpYXCSTVfEYf+tq2traytrezs7eC0tLSCK2tbW2ijgtra2gjWKC3t4IlWKGCGJVjiijVY441VEUKAK/wUxmMxmY4zF5hmGLxOPx+PxNfGY7HYyvVxWMxmMxVWdfE4vF4mvOpWxGJxFac61evWnOrWqznUqTlOTb/ANHqFChhaFHDYajSw+Gw9KnQw+HoU4UaFChSgqdKjRpU1GnSpUqcYwp04RjCEIqMUkkieuY1CgAoAKACgDjPiB8QPCXwv8Jav438b6xb6J4e0W3M13dzktJNI3y29jY265mvdRvZitvZWVurz3M7qiL95h9hwHwFxX4mcV5TwXwXlOIzriDOa6o4XC0Vy06VOPvYjG43EStRweX4OlzV8bjcRKFDD0YSnOWyfw/iP4j8G+E3Budce8e51hsh4ayLDOvjMZiHzVK1SXu4bAYDDRvWx+Z4+s44bAYDDRniMViJxp04O7a/m+/ai/ai8W/tI+LftV19o0TwFolxMPCHhATbo7WNt0R1nWTExhvfEN7CSJphvg06Bzp+nkx/abm+/wCiT6Mv0ZOFPo7cK/V8N7DOeO85w9F8W8Wui1PEzTVX+yMnVWKq4PIMHVSdKk+SvmFeCx+PSqfVsNgv+X76Wn0teMvpQ8Y/WsV9ZyHw7yLE11wXwVGvzU8LTkpUnneeOlJ0cfxJjqLarVk54fLMPUll2XN0ni8Xj/luv6cP5ICgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/07P2Bf8AkxT9iv8A7NL/AGcv/VPeDa9yl/Cpf9e4f+ko/wBc/Db/AJN1wD/2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH75/wDBKn/glPcfGyfQf2j/ANpDQprX4OW8sGqfD34eanC8Fx8V54XEltr+v20gWWH4cRyKslnaOFk8cOodtvhQbvEf+b300/prYTwkw2P8MfDDHYfHeKGMoOjnGc0nTxOE4Bw2Ip3u01OjiOKa1Kang8FUU6WVQlDH5hCU3hcHif6o8BPAOtxrWw3F3F2Hq4fhGhV9pgcDPnpV+JKtKXf3Z0snpzi418RFqeMkpYfDSUfbV6X9a9vb29pbwWlpBDa2trDFb21tbxJDb29vCixwwQQxqscUMUarHFFGqpGiqqqFAFf4MYvF4vMMXisfj8ViMbjsbiK2LxuNxdapicXi8Xiakq2IxWKxFaU62IxGIrTnVrVqs51KtScqlSUpSbf+jdChRw1GjhsNRpYfD4elToUKFCnGlRo0aUVClRo0oKMKdKnCMYU6cIqMIpRikkkTVzmoUAFABQAUAcX8QfiD4S+F3hLV/G/jfV4NF8PaLAZrq6mO6WeVvlt7Gxt1/e3uo3su2CysoFaa4mYKoxuZfseAuAeK/E3ivKeC+C8pr5zxBnNdUcNhaK5adGnH3sRjsdiJfusFl+Dpc1fGY2vKNHD0YylOV+WMvhvEjxI4N8JeDc64949zrD5Fw1kWHdfF4uu+arXqy93DZfl+Gj++x+Z4+ty4fA4HDxnXxNecYwjbmlH+b39qH9qHxd+0j4t+13f2jRfAmi3Ew8IeD1m3RWkbbozq+rmM+VfeIL2I4nn+eGwhc2FgRF9ouLz/AKJfozfRm4U+jtwp9VwvsM545zmhRlxbxbKjy1MXUjaospylVF7XBZBg6utGi+Wtj60Fj8eva/V8Pg/+Xz6Wf0s+M/pQcZfXMZ9ZyLw9yLE148FcFRr81LBUpc1J51nTpSdHH8SY+i7YjEe/Qy6hN5dlz9j9ZxOO+Xa/po/ksKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBOz9gX/kxT9iv/ALNL/Zy/9U94Nr3KX8Kl/wBe4f8ApKP9c/Db/k3XAP8A2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfvx/wSp/4JTT/GmbQP2kP2ktBmtvg9bywar8Ovh1qkLwz/FWaJxLa+IfEVrIFki+HEcirLY2Miq/jllEsgHhED/hJv8ANn6an018L4UYbH+F/hdj6GM8TcXQlQzzPKDp4jC8BYevT1hD46VfiqtSnzYbCzU6eUQlHGY2Eq7w+Fqf1X4CeAVbjOrhuL+L8NUocJUaiqZfl9RSpVuI6tOXxS2nTyaE1arWi4zx0k6GHkqaq1o/1pwQQWsENrawxW1tbRRwW9vBGkMEEEKCOKGGKMLHFFFGqpHGiqiIoVQAAK/wcxWKxWOxWJxuNxNfGY3GV62KxeLxVapiMVisViKkq2IxOJxFaU6tevXqznVrVqs5VKtSUpzlKUm3/ozRo0sPSpYfD0qdChQpwo0aNGEaVKjSpRUKdKlTgowp06cIqEIQSjCKUYpJJEtYGgUAFABQAUAcV8Q/iH4R+FnhHV/HHjjV4NF8PaNAZbm5lO6a4mbIttP0+2U+bfaley4gs7KANLPKwAAQO6/ZcAcAcWeJ/FeU8F8FZTXznP8AOK6pYfD0ly0qFKNniMfj8RJeywWXYKlevjMZXcaVClFttycYy+E8SvErgzwj4Mzrj7j7OsPkXDeR4d1sViqz5q2JryusLl2XYWL9tj80x9a2HwOBw6lWxFaSSSgpzj/N5+1B+1B4u/aQ8W/bL0z6L4F0W4mXwh4PWbdFZRNmM6tqxjPlX3iC9i/4+Lj5obGJjYWBEImmuv8Aom+jP9GfhP6O3Cn1TB+wznjfOaFGXFvF0qPLVxlWNqiyrKlUXtcFkODq/wADD3jVxtaKx2OvWdGjhf8Al9+ll9LLjP6UHGX13HfWMi8PsixFePBXBUa/NRwNGXNSec5y6cvY4/iTH0f95xPvUcBRm8vy61D6xXxny/X9MH8mBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/p2fsC/wDJin7Ff/Zpf7OX/qnvBte5S/hUv+vcP/SUf65+G3/JuuAf+yK4V/8AVFgD/FbrQ+1CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD9/f+CVH/AASmm+Mk2gftJftJ6BLb/CKCWDVfhz8ONVgeKf4pzRMJbTxJ4ktZQskXw5jkVZdP06VQ/jlgs0wHhEAeKP8ANT6av02MP4V0Mw8LfCzMKGL8SsVRlh8/z/Dyp4jDcCUK0LSoUX79KvxXVpyvSoyUqeSwlHE4qMsW6OHj/V3gH4BVeMKuG4w4xw1SjwpSmquW5bVUqdXiOpTlpUqLSdPJoSVpzVpY+SdKi1QVSrL+s6CCG1hhtraGK3t7eKOC3t4I0ihghiQRxQwxRhUiiiRVSONFVERQqgAAV/hDicTicbicRjMZiK+LxeLr1cTisViatSvicTia9SVWviMRXqynVrV61WcqlWrUlKpUqSlOcnKTb/0WpUqVClToUKdOjQo04UqNGlCNOlSpU4qFOnTpwShCnTglGEIpRjFKMUkrEtYGgUAFABQAUAcV8Q/iH4Q+FnhLVvG/jnWLfRPD2jw+ZcXM2XmuJ3yLaw0+2TM19qN7JiGzsrdXmmkPACK7r9nwB4f8WeJ/FeVcF8FZRXznP83rezw+GpWjSoUYWeIx+PxM/wBzgsvwdO9bF4yvKNKjTWrc5QhL4TxK8SuC/CPg3OePeP8AO8PkPDeSUPa4nFVm51sRXneOFy7LsLC9fH5njqtqGCwOGjOtXqvRKEak4fzd/tPftP8Ai79pDxcb6+M+jeB9GnnXwh4PSfdDYwtmM6rqpjPlX3iC9ix9pusNFZxN9hsdsAlluf8Aom+jR9GjhP6O3Cn1PBKjnHG2cUKMuLeLZ0eWtja0bVFleVqova4LIcHVv9Xw11VxlWP17Hc1d0qWG/5e/pYfSw40+lBxm8fj3iMj4AyPEV48FcFQr89DL6E703m+bum/ZY/iPH0f96xVpUcFSk8vy/lw6rVcV8wV/S5/JwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+nZ+wL/wAmKfsV/wDZpf7OX/qnvBte5S/hUv8Ar3D/ANJR/rn4bf8AJuuAf+yK4V/9UWAP8VutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/oC/wCCU/8AwSlm+L8vh/8AaT/aV0CWD4TQSW+rfDf4batA0U3xPljYTWnibxPaSgPH8O0dVm03TJlV/HDBbi4UeEAieKP8z/prfTZw/hhQzHwq8Kcxo4rxHxFKeF4i4iw04V8PwLRrQtPC4WS5qdbiupCWkXzU8ii1WrqWYOnRw/8AWPgF4A1eLqmF4x4yws6PC1Kca2V5XVUqdXiKpCV41qy0nDJoyWstJZi04U2sNz1Kv9ZUMMNvDFb28UcEEEaQwQQosUMMMShI4oo0CpHHGihERAFRQFUAACv8JcRiMRi8RXxeLr1sVisVWq4jE4nEVZ1sRiMRWnKpWr161SUqlatWqSlUq1akpTqTlKc5OTbf+ilKlToU6dGjThRo0YQpUqVKEadOlTpxUKdOnTglGEIRSjCEUoxikkkkSViWFABQAUAFAHEfEX4i+EPhT4Q1fxx441eHRvD+jw+ZPPId9xdXDgi207TrYES32p30g8mzs4QZJZCSdkSSSJ9p4feH3FvijxZlXBXBWU1s4z7N63s6NCneNDDUItPE5hmGJadPBZdgqbdbF4ys1TpQVlz1Z06c/gvE3xN4L8IOC854+4/zqhknDmSUPaYjEVXz4nF4maawuWZZhU1Vx+a4+qvYYLBUE6lao23yUoVasP5uv2nf2nfF/wC0h4uN/qBm0bwTo086eEPB6T74NPgfKHU9TKER32v3sWPtV2QY7aM/YrEJbq7z/wDRP9Gr6NXCf0duE1gcAqOccaZxRoz4t4unR5K+PrxtNZblqmnUwWQ4Krf6rheZVMTUX13GueIlCFD/AJevpX/Sv40+lBxm8xzF18j4CyTEV4cFcFU8Q54fLcPO9N5rmrpv2WP4jx9K31vGNOlhKUvqGA5MNGpPEfMdf0qfyiFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/ACYp+xX/ANml/s5f+qe8G17lL+FS/wCvcP8A0lH+ufht/wAm64B/7IrhX/1RYA/xW60PtQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP6B/wDglP8A8EpZvi3L4f8A2lP2lvD8kPwqhkt9X+Gvw01e3aOb4mSxss1n4n8UWcyh4/h6jhZtL0qZQ/jdgl1cqPCPlp4o/wAyfpsfTZoeGlHMvCjwozGliPESvTnhOJeJcLOFahwPSqxcamBwU1zU6vFdSEmpP3oZDGXPUUszcKWD/rTwC8AKnFk8LxnxnhZ0uGKc41spymtGUKnEM4O8cRiIu0oZLGS0WksyatBrCKU8R/WLDDFbxRQQRRwQQRpDDDCixxQxRqEjiijQBI440AREQBVUBVAAAr/CqvXr4qvWxWKrVcTicTVqV8RiK9SdavXr1pyqVa1arUlKpVq1akpTqVJylOc5SlKTk2z/AERp06dGnClShClSpQjTpUqcYwp06cIqMIQhFKMIQilGMYpRjFJJJIkrIsKACgAoAKAOH+I/xH8H/Cjwhq/jnxzq8Oj6Bo8O+aZ/nuby5cEWum6bagiS+1O+kHlWlpDl5Hy7mOGOWWP7Xw88POLfFPi3KuCeCcprZvn2b1uSlRheGHwmHg08TmOY4lp08FluCpv2uLxda0KcLRip1qlKlU+B8TvE7grwf4Lznj/j/OaGScOZJQ9pXr1LTxONxU1JYTK8rwikquYZrj6q9jgsFRvUqzbnJ06NOtWp/wA3P7Tn7TnjD9pDxedR1EzaP4L0eadPCHg+Ocvb6bbuSh1HUSmI77X76MA3l4VKQIRZWQjtkJm/6KPo1/Rr4S+jvwmsvy9Uc34yzejRnxbxdUoKGIzGvG01l2XKd6mCyLBVG1hMIpKeImnjca54maVL/l5+lb9K7jX6UHGjzLM3XyTgTJK9enwVwVTxDnhsrw826cs0zSULUsw4ix9JJ43GuLp4am/qGAVPCwlKv8y1/SZ/KYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/Jin7Ff/Zpf7OX/AKp7wbXuUv4VL/r3D/0lH+ufht/ybrgH/siuFf8A1RYA/wAVutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP6DP8AglP/AMEpJfivJ4f/AGlf2l/D8kPwuie31j4afDLV7do5fiVIjLNZ+KfFVlMoeP4fowSfSdInQN42IS7u0HhHy4/FP+Yv02Ppt0fDelmfhN4S5lSxHiDWp1MHxRxRhJxq0eCKVSLjVy/LqsXKFXiycJONWonKHD0Xd3zi0Mu/rfwB8AJ8VTwnGnGmEnT4ahKNfKMorxlCfEE4vmhisVBpSjkqavCLs8za0/2G7xX9YEUUUEUcEEccMMMaRQwxIscUUUahI4440ARI0QBURQFVQFUACv8AC6tWrYmtVxGIq1a+Ir1alavXrVJ1a1atVm51atWrNynUq1JylOpUnKU5zk5SbbbP9D4QhShCnThGnTpxjCnThFQhCEEoxhCMUoxjGKUYxikkkklYkrIoKACgAoAKAOG+JHxI8H/Cfwfq3jnxzq8Oj6BpEW6WV8PdXt04b7LpmmWu4SX2p30i+VaWkXzO26SRo4I5po/t/Dvw74u8VOLcq4J4Jyqtm+e5tV5adKN4YbB4aDj9ZzLMsU4yp4LLcFCXtcXi6vuwjywgqlepSpVPz/xQ8T+CvB3grOeP+P8AOaOS8O5LR56tWdqmLx2LqKX1TKsqwilGpmGa5hUj7HB4Oj71SfNUqSpYelXr0v5uP2mv2m/GH7SHi86lqZm0fwZpE08fhDwfHOXttMtnJQ6hqBXEd9r99GFN7fFdsSkWdkI7WPEn/RR9Gz6NnCX0d+Ell2Wqlm/GOb0aNTi3i6pQUMTmWIh76wGAU+apgciwVRyWDwalz1pp4zGOpiqn7v8A5efpWfSs41+k/wAavNM1dbJOBslr4inwVwVSxDnhcqws24SzLMpQap5hxFj6Si8dj3HkoQtgcDGlhKf7z5nr+kT+VQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9Oz9gX/AJMU/Yr/AOzS/wBnL/1T3g2vcpfwqX/XuH/pKP8AXPw2/wCTdcA/9kVwr/6osAf4rdaH2oUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/Ql/wAEpv8AglJJ8UZPD37S37THh54fhnE9trHwy+GOsW7JL8RpEZZ7LxX4ssplDJ4CVgk+j6POgbxqQl5eoPCXlReKP8wPps/TcpeHlPNPCTwjzOnX4+qwq4LivivB1I1KXBNOpFwrZXldaDlCpxbKLccRXi3DhxNxTed3WU/114A/R/nxPLCca8bYSVPhuEo4jJsmrxcJ5/KL5qeMxkHaUclTXNSpNKWaO0nbL/8Aff6u4oo4Y44YY0ihiRIoookWOOKONQqRxooCoiKAqIoCqoAAAFf4ZVatXEVatevVqVq9apOrWrVZyqVatWpJzqVatSblOpUqTk5znOTlKTcpNttn+hkIQpwhTpwjTp04xhCEIqMIQilGMIRilGMYxSUYpJJJJKw+sygoAKACgAoA4T4k/Enwf8JfB+reOfHOrRaRoOkRbpJGxJd313IG+y6XpdruWS+1S+kUxWlpEcsd8srxW0U88X3Hh14dcXeKvF2VcE8E5VVzbPc2q8tOnG8MLgsLBx+tZnmeK5ZU8FlmChJVMVi6ukU406cauIq0aNT8+8UfFHgnwb4JznxA8QM5o5Lw7k1HmqVJWqYzH4yopfVMpynCc0amYZtmFSLpYPB0vem+erVnRw1GvXpfzc/tM/tM+MP2j/GB1TVDLpHg7SJZ4/B/g+Kcva6VaudrX18y7Y77Xr6NVN9fsu1FxZ2axWkSq3/RT9G76NvCP0d+ElluWKlm3F+bUqFTi3i6rQUMVmmJgudYHAxlzTwOR4KpKSwWBjPmqSvi8ZKri6kpR/5ePpVfSr42+k/xq82zZ1sl4IyWtiKXBXBNLEOphMowlR8kswzCUOWnmHEOPpxjLMMwlDlpxtgsDGjg6UYz+aK/o8/lcKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/07P2Bf+TFP2K/+zS/2cv/AFT3g2vcpfwqX/XuH/pKP9c/Db/k3XAP/ZFcK/8AqiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf0L/APBKb/glHJ8TX8PftL/tM+Hni+G8b22sfDH4X6xbMknxBdCs9j4s8W2U6hk8CqwSfRtFnQN4yIS9vkHhTyYfE/8Al39Nr6blPw/hmnhF4Q5pTrcd1YVcDxbxdgqsalPgqE04V8oyivBuM+LZRcoYvFwbjw0m6VNvPuZ5L/XvgB9H+XEssHxtxvhJQ4djKGIyXJcRFxnn0otSp47G05JOOSp2lRoys81aU5L+zrfX/wCraOOOGNIokSKKJFjiijVUjjjRQqIiKAqIigKqqAqqAAABX+G1WrUrVKlatUnVrVZzq1atWcqlSrUqSc51Kk5tynOcm5TnJuUpNttttn+hMIRpxjCEYwhCKhCEEoxjGKtGMYqyjGKSSSSSSsh9QUFABQAUAFAHB/Ev4l+DvhJ4O1bxz461aLSdB0mLLMcSXd/eSBvsul6Xa7lkvtTvnUx2trGcnDzTPDbQzzxfc+HPhzxd4rcXZVwTwRlVXNs9zWraEFeGFwOEg4/WszzPFcsoYHLMFCSqYrFVE1G8KVKFbE1qFCr+e+KXilwT4NcE5x4geIGc0cm4eyalec5WqYzMMbUU/qeUZRg+aNTMM2zCpF0sHg6TTk1OtWnRwtDEYil/Nx+0x+0x4x/aP8YNqurNLpHhDSJZ4vB/g+KdpLTSLSQ7Te3rLtS/16+jVG1DUGQBQFtLNYbOKOOv+ir6N/0b+Efo78IrK8rjSzbi3NqVCrxdxdVoKGLzbFQXMsHg1LmngcjwVSU44DARm3J82Lxcq2Mq1Kh/y7fSp+lTxt9J/jaWcZxKrk3BWTVsRS4K4Ko4iVTB5Ng6j5ZY/Hyjy08w4hzCnGEsxzGUEopRwWCjRwVGnTfzVX9Gn8tBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/p2fsC/8AJin7Ff8A2aX+zl/6p7wbXuUv4VL/AK9w/wDSUf65+G3/ACbrgH/siuFf/VFgD/FbrQ+1CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+h3/glL/wSik+I7+Hf2mP2mvDrR/DyNrbWfhh8LtZtmSTx66Ms9j4u8X2M6hk8EqwS40XQ7hA3jAhL+/jHhXyIPE3+W/02/pu0+A4Zr4QeEGaxqcc1I1cBxhxhgaqnT4MhNOniMlyXEQbjPi2UXKnjsbTbjwynKhRk+IeeWRf1/4AfR/lxHLB8b8cYOUOHouGJyPI8RBxlnzTUqWPx9OSvHJU0pYfDySeb6VJpZZyrMf6sI444o0iiRI4o0WOOONQkccaAKiIigKqKoCqqgBQAAMCv8OqlSpVqTq1ZzqVak5VKlSpJzqVKk25TnOcm5TnOTcpSk25Nttts/0IjGMIxhCKjGKUYxilGMYxVlGKWiSWiS0S0Q+oGFABQAUAFAHBfEz4meDvhF4O1bx1461aPSdC0qPljiS81C8kDfZNK0q03K99ql86mO1tYyOkk88kFrBcXEX3Xhx4ccX+LHF+V8EcEZVUzXPM1qe7FXp4TAYSm4/Ws0zTFcsoYHLMFCaqYrFVE7XhRowrYqvh6FX878VPFTgjwY4IzjxA8QM4pZPw9k9K8pPlqY7MsdUjP6nlGUYPnhUzDNswqQdPCYSm1e1SvXqUMJQxOJo/zbftLftLeMv2j/GLavrDS6T4S0mWeLwh4PhnaSz0ezkbabu7I2pfa7fIqNqGosg6La2iw2UMUQ/6Kvo4fRw4Q+jvwisqymNPNeLM1p0KvF3F1WgqeLzfFwXMsJhIy5p4HI8FOU45fl8ZveWKxc62MrVar/5d/pT/AEqON/pP8bSznOpVcm4MyetiKXBXBNHESqYLJMFUfLLG42UeSGYcQ5hTjCWZZnOC2jg8HChgaFGivm2v6MP5bCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/07P2Bf8AkxT9iv8A7NL/AGcv/VPeDa9yl/Cpf9e4f+ko/wBc/Db/AJN1wD/2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9Ef/BKX/glE/wAQX8O/tM/tN+HWj8Ao1trPwu+Fus2zJJ44dSs9h4v8Y2E6hl8GqQlxoehXCBvFh8u/1GMeGPIt/En+WX02/pvQ4Hjmvg/4PZrGpxtNVsv4x4ywFVShwdGSdPE5HkeIptqfFjTlSzDMKcnHhlOWHw8nxF7SeQf2D4AfR+fELwfHHHODccgThiciyLEwalnrTU6WYZjSkrrJb2nhsNJXzfSrVX9l8scz/qrREiRI40WOONVSONFCIiIAqoiqAqqqgKqqAAAABiv8PpznUnOpUnKpUqSlOc5ycpznJuUpzlJuUpSk25SbbbbbbbP9BoxjGKjFKMYpRjGKSjGKVkklokloktEh1SMKACgAoAKAOA+J3xO8G/CHwbqvjrx1q0elaHpUf+zJe6leyK5tNJ0m03K99ql86FLa2QgALJcXEkFpBcXEX3fht4bcX+LPF+V8EcEZVUzXPM0qaL3qeDy/B05R+t5rmuL5ZwwOWYKE1PE4mom7yp0KFOvi6+Hw9X868VfFXgjwX4IzjxA8Qc4pZPw/k9LV+7Ux2Z46pGbweT5Pg3OE8wzbMJwlTwmEpyirRqYnE1MPg8PicTR/m2/aV/aV8ZftHeMW1jWWk0rwppUk8PhDwfDO0llo1nI2DdXRARL7XL5FRtR1J0BbCWtqsNlBDCv/AEV/Rx+jjwh9HjhFZTlEaeacVZrToVeLeLq1BQxmcYymm44XCqTnPA5JgpynHL8ujUaV54rFSr42vWrS/wCXX6U30puN/pPcbyzvO5Vcn4OyeriKPBXBVDESqYHI8DUlyyxmLkuSGYcQZhCFOWZ5nOnHmahhMJDD4ChQoR+b6/oo/l4KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/07P2Bf8AkxT9iv8A7NL/AGcv/VPeDa9yl/Cpf9e4f+ko/wBc/Db/AJN1wD/2RXCv/qiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf0U/8ABKX/AIJRP48fw7+01+054daPwOjWutfC74Wa1bFX8aMpW4sPGPjKwnUMvhEHZc6FoNymfFJEeoalGPDf2e38Q/5XfTc+m9HguObeD3g7m0Z8ZSVbL+M+NMvrKUeEE06WKyHIcTTk1LitpypZlmVKTXDHv4XCy/1k9rV4e/sX6P8A9H5588FxzxzgmsiThishyHFQs87ek6WZZlRmrrJr2nhcLNXzf3a1Zf2VyRzP+qVESNFjjVUjRVRERQqIigKqqqgBVUABVAAAAAGK/wAQpSlOUpzlKc5ycpzk3KUpSbcpSk23KUm22222223c/wBAklFKMUoxikkkrJJaJJLRJLRJbDqkYUAFABQAUAef/E/4n+Dfg/4N1Xx1461WPS9D0uPAA2yX2p30iubTSdJtC6PfapfOjJbWyMqhVkuLiW3tILi5h+88NfDXjDxa4vyvgjgjK6maZ3mc29XKngsuwVOUVis1zXF8s4YLLMFGanicTOMm5Sp4fD06+Lr4fD1vzjxX8V+B/BXgfOPEHxBzinlGQZRTtpy1cfmmPqxm8Hk2TYNzhPMM2zCcJQwuFpyilGNXE4mrh8Fh8ViqP82v7Sf7SnjP9o7xk2ta20mleFtKknh8IeEIJ2ksdEspGAa4uGARL7XL5EjbUtTeNS5VLa2S3sYILdP+iz6Of0cuD/o8cIRyfJ40804ozSnQrcW8W1qEaeNznG04trD4aLc54HJMFOdSOXZdGpJQUp4nFTr42vXrz/5dfpSfSk44+k9xvPPc9nUyjhDKKuIocFcFUMRKrgMhwFSSUsVipJQhmGf4+EKc80zSdOLnJQwuEhh8Bh8Ph6fzjX9EH8wBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/p2fsC/wDJin7Ff/Zpf7OX/qnvBte5S/hUv+vcP/SUf65+G3/JuuAf+yK4V/8AVFgD/FbrQ+1CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+jD/glJ/wShfxs3hz9pr9p3w4yeDFNtrXws+FetWxV/F7KVnsPGXjOwnUMvhUEJc6D4fuUB8THy9S1OP/AIR37Pa6/wD5VfTd+m/Hg5Zt4O+DmbKXF8lXy7jXjbL6ylHhJNSpYrh/h/E021Lippzo5nmlKVuGPfwmEm+JPa1uHf7H+j/9H5568Fxzx1gmskThisgyDFQaedPSdHM8zpTV1k60qYTBzV8392vXX9lclPNP6nkRY1VEVURFCIiAKqqowqqowFVQAAAAABgV/iLKUpSlKUnKUm5SlJtylJu7lJu7bbbbbd29Wf6AJJJJJJJJJJWSS2SXRLoh1IYUAFABQAUAeffFD4oeDPg94M1Xx1461VNL0TS0wqjbJf6pfyK5tNI0i0Lo99ql86MlvboyqqrLc3MtvZ29zcw/e+Gnhpxh4ucYZXwRwRlc8zzrM53bfNTwOW4GnKCxebZti1CccFlmCjOM8RiJxlKUpU8NhqeIxmIw2Grfm/ix4s8D+CnA+b+IPiDnFPKMgymnaKXLVzDNswqxm8Hk2S4JzhPMM2zCcJQw2GhKMYxjVxWKrYbA4bFYqh/Np+0l+0l4z/aO8ZPrmuPJpfhfS5J4fCHhCC4aSx0OxkYBp52ARb7W75EjfU9TeNWlZUtrZLexgt7aP/ot+jp9HTg/6PHCEclyWMM04nzOFCtxbxbWoRp47O8bTi2qFBNzngclwU51IZblsKko04ynicTPEY6viMRU/wCXT6Uf0o+OPpPccTz7Pp1Mo4SyipiKHBXBVDESq5fkGAqySliMRJKEMwz/ADCEKc81zWdOMqsowwuFp4bL8PhcLS+c6/oY/mIKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/Ts/YF/5MU/Yr/7NL/Zy/wDVPeDa9yl/Cpf9e4f+ko/1z8Nv+TdcA/8AZFcK/wDqiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9G//AASk/wCCULeL28OftN/tPeHGTwkptdb+Ffwp1u1Kt4qYFbjT/GfjXT7hMr4ZB8u58P8Ah65QHxGfL1PVI/7A+zWuu/5TfTd+m+uElm/g54N5tfix+2y7jbjfLq91wqnzUsVw9w7iqUteKH71HNc1oytwz7+DwU3xJ7etw5/ZP0fvo/POnguOuO8E1kv7vFZBw/iqdnnO06OaZpRmv+RRtUweDqL/AIVvdr11/ZXs6eaf1LqqoqoiqiIoVVUBVVVGFVVGAFAAAA4A4Ff4ktuTcpNylJtyk22227ttvVtvVt6t6s/v9JJJJWS0SWyXZDqQBQAUAFABQB558Ufil4M+DvgzVPHXjvVU0zRdMTaiLtlv9Vv5Fc2mkaRaF0a+1S9ZGWC3VlREWW6upbeyt7m5h+/8MvDLjHxd4wyvgjgfK55nnWZTcpyk5U8DlmBpygsXm2bYtQnHBZZgozjKviJRlOc5UsNhqWIxuIw2Grfm3i14tcDeCXA+b+IPiDm8MpyHKocsIRUauY5vmNWM3g8lyXBOpTnj82x84SjhsNGcIQhGti8XWw2Bw2KxVD+bX9pH9pHxn+0b4zfXddd9L8M6ZJcQeEPCEE7SWGhWEjANNMwCLfa3fKkb6pqjxq8zqlvbpb2Fva2sP/Rb9HX6OvB/0eeD45JkkIZnxLmcKFfi3i2vQjTx2eY6nFuNGim5ywWTYKU6kMtyyFSUaUZTxGIniMdiMTiav/Ln9KH6UXHP0neOJ8QcQTqZTwplNTEUOC+CsPiZ1cu4fy+rJKVevJKnDMM+x8IU55tm1SlCdecYYbDU8Nl+GwmEo/Olf0IfzIFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf6dn7Av/Jin7Ff/AGaX+zl/6p7wbXuUv4VL/r3D/wBJR/rn4bf8m64B/wCyK4V/9UWAP8VutD7UKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+j3/glH/wSgbxS3hz9p39p/w4V8MA2ut/Cr4T63akN4lIK3GneNfG+n3CZHh0HZdeHvDl1HnXz5WqarENC+y2mt/5Q/Td+nAuFlm/g34NZunxQ/b5bxvxxl1a64YT5qWL4d4cxdKTvxL8dHNs2oy/4xz38FgZ/wCsXt6/D39l/R++j684eC4647wX/CR+7xXD/D2Kp65u9J0c0zSjNf8AIp2qYPBVI/8ACr7uIxEf7L5KeZ/1JKqqoVQFVQFVVACqoGAABwABwAOAK/xObcm22222227tt6ttvVtvVt7n9+7aLRLYWkAUAFABQAUAed/FP4p+DPg34M1Tx1461RNN0bTU2xxrtk1DVtQkVzaaPo9oXRr3U71kZYIFZUjRZbu7lt7K2ubmH9A8MfDHjHxe4xyzgjgfK55lnOZTc6k5OVPAZXgKcoLF5vm+LUJxwWWYKM4yr15RlOc50sLhaWIxuIw2Grfmni34t8DeCPA2b+IPiDm8MqyLKoclOnHlq5lnGZVYzeCyXJME5055hm2PlCUcPh4yhTp04VsZjK2FwGFxeLofzaftH/tH+NP2jPGb69r8j6b4b0154PCPhCCd5NP0GwkYZkkOEW91q9VI31TVHjWS4dUggS3sLe1tYP8Aou+jv9Hbg76PPB8MjyOEMy4jzKFCvxZxbXoRp4/PcdTi2qdNXnLBZPgpTqQyzLIVJQoQlOvXniMdiMViq3/Ln9KD6UHHP0neOZ8Q8RTnlXC2VVMRh+C+C8PiJ1cu4dy6rJc1WrK1OOYZ7j406VTN83qUoVMRUhDD4anhcuw2DwdD52r+gj+ZgoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/07P2Bf+TFP2K/+zS/2cv/AFT3g2vcpfwqX/XuH/pKP9c/Db/k3XAP/ZFcK/8AqiwB/it1ofahQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/SH/AMEo/wDglA3iM+G/2nf2oPDZHh0G11z4UfCfXLXB8QH5bjTvG3jjTrhMjQQfLuvDnhu6jB1z91qurxf2L9ls9Y/yd+m99OD/AFa/tfwb8Gc3vxL+/wAt4445y6vdcOfFSxfDnDeLpS14i+Ojm+cUZf8AGPe/gcBUfEHt8Rw//Z30fvo+vNnguO+O8F/wlfu8Xw/w9iqeuavSdHNM1ozX/Ir2qYLA1F/wqe7iMTH+zPZ0sz/qMACgKoCqoAAAwABwAAOAAOAB0r/FJttttttu7b1bb3bfVs/vrbRaJbC0gCgAoAKACgDzr4qfFTwX8GvBeqeOvHeqLpujaauyKJNsmo6vqMiu1po+j2jOjXup3rIwhhDJHHGk13dzW1jbXNzD+heF/hfxj4wcY5ZwPwRlk8xzjMZ89WrPmp5flOX05QWLzfN8WoTjgstwUZxlWrSjOpUqTpYTC0sRjcThsNW/M/F3xd4F8D+Bs28QfEHN4ZXkeVw5KVKHJVzPOcyqwqSweSZJgpVKcsfm2PlTmqFCM4U6VOFbGYythcBhcXi6H82f7R37R3jT9ozxo/iDxBI+m+HNNeeDwj4Qt7h5dO8P6fIwy7nEa32s3qpG+q6q8SSXMiJBAltYW1nZ2/8A0XfR4+jxwd9Hrg6GRZDCGY8Q5lGhX4s4sr0IU8wz7H04u0IK85YLJ8FKdSGV5XCrOnh4TqV608Rj8Ti8XX/5cfpPfSe46+k5xzU4j4jqTyrhjK54jD8GcF4fEzq5Zw5ltSavUqO1OGPzzHxhSqZvnFSlCriqkKeHoU8Nl2FwWCw/zxX9AH80hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+ov8AwTL/AOCYvxO/4KC/Ebzn/tLwT+z74O1O3T4m/FL7KN0siiK6bwR4FF1E9rq3jbUrWSN5ZHS40zwnp9zFrWuJPJPouia/vQoSrS6qC+KX6Luz9g8JPCPN/E3NeZ+1y/hnAVorN845NZPSby/LudOFfMKsGnJtSpYKlOOIxCk5YfD4n+0fRv8Agkv/AME5tF0jS9Hi/ZM+Fd/FpWn2enR32s6dfatq96llbx2y3eqape38t5qOo3AjE17fXUj3F1cPJPK7O7GvTVCilb2cdO6u/m3qz+98P4LeFuHoUaEeCsmqqjSp0lUxFGdevUVOKip1q1SpKpVqytzVKk5OU5NybuzS/wCHVX/BOr/o0H4M/wDhPS//ACZT9hR/59w+5G3/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBuH/Dqr/gnV/0aD8Gf/Cel/wDkyj2FH/n3D7kH/EHPC7/oh8g/8JP/ALcP+HVX/BOr/o0H4M/+E9L/APJlHsKP/PuH3IP+IOeF3/RD5B/4Sf8A24f8Oqv+CdX/AEaD8Gf/AAnpf/kyj2FH/n3D7kH/ABBzwu/6IfIP/CT/AO3D/h1V/wAE6v8Ao0H4M/8AhPS//JlHsKP/AD7h9yD/AIg54Xf9EPkH/hJ/9uH/AA6q/wCCdX/RoPwZ/wDCel/+TKPYUf8An3D7kH/EHPC7/oh8g/8ACT/7cP8Ah1V/wTq/6NB+DP8A4T0v/wAmUewo/wDPuH3IP+IOeF3/AEQ+Qf8AhJ/9uH/Dqr/gnV/0aD8Gf/Cel/8Akyj2FH/n3D7kH/EHPC7/AKIfIP8Awk/+3D/h1V/wTq/6NB+DP/hPS/8AyZR7Cj/z7h9yD/iDnhd/0Q+Qf+En/wBufcXhLwp4c8CeFfDPgfwfo9l4e8JeDfD+jeFPC2gabGYdO0Pw54e0620jRNHsIizGKy0zTLO1srWMsxSCFFLHGa0SSSS0SVkuyR+hYLB4XLsHhMvwNCnhcFgMNQweDw1JctLD4XC0oUMPQpx15adKlCFOC6Rikf4Y9M6QoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+kr/glH/wSgOunw3+09+1B4bI0IfZNc+E/wm1y151s/Jc6d438c6dcLkaL/q7rw54Zu486x+61bWIv7I+yWerf5M/Te+nAuHv7X8GvBnN78Q/v8t4445y2v/yT/wAVLF8OcN4ulLXP/ioZvnFCf/CD+8wOAn/bvt8RkX9o/R++j7/af1LjvjzBf8Jn7vFcP8O4qnrme06Oa5tRmv8AkW/DUwOBqR/4UvdxOJj/AGd7OlmP9QoAAAAAAAAAGAAOgA7AdhX+Krbbbbu3q29W2+rP73FoAKACgAoAKAPOPit8VvBfwZ8F6p468daounaPpy7IYU2SajrGoyI7Wmj6PaM6Ne6nesjCGEMkUUaTXd5NbWNtc3MP6H4XeF3GXjDxjlnBHA+WSzHN8wlz1q0+anl+U5fTnBYvN83xahOODy3BqpF1qzjOpVqTo4TCUsTjsThsNW/MfF/xf4F8DuBc28QfEHNoZZkmWR9nQoU+SrmedZnVhUlgskyTBSnTlj81x8qc1RoKcKVGlCvjcbXwuX4XF4uh/Nn+0Z+0b40/aM8aP4h8QyNp3h7Tmnt/CPhG3neTTvD+nSOMsxIQXusXqxxSarq0kSSXUiJDDHbWFtZ2dv8A9F30efo88HfR64OhkGQQjmPEGYxoV+K+LMRQjTzDPswpxdoxV5ywWUYOU6kMryuFSdPDU5zrVqmJx2IxeLxH/Lj9J36TvHX0nOOanEnEtSeWcNZXPEYfgzgzDYidXLOG8tqzjzSk7U44/O8fGnSqZxnFSlCri6sKdChTwuXYXBYHDfPVfv5/NYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfq9/wS//AOCW/wAR/wDgoD4/Gtat/avgf9m/wdqsMXxD+JKW6x3es3UQiuZPAfw+N3FJa6j4rvbeSI6hqLxXOleD9PuYtU1aK7vLnRdC1zooUJVnd6QT959/7q8336bvon+0eEPg9mviZmX1iv7bLuFcBWjHM81ULTxE1yzeW5ZzxcKuNnFp1arU6OBpTjWrKdSeHw+I/v7+Efwi+HHwI+HXhX4T/CXwnpfgnwB4M02PS/D/AIe0iJkgtoVZpJ7q5nlaS71LVdRupJr/AFfWNRnutT1fUrm61HUbq5vLmaZ/XjGMUoxVktl/WrfdvV7s/wBK8kyTKuHMrweS5LgqOX5ZgKKo4XC0U1GEU25TnKTc6tarNyq169WU61erOdWrOdScpP0emeqFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf4U9ABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf0of8Eov+CT51c+G/2nv2ofDZGk5tdc+E/wAJNctSDqpG2407xx46065TP9l/6u78NeGbuP8A4mn7rV9Zi/sz7JZal/kp9N76cCyP+1/BrwYze+efv8t4446y2vpkfxUsXw3w1i6Utc7+KhnGc0J2yX95gMvqPOvrGIyX+0/o/fR9/tD6lx3x5gv+E/8Ad4vh7h3F0/8AkYbToZrm1Ga/5F+1TBYCpH/b/dxGJj9Q9nSx/wDUAAAAAMAcADoB6Cv8Wt9Wf3qFABQAUAFABQB5v8V/iv4K+DHgvU/HXjvVF0/SNPXy4II9kupazqUiO1po2jWjPGb3U70xsIog6RQxJNeXk1tY21zdQ/onhb4W8ZeMXGWW8EcD5ZLMM2zCXtK9epz08uyjLqc4RxecZxi1CccHluDVSLq1XGdWtVnRwmEo4nHYnDYat+YeMHjBwL4G8C5t4g+IObRyzJctj7PD4enyVc0zvNKsKksFkeR4KVSnLHZrjpU5qjRU4UqNGnXxuNr4XL8Li8XQ/my/aL/aM8a/tGeNJPEXiKRtO0DTmuLfwj4Rt53l03w9p0rgkkkRi91i9WOKTVtWkiSW7lSOKGO10+2srK2/6L/o9/R64N+j1wdDh/h+nHMM+zCNDEcWcWYihGnmPEGYU4yskrzeCynByqVaeV5VTqTp4WnOdatPE4/E4zGYn/lw+k39Jvjr6TfHVTiXiapLLOHMsniMNwbwZhsROrlnDeWVZq7bapxx+dY6NOlUzjOKlKFXGVYU6NGnhcuwuBwOF+fK/fT+bQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/Xj/AIJZf8EqvH/7fXjdPF3itdW8EfsyeD9WSHxr48ji+z6j4vv7YxzT+A/h69zFJBda1PEyLrevNFc6Z4Ssp1uLqO81WfTNHv8AooYd1nd6QT1fV+S8+76euj/b/B3wbzLxKzBY7G+3y/hHA11HMMxUeWrjqsLSlluVuacZ4iSaWIxLjOlgqclKaqVpUaFT++z4YfDDwB8F/APhb4XfC3wrpPgrwD4L0qHRvDfhrRYPIsdPsoSzsSztJcXl7eXEk19qep301zqWq6jc3Wpald3V/dXFxJ68YqKUYqyWiS/r8er1Z/pPlGUZbkOW4PJ8nwdDL8twFGNDCYTDx5adKnG7e7cqlScnKpWrVJTq1qs51a051Jzm+8pnpBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/hT0AFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9K3/BKL/gk+dSPhv8Aaf8A2ofDf/Eu/wBE1z4TfCTXbTnUD8lzpvjnx3ptyn/IP/1d34a8MXcf/Ew/c6vrMX2D7HY33+R/03vpwf2P/a/g14MZv/wsfv8ALeOOO8tr/wDIo+Kli+G+GcZRl/yN/joZxnNCf/CT+8wGX1P7W9vicp/tb6P/ANHz699R4748wX+w/u8Xw9w7iqeuO2nQzXNqE1/uXw1MDgKi/wBt93E4qP1L2dLG/wBPXTgcAdK/xfP7yCgAoAKACgAoA82+LHxY8FfBfwVqfjvx1qa6fpOnr5dvbx7JNS1rUpEdrTRtGtGeM3upXpjfy496QwQpNe3k1tY211dQ/ovhZ4WcZeMfGWW8EcD5ZLMM1x8vaYjEVOenl2T5dTnCOLzjOMXGFRYPLcGqkXUquM6tarOjg8HRxOOxOGw1b8u8YvGLgTwL4FzXxB8Qc1jl2TZdH2WGw1Lkq5pnmaVYVJYLI8jwUqlOWOzTHSpzVKkpwo0KMK+Ox1fC5fhcVi6P82P7RX7RXjX9ovxpJ4k8SSNp+hae09v4S8JW07yab4d02RwSASsYvdWvBHFJq2rSRJLeSokUUdtYW1lZWv8A0YfR8+j5wb9Hvg6nw9w9Tjj88zCNDEcV8V4ihCnmPEGY04NJtJ1Hg8qwcp1aeVZVTqzpYSlOdWrUxOPxOMxmJ/5cPpNfSa46+k1x1V4n4nqyy3h7LZYjDcG8G4bETq5XwzldWabSbjTjjs5xyp0qmc5zUowrY2tCnRpU8Ll2FwGAwvz9X72fzeFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH7Jf8Ep/+CT3jj9vHxfb/ELx/Fq/gv8AZc8KasIvEviuNWs9X+IupWMga58EfD+WaJld9wEHiTxQElsvD0LvbW/2vXGjtIenD4d1nd3UFu+sn2X6vp6n7v4NeC+YeI+OjmmZxr4Dg/BV7YvGJOnXzWrTd55fljknd393F4yzp4WLcY8+ItCP97nw8+Hngj4TeCfDPw3+G3hjSPBngXwdpNtonhnwzoVqtppmk6bag7IYYwWeWWWRpLm8vLmSa91C9muL6/uLm9uZ55PWSUUklZLRJH+kWWZZl+S5fhMqyrCUMBl2Aoxw+EwmGgoUaNKG0YrduTbnUqScqlWpKdSpKdScpPsqZ3hQAUAFAH8/f7c//BxP+yH+xd8cNV+Aun+EPG37QHi3wjHJbfEPVPhnqvhq38M+C/FCTtHN4Kn1fWLxE1fxLpaIx8R2+lxy2egXjpot5eNrtrrGmaT+6cF+AfFfF+TUs8nisHkeFxTUsBTzGniXicZhnFNYyNKlBulh6rf+zyqtTrwTrQh7CVKpV/LuJfFfIeHsynlkaGJzSvQTWLng50VRw9a9nh3UqTXtK0P+XygnGlJ+zlL2sakKfxp/xFrfszf9Go/HX/wqfAH/AMmV9b/xK5xH/wBFPkn/AIT4/wD+Vnz/APxHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJh/xFrfszf9Go/HX/AMKnwB/8mUf8SucR/wDRT5J/4T4//wCVh/xHLJv+hJmf/g7C/wDyYf8AEWt+zN/0aj8df/Cp8Af/ACZR/wASucR/9FPkn/hPj/8A5WH/ABHLJv8AoSZn/wCDsL/8mH/EWt+zN/0aj8df/Cp8Af8AyZR/xK5xH/0U+Sf+E+P/APlYf8Ryyb/oSZn/AODsL/8AJh/xFrfszf8ARqPx1/8ACp8Af/JlH/ErnEf/AEU+Sf8AhPj/AP5WH/Ecsm/6EmZ/+DsL/wDJn9OvwC+LemfH74FfBX476LpN/oOjfGv4S/Dj4t6ToeqS28+p6LpnxH8HaN4xsNJ1Ge0ZrWa/0611mKzvJbZmt5LiGR4WMZU1/OWeZXUyPOs4yStVhXrZPmmYZXVrUlJU61TL8XWwlSrTU7SUKkqLnBSSkotX1ufseV46GaZbl2Z04SpU8xwOEx1OnNpzpwxeHp4iMJuN4uUI1FGTTabTtof4gNeWdwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/S5/wAEov8Agk/9s/4Rv9p/9qLw1/of+ia58JfhHrtr/wAfn3LnTfHXjzTbhP8Aj0/1d34Z8L3cf+l/utY1qH7L9jsbr/Ir6b304f7M/tfwa8GM4/4U/wB/lvHPHeW1/wDkWfFSxfDXDOLpS/5Gfx0M4zqhP/hN/eYDLqn9p/WMTln9s/R++j79b+o8d8eYL/ZP3eL4e4dxdP8A3vadDNc2ozX+6bVMDgKkf9r93E4qP1T2dLF/06V/jIf3eFABQAUAFABQB5p8Wvi14K+CvgrU/HXjrU1sNKsF8q1tYtkmp63qciO1po2jWjPGbzUrwxvsTckMEKTXt7NbWNtc3MX6P4VeFXGfjJxllvBHBGWyx+a46XtcTiavPTy3JstpzhHF5xnGLjCosHl2EVSPtKnLOtXrTo4PB0cTjsThsNV/LfGPxj4E8CuBc18QfEHNY5dk+XR9lhcLS5Kua57mtWFSWDyPI8FKpTljs0xrpz9nT54UcPQhXx2Or4XL8LisVR/mx/aI/aJ8a/tFeNJPEniWVrDRLBp7fwn4Stp3k0zw5psjqSqkiMXmq3gjik1bVpIkmvZkjjjjtrC2sbG1/wCi/wCj79H3g36PnBtPh3h2nHH53j40MRxVxXiKEKeZcQ5jThJKUknUeDyrBupVp5VlVOrOlg6U6lSpUxOPxONxuK/5b/pMfSY47+k1x3V4o4oqyy7IMuliMNwbwbhcROrlXDGV1ZxbjFuNNY7OMcqdGrnOc1aUK+OrU6dKlTwuXYXAYDCeAV+9H84hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+13/BJz/gkl4v/AG5vE9p8Vfirb6v4O/ZY8MaqY9S1ePztP1v4tarp0+Lvwf4HuGVZINIgmRrXxZ4xhzHpp87RdEebX/tk+gdWHw7qvmlpTT+ctdUv1f3a7fv/AIL+CeO8Q8XTzrOYV8DwdhK1qtdXpYjO61KXv4HL5P3o0IyThjcfHSl72Hw7lieeWG/vJ8EeCPCHw18IeHfAPgDw5pHhDwX4R0m00Pw14a0Gyi0/SNG0mxjEVtZ2VpCqpHGigs7ndLPM0k88ks8skjeqkkkkrJKyXZH+juX5fgcqwOFy3LcLQwWAwVGGHwmEw1ONKhQo01ywp04R0SS3erlJuUm5Nt9TTOwKACgAoA/ke/4Lqf8ABdyP4Qp4u/Yx/Yr8XxzfFmRb3w78bfjh4dvFki+FqsHtdU+H/wAPdUtnKSfEpgZbTxN4ntJGT4egy6VpMp8d/arvwT/Uvgt4KPNnheLuMMK45WnDEZPk2IhZ5m9JU8dj6ctVl20sPh5K+O0q1V9S5Y4z8L8SvEv+z/b8P8PV74/3qWY5lSldYLdTwuFmnrjN41q0X/smsIP605Sw38Lksss8sk88kk000jyzTSu0ksssjF5JJJHJd5HclndiWZiWYkkmv7SSUUoxSSSSSSsklokktEktkfzY2222222223dtvVtt6tt7sjpiCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/1t/8Agm7/AMo7v2Cf+zLv2W//AFR3gWv8tvEH/kveN/8Asr+Jf/V1jT+5+Ef+SU4Y/wCyeyX/ANVuGP8AF4r5A+hCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/pi/4JRf8En/ADv+Eb/ag/ai8Nfuf9E134S/CPXbT/XfcudN8dePNNuU/wBT/q7vwz4WvI/337rWNbh8n7HYzf5C/Te+nD9R/tfwa8F84/279/lvHPHeWV/9x+Kli+G+GMZRl/v3x0M4zqhP/Yf3mAy6p9e+sYnA/wBufR++j79Y+pcd8eYH/Z/3eL4e4dxdP/eNp0c1zehNf7v8NTA4Cov9o93E4qP1f2dLEf03V/jSf3YFABQAUAFABQB5n8W/i34J+CngrUvHXjrUxY6XYjyrS0i2SaprmpyI7WmjaLaNJGbzUbwo2xN6QW8KT3t7PbWNtc3MP6R4U+FPGfjNxnl3BHBGWyx2Z46XtcViqvPTyzJctpzhHF5xnGLjCosJl+EU488+WdbEVp0cHg6OJx2Jw2Gq/lfjL4y8CeBHAma+IPiDmscvyjL4+xweDo8lXNs+zWpCpLB5HkeCnUpvG5njXTnyQ56dDDUIV8dj6+Fy/C4rFUf5sf2h/wBofxt+0V41l8S+JpWsdFsWnt/CfhO2neTTPDmmSODsTIjF5qt4I4pNW1aSJJr6ZERI7awtrKxtf+jD6P8A9H7gz6PnBtPhzhuksdnOOjQxHFXFWJoQp5lxDmVOElzzs5vCZXhHUqwyrKqdWdHBUZznUqYnHYjG43Ff8t/0l/pL8d/SZ47q8U8VVpZfkWXyxGF4O4OwuIqVcq4YyqrNNwhdU1js3xqp0quc5zVpQr4+vCnTp08Ll2Fy/L8H4DX7wfzmFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH7lf8ABJP/AIJD+J/219f0/wCM/wAaLLVvCn7LPh3VD86m40zXPjRqum3BS68M+E7keXcWXhS1uYntPFvjK2KuJEuPDnhqY64uqap4X68PhnVfNO6p/O8n5eXeXdWWt7f0P4KeCGM4/wATSz/P6dbBcHYWtv71LEZ/WpStPCYKWk6eChNOGNx8bO6lhcJJ4j21bB/3a+EvCXhjwF4Y0HwX4K0DSfCvhLwtpVlofhzw5oNjb6Zo+i6Rp0CW1jp2nWFqkdva2ttBGsccUaAADJyxJPqJJJJKyWiS6H+i+CwWEy3CYbAYDDUcHgsHRp4fC4XDU40qGHoUoqFOlSpwSjCEIpJJL8ToaZ1BQAUAFAH8h/8AwXV/4Lup8Ll8X/sXfsU+L1k+JzLe+G/jj8dfDl6GT4bKwe11X4d/DjVbVyH+IZBls/FXi2ylK+AszaPokx8cfbL3wV/VPgt4KPM3hOL+MMK1lqcMRk2S4iGuY7Sp4/MKctVgL2nhsLNXx2lWsvqfLDGfhPiV4lrBfWOHuHa98Y+ajmWZ0ZX+p7xqYTCTj/zF7xr14v8A2XWnTf1nmnh/4apJHld5ZXeSSR2kkkkYu8juSzu7sSzOzEszMSWJJJJNf2ckkkkkklZJaJJbJLokfza22227t6tvVtvqxlMAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/AOCbv/KO79gn/sy79lv/ANUd4Fr/AC28Qf8AkveN/wDsr+Jf/V1jT+5+Ef8AklOGP+yeyX/1W4Y/xeK+QPoQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/pn/4JRf8ABJ858NftQftReGxjFrrvwl+Eeu2nOf3dzpnjvx5p1yvT7l34Z8LXkRzmHWNaiGLOwP8AkB9N76cP1f8Atfwa8GM3/wBp/f5bxzx5llf/AHb4qOM4a4ZxdF/7z8VDOM7oT/2b95gMuqfWfb4rDf2/9H76PntfqXHfHmB/dfu8Vw9w5i6f8XadDNc2oVF/C2qYHAVI/vfdxOKj7L2dKt/TTX+Nx/dIUAFABQAUAFAHmPxd+Lvgn4J+CtS8deOtSFlplkPJs7OHZJqmu6pIkj2mjaNaO8Zu9RuzG5VdyQW0CT3t7PbWNtc3MX6T4UeFHGfjPxnl3BHBGXPG5ljX7XF4urz08syXLKc4RxecZxi4wqLCZfhFOPNLlnWxFadHB4OjicbicPh6v5V4zeM3AfgPwJmniB4g5osBlOAXscFgqPs6ubZ/m1WnUng8jyPBTqU3jczxjpzcIc9OhhqFOvjsdXwuAwuKxVH+bH9ob9obxt+0T41l8T+J5TY6PYme28KeFLad5NL8N6ZI4Plx5CC71S8EcUmratJEk9/OkaIltYW1jY2n/Rj4AeAHBn0feDaXDfDVJY3N8aqOJ4p4pxNGFPM+Isypwa9pUs5vCZZhHOrTyrKqdWdHA0ZznOpicdicbjcV/wAt30lfpK8efSY47rcV8V1pYDJMBLEYXg/g7C4ipVynhfKatRP2VK8aaxubY1U6VXOc5q0aeIzCvCnCFPC5dhcvy/B+B1+7H86hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+9H/BIr/gj9r/7Yus6V8ePj3puqeGv2XtD1Ey6ZpzG50vW/jhqenXBSfR9DnQw3mn+A7S6ie18TeLLVo57+WO48OeF511JdW1nwz14fDOo+eekE9F1n/wO7+S62/pDwR8DcTx1Xo8ScS0q2E4Qw9W9Gk+ejiOIa1KVpUMPJOM6WWwmnDF42DUqslLC4OXtfb18J/c94c8OaB4P0DRfCvhTRdL8N+GfDmmWWi6BoGiWNtpmj6NpGm28dpp+maZp1nHDa2VjZW0UcFtbW8UcUMSKiKFAFertotEtj/RHC4XDYHDUMHgsPRwuEwtGnh8NhsPThRoYehSioUqNGlTUYU6dOCUYQilGMUkkbVBuFABQAUAfyAf8F1v+C7yfDxfGH7Fn7E3jBX+ILC+8NfHP48eGr4MngJWElpq3w4+GmrWrkN45OZbLxb4xspSvgr9/omgzHxl9t1Dwf/Vngt4J/wBofVOL+McK1gPcxGTZJiIWeO2nSx+Y0pq/1LaeFws1fGaVq6+qclPF/g/iX4l/VPrHD3DuI/2v3qOZZnRl/uu8amEwdRP/AHreNfERf+za06T+sc08P/Dw7s7M7szu7F3dyWZ2YkszMSSzMSSSSSSSSc1/ZiSSSSsloktkuyP5v31Y2gAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/1t/wDgm7/yju/YJ/7Mu/Zb/wDVHeBa/wAtvEH/AJL3jf8A7K/iX/1dY0/ufhH/AJJThj/snsl/9VuGP8XivkD6EKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/pu/wCCUX/BJ/yf+Eb/AGn/ANqPw1++/wBE1z4S/CPXbT/U/cudN8dePNNuU/13+ru/DPha8j/c/udY1qHzvsdjD/j19N76cPJ/a/g14MZv7/7/AC3jnjvLa/wfFSxnDXDOMoy+P46GcZ1Qn7n7zAZdU5/rGKh/cX0fvo+831LjvjzA6fu8Xw9w7i6fxbToZrm9Ca+H4amBwFRa+7icVHl9nSl/TFX+Oh/cwUAFABQAUAFAHmHxe+L3gn4I+CdS8deOtSFnptmPJsrKHZJquu6rIjvaaNo1o7obvULsoxALJBawJNe3s1tY21xcRfpfhN4TcZ+NHGeXcEcEZc8ZmOMftsZjK3PTyzJMspzhHFZxnGLjCosLgMKqkbtRnXxNedHB4KjicbiMPh6v5R40eNHAfgNwJmniB4g5osDlWBXsMDgaHJVzfiDNqlOpPB5HkeDnUpvGZljHTm4xc6eHwuHhXx+Pr4XAYXE4ql/Nh+0J+0J42/aI8bTeKPFExstJsjPbeFPCltO8ml+GtLkcN5MO4Ri71K7CRSatq0kST6hOiKqW1jbWNjaf9GPgD4A8GfR+4NpcNcM0VjM2xqo4nininE0YQzPiPM6cGva1bOo8LluFc6tPKsqp1Z0MDRnOcp4jHYnHY7F/8t30k/pKcefSY47rcWcWVngMmwLxGF4P4PwuIqVcp4WymrUjL2NHmVNY3NcYqdKrnOc1aNPEZjiKdOMaeFy/C5fl+D8Fr90P53CgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+gf/gkF/wAEddY/aw1LRf2iP2jNJ1HQf2adKvluvDXhmf7Tpus/HK/sZyGgtZEMN5p3w3trmJoNa8QW7xXWvyRzaJ4cnjYalrOkdmHwzqNTn8HRdZv/AC7vrsurX9M+B/gXX4zq4finiqhVw3CdGpz4PCS5qVfiGrTlrGD92dLKoTTjXxMWp4qSlh8LJWq16P8AcNomiaN4a0fSvDvh3StO0LQNC06y0jRNE0eyttN0nSNK063jtNP03TNPs44bSxsLG1iitrS0too4LeCNIoo0RFUeptotEtj/AEJw+HoYShRwuFo0sNhsPSp0MPh6FOFKjQo0oKFKjRpU1GFOlThGMIQhFRhFKMUkkjToNgoAKACgD+PH/guv/wAF318Fjxh+xV+xJ4xDeMyL/wAM/Hb49eGb7K+Ds+ZZ6v8ADX4Y6xaOQ3i//XWPi/xnYTY8Jfv9C8P3H/CV/b9R8Lf1f4K+CjxjwnGHGOE/2P3MRkuSYmGuLekqWYZjSmv902nhcJNf7VpWrx+rclPE/gniV4mfVvrHDvDuI/2j3qOZ5nRlf6vvGpg8HUT/AN43jiMRFv2GtKk/b806H8QzMzMWYlmYlmZiSzMTkkk8kk8knknk1/ZJ/OW+rEoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/wCCbv8Ayju/YJ/7Mu/Zb/8AVHeBa/y28Qf+S943/wCyv4l/9XWNP7n4R/5JThj/ALJ7Jf8A1W4Y/wAXivkD6EKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+nT/AIJRf8En/sf/AAjf7UH7UXhr/TP9E134S/CPXbT/AI8/uXOm+OvHmm3Kf8fn+ru/DPha7j/0T9zrGtQ/bPsdjaf46/Te+nDzf2v4NeC+ce7+/wAt4547y2v8XxUcXw1wxjKUvh+OhnGdUJ+97+Ay6py/WMTL+5fo/fR9t9R4748wXvfu8Xw9w7i6e206Ga5vQqLf4amBwFRae7icVG/s6S/pcr/Hk/uMKACgAoAKACgDy74wfGDwT8EPBOo+OvHWoi0060BhsLCEpJquvarIjvaaNo1o7obq/uijHlkt7W3Sa9vZreyt7i4j/TPCXwl408aeM8v4J4Iy54vMMW/bY7HVuenleR5ZCcI4rOM4xUYTWGwOFU4p2jPEYqvOjg8FRxGNxFChU/JvGnxp4D8BeA8z8QPEDNFgsswSdDAYCg6dTN+Ic3qU6k8HkeR4OdSm8ZmOLdObScoYfCYenXx+Pr4bAYXE4ml/Nh+0H+0H42/aI8bTeKfFMxs9KszPbeFfCttPJJpXhrS5HVvIg3LGLrUboRxSatq0sSXGoTogCW1jbWNjaf8ARj4B+AXBn0fuDKXDPDFFYzNMYqOJ4o4oxNGFPNOI8zpwa9tW5XN4XLsK51aeVZVTqzoYChOcpTxGOxGOx2L/AOW76SX0kuPfpL8d1uLeLq7wWT4F4jC8IcIYWvUqZRwtlNWpGXsKCkoLGZpjFTpVc5zmrShiMxxFOnGMMLl+FwGX4Pwev3M/nkKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP6Jf+CPn/BG7U/2mb3QP2lv2m9DvdI/Z3sbmLUvBHgW9W4sNV+Nt1ayhor27X91daf8ADCKZCLi9RorvxeyNZaTJFpJuNTfsw2G57VJr3Oi/n/8AtfzP6i8DfAmrxbUw3FnF2HqUOF6c1Vy/LqnNTrcQThK8ak/hnSyhSXvVE4zxzThRaoc1Z/226Zpmm6Jpun6No2n2Ok6RpNjaaZpWlaZaQWGm6ZpthBHa2On6fY2scVtZ2VnbRRW9ra28UcFvBHHFDGkaKo9Q/wBAqVKlh6VKhQpU6NCjThSo0aUI06VKlTioU6VKnBKFOnTglGEIpRjFKMUkrF6g0CgAoAKAP44v+C7H/Bd//hGf+Ex/Yp/Yj8ZA+Jj9v8M/Hf4+eGL/ACPDf37TV/hp8LtYs5OfEn+usfGXjawlP/CO/v8AQPDlx/wkX2/U/D39ZeCvgm8T9U4w4xwn+ze5iclyPEw/3nadHMcypTX+77VMJg5r/aPdr4iP1fkp1/wHxK8TPY/WOHeHMR++96jmeaUZaUd41MHgqkX/ABt44jERf7rWlSftuedL+JUkkkkkkkkknJJPJJJ5JJ6mv7GP51EoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBbf/gm7/yju/YJ/wCzLv2W/wD1R3gWv8tvEH/kveN/+yv4l/8AV1jT+5+Ef+SU4Y/7J7Jf/Vbhj/F4r5A+hCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD+nn/glF/wAEn/7P/wCEb/ag/ai8Nf8AEx/0TXfhL8I9dtP+Qd9y503x148025T/AJCP+ru/DPhe8j/4l/7nWNah/tH7HY6d/jh9N76cPt/7Y8GvBjN/3H7/AC3jnjvLa/8AH+KljOG+GMZRl/A+OhnGdUJ/7R+8wGXVPYe3xVf+6Po/fR89n9S4748wP7z93i+HuHcXT/h7To5rm9Ca/ifDUwOAqL937uJxUfaezpU/6V6/x+P7gCgAoAKACgAoA8t+MPxh8E/A/wAE6j458c6iLSwtAYNP0+ApJquv6rIjva6No1q7obm+uSjEkslvaW6TXt7Nb2dvPPH+m+EnhJxp41cZ5fwTwTl7xePxb9tjsfW54ZXkWVwnCOKzjOMVGE/q2Bw3PFaRniMViJ0cHgqOIxmIoUKn5L41+NfAfgHwHmfiB4gZmsHluDToZdl1B06mb8RZvUp1J4PI8jwc6lN4vMMW6cnrKGGweGp18fj6+GwOGxGIp/zX/tA/tA+N/wBofxtP4q8VTm00y0M9t4W8K208kmleGdKkcMLe3DCMXWoXQSKTVtWliS41G4RPkt7K3sbG0/6MvATwD4L+j/wZR4Y4XorF5ni1RxPE/FGJowhmnEmaQg4uvX5XN4bL8M51aeVZVTqzoZfQnNueIxuIxuNxf/Lb9JD6SHHv0luPK/F3F9d4LKsE6+F4R4RwlepUyjhXKKlRS+r4ZSUPreZYtU6VXOM4q0oYnMsTCCUMNgMNl+AwXhNfuJ/PYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf0i/8Ed/+CNF58f7jw5+1F+1T4eudP+BdvLb6x8N/hlqkMtrffGSaJxLaeIPENu4Se0+F6OqyWlswjuPHhAZDH4UPneIO3DYbntUqL3d4x/m83/d/P03/AKs8DPAefEssLxfxlhZ0uHYyjXynKKylCpnsk1KGKxUXaUMoTScIO0sy3VsF72K/tSsrKz02ztNO060trDT7C2gsrGxsoIrWzsrO1iSC1tLS1gSOC2traCNIYIIUSKGJEjjRUUAemf3zTp06VOFKlCFOlThGnTp04qFOnTglGEIQilGMIxSjGMUlFJJKyLNBYUAFABQB/Gl/wXY/4Lwf2R/wmP7FH7EXjL/ib/6f4Y+PPx+8MX//ACCf9ZZ6x8MvhbrNnJ/yFh++sPGfjiwmP9kn7R4f8N3H9rf2hqmlf1t4K+CftvqnGHGWE/de5ickyPEw/i7To5jmVKa/hbTwmDqL977tfER9lyU6n8/+JXiZ7P6xw7w5iP3nvUczzSjL+HvGpg8FUi/4m8cRiYv93rSov2nNOn/FJ15Nf2EfzuFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/8ABN3/AJR3fsE/9mXfst/+qO8C1/lt4g/8l7xv/wBlfxL/AOrrGn9z8I/8kpwx/wBk9kv/AKrcMf4vFfIH0IUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAB1oA/p//AOCUX/BJ/wDsj/hG/wBp/wDai8Ng6t/omu/Cb4R67aZ/sn7lzpvjnx5ptyn/ACFv9Xd+GfC93Gf7K/daxrUX9q/Y7HS/8a/pvfTh+uf2v4NeDGb/AOx/v8t4547y2v8A758VHF8N8M4yjL/c/joZxnVCd8Z+8wGXVPqn1jE4z+6/o/fR89h9S4748wP7/wDd4vh7h3F0/wCBtOjmub0Jr+P8NTA4CpH9x7uJxUfb+zpUP6Ua/wAgz+3QoAKACgAoAKAPK/jH8Y/BHwO8E6h458c6iLWxtswadp0BSTVvEGrSRu9ro2jWruhub252MzMzJb2luk17ezW9nbzzp+n+EXhFxp418aZfwTwTl7xWOxTVfMMwrqcMryLK4ThHFZxnGKjCf1fBYfniklGeIxeInRweDo4jGYijRn+R+NnjbwF4BcB5lx/4gZmsJl2ETw+W5bh3TqZxxHnFSnOeEyTJMJOdN4rH4pwlJuUqeGweGhXx2Pr4bBYeviKf82H7QH7QHjf9obxtP4r8Vzm1061M1t4X8LWs8kmk+GdKeQMLa2DBBc39yEik1bVpYkuNRuET5Leyt7GxtP8Aoy8BfAXgv6P/AAZR4X4Wo/WsxxSo4nififE0YQzTiTNIQcXiMQ4uf1bAYZzq08ryunUnQy+hOfv4jGYjG43F/wDLb9I/6R/Hv0lePMRxfxhiHg8rwbr4XhLhLCV6lTJ+FcoqVFJYbCqSh9azHF+zpVc4zirShiczxMIe5hsBhsvy/BeFV+3n8+hQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/TN/wAEc/8AgjFcfGSTwz+1V+1r4ZmtfhFG9rrnwq+EetWzw3PxUZCtxYeLvGdjMqywfDYMEuNG0SdFm8e/Je3iJ4LMMfi7uw2G5rVKi93eMX9rzflr8/Tf+tfArwGlnrwnGXGuElDJE4YjJskxEHGWcte9Tx2PpyV45VtPD4eSTzLSpUSwHKsd/Zpb28FpBBa2sENta20MdvbW1vGkMFvBCgjhgghjCxxQxRqsccaKqIihVAUAV6R/eMYxhGMIRjCEIqMYxSjGMYq0YxirJRSSSS0S0RNQUFABQAUAfxif8F1/+C8K26+Mv2KP2IvGRN0TfeGfjv8AtAeF9Q4th+9s9a+GPws1izfJuSfMsfGPjrT7gC2AufD/AIana4a+1a1/rrwV8E+b6pxjxlhPd9zE5JkWJp/FtOjmWZ0p/Z2nhMFUj7/u18QuXkpT/n3xL8TLfWOHeHMR73vUczzSjP4d41MHgqkftbxxGJg/d1pUXzc84/xZV/Xx/PIUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/wDBN3/lHd+wT/2Zd+y3/wCqO8C1/lt4g/8AJe8b/wDZX8S/+rrGn9z8I/8AJKcMf9k9kv8A6rcMf4vFfIH0IUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAHJOByT07kk0Af1Df8Eov+CUA0L/AIRv9p/9qHw3nXD9l1z4T/CTXbTP9h/cudN8c+OtOuU/5Df+ru/DXhi7j/4kv7nV9Yi/tn7JZaP/AIzfTe+nB/aP9r+DXgxm/wDwnfv8t4546y2v/wAjH4qOM4b4ZxdKX/Iv+OhnGc0J/wDCh+8wGX1P7P8Ab4nMP7u+j99Hz6r9S4748wX+1fu8Xw9w7iqf+67To5rm1Ga/3r4amBwFRf7N7uJxUfrXsqWG/pKr/Ig/toKACgAoAKACgDyr4yfGTwR8DfBOoeOfHOoC2srbNvpumwGN9W8Qas8bva6Po9q7obi8uNjM7MyW9nbJNe3s0FpBNMn6h4Q+EPGnjZxpgOCeCcA8TjcS1XzHMa6qQyvIcrhUhDFZvm+KhCf1fB4fnjGMYxniMXiJ0cHg6VfF16NGf5D43eN3AXgBwHmXH/H+ZLC4DC3w+WZZh3TqZxxHnFSnOeEyTJMJOcPrWOxPJKUpSlDDYLDQr47HV8PgsPXrw/mv+P3x/wDG/wC0L42n8WeLbj7NYWvnWvhfwvazSPpPhnSpJAwtbVWCfaL652RSarqssaXOpXCISsFnb2NjZ/8ARl4D+A3BfgBwZQ4W4VoLE5hiVRxPE3E2Jo04ZrxJmkIOLxOKlFz+r4HDudSnleV06k8Pl2HnO0q+MxGNxuL/AOW36Rv0jePfpKceYjjDjHEPCZbhHXwvCfCWEr1KmT8K5PUqKSwmEjJQWKzDFKFKrnGcVaUMTmeJhBuGHwWGwGAwfhlftx/P4UAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf1Ef8Ecf+CL0nxFPhb9rD9rzwu8Xw9BtNf+EPwY160ZJfHxBS503xx8QNNuUDR+Bv9Xd+HfC90gfxoPK1TWIl8IfZrPxb34bDXtUqLTRxi+vnJPp2XW99t/6/8CvAR5p9T4044wbjlnuYnI8hxMGnmW06OY5nSmrrL/hqYXBzV8f7tauvqPJDG/2KxxxxRpFEiRRRIscccahI440AVERFAVERQFVVACgAAACvRP7oSUUoxSSSSSSsklokktEktkPoGFABQAUAfxZ/8F2P+C8H2j/hMf2KP2IfGX+j/wCn+GPjz8f/AAvf/wDHx9+z1j4ZfCzWbOT/AI9/9dYeM/HFhN/pH+kaB4auPs/2/Vbj+vvBXwT5fqnGHGWE973MTkmR4mG206WY5nRmvi2nhMFUXu+7XxEb8lKP89eJfiZf6xw7w5iNPeo5nmlGe+8amDwVSPTeOIxMXr71Ki7c03/GHX9dH8+BQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/APBN3/lHd+wT/wBmXfst/wDqjvAtf5beIP8AyXvG/wD2V/Ev/q6xp/c/CP8AySnDH/ZPZL/6rcMf4vFfIH0IUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAoBYgAEkkAADJJPQAdSSeg70ebA/qL/AOCUf/BJ8eHP+Eb/AGnv2ofDYPiL/Rdc+FHwl1y1BHh7Oy507xx45064Ug6//q7rw34Zu4/+JD+61bWIv7c+yWeif4xfTe+nB/a39r+DXgxm/wDwlfv8t4546y2v/wAjX4qOL4b4ZxdKX/Ir+OhnGc0Jf8Knv4DL6n9l/WMTmf8AeP0fvo+/U/qXHfHmCvjX7PFcPcO4un/ue06Oa5tQmv8AfNqmBwNRf7H7uJxMfrfs6WE/pEr/ACMP7XCgAoAKACgAoA8p+Mvxl8EfAvwRqHjjxzqH2ezt82+maZbmN9W8Q6s8bva6Po9q7p9ou59jNI7MltZWyTXt7NBaQTSr+o+EHhBxp428aYDgrgrAfWMZiLV8yzKuqkMqyDKoVIQxOb5viYQn7DCUOeMYQjGeIxmInSweDpV8VXpUp/kHjh44cBfR/wCA8x4+4/zL6tgcNfD5VlWHdOpnHEmcTpznhckyTCTnD6zjcRySlOpKUMLgcNCtjsdWw+DoVq0f5rvj78ffG/7Qvje48W+Lbg29jbGa28MeGLWaR9J8MaS8gYWlorBPtF7cbI5NV1WWNLnUrlFZlgtILKytP+jPwI8COC/AHgyhwrwph/rGOxCo4niXiXE0oQzXiXNYQcZYvFyi5/V8HQ56lPLMrp1J4fLsPOSUq+Lr4zG4v/lt+kZ9Izj36SnHuI4y4yxDwuX4V18JwpwphK9SeT8KZPUqKccHg4yUPrOPxPJSq5vm9WlDFZpioQk4YfBYfA4DB+HV+2H4AFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH9Vf/AARv/wCCLjeJP+EV/a0/bB8KFfDmbPxD8Hfgj4hsyG8RYKXWl+PfiTpV0gK+Hv8AVXnhnwdexZ8QDydX8QQDQDa6br3fhsNe1SotN4xfX+9Jduy67vTf+yvArwDeK+p8a8c4L/ZfcxWRcPYqn/vW06OZZtRnqsNtUwmAqRvifdr4qP1bko4n+vQAKAqgKqgAADAAHAAA4AA4AHSvRP7eFoAKACgA9zQB/FJ/wXX/AOC8J1f/AITH9in9iHxkRpH+n+GPjx8fvDF/zq/+ss9Y+GXws1mzkz/ZP+usPGXjiwmB1b/SNA8N3H9k/wBoapq39g+Cvgn7P6pxjxlhP3vuYnJMjxMP4e06WY5nSmv4u1TB4Oa/d+7XxEfa8lKn/PHiX4mc/wBY4d4cxHue9RzPNKMv4m8amDwVSL+DeOIxMH7+tKjLk551P40q/rc/n4KACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/wCCbv8Ayju/YJ/7Mu/Zb/8AVHeBa/y28Qf+S943/wCyv4l/9XWNP7n4R/5JThj/ALJ7Jf8A1W4Y/wAXivkD6EKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgBQCxCqCzMQAACSSTgAAckk8ADkmjzYH9SX/AASj/wCCUA8Lf8I3+07+1B4cDeKCLXXPhT8Jtbtcjwznbcad428cafcLg+JMeXdeHPDd1Hjw9mLVdXi/4SD7JZ6D/i99N76cH9s/2v4NeDOb/wDCN+/y3jjjrLa//I4+Kli+G+GsXSl/yKPjo5xnNCf/AAr+/gMvqf2R9YxGbf3n9H76Pv1H6lx3x5gf9v8A3eK4e4dxdP8A3HadHNc2oTX+/bVMDgai/wBi93E4mP132dLBf0e1/kgf2qFABQAUAFABQB5P8ZvjP4I+BXgi/wDHHji/8i0g3W+l6XbmN9W8Ras8bvbaPo9s7p591PsLSSuUtrK2WW8vZobWGSRf1Pwe8HuNPG/jTA8FcFYD2+Lr2xGZ5niFUhlWQZVGpGGJzfN8TCM/Y4WhzqNOnFTxOMxEqWDwdKtiq1OnL8e8cvHLgL6P3AWY8fcfZj9XweHvhspynDOnPOeJc5nTnPC5JkmFnOH1jGYjklOrVnKGFwOFhWx2OrUMJQq1Y/zXfHz4+eN/2hPG9x4u8XXHkWdv51r4Z8M2s0j6T4Y0l5A62dmrBPPu59kcmqarLGtzqVyis6w2sFlZWn/Rl4E+BPBfgFwXQ4U4Uw/t8biPY4niTiTE0oRzXiXNYU3GWLxcoufsMJQ56lPLMsp1J4fLsPOUYyrYqtjMZiv+W36RX0iuPfpJ8e4njLjPEvDYDDOtheFeFMJXqTybhTJ51FKOCwUZqH1jG4jkp1c3zerShis0xUIylGhg6GBwOD8Qr9rPwIKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP6z/APgjf/wRbJPhX9rT9sTwpx/ofiH4OfA3xHZck/JdaV4++JmlXadP9Xe+GPBN9Fkkwav4kgAFtpLehhsNtUqLzjF/+lSX5L5s/tTwJ8A/9z4146wX/PvFZFw7iqfpOjmebUZr0nhMvqR/lr4uPwUT+tGvQP7UCgAoAKAEJABJIAAJJJwABySSegHc0AfxK/8ABdj/AILwHxN/wmP7FP7EfjLHhr/TvDPx3+Pvhi/58S/6yz1j4afC7WLOT/kXP9bY+MfG1hNnxF+/0Hw5cf8ACPfb9T8Qf2J4LeCf1f6pxhxlhP8AaPcxOS5HiYf7vtOjmOZ0pr/eNp4TBzX7j3a+Ij7fkp0P518SvEz231jh3hzEfuvfo5nmlGX8XeNTB4KpF/wt44jExf73WlRfsuedX+OKv60PwEKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/1t/+Cbv/ACju/YJ/7Mu/Zb/9Ud4Fr/LbxB/5L3jf/sr+Jf8A1dY0/ufhH/klOGP+yeyX/wBVuGP8XivkD6EKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAFVWZgqgszEKqqCWZicAADJJJOAByTQ3a7bslq2/zYbn9TH/BKT/glAPCA8N/tO/tP+HA3i4i11v4VfCjW7UMvhQEJcad418b6fcIQ3ij7l14e8OXSY8NfutU1WI+I/stp4e/xa+m99N956838GvBnN/+EP8Af5bxxxzltfXO371LF8OcN4ulL/kS/HQzfOKE/wDhZ9/A4Cp/Y/t8RnH96/R++j7/AGd9S4747wP/AAofu8Vw9w9iof8AIv2nRzXNaM1/v/w1MFgai/2D3cRiY/X/AGdLA/0b1/kof2mFABQAUAFABQB5N8aPjR4I+BPgi/8AG/ji/wDJtod1vpWk27Rvq/iLVmjZ7bSNHtndPOuZtpeaZyttY2yy3l5LDbRSSD9U8HfB3jTxv40wPBXBWB9tiq1sRmmaYhVI5Vw/lUakYYjNs2xMIz9jhqPMo0qUVLE43ESpYTCUq2Iqwg/xzxz8c+Afo+cBZhx7x9mPsMJQ5sNk+T4aVOec8TZzKnOeFyXJcLOcPb4uvyudatOUMJgMLGtjsdWoYWjUqL+a749/Hvxv+0H43ufF/i+58i0g8628NeGrWaRtJ8MaS8gZbKyRwvnXU+yOTVNUkjW51K5RXcQ20NnZ2v8A0Z+BfgVwX4B8F4fhPhPD+2xdf2WJ4j4kxNKnHNeJc1hBxljMZKLl7HC0eepTyzLKdSeHy7DzlCEq2JrYvF4r/lt+kR9Ijj36SXHuJ4040xX1fB4f22E4W4WwlapPJuFMmnUU44HAwmoe3xlfkp1c2zarThis0xUIznGhhKGCwWD8Rr9pPwQKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAUAsQqgszEAAAkkk4AAHJJPQdSaA3P68f+CN//AARcHh7/AIRX9rT9sLwoG8Q/6H4h+DvwP8RWQK+H/uXWlePviVpV0hDeIP8AVXvhjwbexY0A+Tq/iGE699l0zQfRw2GtapUWu8Yvp5yT69l0332/t7wK8BFhfqfGvHODviv3eKyLh7FQ0wu06OZZtRmtcTtUwmAqK2G92viovE8lHDf1XV3n9lBQAUAFACMwUFmIVVBZmYgAADJJJ4AA5JPAHJo3A/iG/wCC6/8AwXfbxofGP7FX7EvjEr4MBv8Awz8dfj14ZviG8Ykb7TV/ht8MdYtJMjwhnzrHxh4zsZc+Lf3+heH5/wDhFPt2o+Kf7I8FfBP6p9U4w4xwl8Z7mJyXJMTD/dNp0swzGlNf71tPC4Sa/wBl92tXj9Z5KeH/AJz8SvEz6x9Y4e4cxH+z+9RzPNKMv4+8amEwVSL/AIG8a+Ii/wB/rTpP2PNOt/HjX9YH4GFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+tv/wTd/5R3fsE/wDZl37Lf/qjvAtf5beIP/Je8b/9lfxL/wCrrGn9z8I/8kpwx/2T2S/+q3DH+LxXyB9CFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAOVWdlRFZ3dgqqoLMzMcKqqMlmYkAAZJJwOaTaSbbSSTbbdkktW23sl1Yat2Wre3Vts/qe/wCCUn/BKFPBS+HP2nP2nfDgfxowtdb+Ffwp1u1DJ4PUhbjT/GfjXT7hCG8WkFLrw/4euUx4W/danqkZ8TfZrXw5/iv9N76cD4geb+Dfgzm7WQJ18u4345y2u08+fvUsXw5w3i6Ur/2F8dHNs3oT/wCFz38DgZ/2J7fEZ1/e/wBH76Pv9mfUuOuO8FfMv3eK4f4exVPTLtp0c0zWjNa5jtUwWCqK2X+7iMQnj3Tp4D+jCv8AJg/tEKACgAoAKACgDyT41fGrwP8AAjwRfeN/G9/5VvFut9I0i2aN9X8R6u0bPb6RpFs7L5txLtLzzuVtbG2WS8vJYreJ3r9W8HPBzjTxx40wPBfBeB9tiKvLiM1zXERqRyrh7Ko1IwxGbZtiYRl7LD0ublo0YqWJxuJlTwmEpVcRVhA/G/HXx14C+j3wFmHHvHuYexw1Hmw2TZNhZU55zxNnMqc54bJslws5x9tia3K51683HC5fhY1cbjatHDUZzP5rvjx8efHH7QXje58YeMLnyraLzbbw34btZZG0jwxpLSBksbBHC+bcS7Y5NT1OWNbrU7lRJIIreK0tLX/oz8DPAzgrwD4Lw/CXCWG9tia3ssTxFxFiqVOObcS5rGDjPG42ceb2WHpc06eW5bTnLDZdhpOEHVxFXF4vFf8ALZ9If6Q/Hv0kePcVxrxrivYYWh7bCcL8L4StUnk3CmTTqKcMBl8J8vt8VX5KdXNc1q044rNMVGNSoqOFo4LBYTxOv2c/BgoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgB8ccksiRRI8ssrrHHHGrPJJI7BUREUFnd2IVVUFmYgAEmgaTk1GKcpSaSSTbbbskktW29lu2f2Jf8EcP+CLyfD0eFv2sf2vfCyy/EAiz1/4Q/BfX7VXi8BghLnTPHPxB0y5QrJ43/1d34c8LXaFPBh8rVNahbxeLWy8JejhsNa1Sotd4xfTrzPz7Lpvvt/dHgV4CLK/qfGnG+ETzP3MTkeQYmF1lu06OY5nSle+YbTwuDmrYDStiIvHctPBf1F13n9fhQAUAFADXdUVndlREUu7uQqqqglmZiQFVQCSSQAMkmjVuy1b26tth5s/h4/4Lrf8F3n+IbeMP2LP2J/GDJ8PlN94a+Ofx48NXxV/HxBktNW+HHw01e0kyvgXPm2Xi3xjZShvG37/AETQZ/8AhDPtuoeMP7L8FvBNYH6pxhxjhL458mJyXJMRDTBbTpZhmNKS1xu08LhJq2E92tXX1vkhhP5x8SvEx4r6xw7w7X/2X3qOZZnRlrit41MHg5xemG3jiMRF/wC060qT+r888R/IBX9Wn4KFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/wDBN3/lHd+wT/2Zd+y3/wCqO8C1/lt4g/8AJe8b/wDZX8S/+rrGn9z8I/8AJKcMf9k9kv8A6rcMf4vFfIH0IUAFABQAUAFABQAUAFABQAUAFABQAUAFADkR5HWONWkkkZUREUs7uxCqqqoLMzMQFUAkkgAEmk2opyk0kk223ZJLVtt6JJatsaTbSSbbdklq230Xds/ql/4JS/8ABKFPAa+Hf2mv2nfDqyeOXFtrXwt+FWtWoZPBSsFnsPGXjOwnUhvGDAx3Og6BcpjwoPL1LUoz4n+z23hr/FP6bv04HxI838G/BrNnHhxOvl3G3HGXV7PiJrmpYvh3hzFUndZAnzUc1zajO+fe/gsDNZH7evnf98/R++j7/ZX1LjrjvBJ5o/Z4rh/h/FQv/Zm06OaZpRmrPMtqmCwU1/wne7iMRH+0fZ08v/opr/Jw/s4KACgAoAKACgDyL41/GvwR8B/BF9428b33lwx7rbR9HtmjbV/EertGz2+k6Tbuy+ZNJtL3Fw5W2sLZZLu7ljhjJP6v4NeDXGvjlxpguC+C8F7WvU5cRm+b4mNSOU8PZVGpGGIzXNsRCMvZ0KfMoUKEFLE47Eyp4TCU6lepGJ+M+O3jtwD9HrgLH8e8e5h7LD0ubC5LkuFlTlnXE+cypynhsmybDTnH2uIq8rniMRNxwmX4WNXG42rSw9KUj+a747/Hfxv+0D43uvGPjG58qCLzbXw54ctZZG0jwxpDSb0sLBH2+bPLtSTUtSkRbrU7pRLKIoIrW0tf+jPwN8DeC/ATgvDcJcI4b2uIq+yxPEXEWJpQjm3Euaxp8s8djpx5vZUKXNOnl2XU5yw2XYeTp03Ur1cVisT/AMtn0hfpDcffSQ49xXGvG2L9lhqPtsJwxwxhK1SWTcKZNKopwy/L6c+X2uJrctOrmmaVYRxWaYqKq1VSw9LB4TC+K1+zH4QFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUATW9vPdzw2trDNc3NzNHb29vbxvNPcTzOI4YYYYw0ks0sjKkcaKzu7BVBYgEKjGU5RhCMpznJRjGKcpSlJ2jGMVdylJtJJXbbstT+zL/gjn/wRig+DyeGP2q/2tfDUV18W5Etdc+FPwi1q2Sa2+FquFuNP8X+NLGZWjuPiOVKXOi6JOrReAspe3qP40EMXhH08NhuW1Sovf3jF/Z83/e/Lf4tv7x8C/AWORLCcZca4RTztqGJybJMRFSjk17SpY3H05JqWa7ToYeV1lulSonj+VYL+miu0/rUKACgAoAZJIkSPLK6Rxxo0kkkjBEjRAWd3diFVFUFmZiAoBJIFNJtpJNtuyS1bb6Lu2DdrtuyWrb/Nn8NX/BdX/gu7J8UX8X/sXfsU+L3j+GSNe+G/jh8dPDl6ySfEhlL2uq/Dz4carauGT4egiWz8VeLLOQN49PnaPosw8EfbL3xp/Z3gr4Kf2b9U4v4wwt8xfJiMmyXEQ/5F+0qWPzClJa4/aeGws1/sXu1qy+ucsMH/ADd4leJjxv1jh7h2u1g7yo5lmdGVni+k8Jg5xd/qu8a9eL/2rWnTf1bmlif5D6/qk/CAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/AOCbv/KO79gn/sy79lv/ANUd4Fr/AC28Qf8AkveN/wDsr+Jf/V1jT+5+Ef8AklOGP+yeyX/1W4Y/xeK+QPoQoAKACgAoAKACgAoAKACgAoAKACgAoAciPK6Rxo0kkjKkcaKXd3chVRFUFmZmICqASxIABJpNqKcpNRjFNyk3ZJLVtt6JJatvbdjSbaSTbbsktW29klu22f1V/wDBKX/glEnw9Tw7+0z+034dWTx/IttrPwu+Fms2wdPAyMFnsfGHjGxnUq/jRlKXGh6FcIV8Igx6hqMZ8UmC38M/4n/Td+m/Lih5t4OeDebOPDKdbLuNeN8urNS4kabpYrh/h3FUmmuHU+ejmma0ZX4g9/B4KayL21bPP77+j/8AR9WULBcdcdYK+bP2eKyDh/FQusqWk6OZ5pRmtc0elTB4Oa/4Tfdr4hf2jyU8v/ojr/KA/swKACgAoAKACgDyH42fGzwR8BvBF7428bX2yJN9to2jWzxtq/iTVzGzwaVpNu7DzJXxvubl8WthbCS7u5EiT5v1nwZ8GeNfHPjTBcGcGYL2lWfLiM4zjExqLKeHcpVSMK+a5riIRfJShfkw+HhzYnHYlwwuFp1Ks9Pxfx48eOAfo88BY/jzjzH+zo0+fC5HkeFnTlnXFGcypynh8nyfDTlH2lapb2mKxNTlwmXYVVMZjKtOjTvL+a747fHbxx+0B43uvGPjK68uGPzbXw74dtZZDpHhnSWk3pp+nxvjzJpNqSajqMqC61K5USy7IY7a2tv+jPwP8D+C/AXgvDcI8IYb2lap7PE8QcQYqnTWbcS5sqfJPH4+pC/s6NPmnTy7L6c3hsuw0nSpe0rVMTicT/y2fSD+kJx99I/j3F8b8b4v2dCl7bCcM8M4SrUlk3CmTSq89PLstpz5faV6vLCrmeZ1YRxeaYqKrVvZ0KWEwuF8Wr9kPwoKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAs2VlealeWmn6faXN/qF/cwWVjY2UEt1eXt5dSrBbWlpbQLJPc3NxPIkMEEKPLNK6xxqzsATcunTqVakKVKE6tWrONOnTpxlOpUqTkowhCEU5TnOTUYxinKUmkk2z+1D/gjv8A8EZ7P4AweHP2ov2qvD1tqHxzuIrfWPht8MdUhiurH4OxTIJbTxD4it3EkF38T3Rlks7VhJb+A8h18zxXibw96eGw3JapUXv7xi/s+b/vfl67f314GeA8OGo4XjDjLCwq8RTjGvlWUVoqdPIoyXNDFYqLvGecNNOEHeOW3ur433sL/SPXaf1WFABQAUARyyxQRSTzyRwwwxvLNNK6xxRRRqXkkkkchUjRQWd2IVVBZiACaaTk0km22kkldtvZJbtt7LqDaSbbslq29kurbP4XP+C6n/BdyX4vSeL/ANjH9izxdJD8JY3vfDnxs+OHh28aOb4pujPa6p4A+Hup2zh4/hqrCS08S+J7SQP8QiJdK0mUeBPtV342/tHwW8FFlSwnF/F+FTzRqGIyfJsRC6y29pUsdj6clrmFrTw+GkrYG6q1V9d5Y4T+bPErxLeOeI4e4ertYH3qOY5lSlrjd1UwuFmtfqm8a1eL/wBq1pwf1a8sR/I9X9TH4UFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+tv/wAE3f8AlHd+wT/2Zd+y3/6o7wLX+W3iD/yXvG//AGV/Ev8A6usaf3Pwj/ySnDH/AGT2S/8Aqtwx/i8V8gfQhQAUAFABQAUAFABQAUAFABQAUAFAD445JpEiiR5ZZXWOOONWeSSR2CoiIoLO7sQqqoLMxAAJNKUowjKc5KMYpylKTSjGKV3KTeiSV223ZLVjScmkk220kkrtt7JLdtvZdT+rH/glN/wSij+HCeHf2mP2mvDyS/ESRbbWfhh8LtYtg8fgFGCz2Pi7xfYzqVfxuwKXGi6JOhXwePLv79D4qMMHhn/Ev6bn035cVvNvB3wbzZx4WTrZdxrxtl9ZqXE7TdPFZBw/iqbTjw4nzUc0zSlK/EPv4TCTWQ+2q55/fv0f/o/LJlguOeOsEnnDUMTkOQYqndZRe06WZZnRmtc12nhMJNf8Jmleuv7R9nDL/wChyv8AKQ/soKACgAoAKACgDyD42/G7wR8BfBF7418bXuyNd9toui2zxnV/EmrmMvBpWlQOw3yNjfdXT4tbC2D3V1IkagN+teDHgvxr46caYPg3gzBc9SfJiM5znEwqLKeHcqVRQr5pmleEXyU43cMNhoXxOPxLhhcLCdSTcfxbx58eeAfo8cBY7jvjzH8lOHPhciyLCzpvOuKM5dOU8PlGT4eclz1Z29pi8XUthMuwqqYvF1IU4JT/AJrfjr8dfG/x/wDG914y8ZXeyNPNtvD3h62lkOkeGdIMm+PTtOjfG+R8JJqGoSILrUrkedNsjS3t7f8A6NPBDwQ4L8BuC8LwhwhheerP2eJz/P8AE06azbiTNlT5KmYZhUjfkpwvOngMBTm8Nl2GfsqXPVniMRiP+Wv6QX0guPvpHce4vjfjfF8lKn7XCcNcNYSrUeTcKZNKpz08ty2lO3tKtS0KuZZlVgsXmmLXtq3JSp4bDYbxiv2I/DAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgC7pum6lrWpafo+j6fe6tq+rXtrpul6VptrPf6lqepX88drY6fp9jaxy3V5e3l1LFb2trbxST3E8kcMMbyOqk1bstW9urbZpSpVcRVpUKFKpXr16kKVGjShKpVrVaklCnSpU4KU6lSpOSjCEU5Tk1GKbdj+2z/gj5/wRu079me00D9pf9pzRLLV/wBoa+tYtS8DeA71Le/0v4JW11EHiv7wfvbbUPifLC+JryNpLTwervZ6W82rifUofUw2G5LVJr33sv5f/tvy9T/QLwN8CaXCUMNxZxdh6dfiipBVcuy6oo1aPD8Jq6qVPihVzdxfvVE5QwN3Ci5V+asv6J67D+ogoAKACgCKeeG2hmubmaK3t7eKSeeeeRYoYYYlMks00shVI4o0VnkkdgqKCzEAE00pSkoxTlKTSjFJuUpN2SSWrbeiS1bE2km20kk223ZJLVtt7Jbts/hO/wCC6f8AwXam+Nkvi39jT9i/xbLb/BqGS88PfGj42eH7x4p/i5JGz22peBfAWpW7rJF8LkcSWviHxFayK/xGIk07TZR4C+0T+N/7U8FvBRZOsLxdxfhVLNmoV8nyevBNZWmlKnjcdTlo8yaalQw8lbAaVKi+u8scH/NfiV4lvMHX4f4ertZeualmOY0pWePeqnhsLNarBrWNWtF3xesIP6rzPE/yW1/UZ+GBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62//BN3/lHd+wT/ANmXfst/+qO8C1/lt4g/8l7xv/2V/Ev/AKusaf3Pwj/ySnDH/ZPZL/6rcMf4vFfIH0IUAFABQAUAFABQAUAFABQAUAFAD445JpI4YY3lmldY4oo1aSSWSRgqRxooLO7sQqqoLMxAAJNTKUYRlOcowhCLlOcmoxjGKblKUm0lFJNtt2Su2xpOTUYpylJpJJNttuySS1bb2W7Z/Vv/AMEpv+CUcfwyj8PftLftM+Hkm+JMqW2sfDL4YaxbK8fw9Rws9l4s8W2UylZPHLqUn0fRp0K+DgUvb1G8VGGHwz/iR9Nz6b8uL5Zt4PeDmbShwlGVbL+M+NMvrOMuKmm6eJyHIMTTaceGk+almWZUpJ8RNSwuFksh9rVz3/QD6P8A9H5ZIsFxzxzg1LOpKnishyHEwusnTtOjmWZUpqzzZ6TwuFmv+EvStWX9pcsMv/oWr/Kg/scKACgAoAKACgDx743/ABv8EfATwReeNPGt7tUb7bRNDtnjOr+JdXMZeHS9Lgcjcx4e7u3AtdPtt9zdOqhVk/XPBbwW418deNMHwbwbg+aT5MTnWdYmFRZTw5lXtFCtmeZ14J8sVrDC4WDeJx+JcMNhoSnKUofinj3498BfR34Cx3HXHeO5YR9phchyHCzpvOuKc5dOU6GU5RQnJc05WVTGYyolhMtwvPisXOMIxjU/mu+Ofx08b/H7xvd+MvGd3tRfMtvD/h+2kkOkeGdIMm+LTtOifG52wsl/fyKLrUbkGecqiwQQf9Gfgj4I8FeA/BeF4Q4PwvNOXs8Tn+fYmnTWbcSZsqfJUzDMasb8sI3lTwOBpyeGy/DP2NFSnKvXr/8ALZ9IH6QPH30jePcZxxxxjOWnH2uF4c4cwtWo8m4VyZ1OellmWUptc1SdoVMxzGpFYvM8WnXruMI4fD4fxmv2E/DQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgDT0XRdY8SaxpXh7w9pWo67r+u6jZaRouiaPZXOpatq+rajcR2mn6ZpmnWcc13fX99dzRW1pZ2sUtxc3EscMMbyOqk1bstW9urbZth8PXxdejhcLRq4nE4irToYfD0Kc6tavWqyUKdKlSgpTqVKk5KEIQi5Sk0optn9wn/AASC/wCCOmj/ALKOm6L+0T+0bpGna9+0tqtkt34Z8L3H2bUtG+BthfQEGG2kQzWepfEm6tpWh1nX4HmtfD0ck2ieHJ5M6lrOrerh8Mqdpz1m9luo/wD23nfTp5/6E+B/gXQ4MpYfijimjSxPFlamqmEwkuWrQ4dp1Y6xg/ehVzWcZOOIxMW4YZOWHwsn+9xFf+gmus/pkKACgAoAhuLi3tLee7u54bW1tYZbi5ubiVIbe3t4UaWaeeaVljihijVpJZZGVI0VndgoJpxjKcoxjFylJqMYxTlKUpOyjFK7bbdkldtvQTainKTSSTbbdkktW23oklq2z+EH/guj/wAF2Lj48TeLP2N/2M/Fc1r8D4JbvQPjH8ZtBunhuPjJLEz2+o+CvBOoQMskPwpR1e31vXLd1k+JDK9lZuvgPzZPGf8Aa/gt4KrJVheLeLsKpZxJQr5RlFeN1lKa5qeMxsJKzzNpqVGjJNZfpOaeNssJ/NHiV4lyzN1+H+H67jlycqWYZhSk1LMGm1PD4aad1gb6Vai1xesYv6rf6x/J3X9QH4eFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62/8AwTd/5R3fsE/9mXfst/8AqjvAtf5beIP/ACXvG/8A2V/Ev/q6xp/c/CP/ACSnDH/ZPZL/AOq3DH+LxXyB9CFABQAUAFABQAUAFABQAUAFAEkUUs8scEEck000iRQwxI0ksssjBI4440DPJJI7BURQWZiFUEmpnOFOEqlSUYQhGU5znJRhCEU5SlKUmlGMUm5SbSSTbdhpSlJRinKUmlGKTcpSbskktW29Elq2f1ef8Ep/+CUkXwtj8PftK/tL+H45vibKlvrHw0+GWr26yRfDpHCzWXinxVZTKVk8eOpSfSdInUr4NBW7u0Pivyo/DX+In03PpvT4zlmvg/4PZrKHB8JVsBxlxnl9ZxnxZJN08TkeRYmm048MJ81LMMxpSvxE+bDYaSyH2lTPP9Avo/8A0flkKwfHHHODUs8koYnIshxMLrJU/fpZhmNKV083atPDYWatld1Vqr+0uWOX/wBCFf5WH9ihQAUAFABQAUAeO/HD44eCPgH4IvPGnjS9wP3ltoWhWzx/2v4l1fyy8OmaZC55J4e8vJB9l0+23XNy4GxJP13wV8FeNfHXjTCcHcG4O7fs8Rned4iFT+yeHMp9ooVszzOtFaJawwmEg/rWYYnlw+Gi25zp/iXj54+cA/R24CxvHPHWOsv3mF4fyDC1KbzrinOfZOpQynKaE3q3pPG42olhMtwvNisVNR5IVf5rvjl8cvHHx98b3fjPxneYA8y20HQbaST+yPDWkGQvFpumxOeWPyyX19Iv2rUbkGe4YAQxQ/8ARn4J+CXBfgRwXhOD+D8JeT9nic9z3E06f9rcSZsqfLVzHMasFpFXlDBYKnJ4bL8M1QoJydWrW/5bPH/x/wCPvpGce4zjjjnG2ivaYXh3h3CVKv8AYvCuTOo50sryujN6ylaNTMMwqR+t5nik8RiJKKo0aHjVfsB+HhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAGz4d8O6/wCL9f0bwr4V0XVPEfibxHqljougaBoljc6nrGtaxqVxHaafpmmadZxzXV7fXt1LHb21rbxSTTTSKiKzMBTSbdkm29ktW2dGFwuJx2JoYPBYeti8Xiq1PD4bDYenOtXxFerNQpUaNKmpTqVKk5KMIRTlKTSSbZ/c7/wSK/4I+6B+x5o+lfHr4+abpfiX9qHW9O83S9MY22qaJ8D9N1G3KT6Rok6Ga01Hx7d2sz2vibxXbPLb6fFJceHfC87ac2raz4l9TD4ZU/fnrUfndR9PN9Xr5db/AOiHgj4G4bgWhR4k4lpUcXxfiKXNRpPlrYfh6lVjaVDDyXNCrmU4ScMXjYNxpJywuDk6TrYjF/vVXWf0gFABQAUAQXV1a2Nrc3t7cwWdnZwTXV3d3U0dva2trbxtNcXNzcTMkUEEESPLNNK6xxxqzuyqpNVGMpyjCEZTnOSjGMU5SlKTtGMYq7lKTaSSu23ZailJRTlJqMYpylKTSSSV223oklq29EtWfwZ/8F0P+C611+0LceKv2O/2OPFU9n8BLaa60L4vfGDQ7mS3u/jXPC7QX/hHwhfQsktv8JopFeHVdVhZZfiO6tBCy+Bw3/CXf214L+CscijhuLOLcKp53JRr5VlNaKlHJ4tc0MXi4O6lmjWtKk7rAJ80k8bb6r/M3iT4lvNXXyDh+u45YnKlj8wpu0sxadpUMPJO6wN7qpUVni9Uv9m/j/ylV/Th+IhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/APBN3/lHd+wT/wBmXfst/wDqjvAtf5beIP8AyXvG/wD2V/Ev/q6xp/c/CP8AySnDH/ZPZL/6rcMf4vFfIH0IUAFABQAUAFABQAUAFABQBJDDNcTRW9vFJPPPIkMMMKNLNNNKwSOKKNAzySSOwREQFnYhVBJAqZzhThOpUnGnTpxlOpUnJRhCEU5SnOUmoxjGKcpSk0kk23YcYynKMYxcpSajGMU5SlKTsoxSu223ZJXbb0P6wv8AglP/AMEpIvhPF4f/AGlP2lvD8c/xSmjt9X+G3w01e3WWL4bRyKs1n4n8UWcysknj91KzaXpcysvgtStzcqfFexPDX+H/ANNv6btTjaea+EHg/msqfBkJVcBxhxjgKrjU4unFuniMkyTEU2pR4XTUqePx9OSlxG+ahQksh5551/oL4AfR+XD6wfG/HGDUs+koYnI8ixEVKOSJrmpY/MKUk1LN2mp4fDTTWV3VSqv7S5Y5f/QVX+WJ/YQUAFABQAUAFAHjfxy+OXgj4BeCLvxn40vOT5ltoOg20kf9r+JdX8svFpumwueAPlkvb6QfZdOtiZ7hsmKKb9f8FPBPjXx340wnB3B2E0/d4nPM8xMKn9k8OZV7Tkq5lmVaC3fvQwWDhL6zmGJtQoL+LVpfiHj94/cBfR14CxvHPHON1/eYXh7h7C1Kf9tcU5z7Nzo5VlVGb2+Gpj8fUj9UyzC82JxMr+ypVv5rvjj8cfG/x88b3njTxpef89LbQtCtpJP7I8NaR5heHTNMhc/SS9vZB9q1G53XFw3+rji/6M/BTwU4K8CeC8JwdwdhP5MTnmeYmFP+1uI829moVcyzKtBf4oYPBwf1bL8Ny4fDx/iVKv8Ay2ePvj9x99Irj3G8c8dY3/n5heHuHsLUqf2Nwtk3tXOjlWVUZv8Aw1Mdjqi+t5niubE4mX8KlR8cr9ePxEKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA6Hwn4T8T+PPE2g+C/Begat4p8W+KNVstD8O+HNCsbjU9Z1rWNRnS2sdO02wtUkuLq7up5EjiiiRmJOTgAkNJtpJNtuyS3bOrBYLGZli8NgMBhq+MxuMrU8PhcLhqcq1fEV6slCnSpU4JynOcmkkkz+7P8A4JJf8Eh/DH7FWgaf8afjVZaT4r/am8Q6YcMDb6nofwX0vUrcpdeGvCtwvm2174turaV7TxX4xtmZPLe48O+Gphojapqnij1cPh1SXNLWo/ujfovPu/ku7/0X8FPBDCcAYaln2f06GN4xxVLf3a2HyGjVjaeEwU9YVMbOLcMbj4XTTlhcJL6v7atjP3MrqP6HCgAoAKAK15eWenWd3qGoXdtY2FjbT3l9fXk8VrZ2dnaxPPc3d3czukNvbW8KPNPPM6RRRI8kjqikioQnUnGnTjKdScowhCEXKc5ydoxjFXlKUm0oxSbbdldsUpRhGU5yUYxTlKUmlGMUruUm9EkrttuyWrP4J/8Agub/AMF1bz9pG68U/sgfseeKLmw/Z7s57nRPix8WtGuJbW9+OVzbyNDeeGPC93EyT23wjhlRor+9QpN8R3UrlPBYCeKP7c8F/BaPD8cNxXxZhozz2SjWyvK60VKGTRa5oYnEwd1LNGnenB3WAWuuLd8N/MniT4lSzd1shyCu45Um6eOx1NtSzJp2lRoyWqwKek5Kzxeq/wB3/jfysV/TR+JhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62//AATd/wCUd37BP/Zl37Lf/qjvAtf5beIP/Je8b/8AZX8S/wDq6xp/c/CP/JKcMf8AZPZL/wCq3DH+LxXyB9CFABQAUAFABQAUAFABQBLBBNczQ21tDLcXFxLHBBBBG8s080rhIoYYkDPJLI7KkcaKzu7BVBJAqKlSnSpzq1Zwp0qcJVKlSpJQp06cE5TnOcmoxhGKcpSk0opNt2RUYynKMIRlOc5KMYxTlKUpO0YxiruUpNpJK7bdlqf1j/8ABKj/AIJSw/CKHQP2kv2lNAjuPivPHb6t8OPhvq0Cyw/DOKRRLaeJPEtpKGSX4gSIyzadp0qlPBalZ51PiraPDn+Hn02vpu1OOp5p4Q+EGaTpcE051cDxdxfgasoVOMZwbp4jJ8mxFNqUOFVJSp4zGwalxG06NJ/2FzSzn/QbwA+j/Hh2OD4243wcZ8QSUMRkmSYiClHI00pU8djqck1LOGrSoUJJrK9Kk08wssD/AECV/lqf2AFABQAUAFABQB418c/jn4I+APgi78Z+M7vc7eZbaBoFtJGNX8TauI98Wm6bE5OFXKyX9/IpttOtiZ5yWaGGf9g8EvBLjXx440wvB/B+EtFezxOfZ9iadR5Tw3lTqctXMcxqwWspWlDA4GnJYnMMSlQoJRVatR/D/H/6QHAX0c+AsZxxxzjOaUvaYXh3h3CVKf8AbXFWcqm50sryujN6RjeNTMMwqR+qZZhW8RiHKcqFCv8AzW/G/wCN/jj49+N7zxp41vcsd9toeh2zyDSPDek+YXh0vS4HJwo4e7u5AbrULndc3LliiR/9Gfgr4K8FeBXBeE4O4Nwdl7mJzvO8TCm824jzb2ahWzPM60Fq3rDCYSD+rYDDcuHw8UlOdT/lr8fPHzj76RPHuO4646x15P2mFyDIMLOqsm4Wyb2jnQynKaE5O0VpUxuNqJ4vMsXz4rFTcnCFPx6v10/EwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgDqfBPgjxf8SfF3h3wF4B8Oav4u8Z+LdWtND8NeGtBsptQ1fWdWvpRFbWVlaQK0kkjsdzsdsUMSyTzvHDHJIrScmkk23slq2dmX5fjs2x2Fy3LcLXx2PxtaGHwmEw1OVWvXrVHaFOnCN223q3tGKcpNRTa/vH/wCCTf8AwSR8I/sN+GbP4r/Fa30nxh+1P4n0ox6jqsfk6honwl0vUYMXXg/wTcFWjn1ieF2tfFfjGHEmogzaLojw6B9suNf9bD4dUlzS1qPd7qPkvPu/ktL3/wBHfBfwUwPh5hIZznMaGO4xxlG1WuuWrh8ko1Y+/gcvk7qVeUXyY3Hx1q+9h8O1hueeJ/bCuk/fwoAKACgCpf39jpdje6nqd7aadpunWlxf6hqF/cQ2djYWNnC9xd3t7d3DxwWtpawRyT3FxPIkMEKPLK6orMKhCdWcKdOEqlSpKMKdOEXOc5zajGEIxTlKUpNKMUm5NpJNsmUowjKc5RhCEXKc5NRjGMU3KUpNpKKSbbbsldtn8Df/AAXL/wCC6V9+05eeJ/2RP2QfE13pv7OVjc3Gj/FH4paTPNZ3/wAd7u1laK60Dw/coY7i0+EUEyMs8oMc/wAQ5EEkwTwmIrfXv7g8GPBeHDkcNxVxXh41M/nGNXLctqxU4ZLGSvGvXi7xnmkou8VrHAp2V8TeVH+YvEnxJlnMq2RZDWlDKYt08bjabcZZnJO0qVJ6NYBPd6PFvf8A2eyq/wAs1f0sfiwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/wDBN3/lHd+wT/2Zd+y3/wCqO8C1/lt4g/8AJe8b/wDZX8S/+rrGn9z8I/8AJKcMf9k9kv8A6rcMf4vFfIH0IUAFABQAUAFABQAUASwQT3U8Nrawy3NzcyxwW9vBG809xPM4jihhijDSSyyyMqRxorO7sFUFiBUVatOjTqVq1SFKlShOrVq1ZxhTp04Rc51Kk5tRhCEU5TnJqMYpttJNlQhOpOMIRlOc5KEIQTlOc5O0YxiruUpNpJJNtuyuz+s3/glT/wAEpoPg7DoH7SH7SWgxXPxbuIoNV+Hfw51WBJoPhhDKqy2viLxFayho5fiFIjCSxsJFZPBasssoPigj/hH/APDf6bX03KvH1TNPCPwhzSdHganOrgeLeLcFUlCrxnUhJ06+UZRXg1KHCkZJwxeLg1LiN3p02sjTecf6E+AH0f48NxwfG3G+DjPiGShiMlyXERUo5FGS5qeNxtOV4yzlpqVGjK6ytNTl/wAKFvqX7/V/l0f16FABQAUAFABQB4x8dPjr4I+AHgi68ZeMrvfI/mWvh/w/ayRjV/E2riPfHp2nRvnZGmUk1DUJFNtp1sfOmLyPb29x+xeCPgjxr488aYXhDg/C8tOPs8Tn+f4mnUeU8N5U6nLUzDMKkLc1Sdp08BgKclicxxK9jR5YRr16H4b9IH6QPAP0cuAsZxxxxjOepP2uF4b4bwlWms64qzlU+ellmWUp35KcLwqZjmVWDwmWYV+3rudSeHw+I/mu+N3xu8b/AB68b3vjXxre7nbfbaJols8g0jw3pAkZ4NL0qB2O1Fzvurp83WoXRe6unZ2AT/oz8F/BbgrwL4LwfBvBuD5YLkxOdZ1iY03m3EebOmoVszzSvBLmnKzhhcLBrDYDDcmGw0IwjKU/+Wzx68euPvpEce47jvjvHc05c+FyHIcLOosm4XyZVZToZRlGHnJ8sI3VTGYypfF5jinPF4upKcoxh4/X62figUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB2Xw9+Hvjf4seNvDPw4+G/hjV/Gfjnxlq1tonhnwzoVq15qmrandMRHDBEuEjijRZLm8vLiSGysLKG4vr64t7O3nnjaTk0optvZLdnfleV5hnWYYTKsqwlfH5jj68MPhMJh4OdavWntGK2SSTnUqTcadKnGdWpONOEpL+9n/AIJTf8EnfBP7B/hGD4h/EGLSPGv7UfizSRF4j8UxIt5pHw402+jVrrwR4AmmjDDKkQeJfFSpFeeIJUa1tRaaGq21z62Hw6pLmlrUe7/l8l+r6n+kXgz4L5f4cYGOaZmqGYcYY2hbF4xJVKGVUqivPL8tlJX/ALuLxllUxUk4Q5MOlCf7K10n7uFABQAUAU9R1Gw0iwvtW1a+s9M0vTLO51HUtS1G5hsrDT7Cyhe5vL6+vLl47e0s7S3jkuLm5uJI4YIY3lldUVmF06dSrUhSpQnUq1Jxp06dOLnUqVJtRhCEIpynOcmoxjFNybSSbZM5xpxlOcowhCMpznOSjGMYpuUpSbSjGKTcpNpJJts/gO/4Lk/8Fz9Q/anvvEv7JP7I/iS90r9mnTrubSviT8S9LmnstR+Pl7ZzFJ9I0iZTFc2Xwit54/ljPl3Hj6VFu75Y/DYtrLUv7i8GfBenw1DD8U8VYeFXiGpCNXL8uqpTp5HCauqtVO8Z5rKL31jgU3GDeIcp0/5g8SPEmedyrZFkVWUMohJwxmMg3GeaSi9adN6OOATXrimuaVqNoz/l0r+kz8YCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/AOCbv/KO79gn/sy79lv/ANUd4Fr/AC28Qf8AkveN/wDsr+Jf/V1jT+5+Ef8AklOGP+yeyX/1W4Y/xeK+QPoQoAKACgAoAKACgCa3t7i7uILS0gmurq6mit7a2t4nmuLi4mdY4YIIY1aSaaaRljiijVnkdlVVLEA51atKhSqV69SnRo0ac6tatVnGnSpUqcXOpUqVJtRhThFOU5yajGKcpNJNlQhOpONOnGVSpUlGEIQi5TnOTUYxjGKcpSlJpRik220km2f1pf8ABKn/AIJT2/wYg0H9o/8AaQ0KG6+L9zFBqnw9+HeqQJNb/C2GZBLa6/r9tIGjm+IckbCSzs5AyeC1YOwbxOSdB/w1+m19Nyr4g1c08JPCPM6lHgOlOpguK+LMFUlTq8aVIScK+V5VWg1KHCkJJwxGIi1LiKScYtZL/wAjX/QzwA8AIcMwwnGvG2EjU4jnGNfJsmxEVOGQQkrwxmMhK8ZZzJNSpUndZWnd3zD/AHP9+K/y9P66CgAoAKACgAoA8W+O3x28Efs/+CLrxj4xuvMmk8228O+HbWSMav4m1cR749P0+N87IY9ySajqMiG2062bzZd8r29vcfsfgh4Ica+PPGmG4Q4QwvJSh7PE8QcQYmnUeU8N5VKpyVMwzCpG3PVnadPL8vpzWJzHEr2VLkpQxGIw/wCF/SC+kFwF9HHgLF8b8b4z2lap7XC8NcNYSrTWdcV5zGnz08ty2lPm5KNPmhVzPM6sJYTK8LL21b2laphcNif5rvjZ8bPHHx58b3vjXxtfeZK++20bRrZ5BpHhzSRIzwaVpVu7HZEmd9zcvm5v7kvdXUjyv8v/AEZ+DHgxwX4GcF4PgzgzBclOHLiM4zjExpvNuIs2dNQr5pmteEVz1J25MNhoWw2AwyhhcLCFOHvf8tnjz488ffSH49x/HnHmP56s+fC5HkeFnUWTcL5Mqkp4fJ8nw85S5KUL+0xWKqc2LzHFupi8ZUnVn7vkNfrR+LBQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHefDH4Y+P/AIz+PvC3wv8Ahd4V1bxr4+8aarBo3hvw1osHn32o30wZ2JLtHb2llaQRzXup6nfTW2m6Vp1vdalqV3a2FrcXEbjGUmoxTbeyX9fj956WUZRmefZlg8nyfB18wzLH1o0MLhMPHmqVakrt6tqMKcIqVStWqSjSo0ozq1ZwpwlJf31f8Esf+CVPgD9gXwTH4v8AFi6T44/ab8YaSkPjTx1FD9o03wdYXQjmn8BfD17mJJ7bRoZVRdc19orfU/Ft7As9zHZ6VBpmkWPr0MPGkrvWb3e9vKPl57vqf6T+Dvg3lvhrl6x2N9jmHFuOoqOYZio81LA052lLLcsc0pQoRkl9YxLUa2NqR5pqFGNKhT/Xqug/bwoAKACgCjqmqaZoemajrWtajY6Po2j2N3qmratql3b6fpml6Zp9vJd3+o6jf3ckVrZWNlaxS3N3d3MsVvbW8Uk00iRozC6dOpWqU6NGnOrWqzjTpUqcZVKlWpUkowp04RTlOc5NRjGKcpSaSTbJnOFOE6lScadOnGU6lSclGEIRTlKc5SajGMYpylKTSSTbdj+AL/guN/wXM1P9rLUfEf7J/wCyb4hvtH/Zh0u9l0z4g/ELTpLjT9U+P9/ZT7ZLGyf91dWPwktbmLfaWUgiufG8scepavHHo62Wmyf3L4M+C9Phenh+KOKMPCtxHVgquAwFRRqU8jhNXU5rWM81lF+9Nc0cGm6dJutz1F/L3iR4kzzydbI8jqyp5NCThi8XBuM80nF6xjtKOBi17sXZ4l+/UXs+WD/mCr+kD8aCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/1t/8Agm7/AMo7v2Cf+zLv2W//AFR3gWv8tvEH/kveN/8Asr+Jf/V1jT+5+Ef+SU4Y/wCyeyX/ANVuGP8AF4r5A+hCgAoAKACgAoAntra5vbm3s7O3nu7u7nitrW1topJ7m5uZ5Figt7eCJXlmnmldY4oo1aSSRlRFLMAc61alh6VXEYirToUKFOdatWrTjTpUaVOLnUq1ak2oU6dOClOc5yUYxTlJpJsqEJ1Zwp04SqVKkowp04Rc5znNqMYQjFOUpSk0oxSbk2kk2z+tf/glV/wSot/gnb6F+0b+0doUN38Y7qGHU/h/8PtThSe2+FUEyCS31zXLeQNHN8RJY2D21s4aPwYjfxeJSz6J/hh9Nn6bdXxFq5n4S+EmZ1KHAFGpUwfFPFWDqSp1uNqtOThWy3LKsWp0+E4STjWrRalxDJdMnSWZ/wCh3gD4AQ4XhhONONcLGpxLUjGvk+T14qcMghNc0MXioO8Z5zKLThB3WWJ9cc28L++Vf5hH9cBQAUAFABQAUAeKfHf48eCP2fvBFz4w8Y3Xm3Evm2vhzw5ayxjV/E2rrHvSwsI33eXBHuSTUtSkRrbTbZhJL5k8tra3P7J4HeB3Gvj1xphuEeEMN7OjT9lieIeIcTTqSynhrKZVOSePx9SPL7StU5Z08uy6nOOJzHExdOk6dGnisThvwn6Qn0hOAvo4cBYvjbjbF+1xFX2uE4Z4ZwlWnHOeK85jT56eXZdTnzeyoUuaFXM8zqwlhcswslVq+1r1cJhcV/Nd8a/jX43+PHje+8beN77zJpN1vo+j2zSLpHhzSRIzwaTpNu7N5cMed9xcPuub+5Ml3dySTSEj/o08GfBngvwN4LwXBnBeC9nRhy4nN83xMacs24izaVOMK+a5riIRj7StUtyYfDwthsDhlTwuFp06VNJ/8tnjv478ffSG49x/HnHuYe1r1OfC5JkmFlUjk3DGTKpKeHyfJ8NOUvZ0afNz4nEz5sXmOKlUxmNq1a9RteR1+sH4wFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAej/AAk+EfxG+O3xF8K/Cf4TeE9V8beP/Gmpx6V4f8PaREJLi5nZWlnubmeRo7XTtL061jm1DV9X1Ce20zSdNtrrUdRuraztpp0qMZTkoxTbfT+vxPVyTJM14jzTB5LkuCrZhmePqqjhsLQjeU5O7lOcm1ClRpQUqlevVlClRpRnVqzhCMpL+/n/AIJff8Etvhz/AME/vAI1vWP7L8cftI+MtKhi+IXxIS3MlpolpMYrmXwF8PjdxR3Wn+FbO5jiOo6k8VtqvjDULaLVNVitLK20XQtE9ahQjRV/im170v0XZfi932X+lnhD4PZV4Z5b9Yr+xzHivH0Us0zZRvDDwlyzeW5ZzxU6WChJJ1arUa2OqxVasoU4YfDYf9YK6D9nCgAoAKAM7V9X0nw/pWp69r2p6fomh6Jp97q2s6zq97babpWk6Vp1vJeahqep6jeSQ2lhp9haQzXV5eXU0Vva28Uk88iRozDSlSq16tOjQp1K1atUhSo0qUJVKtWrUkoU6dOnBOc6k5tRhCKcpSaSTbJqVIUoTq1Zxp06cZVKlSpJQhCEE5TnOcmoxjGKcpSk0kk23Y/z+P8AguJ/wXK1b9r3U/EP7Kv7Kev6hon7LWkX72Hjnx1Yvc6bq37QWo2FxzEmRDeaf8JrO6iEul6TMsV14wnjh1rXYY7JdN0m2/ufwZ8GKfCtOhxNxPQp1uJasFUwWCmo1KWRU5x3e8Z5pOLtUqpuOEi3RotzdSrL+XPEfxJqZ9OrkmSVZ08lhJwxOJi3Gpms4vZbShgYyV4QdpYhpVKqUeSmv5j6/o0/HAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/AFt/+Cbv/KO79gn/ALMu/Zb/APVHeBa/y28Qf+S943/7K/iX/wBXWNP7n4R/5JThj/snsl/9VuGP8XivkD6EKACgAoAKAJ7W1ur66t7Kyt57y8vJ4bW0tLWGS4urq6uJFigt7eCJXlnnnldIoYYkaSWRlRFZmAOVatRw1GtiMRWpUMPQpVK1evWqRpUaNGlF1KtatVqOMKdKnCMp1Kk5KMIpyk0k2XTp1KtSFKlCdWrVnGnTp04ynUqVJyUYQhCKcpznJqMYxTlKTSSbZ/W7/wAEq/8AglTbfA630L9ov9ozQ4Lz4zXcMOpeAvAWoxR3Fr8KLedA8Gs6zA4eK4+Ik0bhoYmDR+DkbbGW8RF5dI/wt+mx9Nqt4kVcy8J/CbMquH8PqFSphOJ+J8JOVKtxvVpTcauX5fUi41KXCcJxaqTTjPiCSvO2UqMMf/of4BeAEOFYYTjPjTCwq8TVIxr5TlFaKnT4fhNXjicVF3jPOpRd4x1jlidlfGtyw372V/mMf1sFABQAUAFABQB4l8efjz4I/Z98E3Pi/wAYXPnXM3m23hvw3ayxrq/ibVlj3pY2KPu8q2i3JJqWpyo1tp1sweQS3Etpa3P7N4G+BnGvj3xph+EuEcN7LD0vZYniLiLE0qksq4ayqVRxnjsdOPL7XEVeWpTy7Lac44nMcTFwpulQpYrFYb8H+kN9IbgL6N/AWK4142xXtsVW9theGOF8JWpxznivOY01OGX5fCfN7HDUeanVzTNKsJYXK8LJVKirYmrg8Hi/5r/jT8afG/x38b33jfxvfmW4l3W+kaRbtIuk+HdJWRnt9J0i3dm8qCLcXnnctc31y0l3dySzys1f9Gfg34NcF+B3BeC4M4LwPsqFPlxGbZtiI05ZtxDmsqcYYjNc2xEIx9rXqcvJQox5cNgcNGnhMJTp0KcYn/Lb47eO3Hv0hePcw4949zD2uIq82GyXJcLKpHJuGcmjUlPDZNk2GnOXssPS5nPEYibliswxUquNxtWriKspHklfq5+MhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAHqfwW+CvxO/aG+JnhX4QfB/wnqXjTx/4y1BbDRtF01BwADJealqV5KUtNJ0XSrVZb/WNZ1Ca307S7CCe8vLiKGJmqoxlOSjFXb6fq/Luz2MgyDN+KM2weR5HgquPzPH1VToYekl6zq1akmoUcPRgpVK9erKNKjTjKdSSimz/AEBv+CZf/BMT4Zf8E+vh19ok/s3xt+0H4y0y3T4mfFH7KSkMbNFdN4I8Ci6iS70nwVp13HG00rpBqfizUbaLWtcSCODRdE0D16FCNGPeb+KX6Lsvxe76W/0w8JfCLKPDLK+Z+yzDibHUYrN845HZJtTeX5dzpTo5fSmk5NqNbG1YRxGJUVHD4fDfqPW5+wBQAUAFAGZrWtaP4b0fVvEXiLVtN0HQNB02+1nXNc1m+ttM0jR9I0y2lvdS1TVNSvZYbPT9O0+zhmu729u5ora1topZ55EjRmGlGjWxFalQoUqlevXqQpUaNKEqlWtVqSUKdOnTgpTqVKk2owhFOUpNJJtkVKlOlTnVqzhSpU4SqVKlSShTpwgnKc5zk1GMIxTlKUmlFJtux/n2f8Fwf+C42sftj6tr37Ln7LWu6loP7Kmi6g1n4x8Y2hudM1j9oPVNOuQVkmRhDe6d8KrK7hWfRNBuFiuvFM8UHiDxHBHGmk6Ppf8AdXg14MUuE6dDiXiajTr8TVYKeDwkuWpSyGnUjrZ+9Cpmc4yarV4uUcNFuhh5NurWqfy34j+JFTP51clyWrOlklObjiMRHmhUzWcH12lHAxkr06TtKu0qtZJKFOH8ztf0UfjwUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62//AATd/wCUd37BP/Zl37Lf/qjvAtf5beIP/Je8b/8AZX8S/wDq6xp/c/CP/JKcMf8AZPZL/wCq3DH+LxXyB9CFABQAUAWLS0u7+7trGxtri9vr24htLOztIZLm7u7u5kWG3tra3hV5p7ieZ0ihhiR5JZHVEVmYA5V69DC0K2JxNalh8Nh6VSviMRXqQo0KFClB1KtatVqONOlSpU4ynUqTlGEIRcpNJNl06dStUhSpQnVq1Zxp0qVOMp1KlSclGEIQinKc5yajGMU5Sk0km2f1xf8ABKz/AIJVWvwJttE/aI/aK0S3vfjXeQRaj4F8C6hFHc2nwmtriMPDquqwsHhuPiJNE4Kqd8Xg9GMUBbXjLPpn+FP02Pps1/EytmXhR4UZjVw3h3QqzwnEvEuFnOjX45rUp8tTBYKouWpS4Tpzjq/dnn8kp1UsrUKWM/0S8AvAGnwnDC8Z8Z4WFXiipCNbKsqrKM6fD0Jq8cRiIu8Z51KL03jlqfLBvGOU6H7zV/mUf1mFABQAUAFABQB4j8e/j34I/Z88E3Hi7xfc+fdz+bbeGvDVrNGureJtWWMMtlZKwfybWHdHJqepyRtbadbMHcTXMtpaXX7R4GeBfGvj5xph+E+EsP7HC0fZYniPiPE0qksq4ayqdRxljcbKLj7bE1uWpTy3LadSOJzHERlCDpYeli8Xhvwb6Q/0iOAvo3cBYrjTjTFe3xdf22E4W4WwlanDOeK85hTU4YDAQmp+wwtDnp1c1zWrTlhcrws41KirYqtgsHi/5rvjR8aPG/x28b3/AI38b3/nXM2630rSbdpF0nw7pKyM9tpGkWzu/k20W4vNM5a5vblpby8lluJXc/8ARn4OeDnBfgfwXgeC+C8D7HD0rYjNc1xCpyzbiHNZU4wxGbZtiYRj7XEVeVRo0YKOGwWGjTwmEpUsPShA/wCWzx08dOPfpB8e5hx7x9mPt8VX5sNk+T4aVSGTcM5NGpOeFybJcLOcvY4ajzOdavNyxWPxUquNxtatia05vyav1U/GwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD2D4D/Af4qftLfFLwt8G/gz4UvvGHjzxbeC30/TrQbLWytI8Nf65rmoOPsujeH9Hty15q2r3zx2llbIWd2keKOSoQlUkoxV2/wCrvy7s9zhzhvOeLM4weRZDgqmOzLG1OWnShpCnBa1cRiKr9yhhqEbzr16jUKcVdttpP/QR/wCCbH/BNP4V/wDBPn4Z/ZNP+w+M/jn4v0+1HxS+K8lpsnvpAY7k+EvCC3CC60XwLpd2itBbHyr7xDewR63ro85NM07RvXo0I0Y95v4pfovK/wA3u+h/pp4U+E+TeGWU+zpezx/EOOpQ/tnOnC0qrup/UsCpLnoZdRnZxhpUxNSKxGJ95UqVD9L63P1gKACgAoAyte17RPC+iax4l8S6xpnh7w74e0y/1vXte1u/tdL0fRdH0u1lvdT1XVdTvZYLPT9O0+zgmu729u5ora1topJ55EjRmGtGjWxNalh8PSqV8RXqQo0KFGEqtWtVqSUKdKlTgnOpUqTajCEU5Sk0km2RVq06NOpWrVIUqVKEqlWrUkoU6dOCcpznOTUYwjFOUpSaSSbbsf58P/Bb3/guJrf7aGsa5+zF+zDreqeHv2TtC1I2vinxRbm60vWv2hNW0y6DR3t9G4gvdO+Ftjdwpc+HfDV0kN14iuIrfxJ4nt45U0fRNB/u3wb8GqPCFKjxHxJRp4jievT5sNhpctWjkNKpHWMHrCpmc4txxGIi3HDxcsPh5NOtWr/yx4j+I9TiGpUybJqk6WR0p2r11zQqZrUhLSUlpKOCjJc1KlK0qrSrVkmqdOn/ADVV/Qx+QhQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62//AATd/wCUd37BP/Zl37Lf/qjvAtf5beIP/Je8b/8AZX8S/wDq6xp/c/CP/JKcMf8AZPZL/wCq3DH+LxXyB9CFABQBZs7O81G8tdP0+1ub+/v7mCzsbGzglury8vLqVYLa1tbaBZJri5uJnSGCCFHlmldY41Z2AOOIxFDCUK+KxVejhsLhqNXEYnE4irCjQw9CjCVStXr1qko06VGlTjKpVq1JRhCEZTnJRTZdOnUrVKdGjTnVrVZxp0qVOMqlSrUqSUYU6cIpynOcmoxjFOUpNJJtn9c//BK3/glXafAS10T9of8AaH0W2v8A43XtvFqHgjwRfxxXVn8JLW5jDRalqMTB4bn4izwvy3zxeEY3a3tmfWzPc2P+E301/ps1/E+vmPhV4U5hWw3hzh6s8LxHxHhpzo4jjqtRnaeEwk1y1KPCdOpHSPu1M9lFVa6jlyp0cT/op4BeANPhGnheMuMsNCtxTVhGtleV1VGdPh2nUjeNatF3jUzmcXq9Y5cm4U28Vz1KX7wV/mcf1iFABQAUAFABQB4f8fPj54I/Z78EXHi7xdcfaLy48618M+GbWaNdW8TaskYZbOzVg/kWkG+OTVNUkja2022dWZZrqazs7r9p8CvArjTx940w/CfCeH9hhKHssTxJxJiaVSeVcNZVOo4yxmMlFx9viq/LUp5ZllOpDE5jiYyjGVHDUcZjML+CfSJ+kTwF9G3gPE8Z8aYn6xjcR7fCcK8K4StThnPFecQpqccDgYzU/q+Doc9Orm2bVac8LleFnGc418XXwWCxn81/xm+M3jf46+N7/wAceOL8z3U5a30vS7dpF0nw9pKyM9tpGkWzu/kWsO4tLKxa5vblpby8lmuZpJD/ANGfg74O8F+CHBeB4K4KwPsMNRtiM0zTEKnLNeIM1lCMMTm2b4mEY+2xNblUaVKKjh8Fh40sHg6VHDUacF/y2eOfjlx79ILj3MePuPsx+sYvEc2GyjKMNKpDJuGsmhUnPC5LkuFnOfsMJQ53OrVm54rH4qdbHY6tXxVepUl5PX6ofjoUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB7b+zz+zx8Wv2pvix4Z+DHwV8K3Xivxv4nuMRwpuh0vRNKheMan4l8TaoUeDRPDejRSrPqeqXIKpuhtbWK71G7srK5uEJVJKMVdv8ADzb6JdWfQcMcL53xjnWEyHIMHPG5hjJ6JXjRw9GLXtcXi61nHD4WhF81WtPa6hBTqzp05/6Df/BOf/gnF8Jf+CfXwtGheHVtfFvxf8V2dnJ8VfizdWSw6l4hvYtsw0DQEl3z6H4H0m5Lf2Xo6SedezINW1mS61F0+zevRoxoxstZP4pd/Jdl/wAOz/Tbwt8K8k8Msn+r4VQxueY2nTec51Ony1cVUj731bDJ3lh8vozb9jQT5qkv32IlOq1yfoxWx+pBQAUAFAGP4h8Q6D4S0HWvFPinWtK8OeGfDmlX+ueIPEGu6ha6Vouh6LpVrLfanq2ranfSwWen6bp9nBNd3t7dzRW9tbxSTTSJGjMNaFCviq9HDYajVxGIxFWFGhQowlVrVq1WShTpUqcFKdSpUnJRhCKcpSaSTbM6tWlQpVK9epClRpQlVq1as1CnTpwi5TqVJyajGEYpylKTSSTbdj/PX/4Ldf8ABcDX/wBtnW9a/Zp/Zo1nVfDf7JOgan5HiDxBD9q0rW/2hNX0u6Dw6pq0LiG9074ZWN5Cl14X8KXSRXOt3EVv4o8V26Xq6Lofhr+7/Bzwao8HUaPEXEVKliOKa9PmoUHy1aORUqkbOnSkuaFTMZwbjicVFuNGMpYbDScfbV8T/LHiN4j1OIqlTJ8nqTpZFSnatWXNCpmtSErqc07ShgoySlQoSSlUaVevFSVOlR/m3r+hD8iCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9bf/AIJu/wDKO79gn/sy79lv/wBUd4Fr/LbxB/5L3jf/ALK/iX/1dY0/ufhH/klOGP8Asnsl/wDVbhj/ABeK+QPoQoAtWVle6ne2mnadaXWoahqF1BZWFhZQS3d7e3t3KkFraWlrAkk9zdXM8iQwQQo8s0rpHGjOwBxxOJw+Dw9fF4uvRwuEwtGricVisTVhQw+Gw9CEqtavXrVZRp0aNGnGVSrVqSjCnCMpzkops0pUqterTo0KdStWrVIUqNKlCVSrVq1JKFOnTpwTnOpObUYQinKUmkk2z+uv/gld/wAErLL4AWmi/tCftC6Na6h8cb62jv8AwZ4Lvo4ruy+EdpdRbo769Q+ZBdfES4hfEsw3w+FI3a0smfVjc3dv/hH9Nb6bGI8Ua+Y+FfhXmFbC+G+GrTw3EPEOHlOhiOO69GdpYfDy92rR4UpVI3hTfLUzucVXxKWBVGhV/wBFfALwCpcIU8LxjxjhoVuKqsFVyzLKqjUpcO06kbqrVWsamczi/elrHL4t06TeI9pUh+7df5on9XhQAUAFABQAUAeHfH34++CP2e/BM/i3xbcfaL658628M+GLWaNNW8TaqkYYWlorB/s9nb745NU1WWN7bTrd1Zlnu57Ozu/2vwJ8CONPH7jShwpwph/YYLD+xxPEvEuJpTnlXDWVTqOMsXi5RcPb4yvyVKeWZXTqRxOY4iElGVHC0MZjML+BfSL+kXwF9GzgPE8ZcZYn6zj8T7bCcK8KYSvThnHFecQpqccFgozU/q2Cw/PTq5vm9WlPDZZhpxlKOIxlfA4HGfzXfGX4y+N/jp43v/HHjjUPtF5cZt9M0y3MiaT4e0lJHe20fR7V3fyLSDeWkkZnub25eW9vZp7qeWVv+jPwf8H+C/BHgvA8FcFYD2GEoWxGZ5nXUJZrn+azpwhic3zfEwjH2+Kr8ijTpxUcPg8PGlg8HSo4ajTpx/5bfHHxx49+kDx7mPH3H2ZPE43EuWHynKcNKpDJuGsnhUnPC5LkmEnOf1fB4fncqlSUp4rHYqdbHY6tiMXXq1ZeU1+pH48FABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAe9/s1fs0fF/8Aa1+Lnhz4LfBPwxN4k8X+IJDLPPIZLfQvDGhwSwpqnivxZqyxTR6L4b0dJ43vb6RJZp55bXTNMtdQ1nUNO067uEJVJKMVdv7kurb6Jf1qz6XhPhLPONs7wuQcP4SWKx2JfNOTvHDYPDRlFVsbja1pKhhaCknUqNSlKUoUaMKterSpT/0I/wDgnt/wTx+EH/BP34UJ4R8GRQ+JviX4lt7K5+KnxYvrGODXPGWrwIXWwsULTSaH4N0iaSZPD/huG4kSBXk1DUp9Q1q8vtQn9ejRjRjZayfxS6t/ol0X63Z/pv4Y+GGR+GeSrA4BRxebYuNOec51UpqOIx9eKuqdNXk8PgKEnJYXCxk1FN1asquIqVKsv0CrY/TAoAKACgDF8SeJPD3g3w9rni3xbrmk+GfC3hnSdQ17xF4i17ULXSdE0LRNKtZb7U9W1bU76WCz0/TtPs4Zrq8vLqaK3t4IpJZZFRWYbYfD18XXo4XC0auJxOJqwoYfD0ISq1q1arJQp0qVOClOpUqTkowhFOUpNJJtmdatSw9KrXr1YUaNGE6tatVnGFOlThFynUqTk1GMIxTlKUmkkm2z/PL/AOC2/wDwW88Q/txa9q/7OH7N+rat4Y/ZF8Oap5WsavGLrSdc/aC1nSroSW+ua7buIb3TvhxYXkCXnhLwheJFcanPFbeKvFlsmqLouh+Ff7x8HPBuhwbQpcQcQ0qWJ4qxFO9Gk+WrRyKlUi1KjRkuaFTMKkZOOKxUW404uWGwsnT9tWxP8reI3iNV4jq1MoyipOjkVKdqlRc0Kua1ISuqlRO0oYOMkpUKErSqNKvXSn7OlQ/nFr+gT8kCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/4Ju/8AKO79gn/sy79lv/1R3gWv8tvEH/kveN/+yv4l/wDV1jT+5+Ef+SU4Y/7J7Jf/AFW4Y/xeK+QPoS3Y2N9ql9Z6Zplnd6jqWo3dvY6fp9jbzXl9f315MlvaWdnaW6SXF1d3VxJHBb28EbzTzOkcaM7KDhicThsFhsRjMZiKGEwmEoVcTisViatOhhsNhqFOVWviMRXqyhSo0KNKEqlWrUnGnTpxlOclFNmlKlVxFWnQoU6lavWqQpUaNKEqlWrVqSUKdOnTgpTqVKk5KMIRTlKTUYpt2P68f+CV/wDwStsf2e7LRv2gf2gtHtNR+OuoWqXvhDwfepDeWPwhs7qLK3Vyv7y3uviHcwSbbm6QvD4Xid7DT3fUWu71f8IPprfTXxPipiMw8LfCzMK+E8NcLWlh8+z7DyqUMTx5iKE/eo0n7lWjwpSqRvRoSUamczjHFYuKwioYd/6L+AfgFS4OpYbjDjDDU6/FdamquXZdVUatLhylUjdTnvCpnM4u1SorxwEW6VGTrOpVP3Ur/NU/q0KACgAoAKACgDwz4/8Ax/8ABP7PXgmfxZ4sn+1ahdedbeGPDFrNGmreJtVRAwtrVWD/AGext98cuq6rLG9vp1u6krPeXFlZXf7b4D+A/Gnj9xpQ4V4VoPD4HDOjieJeJcTSnPKuG8rnUcXisU4uH1jG4jkqU8ryunUhiMxxEJpSoYShjcZhfwD6Rn0jOAvo18B4jjHjLE/WsxxSr4XhPhPCV6cM44rzinTUo4PBxkp/VsBhnOlVzfOKtKeFyzDTg3DEY3EYDAY3+a/4x/GPxv8AHLxtqHjjxzqBub25Jg03TYDImleH9KSR3tdH0e1d3+z2dvvLOzM9xeXDy3l7NPdzyzP/ANGXhB4QcF+CXBeA4K4KwH1fB4e1fMcxrqE81z7NZwhDE5vm+JhCH1jF1+RRhCKhh8Jh40sHg6VHC0aVKP8Ay2+N/jfx79IDj3MuP+P8yeKx2Jbw+V5Xh3Uhk/DeTwqTnhMkyTCTnP6tgsNzylOcpTxWOxM62Ox9fEYzEVq0/Kq/UT8gCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAPov9lr9lj4x/th/F/QPgv8FPDj634k1hvtWqapdedb+G/Bvh2GaGPU/Fvi7VY4Z10rQNLE8fmy+XNeX93LaaRo9nqOs6hYafc3TpyqyUYq76vol1bf8AVz6ng7g7PeOs8w2Q8P4V4jF1/frVp80cJgMLGUVWxuOrKMlRw1HmV3aVSpUlChQhVr1adKf+hR+wL+wD8HP2AvhHD4C+H1smveN9ejsr34o/FTUbKGDxH49163jbbkK07aP4V0l5riHwz4Wt7ma10q3lmubqfUtcv9X1jUfYo0Y0Y2WrespPdv8Ay7L83qf6deG3hrkXhrkkctyyKxOYYlU6mcZzVpqOKzLExT6Xk6GDouU44TBxnKFGMpTnKriKlevV+7a1P0UKACgAoAwfFPinw34I8N694x8Za9pHhbwn4W0jUNf8SeJPEGoWuk6HoOiaVay3up6tq2p30sNnYafYWkMtzd3dzNHDBDG8kjqqk1vhsNiMZiKGEwlCricViatOhh8PQpyq1q9arJQp0qVOCc51JzajGMU3JuyVzKtWo4ajVxGIq06NCjTnVrVqs4wp0qcE5TqVJyajGEYpylKTSSTbZ/nhf8Fs/wDgtx4k/bq8Q6r+zx+ztqmr+Fv2QvDWrBNQv1F1pOu/tAazpV0JLbxF4ktnEN5p/wAPrG7hS88H+DLxIri8mjtvFXi22XWV0bRPCX95+Dvg5h+C6FLP8/pUsVxViKV6cHy1aORUasbSoYeWsamPqQk4YvFxbUIuWFwsvZOtWxX8qeIviNW4lqzynKZ1KGQ0anvz96FXNKkJXjVrRdpQwsZJSoYeWsmlXrpVPZ06H86Vfvx+ThQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAf62/8AwTd/5R3fsE/9mXfst/8AqjvAtf5beIP/ACXvG/8A2V/Ev/q6xp/c/CP/ACSnDH/ZPZL/AOq3DH+MHp+n3+rX9lpel2V3qWp6ld22n6dp2n2015f6hf3kyW9pZWVpbpJcXV3dXEkcFtbQRyTTzSJFEju6qficVisNgcNicbjcTQweDwdCtisXi8VWp4fDYXDYenKtiMTicRWlClQoUKUJ1a1arONOnTjKc5RjFs+lo0auIq0qFClUr169SFGjRowlVq1qtWShTpUqcFKdSpUnJQhCCcpyajFNtI/r2/4JY/8ABK6w/Z3sdH+P/wC0BpFpqfx41G1S88J+E7tYbyw+EFleQ8Syj95b3fxCuYJCl9fIZIPDUTyabpcj3jXt/J/g59NX6a2J8V8Tj/C/wux9fB+GWEryoZ5nlB1MPiuPcTQqawg/cq0OFaNWClhsNJQqZvOMcZjYRoLDYaH+jPgH4B0uDKWG4v4vw1OvxbWpqpl+X1OWrR4cpVY/FLeFTOZwk1Vqpyjgot0KDdV1a0v3Pr/Nk/qsKACgAoAKACgDwr9oD9oDwT+zz4Jn8V+K5/teo3QmtvC/he1njTVfE2qpGGFtbBg5trC2LxSarq0sT2+nW7p8lxeXFlY3f7f4DeAvGnj/AMZ0eFuFqDw2X4V0cTxNxNiaNSeV8N5XObi8RiXFw+s47E8lSnleV06kMRmGIhP36GDoY3G4T+fvpHfSO4C+jVwHiOMOMcQsXmeLVfC8JcJYSvTp5xxVnFOmpLC4VSVR4XL8K6lKrnGcVaVTDZZhqkPcxOOxOAwGN/mw+MXxi8b/ABx8bah458c6ibq+uiYNO06AyJpWgaUkjva6Po1q7v8AZrK33szMzPcXdw817ezXF5cTTv8A9GPhD4Q8F+CfBmA4J4Jy/wCrYLDWr5hmFdQnmufZpOEI4nN83xUYQeIxmI5IxjFRhh8Jh4UcHg6VDCUKNGH/AC2+N3jdx74/ceZlx/x/mTxeYYpvD5ZlmHdSGT8OZPCpOeEyTJMJOc/quBw3PKUpSlPE43EzrY/H18TjcRXrz8rr9QPyIKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA+m/2Sv2SPjL+2j8YdE+DXwX0A6jq98VvfEXiK+E8PhbwJ4Zjnji1DxX4t1OKGYWGlWXmKkUUaTahq1/JbaTo9pfapeW1rJpTpyqyUYrXq+iXVt/1fpdn13BXBOfcfZ7h8iyDDe1r1LVMViqnNHB5dhFJRq43HVoxl7OjT5kopKVWvUcKNCFStUhB/6Fv7Cv7CXwa/YK+EFr8NfhjZDVPEeqrZ6h8SviZqdpDD4o+IniSCF0N7fsjTHTdB05prmDwz4XtriWw0KymmYy3+sX+s6zqnr0qUaUeWO71lJ7yf8Al2XQ/wBO/Drw5yHw3yOGU5RT9tiqyhVzbNqsIxxmaYuMWvaVGub2WGpOU44TBxnKnhqcpe9Ur1K9et9sVqffhQAUAFAHPeLPFnhjwH4Y8QeNfGviDR/CnhDwpo+oeIPE3ibxBqFtpOh6Doek20l7qeratqd7JDaWNhY2kMtxc3NxKkUUSMzMAK3wuFxONxNDB4OhVxWKxVWFDD4ehCVWtXrVZKFOlSpwTlOc5NRjGKbbZlXr0cNRq4jEVadChQpzq1q1WahTpU4JynUqTk1GMYxTcpNpJas/zuf+C1//AAW08Uft5eJNU/Z//Z/1LWPCf7H/AIX1cCebF1pOvfH3WtJuhJa+KfFdq/k3dh4FsruFL3wX4JvEjmlljtvFfiy3GvDR9F8If3t4PeDuG4Jw9LPc9p0sVxXiaXux92rRyOjVjaWGw0leNTGzi3DF4yDaScsNhZex9rWxX8peIniLW4mqzyvK51KGQ0amr96FXNKkJXjWrxdpQw0ZLmw+GkrtpV669r7OnQ/ner97PykKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/1t/8Agm7/AMo7v2Cf+zLv2W//AFR3gWv8tvEH/kveN/8Asr+Jf/V1jT+5+Ef+SU4Y/wCyeyX/ANVuGP4f/wDgll/wSv0/9nKx0j4+fH3SbPVPj3qVot14W8LXSw3un/B+xvIfvE/vLe8+IN1byGPUdSiMkHh2GSXStJke4a/1G6/54/pqfTWxXi1icf4YeGGOxGC8McHXdHOs6ouph8Xx7icPU2+xVocLUasFPCYOajUzapGGOx8I01hsJQ/1A8BPAOjwXRw3F3F2Hp1+Lq9NVMBgKnLVo8N0qsf+3oVM4qQk1Xrpyjg4uWHw0nN1q1T9ya/zcP6pCgAoAKACgAoA8I/aC/aC8E/s8eCZ/FXiqcXep3YntvC3ha2nSPVfE2qxoGFvbhhIbbT7YvFJq2rSxPb6fbugCXF7cWNjd/uHgL4CcafSA4zo8L8L0Xhctwro4nififE0ZzyvhvK6k3F4jEOLh9Zx+J5KtPK8qp1IYjMK8JvnoYPD43G4T+fPpH/SQ4D+jTwHX4v4vxCxma4xV8Lwjwjha9OnnHFWb06aksNhlJVHhMtwjqUqucZxVpTw2W4epBcmJx+JwGAxv82Hxh+MPjf44eNtR8c+OtRN3qF2TBYWEBePStB0pHd7XR9GtHdxa2NtvZiSz3F3cPNe3s1xeXE88n/Rl4R+EfBfgpwZl/BPBOXrC4HCpV8fj6/JPNM9zScIRxWb5vioQg8TjcS4RSSjChhaEKWDwdGhhKFGjD/lt8bPGzj3x948zLj/AMQMzeMzHFt4fLcuw7qU8o4dyinUqTwmSZHg5zqLC4DC+0lJuUqmJxmJnXx+Pr4nHYnEYip5bX6cfkgUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9W/sdfscfGf8Abe+MWk/CD4OaL507+VqHi/xhqMc8fhT4feF/PWK78S+J7+KN/JgTLRadpsO/U9c1Ax6dplvNO7GPSnTlVlyx+b6Jd3/Wp9nwLwJn3iDntHI8ioc0ny1cdjqqksFlmD5rTxeMqxT5YrVUqUb1sTVtSoxlJu3+hd+xJ+xF8Gf2E/g9YfCv4T6b9p1G7FrqPxB+IWp20C+K/iN4nigMcmsa1PFv+zafamSeDw94dtpX03w/YSPDAbm+utT1PUfYpUo0o8sd/tSe8n3f6Lp6tt/6d+H/AIfZD4dZHTybJaXPVnyVczzOtGKxua4tRadfESjflpwvKOFwsZOlhaTcY89SdatV+xa0PugoAKACgDm/GHjHwr8PvCviLxz458RaN4R8G+EdH1DxD4n8UeIdQttK0PQdD0q2kvNS1XVdSvJIrWzsrO2ikmnnmkVERSSc4B6MJhMTj8TQwWCw9bFYvFVYUMNhqFOVWtXrVZKFOlSpwTlOc5NKMYpttmOIxFDC0KuJxNWnQw9CnOrWrVZqFOlTgnKc5zk0oxik222f51//AAWs/wCC2Hiv9vnxPqPwH+A2oaz4Q/Y+8K6wCFIudJ134863pVyHtPF3jG0byrqx8G2V3Cl54K8EXio6ypb+KfFVv/wkI0jSfCH98eD3g9huCMNTzzPKdLFcV4ql/dq0cko1Y2lhcLLWM8ZOLcMZjINqzlhsNL2HtauK/lHxD8RK/E9aeWZZKpQyGhU84VczqQd1XxEdHHDxkubD4aWt0q9de15KeH/nrr94PysKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9bf/gm7/wAo7v2Cf+zLv2W//VHeBa/y28Qf+S943/7K/iX/ANXWNP7n4R/5JThj/snsl/8AVbhj8SK/4wD/AHpCgAoAKACgAoA8H/aD/aD8E/s8eCpvFPimYXmq3gntvCvhW1nSPVfEuqRoGEEG4SG1061LxSatq0sTwafA6AJc31xY2N3+5eAfgHxp9IHjOjwxwxReEyzCOjieKOKMTRnPK+HMsqTcXXrtOCxWYYpQqU8qyqnVhiMfXhNueHwWHx2Nwn89fSR+kjwH9GngOvxdxdXWNzfGqvheEeEcLXp0834qzenTUvq+HUlUeEyzCOpSq5xnNWlPDZbh6lNKGJx+Ky/L8b/Nh8X/AIweNvjf421Hx1461E3mo3Z8ixsYPMj0rQdKjd3tdG0a0d5Pstha72PLPcXVw817ez3N7c3FxJ/0Y+EnhJwX4K8GZfwTwRlywmX4Ve3x2OrclTNM9zScIRxWcZxiowg8TjsS4RWkYUMLQhRweCo4fB4ehQp/8t3jV418eePnHmZ+IHiBmbxuZYxuhl+X0HUp5Rw9lFOpOeDyTI8HOpUWEy/CKpJ6zqYnF4ipXx2Or4nHYnEYip5dX6afkoUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9ffsVfsUfGf9un4xaf8KPhHpflWtv9m1Hx74+1O3nPhT4c+F5JzFNrmvXMW3zrqfZNBoOgW0g1PxBqEbW1oIrW31HULDSlSnVlyxXq3tFd3/lu/vPuOAOAM+8Rc9pZLklHlhHlq5lmVaMngsrwblaWIxM18U5WlHDYaL9tiqq5YcsI1atP/Qw/Y1/Yz+DH7D3wd0v4RfB7R9i/uNQ8aeNNSigbxX8Q/FAgEVz4i8S3sSDc3Lw6TpMBXTNBsCtjp0KKZ5rj2KdONKPLH5vrJ93/AFof6ecCcB5D4e5FRyTIqFvhq4/H1VF43NMZy2nisXUild7xo0Y2pYan+7pRXvSl9ZVofaBQAUAFAHMeNfGvhH4ceEfEnj3x94k0bwf4K8H6NqHiHxR4p8Rahb6Voeg6Jpdu93qGp6pqF28dva2lrbxvJLLI4GBtXc7Kp6cHg8VmGKw+BwOHrYvGYutChhsNQhKrWr1qslGnTp04pylOcmkkkY4jE0MJQrYrFVqeHw+Hpzq169WahTpU4JynOc5NKMYpNttn+dL/AMFp/wDgtZ4u/wCCgHivUPgf8D77WvB37HvhLWA1taOLjStd+OetaVc7rTxn42tSY7mz8K2lzEt54K8EXYU2zLb+JvE8DeI/7M03wp/fXhB4P4XgbDQzrOoUcXxXiqXvSXLVo5LRqxtLCYOesZ4qcW44zGR+K8sNh5fV/aVcV/KHiH4iV+KK88ty2VShkNCpotYVcyqQl7uIxC0caEWubD4aWztWrJ1eSFD+fiv3U/LAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP9bf/AIJu/wDKO79gn/sy79lv/wBUd4Fr/LbxB/5L3jf/ALK/iX/1dY0/ufhH/klOGP8Asnsl/wDVbhj8SK/4wD/ekKACgAoAKAPBf2hP2hPBP7O/gmbxR4omF7q16J7bwp4Utp0j1TxLqkaBvJh3CQ2mm2hkik1bVpIngsIHRVS5vrmxsbv9z8AvALjT6QPGdLhnhmi8HlWDdHEcUcU4mjOplnDmWVJyXtq1nBYrMcUoVaeVZVTqwr4+vCcpTw+Cw+OxuE/nj6SX0k+A/o0cCVuLeLayx2cY5V8LwhwhhcRTp5vxTm1OEZewoc0ajweV4N1KVXOc5q0qmHy7D1KcY08VmGKy/L8b/Nh8Xvi942+N3jbUvHXjrUje6leHybKyh3x6XoWlxu72mjaNaO8gtNPtQ7EAs89zO817ezXN7c3FxL/0ZeE3hLwZ4L8GZfwRwRlyweXYRe2xuNrclTNM8zOpCEcVnGcYqMIPFY/FOEU2owoYahCjg8HRw+Dw9ChT/wCW7xp8aePPHrjzNPEDxAzR47NMc/YYDAUPaU8o4fyinUqTweR5Hg51KiwmXYNVJtJzqYjF4ipXx+Pr4rH4rE4mr5hX6Wfk4UAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB9pfsNfsL/Gb9vP4wWnwx+FtidO0LTTaah8R/iTqdpPL4W+HXhqeZkOo6m8bRf2hrN+IbiDw14YtriLUNfvoZQsllpVjrGsaXrSpSqy5Y7fak9orv5+S6v7z77w88O8+8SM8hlGT0/ZYelyVc1zatCTweV4SUmva1mre1xFXlnHCYSMlVxNSMtadGnXr0f8AQt/ZE/ZD+DP7FXwd0X4OfBjQhY6baeXfeJ/E1+sE3irx94neCOG/8VeLNSiiiN7qV35YjtraJItO0ewS30rSLSz0+1hgX2KdONKPLFW7vq33b6/of6d8EcEZDwDkdDIshw3s6ULVMXi6ijLGZli3FKpjMbWSXtKs7WhFKNKhTUaNCFOlCMV9Q1Z9eFABQAUAcr458c+Dvhl4O8TfEL4heJtF8GeB/BujX/iLxV4q8Rahb6Xomg6JpkD3N/qWpX908cFvbW8KMzM7bmbbHGryOiN04PB4vMcXh8DgcPWxeMxdaFDDYbDwlVrV61SXLCnThFOUpSb6ersk2Y4nE0MHQrYrFVqeHw2Hpyq169WShTpU4JylOcpNJJJXbf5n+c5/wWj/AOC1HjH/AIKDeLr74L/Be81rwZ+x74Q1kPp2mSefpeu/G3WtLuCbTxx46tcpPa+HraeNbvwV4IuhjT8QeIvEcL+JDY2Phn+/PCHwgwnAmFhnGcQo4vizFUbVKi5atHJ6NWPv4PBS1jLEST5MZjI35/eoYeX1fnniP5N8QvEPEcU15Zdl8qmHyGhUvCDvCrmNSD93E4lbqkmubD4eXw6VaydbkjR/ASv3M/LwoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/4Ju/8AKO79gn/sy79lv/1R3gWv8tvEH/kveN/+yv4l/wDV1jT+5+Ef+SU4Y/7J7Jf/AFW4Y/Eiv+MA/wB6QoAKACgDwP8AaG/aG8Ffs7eCpfE/ieUX2s3wntvCnhS2nSPVPEmqRoD5ceRIbTS7QyRSatq0kTwWMDxoiXN/c2Njd/uvgD4AcZ/SC4zpcNcNUngspwTo4nininE0J1Ms4dy2pNr2tW0qaxeZYtQq08qyqnVhXx1aE5znhsDh8bjcL/O30lPpKcB/Rn4ErcV8WV1j86x6xGF4P4PwuIp0824pzanTUvZUbxqPB5Vg3UpVM5zmrSqYfLqFSnCFPFZhisvy/G/zYfFz4ueNvjZ411Lx1461I32p3p8mzs4d8el6Hpcbu1po2jWjPILTT7QSMVXc89zO897ezXN9c3NzL/0Y+E/hPwZ4McGZdwRwRlyweW4Ne2xmMrclTM87zOpCEcXnGcYuMIPFY/FOEVKSjChh6EKODwdHD4LD4fD0v+W/xn8Z+PPHnjzNPEDxBzR4/Ncc/YYHA0faU8p4fymnUqTweR5Hg51KiweW4NVJuMXOpiMViKlfH4+viswxWKxVbzKv0o/KAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD7m/YK/YH+Mn7ffxdg+H3w5tW0XwfokllffE/wCKWpWU0/hr4f8Ah+4lcCWba8I1fxNqqQ3MPhjwrbXMN5rN1DNPPPpuiWGsa1pmtKjKtKy2WspPZL17vouvpdn6J4b+G2feJWdxyzKoPD4HDunUzfOKtOUsJlmGk3rKzj7fF1lGccJg4zjOvNSlKVLD06+Io/6FX7Kn7Kfwb/Y3+D+g/Bj4K+Hl0jw/pQF3rOs3fk3Hibxr4kmhii1Lxb4u1WOGFtT1zUjDGCRHDY6bZRWukaPZ6fo9hY2Nv7FOnGnHlird31b7vuz/AE64N4MyLgTI8NkOQYVUMLR9+vXnyyxeYYuUUq2Nx1ZRi62Iq8qu7Rp0qahQoU6VCnTpx+kKs+qCgAoAKAOR8fePvBXws8FeKPiN8RvFGi+CvAngvRb7xF4r8V+Ir+HTdF0LRdNhae81DUL24ZY4oYo1woy0s0rRwQJLPLHG/VgcDjMzxmGwGAw1bGY3GVoUMNhqEHUrV61R8sIQhG7bbfoleUmkmzDFYrD4LD1sXi61PD4bD05Va9etNQp0qcFeU5yeiS+9vRXbSP8AOV/4LP8A/BaPxr/wUM8Y3nwf+D93rXgr9j7wdrXm6Nokpm0zXPjNrWmTn7J488f2oZZYNIgmQXXgzwVcZj0lfJ1zXI5fEjWtv4f/AL+8IvCHB8CYSGbZtCjjOK8XS/fVlapRyilUXvYLBS1UqrT5cXjI61XejRaw6lKv/JniD4hYnirESwGAlUw+QYepenTd4VcwqQfu4rFLdU0/ew+HfwaVKqdZxVL8D6/cD8xCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/AOCbv/KO79gn/sy79lv/ANUd4Fr/AC28Qf8AkveN/wDsr+Jf/V1jT+5+Ef8AklOGP+yeyX/1W4Y/Eiv+MA/3pCgAoA8B/aH/AGh/BP7OvgqXxN4mlF/rV+J7bwn4Ttp0j1TxJqcaA7EyJDZ6VZmSKTV9XkieGxheONEudQubGxu/3f6P/wBH/jP6QfGVLhvhuk8Fk+CdHEcVcVYmhOplnDuW1Jtc9SzgsXmeLUKtPKsqp1YVsbWhOc54bAYbHY7C/wA6fSW+ktwJ9GbgStxVxXWjmGeZhHEYXg/g7C4inTzbijNacIv2dO6qPBZRgnUpVc5zmrSqUMvoVKdOnTxWY4rL8vxn82Pxb+Lfjb41+NdS8deOtTa+1S+PlWlpFvj0zRNMjd2tNG0a0Z5BZ6daCRyib3muJ3mvb2a5vrm5uZf+jDwo8KODPBng3LuCOCMtWCyzBL2uLxdXkqZnnWZVIQjis4zjFxhTeLzDFOEeafLCjh6MKODwdHD4LD4fD0v+W/xm8ZuPPHjjvNPEDxBzWWPzbHy9jgsFR9pSynIMpp1Kk8HkeR4KdSosFlmCVSfJDnqYjE16lfH4+viswxWKxVbzOv0k/KgoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD77/4J9/8ABPf4wf8ABQH4sx+DPA8Mvhz4d+HJ7K7+KnxXv7KWfQfBOjXEjMtrbLuhTW/GGrxRTx+HPDMFxFNeSpLfX8+naJZajqdptRoyrSstEvil0S/Vvov0uz9K8MvDHPPEzOlgMvjLC5XhZU55znVWm5YbL6E3pCCvFYjHV0pLC4SM1Ko1KpVlSw9OrWh/oS/syfsx/B/9kX4ReHfgt8E/DMXh7wpoSefeXcxjufEHivX54oY9U8WeLdWWGGTWfEertBG11dvHFb21vFa6VpVpp2jafp2nWnrwhGnFRirJfe31bfVv+tD/AE44S4RyPgnJMLkHD+EWFwWGXNOcmp4nG4mSiq2NxtZRi6+KruKc5tRjCKhRowpUKVKlD6Aqz6YKACgAoA434h/EPwP8JvA/in4lfEvxVovgjwF4J0W98Q+K/FniK+i07RtD0bT4jLdXt9dzEKqqAI4okD3F1cSRWtrFNczRRP14DAY3NMbhsuy7DVsZjsZWhQwuFw8HUrVq1R2jCEVq31bdoxinKTUU2ufF4vDYHDV8ZjK9PD4XDU5Va9erJQp06cFeUpSf4LVybSinJpP/ADjP+CzX/BZ7xx/wUR8a3Xwo+FFzrXgj9j/wZrXm+HfDcxl07W/i7rOmzMLX4gfEO2RwyWcci/afB3gyYvb6FE0eq6qk/iORBpH+gHhH4RYLgLCRzTNI0cbxXi6NsRiFapRyqlUV5YHASa1m17uLxas6zvSpWoJur/JfiD4g4nivESwOBlUw2Q4epelRd4VMfUg3y4rFrdRvrh8O9KatUqJ1mvZ/g5X7afmYUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFAH+tv8A8E3f+Ud37BP/AGZd+y3/AOqO8C1/lt4g/wDJe8b/APZX8S/+rrGn9z8I/wDJKcMf9k9kv/qtwx+JFf8AGAf70hQB4B+0R+0T4K/Z18FS+JPEsq3+uX6z2/hPwlbTpHqfiPUo0UlVyJDZ6VZmSKTVtWkieGyheOONLm/ubGxuv3n6P30fuMvpB8ZU+HOHKUsDkuBlQxHFXFWIoTqZbw9ltSclzSs4LF5pi1CrTyrKqdWFbG1oTqVKmGwOGxuNwv8AOX0l/pL8CfRm4Fq8U8U1Y5jn2YxxGG4O4NwuIhSzXifNaUItwg3Go8FlGCdSjVznOatGpQwFGpTp06eKzHFYDAYv+bH4tfFrxr8avGup+OvHWptf6rft5Vraxb49M0XTI3drTRtGtGeQWenWYkfYm55p5nmvb2a5vrm5uZv+jDwq8KeDPBrg3LuCOCMtWByvAr2uKxNXkqZlnOZVIQji84zjFxhTeLzHFuEeefLCjQowo4PB0cNgsNh8PS/5b/GTxk478duO818QfEHNZZjm+YS9lhMJR9pSyrIcqpzqSweR5Hgp1KiwWV4JVJ+zp89Svia9SvjsfXxeYYrFYqt5pX6OflgUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB+iP8AwTs/4Jz/ABc/4KCfFRfDvheO58KfCbwtd2c3xV+LV3ZPNpXhjT5SJRo2ipIY4dd8b6vbq40bQopQsKH+1NXls9Khed9qNGVaVlpFfFLt/m32+bP1Hwu8Lc78Tc5+q4RTwWS4OcJZznU6blRwlKXvLD4dO0cRmFeKfsMOpWiv31eUKMXJ/wCg7+zn+zn8I/2VPhN4Z+C/wV8L23hfwX4agzgbJ9X8QaxOkS6n4o8Uap5cc2teJNZliSbUdSnVRtSCxsoLPS7KwsLX2IQjTioxVkvvb6tvq2f6bcLcLZJwbkuEyHIMHHB4DCR8pV8TXkl7bGYytZSxGLxDSlVqyS2jTpxp0adOnD3KqPoQoAKACgDifiR8SPAfwg8B+K/if8T/ABXovgb4f+B9FvPEPizxZ4ivY7DR9E0ixTfPd3dxJyWYlILW2gWW7vrya3sbKC4vLiCCTsy/L8bmuNw2W5bha2Nx2NrQoYXC4eDqVq1WbtGMYr5ylJtRhFSnOUYRlJc+LxeGwGGr4zGV6eGwuGpyq169WSjTp04q7lJv7kleUpNRinJpP/OE/wCCyv8AwWZ8ef8ABRbxxcfDD4Yz614G/ZC8Fa003hXwnM8lhrPxU1fT5XS1+IfxGt4pOgIM/hLwhK8tp4bt5Fvb0XPiGWSez/0C8JPCPBcA4OOZZkqON4rxlG2JxStOjllKorzwGXya3+zisWrTxElyQ5aCSn/JXiB4g4rizEvBYN1MNkOHqXo0G3GpjqkG+XF4tL76GHbcaK96XNVd4/hRX7UfmgUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/8E3f+Ud37BP8A2Zd+y3/6o7wLX+W3iD/yXvG//ZX8S/8Aq6xp/c/CP/JKcMf9k9kv/qtwx+JFf8YB/vSfPv7RX7RXgr9nTwXJ4j8RyLqGvags9v4S8JW86R6l4i1KNATgkSGy0iyMkUmratJE8VnE8cUUdzqF1ZWV1+9/R9+j5xl9ITjKnw9w7TlgMkwEqGI4r4rxNCdTLeH8uqTaTaTgsZmuMUKtPKsqp1YVcZVhUq1amGwGGxuNw384fSZ+kzwJ9GXgWrxRxRVjmXEGYxxGG4N4NwuIhSzXifNKUE3FNxqPA5NgZVKVTOc5q0alHA0alOlSp4vMcVgMBi/5svix8WPGvxp8a6n478dam2oatqDeXb28e+PTdG02N3a00bRrRnkFlptmJH8uPe808zzXt5Nc39zdXU3/AEX+FfhXwb4N8G5bwRwPlscBlWBj7XE4mpyVMyznMqkIRxecZxi4whLGZjjHCPtKjjClQowo4PB0cNgcNhsNS/5cPGPxj478deO818QfEHNZZjnOYy9lhcLS56WVZFldKdSWDyPI8FKpVWByvBKpP2VLnqV8RWqV8djq+KzDFYrFVvNq/Rj8tCgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP0n/wCCbv8AwTY+K3/BQf4nf2dpIvfB3wT8I39ofip8WZrMyWmlQuEuP+EW8KpOottb8darasGtLEM9polpMmta6UtWsLPVd6NGVaXaKfvS7eS7t/huz9X8KvCjOfE3N/ZUfaYHh/BVIf2znUoXhRi7S+p4JSXLiMxrQd4U9aeHhJYjE2h7OnW/0FfgH8AvhT+zJ8K/C/wa+DHhSy8I+BfCdp5NnZW48291K+lCnUdf1/UXAuda8Q6xOputV1a8Z7i6mIUeXbxQQRevCEYRUYqyX9Xfdvqz/TThvhrJuEsmweQ5Dg6eBy7BQ5adOOtSrUlZ1cTiar9/EYqvL361eo3Kcn0ioxXslUe6FABQAUAcL8Tfib4A+DPgDxb8U/in4t0XwL8PPAui3fiHxZ4s8Q3a2Wk6NpNmoMtxcSkNJLLLI0drZWVrHPf6lfz22n6fbXV9dW9vL25dl2OzfHYXLctwtbG4/G1o0MLhcPBzq1qs3pGK6JK8pzk1CnBSqVJRhGUlzYzGYXL8LXxuNr08NhcNTlVr16suWFOEd23u23aMYxTlOTUIKUpJP/N9/wCCyH/BZTx//wAFG/Hk3w6+HUuteBP2RfA+tPN4N8FzSNZax8StWsXkit/iL8SYIJWSS6dS8vhbwo0k9h4Us5vMka7164vL+v8AQPwl8JMDwBglmGYKjjeKsbRSxeMSU6WXUppOWAy+UldRW2KxSSnipqy5aCjA/knj/wAQMVxbiXhMI6mGyLDVG8Ph2+WpjKkW0sXi0nv1o0LuNCLu+aq5SPwzr9oPzYKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKAP8AW3/4Ju/8o7v2Cf8Asy79lv8A9Ud4Fr/LbxB/5L3jf/sr+Jf/AFdY0/ufhH/klOGP+yeyX/1W4Y/nF/aL/aM8F/s6eC5PEXiKRdR1/UVnt/CXhG3uEi1LxDqMaAk52yNZaRZNJFJq2rSRPFaRPHDDHc6hc2Vlc/8AJ59Hv6PfGX0hOMqfD/D1OWX5Fl8qGI4r4rxFCdTLuH8uqTlbTmprG5tjFCrTyrKqdWFXF1YTq1amGwGGxmMw3+rH0mvpNcCfRl4Fq8TcT1YZlxFmUMRhuDeDMNiIUs04lzOlBX1aqSwOS4GVSlUzjOKlKdHB0p06NGnisxxWBwOK/my+K/xX8a/Gjxpqfjrx3qjahq+oN5dvBHvj03RtNjd2tNG0a0Z5BZaZZCRxFEGeWaV5ry8mub65urqb/ov8LPCzg3wc4Ny3gjgfLY4DKcBH2mIr1OWpmOcZjUhCOLzjOMWoQljMyxjhF1arjClRpQo4PB0cNgsNhsNR/wCXHxh8YeO/HPjrNfEHxBzaeZZzmMvZYbD0+elleR5XSnUlg8kyPBSqVI4HK8Cqk1SpKc61etUr43HV8VmGKxWLr+b1+in5cFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAfp9/wAE0P8AgmT8UP8AgoN8SMp/aPgn4BeDtStl+J/xUNqDhgIrpvBXghbqNrXV/HGp2kiOdyXGm+FrC4i1vXklM+jaNr29ChKtLtBP3pfoujf5Xu/P9e8JvCTN/E3Nbr2uX8NYGrBZvnLh10m8vy/nThXzCrBpvSVLB05xxGJT5qFDE/6BfwR+CHwv/Z0+GPhX4PfB3wnp3gzwD4OsFsdI0ewQs8sjHzL3VdVvZS93q+uatdNJfaxrOoSz3+pX00tzdTPI/HrxjGEVGKsl0/rq+p/pjw/w/lHC2UYPI8iwVLAZbgafs6FCktW3rUrVqjvOviK071K9erKVWtUlKc5Ns9Xqj2QoAKACgDgfil8Uvh58E/h74u+K/wAWPF2i+A/h34E0a68QeLPFniC6FppekaXaABpZXw81xc3EzxWenafZxXOo6pqNxa6bptrd393bW0vdluW4/OMfhcsyvC1sdj8bWjQwuFoR56tarLololGKTnUqTcadOnGVSpKMIykubG43C5dha+Ox1enhsJhqcqtevVlywpwj1b1bbbUYQinOc3GEIynJJ/5uv/BYv/gsd8Q/+Cj/AI/l8BeBJNa8B/skeBtakm8DeA5pTaar8QdUs2lgg+JHxKht5Xin1WaNpH8NeGTLcad4PsJ2SN7zW7nU9Vuf9BfCbwlwHh/gVjscqOO4qxtFLG41LnpYClNKUsvy6UldU07LE4m0amLmtVCjGnTj/I/H/H+K4txTwuFdTDZFhqjeGwzfLUxc43SxeMSbTm9XRo3cKEX9qo5zf4f1+yn5wFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQAUAFABQB/rb/8E3f+Ud37BP8A2Zd+y3/6o7wLX+W3iD/yXvG//ZX8S/8Aq6xp/c/CP/JKcMf9k9kv/qtwx/Cd/wAFP/8Ak4LQv+ya6F/6evEdf5Qfs0P+TCZ5/wBnGzv/ANU3Dx+bftY/+UkOH/8As1vD/wD6veJz84a/0QP8wQoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoA/wBCf/ghV/yjM+Av/YR+J/8A6szxVXsYX+DH5/mf6c/R1/5NJw3/ANfc4/8AVxjT9ea6D9vCgAoAKACgD+cz/g6I/wCUZ1r/ANnE/C3/ANNPjav3/wCjb/ycSf8A2IMy/wDTuDPybxm/5I+P/Y1wf/pGIP8AOqr++T+UAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgAoAKACgD/W3/wCCbv8Ayju/YJ/7Mu/Zb/8AVHeBa/y28Qf+S943/wCyv4l/9XWNP7n4R/5JThj/ALJ7Jf8A1W4Y/9k="
            flag_bytes = base64.b64decode(FLAG_B64)
            flag_pil = Image.open(io.BytesIO(flag_bytes))
            flag_img = ctk.CTkImage(light_image=flag_pil, dark_image=flag_pil, size=(150, 90))
            flag_label = ctk.CTkLabel(banner, image=flag_img, text="")
            flag_label.pack(side="right", padx=(0, 10))
        except Exception as e:
            logger.warning("Flag load error: %s", e)

        # Made in SA text under banner
        ctk.CTkLabel(
            self,
            text="Made in South Africa  ·  for South Africans",
            font=("Arial", 10),
            text_color="#888888"
        ).pack(pady=(2, 6))

        # Settings button (top right)
        settings_btn = ctk.CTkButton(
            self, text="⚙️ Settings", command=self.open_settings,
            width=120, height=30, fg_color="gray30"
        )
        settings_btn.place(x=480, y=10)

        # Credits / About button (below Settings)
        credits_btn = ctk.CTkButton(
            self, text="ℹ️ About", command=self.open_credits,
            width=120, height=30, fg_color="gray30"
        )
        credits_btn.place(x=480, y=48)

        # 1. Task
        ctk.CTkLabel(self, text="1. Select Task:", font=("Arial", 13, "bold")).pack(pady=(10, 2))
        self.task_var = ctk.StringVar(value="Translate to English")
        ctk.CTkOptionMenu(
            self, values=["Original Language", "Translate to English"],
            variable=self.task_var, width=380
        ).pack(pady=5)

        # 2. Language
        ctk.CTkLabel(self, text="2. Audio Language:", font=("Arial", 13, "bold")).pack(pady=(10, 2))
        self.lang_var = ctk.StringVar(value=self.config_data.get("last_language", "Auto-Detect"))
        ctk.CTkOptionMenu(
            self, values=list(LANG_CODES.keys()),
            variable=self.lang_var, width=380
        ).pack(pady=5)

        # 3. Mode — Fast vs Accurate. Two plain buttons rather than
        # CTkSegmentedButton because Phase 1/2 needed to render one segment
        # disabled while the other was live; Accurate is active as of Wave 2.3,
        # but the two-button layout is kept rather than reshuffled, since it
        # already matches the rest of the app's button styling.
        ctk.CTkLabel(self, text="3. Mode:", font=("Arial", 13, "bold")).pack(pady=(10, 2))
        self.mode_var = ctk.StringVar(value="fast")

        mode_frame = ctk.CTkFrame(self, fg_color="transparent")
        mode_frame.pack(pady=2)

        self.fast_btn = ctk.CTkButton(
            mode_frame, text="⚡ Fast (Vinnig)",
            command=self._select_fast_mode,
            width=185, height=40, fg_color="#1f538d", font=("Arial", 13, "bold")
        )
        self.fast_btn.grid(row=0, column=0, padx=(0, 8))

        self.accurate_btn = ctk.CTkButton(
            mode_frame, text="🎯 Accurate (Akkuraat)",
            command=self._select_accurate_mode,
            width=185, height=40, fg_color="gray30", font=("Arial", 13, "bold")
        )
        self.accurate_btn.grid(row=0, column=1)

        # 4. Speaker identification
        ctk.CTkLabel(self, text="4. Speaker Identification:", font=("Arial", 13, "bold")).pack(pady=(10, 2))
        self.diarize_var = ctk.BooleanVar(value=False)

        diarize_frame = ctk.CTkFrame(self, fg_color="transparent")
        diarize_frame.pack(pady=2)
        ctk.CTkSwitch(
            diarize_frame, text="Identify different speakers",
            variable=self.diarize_var
        ).pack(side="left", padx=10)

        # Number of speakers selector
        ctk.CTkLabel(diarize_frame, text="  Expected speakers:", font=("Arial", 11)).pack(side="left", padx=(20, 5))
        self.num_speakers_var = ctk.StringVar(value="2")
        ctk.CTkOptionMenu(
            diarize_frame,
            values=["2", "3", "4", "5", "6", "7", "8"],
            variable=self.num_speakers_var,
            width=70
        ).pack(side="left")

        # Status & progress
        self.status_label = ctk.CTkLabel(self, text="Ready", font=("Arial", 13), text_color="gray")
        self.status_label.pack(pady=(20, 5))
        self.progress_bar = ctk.CTkProgressBar(self, width=480)
        self.progress_bar.pack(pady=5)
        self.progress_bar.set(0)

        # Main button
        self.main_btn = ctk.CTkButton(
            self, text="▶  START PROCESSING",
            command=self.start_transcription,
            height=60, width=340,
            font=("Arial", 18, "bold"),
            fg_color="#1f538d"
        )
        self.main_btn.pack(pady=20)

        # Action buttons (after transcription)
        btn_frame = ctk.CTkFrame(self, fg_color="transparent")
        btn_frame.pack(pady=5)

        self.name_btn = ctk.CTkButton(
            btn_frame, text="🎤 Assign Speaker Names",
            command=self.open_name_assignment,
            width=200, height=40, fg_color="gray30",
            state="disabled"
        )
        self.name_btn.grid(row=0, column=0, padx=10)

        self.summarize_btn = ctk.CTkButton(
            btn_frame, text="🤖 Summarize with Claude",
            command=self.summarize_with_claude,
            width=210, height=40, fg_color="#2d6a4f",
            state="disabled"
        )
        self.summarize_btn.grid(row=0, column=1, padx=10)

        # Output area
        ctk.CTkLabel(self, text="Output:", font=("Arial", 12, "bold")).pack(pady=(15, 2))
        self.output_box = ctk.CTkTextbox(self, width=560, height=180, font=("Arial", 12))
        self.output_box.pack(pady=5)

    # ── MODE SELECTION ──────────────────────
    # Both buttons stay enabled and clickable; the colour swap is the only
    # feedback that a mode is "selected" (there is no separate indicator), so
    # both handlers must repaint both buttons — not just set mode_var — or a
    # click would change behaviour invisibly.

    def _select_fast_mode(self):
        self.mode_var.set("fast")
        self.fast_btn.configure(fg_color="#1f538d")
        self.accurate_btn.configure(fg_color="gray30")

    def _select_accurate_mode(self):
        self.mode_var.set("accurate")
        self.accurate_btn.configure(fg_color="#1f538d")
        self.fast_btn.configure(fg_color="gray30")

    # ── SETTINGS / CREDITS ──────────────────

    def open_settings(self):
        SettingsWindow(self)

    def open_credits(self):
        CreditsWindow(self)

    # ── TRANSCRIPTION ───────────────────────

    def start_transcription(self):
        file_path = filedialog.askopenfilename(
            title="Select Audio File",
            filetypes=[("Audio Files", "*.mp3 *.wav *.m4a *.flac *.mp4 *.ogg *.webm")]
        )
        if not file_path:
            return

        # Validate the selection before doing any heavy work.
        if not os.path.isfile(file_path):
            messagebox.showerror("Error", "The selected file could not be found.")
            return
        if os.path.getsize(file_path) == 0:
            messagebox.showerror("Error", "The selected file is empty.")
            return

        mode = self.mode_var.get()

        # Save last used settings
        self.config_data["last_language"] = self.lang_var.get()
        save_config(self.config_data)

        task_choice = self.task_var.get()
        language_label = self.lang_var.get()
        lang_code = LANG_CODES[language_label]
        whisper_task = "translate" if task_choice == "Translate to English" else "transcribe"
        do_diarize = self.diarize_var.get()
        num_speakers = int(self.num_speakers_var.get())

        self.output_box.delete("1.0", "end")

        # Accurate mode has its own flow — first a possible download (with
        # consent), then transcription. It is never allowed to fall back to
        # the Fast/Medium engine, so this branches out completely rather than
        # sharing run() below.
        if mode == "accurate":
            self._start_accurate_transcription(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )
            return

        model_name = FAST_MODE_MODEL
        self.status_label.configure(text=f"Status: Loading model ({model_name})...")
        self.progress_bar.set(0.1)
        self.main_btn.configure(state="disabled")
        self.name_btn.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")

        def run():
            try:
                self._ui(self.progress_bar.set, 0.3)
                self._ui(self.status_label.configure, text="Status: Processing audio...")

                result = transcribe_audio(
                    file_path,
                    language=lang_code,
                    task=whisper_task,
                    model_name=model_name,
                )

                self._ui(self.progress_bar.set, 0.7)

                if do_diarize:
                    self._ui(self.status_label.configure, text="Status: Identifying speakers...")
                    segments = self._diarize(result, max_speakers=num_speakers, audio_path=file_path)
                    self.diarized_segments = segments
                    transcript_text = self._segments_to_text(segments, self.speaker_name_map)
                else:
                    self.diarized_segments = None
                    transcript_text = result["text"]

                # current_transcript feeds the on-screen preview AND the Claude
                # summary prompt, so it must stay pure transcript. Provenance is
                # attached only at the export layer (below), never here.
                self.current_transcript = transcript_text

                provenance = build_provenance(
                    mode="Fast (Vinnig)",
                    language_label=self.lang_var.get(),
                    task=whisper_task,
                    diarized=do_diarize,
                    num_speakers=num_speakers,
                )
                output_path = self._save_transcript(
                    transcript_text, file_path, provenance=provenance
                )

                preview = transcript_text[:2000] + ("..." if len(transcript_text) > 2000 else "")
                self._ui(self.progress_bar.set, 1.0)
                self._ui(self.status_label.configure, text="✅ Done!")
                self._ui(self.output_box.insert, "1.0", preview)

                if do_diarize and self.diarized_segments:
                    self._ui(self.name_btn.configure, state="normal")
                self._ui(self.summarize_btn.configure, state="normal")

                self._ui(messagebox.showinfo, "DanScribe AI", f"Transcription complete!\nSaved to folder:\n{output_path}\n\nBoth .txt and .docx files created.")

            except Exception as e:
                logger.error("Transcription failed: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"An error occurred:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")

            self._ui(self.main_btn.configure, state="normal")
            self._ui(self.progress_bar.set, 0)

        threading.Thread(target=run, daemon=True).start()

    # ── ACCURATE MODE (Wave 2.3) ─────────────
    # First activation: consent -> cancellable/resumable download -> full
    # verification -> transcription. Subsequent activations: straight to
    # transcription. Both routes converge on _transcribe_with_accurate_model()
    # for the actual engine call, so Wave 2.4's transcription-progress /
    # time-warning UI only has one call site to extend, not two.
    #
    # accurate_model_download is imported lazily inside these methods, the
    # same discipline accurate_engine.py and transcribe_audio_accurate() already
    # follow — Fast mode's runtime must never acquire this import just because
    # the module exists on disk.

    def _start_accurate_transcription(self, file_path, language_label, lang_code,
                                       whisper_task, do_diarize, num_speakers):
        import accurate_model_download as amd

        self.main_btn.configure(state="disabled")
        self.name_btn.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")

        if amd.is_model_ready():
            # Subsequent activation — skip consent AND the download UI
            # entirely, per masterplan 2.3. ensure_accurate_model() still runs
            # (a local stat + JSON read; no network) so the directory handed to
            # the engine is always the one it actually verified, never assumed.
            self.status_label.configure(text="Status: Loading Accurate-mode model...")
            self.progress_bar.set(0.1)

            def already_ready_worker():
                try:
                    model_dir = amd.ensure_accurate_model()
                except Exception as e:
                    logger.error("Accurate model check failed unexpectedly: %s", e, exc_info=True)
                    self._ui(self._on_accurate_setup_failed, e)
                    return
                self._transcribe_with_accurate_model(
                    model_dir, file_path, language_label, lang_code,
                    whisper_task, do_diarize, num_speakers,
                )

            threading.Thread(target=already_ready_worker, daemon=True).start()
            return

        # First activation. Fetch the REAL size/file-count before asking for
        # consent — never a number remembered from an earlier report, which
        # could be stale — off the main thread like every other network call.
        self.status_label.configure(text="Status: Checking Accurate-mode download details…")

        def fetch_info():
            try:
                info = amd.describe_download()
            except Exception as e:
                self._ui(self._accurate_consent_check_failed, e)
                return
            self._ui(
                self._show_accurate_consent, info,
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )

        threading.Thread(target=fetch_info, daemon=True).start()

    def _on_accurate_setup_failed(self, exc):
        self.status_label.configure(text="Status: Ready")
        self.progress_bar.set(0)
        self.main_btn.configure(state="normal")
        messagebox.showerror("DanScribe AI", f"Accurate mode could not start:\n{exc}")

    def _accurate_consent_check_failed(self, exc):
        self.status_label.configure(text="Status: Ready")
        self.main_btn.configure(state="normal")
        messagebox.showerror(
            "DanScribe AI",
            "Could not check the Accurate-mode download details. This usually "
            f"means there is no internet connection right now.\n\n{exc}\n\n"
            "Fast mode works offline."
        )

    def _show_accurate_consent(self, info, file_path, language_label, lang_code,
                                whisper_task, do_diarize, num_speakers):
        self.status_label.configure(text="Status: Ready")

        def on_confirm():
            self._download_then_transcribe_accurate(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )

        def on_decline():
            # The user said no. Accurate mode simply isn't used this time —
            # reset and don't ask again until they click Accurate again.
            self.main_btn.configure(state="normal")
            self.progress_bar.set(0)

        AccurateConsentDialog(self, info, on_confirm=on_confirm, on_decline=on_decline)

    def _download_then_transcribe_accurate(self, file_path, language_label, lang_code,
                                            whisper_task, do_diarize, num_speakers):
        import accurate_model_download as amd

        cancel_event = threading.Event()
        progress_dialog = AccurateDownloadProgressDialog(self, on_cancel=cancel_event.set)

        def worker():
            try:
                model_dir = amd.ensure_accurate_model(
                    progress_callback=lambda p: self._ui(progress_dialog.update_progress, p),
                    cancel_event=cancel_event,
                )
            except amd.AccurateModelDownloadCancelled:
                # Catch this BEFORE AccurateModelDownloadError (its own base
                # class) — a deliberate cancel is not a failure and must never
                # show an error dialog. Whatever chunks landed stay on disk.
                self._ui(self._on_accurate_download_cancelled, progress_dialog)
                return
            except amd.AccurateModelDownloadError as e:
                self._ui(
                    self._on_accurate_download_failed, progress_dialog, e,
                    file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
                )
                return
            except Exception as e:
                logger.error("Accurate model download failed unexpectedly: %s", e, exc_info=True)
                self._ui(
                    self._on_accurate_download_failed, progress_dialog, e,
                    file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
                )
                return

            self._ui(progress_dialog.destroy)
            self._transcribe_with_accurate_model(
                model_dir, file_path, language_label, lang_code,
                whisper_task, do_diarize, num_speakers,
            )

        threading.Thread(target=worker, daemon=True).start()

    def _on_accurate_download_cancelled(self, dialog):
        dialog.destroy()
        self.status_label.configure(text="Status: Ready (download cancelled)")
        self.progress_bar.set(0)
        self.main_btn.configure(state="normal")
        # Deliberately no error dialog — the user asked for this, and nothing
        # already downloaded was lost (see accurate_model_download's
        # cancellation notes: chunks are only ever removed after the guard
        # passes, never on cancel).

    def _on_accurate_download_failed(self, dialog, exc, file_path, language_label, lang_code,
                                      whisper_task, do_diarize, num_speakers):
        dialog.destroy()
        self.status_label.configure(text="Status: Ready")
        self.progress_bar.set(0)

        # AccurateModelDownloadError's message is already written for a
        # non-technical user (masterplan 2.1a.2) — shown as-is, no wrapping.
        can_retry = getattr(exc, "can_retry", False)
        if can_retry and messagebox.askretrycancel("DanScribe AI", str(exc)):
            self._download_then_transcribe_accurate(
                file_path, language_label, lang_code, whisper_task, do_diarize, num_speakers,
            )
            return

        self.main_btn.configure(state="normal")

    def _transcribe_with_accurate_model(self, model_dir, file_path, language_label, lang_code,
                                         whisper_task, do_diarize, num_speakers):
        """Runs on a background thread — the shared tail for both Accurate-mode
        routes above. Everything from here down is what Wave 2.4's duration
        warning / generate()-progress UI will extend; the transcribe_audio_
        accurate() call is the seam (an extra progress_callback= kwarg is all
        it would need)."""
        self._ui(self.status_label.configure, text="Status: Loading Accurate-mode model...")
        self._ui(self.progress_bar.set, 0.3)
        try:
            self._ui(self.status_label.configure,
                     text="Status: Processing audio (Accurate mode)...")

            result = transcribe_audio_accurate(
                file_path, language=lang_code, task=whisper_task, model_dir=model_dir,
            )

            self._ui(self.progress_bar.set, 0.7)

            if do_diarize:
                self._ui(self.status_label.configure, text="Status: Identifying speakers...")
                segments = self._diarize(result, max_speakers=num_speakers, audio_path=file_path)
                self.diarized_segments = segments
                transcript_text = self._segments_to_text(segments, self.speaker_name_map)
            else:
                self.diarized_segments = None
                transcript_text = result["text"]

            self.current_transcript = transcript_text

            provenance = build_provenance(
                mode="Accurate (Akkuraat)",
                language_label=language_label,
                task=whisper_task,
                diarized=do_diarize,
                num_speakers=num_speakers,
                # accurate_engine.transcribe()'s result already carries its own
                # label (ACCURATE_ENGINE_LABEL) — reuse it rather than import
                # accurate_engine here just for a string.
                engine=result.get("engine"),
                # Masterplan 2.6 — likewise sourced from the result dict, not
                # re-derived here, so accurate_engine stays the single source
                # of truth for what backs a given transcription.
                model_id=result.get("model_id"),
                model_revision=result.get("model_revision"),
                model_layout=result.get("model_layout"),
                guard_verification=result.get("guard_verification"),
            )
            output_path = self._save_transcript(
                transcript_text, file_path, provenance=provenance
            )

            preview = transcript_text[:2000] + ("..." if len(transcript_text) > 2000 else "")
            self._ui(self.progress_bar.set, 1.0)
            self._ui(self.status_label.configure, text="✅ Done!")
            self._ui(self.output_box.insert, "1.0", preview)

            if do_diarize and self.diarized_segments:
                self._ui(self.name_btn.configure, state="normal")
            self._ui(self.summarize_btn.configure, state="normal")

            self._ui(
                messagebox.showinfo, "DanScribe AI",
                f"Transcription complete!\nSaved to folder:\n{output_path}\n\n"
                "Both .txt and .docx files created.",
            )

        except Exception as e:
            logger.error("Accurate transcription failed: %s", e, exc_info=True)
            self._ui(messagebox.showerror, "Error", f"An error occurred:\n{e}")
            self._ui(self.status_label.configure, text="Status: Error")

        self._ui(self.main_btn.configure, state="normal")
        self._ui(self.progress_bar.set, 0)

    # ── DIARIZATION ─────────────────────────

    def _diarize(self, whisper_result, max_speakers=4, audio_path=None):
        """Attach speaker labels to a transcript.

        Returns exactly the shape it always has — an ordered list of
        {"speaker": "Speaker N", "text": ...} with consecutive same-speaker
        turns merged — so _segments_to_text(), the name-assignment dialog and
        both exporters are unaffected by what happens inside.

        The speaker timeline comes from pyannote (diarization_engine), computed
        over the raw audio and NOT from Whisper's segment boundaries. That
        distinction is the entire point of masterplan 2.12: the old backend
        averaged one feature vector per ~30s Whisper segment, which made a
        speaker change inside a segment unrepresentable and collapsed a
        confirmed four-voice stretch into a single speaker.

        Falls back to the pause-based heuristic below if diarization is
        unavailable — logged at ERROR, not warning, because silently shipping a
        near-one-speaker transcript is precisely the failure this product
        cannot afford.
        """
        segments = whisper_result.get("segments", [])
        if not segments:
            return [{"speaker": "Speaker 1", "text": whisper_result["text"]}]

        try:
            if audio_path is None:
                raise ValueError("no audio path provided")

            import diarization_engine

            turns = diarization_engine.diarize(audio_path, max_speakers=max_speakers)
            if not turns:
                raise ValueError("diarization returned no speaker turns")

            # Referenced via the class, not self. _diarize() has always been
            # callable with an unbound/None self — the existing test harnesses
            # do exactly that — and _attach_speakers is a @staticmethod, so
            # going through self would silently reintroduce an instance
            # requirement and turn every such call into a fallback.
            labelled = DanScribeApp._attach_speakers(segments, turns)
            if not labelled:
                raise ValueError("no text could be attached to the speaker timeline")
            return labelled

        except ImportError as e:
            logger.error("Speaker-diarization libraries unavailable (%s); falling back to "
                         "pause-based detection, which is much less accurate.", e)
        except Exception as e:
            logger.error("Speaker diarization failed (%s); falling back to pause-based "
                         "detection, which is much less accurate.", e, exc_info=True)
        # ── Fallback: pause-based detection ─────────────
        pauses = []
        for i in range(1, len(segments)):
            gap = segments[i].get("start", 0) - segments[i - 1].get("end", 0)
            pauses.append((i, gap))

        sorted_pauses = sorted(pauses, key=lambda x: x[1], reverse=True)
        num_changes   = max_speakers - 1
        change_points = set(idx for idx, gap in sorted_pauses[:num_changes] if gap >= 0.8)

        result = []
        current_speaker = 1
        for i, seg in enumerate(segments):
            text = seg.get("text", "").strip()
            if not text:
                continue
            if i in change_points and current_speaker < max_speakers:
                current_speaker += 1
            speaker_key = f"Speaker {current_speaker}"
            if result and result[-1]["speaker"] == speaker_key:
                result[-1]["text"] += " " + text
            else:
                result.append({"speaker": speaker_key, "text": text})

        return result if result else [{"speaker": "Speaker 1", "text": whisper_result["text"]}]

    @staticmethod
    def _attach_speakers(segments, turns):
        """Map a speaker timeline onto Whisper's segments.

        `turns` is [(start, end, raw_speaker)], non-overlapping and time-sorted.

        A Whisper long-form segment can be ~30s and span several speaker turns,
        so a segment is split where the timeline says the speaker changed, and
        its words are apportioned across the parts in proportion to how long
        each speaker held the floor. Word-level timings would make this exact;
        segment-level timings are all Whisper gives us here, so this is an
        approximation — but a far better one than handing a whole 30s block to
        whichever speaker happened to dominate it, which is what the old
        backend effectively did.

        Raw pyannote labels (SPEAKER_00, ...) are renumbered to "Speaker N" in
        order of first appearance, matching the previous backend's convention
        so the name-assignment dialog keeps working unchanged.
        """
        raw_to_label = {}
        out = []

        for seg in segments:
            text = (seg.get("text") or "").strip()
            if not text:
                continue

            s0 = float(seg.get("start", 0.0) or 0.0)
            s1 = float(seg.get("end", s0) or s0)

            overlaps = []
            if s1 > s0:
                for t0, t1, spk in turns:
                    o = min(s1, t1) - max(s0, t0)
                    if o > 0:
                        overlaps.append((max(s0, t0), o, spk))
                overlaps.sort(key=lambda x: x[0])

            if not overlaps:
                # No speech detected here by the diarizer (music, noise, a gap).
                # Keep the text with whoever was last speaking rather than
                # inventing a speaker or dropping the words.
                carry = out[-1]["_raw"] if out else (turns[0][2] if turns else "SPEAKER_00")
                parts = [(carry, text)]
            elif len(overlaps) == 1:
                parts = [(overlaps[0][2], text)]
            else:
                words = text.split()
                total = sum(o for _, o, _ in overlaps)
                parts, idx = [], 0
                for n, (_, o, spk) in enumerate(overlaps):
                    if n == len(overlaps) - 1:
                        take = len(words) - idx
                    else:
                        take = int(round(len(words) * (o / total))) if total > 0 else 0
                        take = max(0, min(take, len(words) - idx))
                    chunk = words[idx:idx + take]
                    idx += take
                    if chunk:
                        parts.append((spk, " ".join(chunk)))
                if not parts:
                    parts = [(overlaps[0][2], text)]

            for spk, chunk in parts:
                if spk not in raw_to_label:
                    raw_to_label[spk] = f"Speaker {len(raw_to_label) + 1}"
                label = raw_to_label[spk]
                if out and out[-1]["speaker"] == label:
                    out[-1]["text"] += " " + chunk
                else:
                    out.append({"speaker": label, "text": chunk, "_raw": spk})

        for entry in out:
            entry.pop("_raw", None)
        return out

    def _segments_to_text(self, segments, name_map=None):
        lines = []
        for seg in segments:
            speaker = seg["speaker"]
            name = name_map.get(speaker, speaker) if name_map else speaker
            lines.append(f"{name}: {seg['text']}")
        return "\n\n".join(lines)

    # ── NAME ASSIGNMENT ──────────────────────

    def open_name_assignment(self):
        if not self.diarized_segments:
            return
        # Preserve insertion order (no sorting that breaks numbering)
        seen = []
        for seg in self.diarized_segments:
            if seg["speaker"] not in seen:
                seen.append(seg["speaker"])

        def on_names_confirmed(name_map):
            self.speaker_name_map = name_map
            updated_text = self._segments_to_text(self.diarized_segments, name_map)
            self.current_transcript = updated_text
            self.output_box.delete("1.0", "end")
            self.output_box.insert("1.0", updated_text[:2000] + ("..." if len(updated_text) > 2000 else ""))
            self.status_label.configure(text="✅ Names assigned!")

        NameAssignWindow(self, seen, on_names_confirmed)

    # ── AI SUMMARY ──────────────────────────

    def summarize_with_claude(self):
        config = load_config()
        api_key = _get_api_key(config).strip()

        if not api_key:
            messagebox.showwarning(
                "API Key Missing",
                "Please enter your Claude API key via ⚙️ Settings."
            )
            return

        if not self.current_transcript:
            messagebox.showwarning("No Transcript", "Please complete a transcription first.")
            return

        self.summarize_btn.configure(state="disabled")
        self.status_label.configure(text="Status: Claude is generating summary...")
        self.progress_bar.set(0.3)

        def run():
            try:
                client = anthropic.Anthropic(api_key=api_key)

                prompt = f"""You are a professional minutes writer. The following is a transcription of a meeting or conversation.

Please provide a concise summary that includes:
1. Main points discussed
2. Decisions made (if any)
3. Action items (if any)
4. Participants (if names are available)

Transcription:
{self.current_transcript}

Write the summary in the same language as the transcription."""

                message = client.messages.create(
                    model=CLAUDE_MODEL,
                    max_tokens=2000,
                    # Disable extended thinking: a summary doesn't need it, and it
                    # keeps latency and token cost down (Sonnet 5 thinks by default).
                    thinking={"type": "disabled"},
                    messages=[{"role": "user", "content": prompt}]
                )

                summary = message.content[0].text
                output_path = self._save_summary(summary)

                self._ui(self.progress_bar.set, 1.0)
                self._ui(self.status_label.configure, text="✅ Summary complete!")
                self._ui(self.output_box.delete, "1.0", "end")
                self._ui(self.output_box.insert, "1.0", summary)

                self._ui(messagebox.showinfo, "DanScribe AI", f"Summary complete!\nSaved to folder:\n{output_path}\n\nBoth .txt and .docx files created.")

            except anthropic.AuthenticationError:
                logger.warning("Claude authentication failed")
                self._ui(messagebox.showerror, "Invalid API Key", "Please check your Claude API key in Settings.")
                self._ui(self.status_label.configure, text="Status: API error")
            except anthropic.RateLimitError:
                logger.warning("Claude rate limit hit")
                self._ui(messagebox.showerror, "Rate Limited", "Too many requests. Please wait a moment and try again.")
                self._ui(self.status_label.configure, text="Status: Rate limited")
            except anthropic.APIError as e:
                logger.error("Claude API error: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"Claude API error:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")
            except Exception as e:
                logger.critical("Unexpected error during summary: %s", e, exc_info=True)
                self._ui(messagebox.showerror, "Error", f"An unexpected error occurred:\n{e}")
                self._ui(self.status_label.configure, text="Status: Error")

            self._ui(self.summarize_btn.configure, state="normal")
            self._ui(self.progress_bar.set, 0)

        threading.Thread(target=run, daemon=True).start()

    # ── FILE SAVING ──────────────────────────

    def _get_output_dir(self):
        import platform
        if platform.system() == "Windows":
            base = Path.home() / "Downloads"
        else:
            # Linux/Mac: prefer ~/Documents, fall back to ~/Downloads.
            docs = Path.home() / "Documents"
            base = docs if docs.exists() else Path.home() / "Downloads"
        path = str(base / "DanScribe_Transcriptions")
        os.makedirs(path, exist_ok=True)
        return path

    def _make_docx(self, text, docx_path, doc_type="transcript", provenance=None):
        """Convert text to a formatted .docx file using python-docx.

        `provenance`, when given, is rendered as a small gray audit footer at
        the very bottom (below the transcript body) — additive metadata that
        records which engine produced the document.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title_text = "DanScribe AI — Meeting Summary" if doc_type == "summary" else "DanScribe AI — Transcript"
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
        run.font.name = "Calibri"

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%d %B %Y"))
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_run.font.name = "Calibri"

        doc.add_paragraph()  # spacer

        # Parse and add content lines
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Speaker line: "Name: text"
            if doc_type == "transcript" and ":" in stripped:
                parts = stripped.split(":", 1)
                if len(parts[0]) <= 40:
                    p = doc.add_paragraph()
                    speaker_run = p.add_run(parts[0] + ": ")
                    speaker_run.bold = True
                    speaker_run.font.name = "Calibri"
                    speaker_run.font.size = Pt(11)
                    text_run = p.add_run(parts[1].strip())
                    text_run.font.name = "Calibri"
                    text_run.font.size = Pt(11)
                    continue

            # Bold markdown headings **text**
            if stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                r = p.add_run(stripped.strip("*"))
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
                r.font.name = "Calibri"
                continue

            # Numbered headings like "1. **Main Points**"
            if stripped[:3].rstrip(". ").isdigit() and "**" in stripped:
                clean = stripped.replace("**", "").lstrip("0123456789. ")
                p = doc.add_paragraph()
                r = p.add_run(clean)
                r.bold = True
                r.font.size = Pt(13)
                r.font.color.rgb = RGBColor(0x1f, 0x53, 0x8d)
                r.font.name = "Calibri"
                continue

            # Bullet points
            if stripped.startswith("- ") or stripped.startswith("• "):
                bullet_text = stripped.lstrip("-• ").strip()
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(bullet_text)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.name = "Calibri"
            r.font.size = Pt(11)

        # Provenance footer — small gray audit block at the very bottom.
        if provenance:
            doc.add_paragraph()  # spacer
            sep = doc.add_paragraph()
            sep_run = sep.add_run("─" * 40)
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            sep_run.font.name = "Calibri"

            heading = doc.add_paragraph()
            h_run = heading.add_run("Transcription provenance")
            h_run.bold = True
            h_run.font.size = Pt(9)
            h_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h_run.font.name = "Calibri"

            for line in format_provenance_lines(provenance):
                pp = doc.add_paragraph()
                pr = pp.add_run(line)
                pr.font.size = Pt(9)
                pr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                pr.font.name = "Calibri"

        doc.save(docx_path)

    def _save_transcript(self, text, source_path, provenance=None):
        from datetime import datetime
        base = os.path.splitext(os.path.basename(source_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = self._get_output_dir()

        # The transcript body is written verbatim; the provenance footer, if
        # present, is appended AFTER it, clearly separated — additive metadata,
        # not part of the transcript content itself.
        txt_body = text
        if provenance:
            footer = "\n".join(format_provenance_lines(provenance))
            txt_body = (
                text
                + "\n\n" + ("─" * 40) + "\n"
                + "Transcription provenance\n"
                + footer + "\n"
            )

        txt_path = os.path.join(out_dir, f"{base}_transcript_{timestamp}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(txt_body)
        docx_path = os.path.join(out_dir, f"{base}_transcript_{timestamp}.docx")
        try:
            self._make_docx(text, docx_path, "transcript", provenance=provenance)
        except Exception as e:
            logger.warning("Could not create .docx: %s", e)
        return out_dir

    def _save_summary(self, text):
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = self._get_output_dir()
        txt_path = os.path.join(out_dir, f"summary_{timestamp}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        docx_path = os.path.join(out_dir, f"summary_{timestamp}.docx")
        try:
            self._make_docx(text, docx_path, "summary")
        except Exception as e:
            logger.warning("Could not create .docx: %s", e)
        return out_dir


# ─────────────────────────────────────────────
#  START APPLICATION
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import multiprocessing
    multiprocessing.freeze_support()  # Required for the Windows .exe build
    ctk.set_appearance_mode("dark")
    app = DanScribeApp()
    app.mainloop()
